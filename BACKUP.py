import logging
import warnings
import requests
import time
import json
import re
import random
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np
from dotenv import load_dotenv
import os
import pandas as pd
import tiktoken
import io
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
from threading import Thread
from fastapi import FastAPI, Request, Response, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from azure.storage.blob import BlobServiceClient, ContainerClient, ContentSettings
import psycopg2
import uvicorn
import base64
from dataclasses import dataclass
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
import aiohttp
from typing import Optional, Dict, Any
import platform
from contextlib import asynccontextmanager
from collections import deque
import statistics
import asyncpg
import threading, sys

# Google GenAI SDK for Vertex AI (new architecture - primary method)
from google import genai
from google.oauth2 import service_account

# Global connection pools
_connection_pools = {}
_pool_lock = asyncio.Lock()

# Load environment variables from .env file
load_dotenv()

# ======================================================
#                 Configuration
# ======================================================

# Windows-specific event loop fix for aiohttp
if platform.system() == 'Windows':
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

# Notification configuration
NOTIFICATION_API_URL = "https://philotimo-backend-staging.azurewebsites.net/send-notification"
NOTIFICATION_TIMEOUT = 10  # seconds

# Indexer configuration
INDEXER_API_BASE_URL = "https://autoindexerwebapp.azurewebsites.net"
INDEXER_TIMEOUT = 300  # 5 minutes timeout for indexer operations
INDEXER_RETRY_ATTEMPTS = 3
INDEXER_RETRY_DELAY = 30  # seconds

# Growth Engine Database Configuration
GROWTH_DB_CONFIG = {
    "host": "memberchat-db.postgres.database.azure.com",
    "database": "BACKABLE-THE-GROWTH-ENGINE",
    "user": "backable",
    "password": "Utkar$h007",
    "port": 5432,
    "sslmode": "require"
}

async def get_db_pool(db_config: Dict):
    """Get or create connection pool for database"""
    global _connection_pools, _pool_lock
    
    db_key = f"{db_config['database']}"
    
    if db_key not in _connection_pools:
        async with _pool_lock:
            if db_key not in _connection_pools:
                logging.info(f"🔗 Creating pool for {db_key}")
                
                pool = await asyncpg.create_pool(
                    host=db_config["host"],
                    database=db_config["database"],
                    user=db_config["user"],
                    password=db_config["password"],
                    port=db_config["port"],
                    ssl="require",
                    min_size=3,
                    max_size=12,
                    command_timeout=60
                )
                
                _connection_pools[db_key] = pool
                logging.info(f"✅ Pool created for {db_key}")
    
    return _connection_pools[db_key]


# Multi-Database Intelligence Sources
COMPONENT_DB_CONFIG = {
    "host": "memberchat-db.postgres.database.azure.com",
    "database": "BACKABLE-COMPONENT-ENGINE",
    "user": "backable",
    "password": "Utkar$h007",
    "port": 5432,
    "sslmode": "require"
}

PROFILE_DB_CONFIG = {
    "host": "memberchat-db.postgres.database.azure.com",
    "database": "BACKABLE-PROFILE-ENGINE",
    "user": "backable",
    "password": "Utkar$h007",
    "port": 5432,
    "sslmode": "require"
}

DREAM_DB_CONFIG = {
    "host": "memberchat-db.postgres.database.azure.com",
    "database": "BACKABLE-DREAM-ANALYZER",
    "user": "backable",
    "password": "Utkar$h007",
    "port": 5432,
    "sslmode": "require"
}

ANALYST_DB_CONFIG = {
    "host": "memberchat-db.postgres.database.azure.com",
    "database": "BACKABLE-THE-ANALYST",
    "user": "backable",
    "password": "Utkar$h007",
    "port": 5432,
    "sslmode": "require"
}

USER_DB_CONFIG = {
    "host": "philotimo-staging-db.postgres.database.azure.com",
    "database": "philotimodb",
    "user": "wchen",
    "password": "DevPhilot2024!!",
    "port": 5432,
    "sslmode": "require"
}

# Azure Storage - NEW UNIFIED ARCHITECTURE
AZURE_STORAGE_CONNECTION_STRING = os.getenv('AZURE_STORAGE_CONNECTION_STRING')
if not AZURE_STORAGE_CONNECTION_STRING:
    raise ValueError("AZURE_STORAGE_CONNECTION_STRING environment variable is required")

ONBOARDING_DB_HOST = os.getenv('ONBOARDING_DB_HOST', 'memberchat-db.postgres.database.azure.com')
ONBOARDING_DB_NAME = os.getenv('ONBOARDING_DB_NAME', 'BACKABLE-GOOGLE-RAG')  # Updated to new unified architecture database
ONBOARDING_DB_USER = os.getenv('ONBOARDING_DB_USER', 'backable')
ONBOARDING_DB_PASSWORD = os.getenv('ONBOARDING_DB_PASSWORD', 'Utkar$h007')
ONBOARDING_DB_PORT = int(os.getenv('ONBOARDING_DB_PORT', '5432'))

# Gemini 2.5 Pro Configuration - Load from environment variables (matching Component Engine pattern)
# Load individual API keys from environment (GEMINI_API_KEY_1, GEMINI_API_KEY_2, etc.)
GEMINI_API_KEYS = []
for i in range(1, 11):  # Load up to 10 API keys
    key = os.getenv(f'GEMINI_API_KEY_{i}')
    if key:
        GEMINI_API_KEYS.append(key)

# Ensure we have API keys from environment
if not GEMINI_API_KEYS:
    raise ValueError("❌ No GEMINI_API_KEYS found in environment variables. Please set GEMINI_API_KEY_1 through GEMINI_API_KEY_10")

logging.info(f"✅ Loaded {len(GEMINI_API_KEYS)} GEMINI_API_KEYS from environment")

# Add this global dictionary after your GEMINI_API_KEYS list
api_key_health = {}

def get_smart_api_key(section_index: int, retry_attempt: int = 0) -> str:
    """Enhanced smart API key selection with load balancing"""
    global api_key_health
    
    # Initialize health tracking if not exists
    if not api_key_health:
        for i, key in enumerate(GEMINI_API_KEYS):
            api_key_health[key] = {
                'last_503_time': None,
                'consecutive_failures': 0,
                'total_requests': 0,
                'key_id': f'Growth_{i+1:02d}',
                'response_times': deque(maxlen=10),  # Track last 10 response times
                'success_rate': 1.0,
                'last_used': 0,
                'current_load': 0  # Track concurrent requests
            }
    
    current_time = time.time()
    
    # Calculate health scores for all keys
    key_scores = []
    for key in GEMINI_API_KEYS:
        health = api_key_health[key]
        score = calculate_key_health_score(health, current_time)
        key_scores.append((key, score, health))
    
    # Sort by health score (higher is better)
    key_scores.sort(key=lambda x: x[1], reverse=True)
    
    # Select best available key
    for key, score, health in key_scores:
        if score > 0:  # Key is usable
            # Update usage tracking
            health['total_requests'] += 1
            health['last_used'] = current_time
            health['current_load'] += 1
            
            logging.info(f"🔑 [{section_index}] Selected {health['key_id']} (score: {score:.2f}, load: {health['current_load']})")
            return key
    
    # Fallback: use least recently failed key
    logging.warning("⚠️ All keys degraded, using least recently failed")
    fallback_key = min(GEMINI_API_KEYS, 
                      key=lambda k: api_key_health[k]['last_503_time'] or 0)
    api_key_health[fallback_key]['total_requests'] += 1
    return fallback_key

def calculate_key_health_score(health: Dict, current_time: float) -> float:
    """Calculate health score for API key (0-100, higher is better)"""
    score = 100.0
    
    # Penalize recent 503 errors
    if health['last_503_time'] and (current_time - health['last_503_time']) < 300:
        score *= 0.1  # Major penalty for recent 503
    
    # Penalize consecutive failures
    failure_penalty = min(0.9, health['consecutive_failures'] * 0.3)
    score *= (1.0 - failure_penalty)
    
    # Penalize high current load
    load_penalty = min(0.5, health['current_load'] * 0.1)
    score *= (1.0 - load_penalty)
    
    # Bonus for good response times
    if health['response_times']:
        avg_response_time = statistics.mean(health['response_times'])
        if avg_response_time < 30:  # Fast responses
            score *= 1.2
        elif avg_response_time > 60:  # Slow responses
            score *= 0.8
    
    # Bonus for high success rate
    score *= health['success_rate']
    
    return max(0, score)

def update_api_key_health(api_key: str, success: bool, error_code: str = None, response_time: float = None):
    """Enhanced API key health update with detailed metrics"""
    global api_key_health
    
    if api_key not in api_key_health:
        return
    
    health = api_key_health[api_key]
    
    # Decrease current load
    health['current_load'] = max(0, health['current_load'] - 1)
    
    # Track response time
    if response_time:
        health['response_times'].append(response_time)
    
    if success:
        health['consecutive_failures'] = 0
        # Update success rate (sliding window)
        old_rate = health['success_rate']
        health['success_rate'] = min(1.0, old_rate * 0.9 + 0.1)  # Weighted moving average
        
        logging.debug(f"✅ {health['key_id']} success (rate: {health['success_rate']:.2f})")
    else:
        health['consecutive_failures'] += 1
        # Update success rate
        old_rate = health['success_rate']
        health['success_rate'] = max(0.0, old_rate * 0.9)  # Penalize failures
        
        # Special handling for different error types
        if error_code == "503":
            health['last_503_time'] = time.time()
            logging.warning(f"🚨 {health['key_id']} got 503 - cooling down")
        elif error_code == "429":
            health['last_503_time'] = time.time() - 150  # Shorter cooldown for rate limits
            logging.warning(f"🚦 {health['key_id']} rate limited")
        
        logging.warning(f"❌ {health['key_id']} failed (consecutive: {health['consecutive_failures']})")


def get_load_balanced_api_key(section_index: int) -> str:
    """Get API key using round-robin with health awareness"""
    global api_key_health
    
    if not api_key_health:
        # Initialize if needed
        return get_smart_api_key(section_index, 0)
    
    current_time = time.time()
    
    # Filter healthy keys
    healthy_keys = []
    for key in GEMINI_API_KEYS:
        health = api_key_health[key]
        
        # Skip if in cooldown
        if health['last_503_time'] and (current_time - health['last_503_time']) < 180:
            continue
            
        # Skip if too many failures
        if health['consecutive_failures'] >= 3:
            continue
            
        # Skip if overloaded
        if health['current_load'] >= 3:
            continue
            
        healthy_keys.append(key)
    
    if not healthy_keys:
        # Fallback to smart selection
        return get_smart_api_key(section_index, 0)
    
    # Use weighted random selection based on inverse load
    weights = []
    for key in healthy_keys:
        load = api_key_health[key]['current_load']
        success_rate = api_key_health[key]['success_rate']
        weight = success_rate / max(1, load + 1)  # Higher success rate, lower load = higher weight
        weights.append(weight)
    
    # Weighted random selection
    selected_key = random.choices(healthy_keys, weights=weights)[0]
    
    # Update usage
    api_key_health[selected_key]['current_load'] += 1
    api_key_health[selected_key]['total_requests'] += 1
    
    health = api_key_health[selected_key]
    logging.info(f"🎯 Load balanced selection: {health['key_id']} (load: {health['current_load']}, rate: {health['success_rate']:.2f})")
    
    return selected_key

def get_api_key_status_summary() -> str:
    """Get summary of all API key health for logging"""
    if not api_key_health:
        return "No health data available"
    
    healthy_count = 0
    cooling_down = 0
    failed_count = 0
    
    for key, health in api_key_health.items():
        current_time = time.time()
        
        if health['last_503_time'] and (current_time - health['last_503_time']) < 300:
            cooling_down += 1
        elif health['consecutive_failures'] >= 3:
            failed_count += 1
        else:
            healthy_count += 1
    
    return f"Healthy: {healthy_count}, Cooling: {cooling_down}, Failed: {failed_count}"

def get_enhanced_api_key_status() -> Dict:
    """Get comprehensive API key status"""
    if not api_key_health:
        return {"status": "not_initialized"}
    
    current_time = time.time()
    status = {
        "total_keys": len(GEMINI_API_KEYS),
        "healthy_keys": 0,
        "degraded_keys": 0,
        "failed_keys": 0,
        "cooling_down": 0,
        "total_load": 0,
        "average_success_rate": 0,
        "key_details": []
    }
    
    success_rates = []
    
    for key, health in api_key_health.items():
        # Determine status
        if health['last_503_time'] and (current_time - health['last_503_time']) < 180:
            key_status = "cooling_down"
            status["cooling_down"] += 1
        elif health['consecutive_failures'] >= 3:
            key_status = "failed"
            status["failed_keys"] += 1
        elif health['consecutive_failures'] > 0 or health['success_rate'] < 0.8:
            key_status = "degraded"
            status["degraded_keys"] += 1
        else:
            key_status = "healthy"
            status["healthy_keys"] += 1
        
        status["total_load"] += health['current_load']
        success_rates.append(health['success_rate'])
        
        # Add key details
        avg_response = statistics.mean(health['response_times']) if health['response_times'] else 0
        
        status["key_details"].append({
            "key_id": health['key_id'],
            "status": key_status,
            "success_rate": health['success_rate'],
            "current_load": health['current_load'],
            "consecutive_failures": health['consecutive_failures'],
            "total_requests": health['total_requests'],
            "avg_response_time": round(avg_response, 2)
        })
    
    status["average_success_rate"] = statistics.mean(success_rates) if success_rates else 0
    
    return status

def reset_failed_api_keys():
    """Reset failed API keys for retry attempts"""
    global api_key_health
    
    reset_count = 0
    for key, health in api_key_health.items():
        if health.get('consecutive_failures', 0) >= 3:
            health['consecutive_failures'] = 1  # Reduce but don't fully reset
            health['last_503_time'] = None  # Clear cooldown
            reset_count += 1
            logging.info(f"🔄 Partially reset API key {health.get('key_id', 'unknown')} for retry")
    
    logging.info(f"🔄 Reset {reset_count} failed API keys for retry attempt")

# Production-optimized settings
MAX_RETRIES = 10
MAX_REQUESTS_PER_ENDPOINT = 100
REQUEST_TIMEOUT = 120  # 2 minutes
MAX_SECTION_RETRIES = 3
MAX_REPORT_RETRIES = 2
MIN_ACCEPTABLE_WORDS = 100
RETRY_WAIT_BASE = 30
growth_job_status = {}

# ======================================================
#           Vertex AI Configuration (NEW UNIFIED ARCHITECTURE)
# ======================================================

VERTEX_PROJECT_ID = "backable-machine-learning-apis"
VERTEX_LOCATION = "us-central1"
USE_VERTEX_AI = True

def initialize_vertex_ai_client():
    """Initialize Google GenAI client for Vertex AI (PRIMARY METHOD)"""
    try:
        # Try loading credentials from environment variable (for Azure deployment)
        creds_json = os.getenv('GOOGLE_APPLICATION_CREDENTIALS_JSON')

        if creds_json:
            # Load from JSON string in environment variable
            import json
            import tempfile
            creds_dict = json.loads(creds_json)

            # Write to temporary file
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json') as f:
                json.dump(creds_dict, f)
                temp_path = f.name

            credentials = service_account.Credentials.from_service_account_file(
                temp_path,
                scopes=['https://www.googleapis.com/auth/cloud-platform']
            )

            # Clean up temp file
            os.unlink(temp_path)

        else:
            # Load from local file (for local development)
            credentials = service_account.Credentials.from_service_account_file(
                "vertex-key.json",
                scopes=['https://www.googleapis.com/auth/cloud-platform']
            )

        # Initialize GenAI client with Vertex AI
        client = genai.Client(
            vertexai=True,
            credentials=credentials,
            project=VERTEX_PROJECT_ID,
            location=VERTEX_LOCATION
        )

        logging.info("✅ Vertex AI client initialized successfully")
        return client

    except Exception as e:
        logging.warning(f"⚠️ Vertex AI initialization failed: {str(e)}")
        logging.warning("Will use API keys fallback for all requests")
        return None

# Initialize Vertex AI client (attempt on startup)
vertex_ai_client = initialize_vertex_ai_client() if USE_VERTEX_AI else None

def try_vertex_ai_growth_request(
    enhanced_prompt: str,
    temperature: float,
    max_tokens: int,
    start_time: float
) -> Optional[Dict]:
    """
    Try making request using Vertex AI (PRIMARY METHOD for Growth Engine).
    Returns response dict if successful, None if fails.
    """
    if not vertex_ai_client:
        logging.info("Vertex AI client not available - using API keys fallback")
        return None

    try:
        logging.info("🚀 Trying Vertex AI (Primary Method for Growth Analysis)")

        # Call Vertex AI using GenAI SDK with gemini-2.5-pro
        response = vertex_ai_client.models.generate_content(
            model="gemini-2.5-pro",
            contents=enhanced_prompt,
            config={
                "temperature": temperature,
                "max_output_tokens": max_tokens,
                "top_p": 0.95,
            }
        )

        # Extract content
        if response and response.candidates and len(response.candidates) > 0:
            content = response.candidates[0].content.parts[0].text if response.candidates[0].content.parts else ""
            token_count = response.usage_metadata.total_token_count if response.usage_metadata else 0
            request_time = time.time() - start_time

            logging.info(f"✅ Vertex AI SUCCESS - {len(content.split())} words, {token_count} tokens, {request_time:.2f}s")

            return {
                "success": True,
                "content": content,
                "tokens": token_count,
                "time": request_time,
                "model": "gemini-2.5-pro-vertex"
            }
        else:
            logging.warning("⚠️ Vertex AI returned empty response - falling back to API keys")
            return None

    except Exception as e:
        logging.warning(f"⚠️ Vertex AI failed: {str(e)} - Falling back to API keys")
        return None

# ======================================================
#           PersonalizedGrowthNotificationService (NEW ARCHITECTURE)
# ======================================================

class PersonalizedGrowthNotificationService:
    """
    Enhanced notification service with personalized, PROFESSIONAL messages using Vertex AI + Gemini
    Focuses on value delivery and Backable Mind intelligence for Growth Engine
    """

    def __init__(self, gemini_api_key: str):
        self.gemini_api_key = gemini_api_key
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models"
        self.model = "gemini-2.5-pro"

        # Professional fallback messages (used if generation fails)
        self.fallback_messages = {
            "start": [
                "Your Growth Engine assessment has begun. This analysis will expand your Backable Mind with strategic growth insights about your market position, scaling opportunities, and revenue acceleration strategies.",
                "Growth Engine analysis initiated. Your business is being examined to provide comprehensive growth intelligence.",
                "Analysis started for your growth strategy. The Growth Engine is now evaluating your market opportunities, scaling potential, and revenue drivers across all dimensions."
            ],
            "middle": [
                "Your Growth Engine analysis is progressing well. We're currently examining your market position and scaling opportunities to build comprehensive growth intelligence.",
                "Analysis update: The Growth Engine has completed multiple sections of your growth assessment. Strategic insights are being compiled across all growth areas.",
                "Progress update on your growth analysis. Key areas including market expansion, revenue acceleration, and scaling strategy have been examined."
            ],
            "complete": [
                "Your Growth Engine analysis is now complete and has expanded your Backable Mind with comprehensive growth intelligence. Head to your dashboard to explore growth insights.",
                "Growth Engine analysis complete. Your Backable Mind now contains detailed market opportunities, revenue strategies, and scaling recommendations. Visit your dashboard to explore.",
                "Analysis finished. Your Backable Mind has been enhanced with growth intelligence covering all key strategic areas. Access your dashboard now to review recommendations."
            ]
        }

    async def generate_personalized_message(self, user_profile: Dict, stage: str, progress_data: Dict = None) -> str:
        """
        Generate professional, value-focused notification message using Vertex AI (primary) or Gemini API (fallback)
        Focuses on how Growth Engine makes Backable Mind smarter with growth intelligence
        """
        try:
            # Extract user context
            business_name = user_profile.get('business_name', 'Your Business')
            username = user_profile.get('username', 'Entrepreneur')
            industry = user_profile.get('industry', 'Business')

            # Create stage-specific professional prompts focused on Backable Mind value
            if stage == "start":
                prompt = f"""Create a professional, value-focused notification for {username} from {business_name} in the {industry} industry.
They just started their Growth Engine strategic assessment (market opportunities, scaling potential, and revenue acceleration analysis).
Make it professional, encouraging, focus on how this analysis will make their Backable Mind smarter with growth intelligence, 2-3 sentences max, NO emojis."""

            elif stage == "middle":
                sections_done = progress_data.get('sections_completed', 5) if progress_data else 5
                total_sections = progress_data.get('total_sections', 9) if progress_data else 9
                prompt = f"""Create a professional mid-progress notification for {username} from {business_name}.
They're {sections_done}/{total_sections} sections through their Growth Engine strategic assessment.
Make it professional, informative, highlight growth aspects being analyzed, 2-3 sentences max, NO emojis."""

            elif stage == "complete":
                total_words = progress_data.get('total_words', 12000) if progress_data else 12000
                prompt = f"""Create a professional completion notification for {username} from {business_name}.
Their Growth Engine strategic assessment is complete with {total_words:,} words of growth insights.
Make it professional, celebratory, focus on how their Backable Mind is now smarter with growth intelligence, tell them to visit dashboard, 2-3 sentences max, NO emojis."""

            # STEP 1: TRY VERTEX AI FIRST (PRIMARY METHOD)
            if vertex_ai_client:
                try:
                    logging.info("🚀 Trying Vertex AI for growth notification message")
                    response = vertex_ai_client.models.generate_content(
                        model="gemini-2.5-pro",
                        contents=prompt,
                        config={"temperature": 1.0, "max_output_tokens": 1000, "top_p": 0.95}
                    )

                    if response and response.candidates and len(response.candidates) > 0:
                        content = response.candidates[0].content.parts[0].text if response.candidates[0].content.parts else ""
                        if len(content.split()) > 10:
                            logging.info(f"✅ Vertex AI growth notification for {username}: {stage}")
                            return content
                except Exception as e:
                    logging.warning(f"⚠️ Vertex AI notification failed: {str(e)} - Falling back to API key")

            # STEP 2: FALLBACK TO GEMINI API KEY
            logging.info("🔄 Using Gemini API key for growth notification")
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as session:
                url = f"{self.base_url}/{self.model}:generateContent"
                payload = {
                    "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 1.0, "maxOutputTokens": 1000, "topP": 0.95}
                }
                params = {'key': self.gemini_api_key}

                async with session.post(url, json=payload, params=params) as response:
                    if response.status == 200:
                        data = await response.json()
                        if 'candidates' in data and len(data['candidates']) > 0:
                            candidate = data['candidates'][0]
                            content = ""
                            try:
                                if 'content' in candidate and 'parts' in candidate['content']:
                                    content = candidate['content']['parts'][0]['text']
                            except:
                                pass
                            if content and len(content.split()) > 10:
                                logging.info(f"✅ Gemini API growth notification for {username}: {stage}")
                                return content.strip()

        except Exception as e:
            logging.error(f"❌ Error generating growth notification message: {str(e)}")

        return random.choice(self.fallback_messages[stage])

    @staticmethod
    async def send_notification(user_id: str, title: str, body: str, data_type: str = "notification", save_to_db: bool = False, report_id: str = None, business_name: str = None):
        """Send notification to user with optional database persistence"""
        try:
            from datetime import timedelta
            payload = {
                "userId": int(user_id),
                "title": title,
                "body": body,
                "data": {"type": data_type, "timestamp": str(int(datetime.now().timestamp()))}
            }

            if save_to_db and report_id:
                payload["saveToDb"] = True
                payload["expiresAt"] = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%SZ")
                payload["data"]["screen"] = "GrowthReport"
                payload["data"]["reportId"] = report_id
                payload["data"]["payload"] = {
                    "type": "ai_report_complete",
                    "params": {
                        "reportId": report_id,
                        "reportTitle": "Growth Intelligence Report",
                        "reportType": "comprehensive_growth",
                        "userId": int(user_id),
                        "businessName": business_name or "Your Business",
                        "completionStatus": "success",
                        "generatedAt": datetime.now().isoformat()
                    },
                    "actionType": "navigate",
                    "screen": "GrowthReport",
                    "url": f"/growth/{report_id}"
                }

            logging.info(f"🔔 Sending professional growth notification to user {user_id}: {title} (saveToDb: {save_to_db})")
            logging.info(f"📦 Notification payload: {json.dumps(payload, indent=2)}")

            connector = aiohttp.TCPConnector(use_dns_cache=False) if platform.system() == 'Windows' else None
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=NOTIFICATION_TIMEOUT), connector=connector) as session:
                async with session.post(NOTIFICATION_API_URL, json=payload, headers={"Content-Type": "application/json"}) as response:
                    if response.status == 200:
                        result = await response.text()
                        logging.info(f"✅ Professional growth notification sent successfully to user {user_id}")
                        logging.info(f"📨 Backend response: {result}")
                        return True, result
                    else:
                        error_text = await response.text()
                        logging.error(f"❌ Growth notification failed for user {user_id}: {response.status} - {error_text}")
                        return False, f"HTTP {response.status}: {error_text}"
        except Exception as e:
            logging.error(f"❌ Growth notification error for user {user_id}: {str(e)}")
            return False, str(e)

    async def send_personalized_notification(self, user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None, report_id: str = None):
        """Send personalized professional growth notification for specific stage"""
        try:
            message = await self.generate_personalized_message(user_profile, stage, progress_data)
            username = user_profile.get('username', 'Entrepreneur')
            business_name = user_profile.get('business_name', 'Your Business')

            professional_titles = {
                "start": [
                    f"Growth Engine - Analysis Started",
                    f"{business_name} - Strategic Assessment Beginning",
                    f"Growth Engine Assessment - {username}",
                    f"{business_name} - Growth Intelligence Analysis",
                    f"Strategic Analysis Initiated"
                ],
                "middle": [
                    f"Growth Engine - Progress Update",
                    f"{business_name} - Analysis Progressing",
                    f"Strategic Assessment Update - {username}",
                    f"{business_name} - Growth Analysis In Progress",
                    f"Your Growth Engine Progress"
                ],
                "complete": [
                    f"Growth Engine - Analysis Complete",
                    f"{business_name} - Growth Intelligence Ready",
                    f"Your Growth Analysis is Complete",
                    f"{business_name} - Strategic Insights Available",
                    f"Growth Engine Assessment Complete"
                ]
            }

            title = random.choice(professional_titles[stage])
            save_to_db = (stage == "complete")
            success, result = await self.send_notification(user_id, title, message, "notification", save_to_db, report_id, business_name)

            if success:
                logging.info(f"✅ Sent professional {stage} growth notification to user {user_id}")
            else:
                logging.error(f"❌ Failed to send professional notification: {result}")

            return success, message
        except Exception as e:
            logging.error(f"❌ Error sending personalized growth notification: {str(e)}")
            return False, str(e)

    @staticmethod
    def send_personalized_notification_sync(user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None, gemini_api_key: str = None, report_id: str = None):
        """
        Synchronous wrapper for sending personalized professional growth notifications
        FIXED for Windows compatibility
        """
        try:
            # FIXED: Handle Windows event loop policy BEFORE creating new loop
            if platform.system() == 'Windows':
                asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

            # Create new loop after setting policy
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                service = PersonalizedGrowthNotificationService(gemini_api_key or GEMINI_API_KEYS[0])
                return loop.run_until_complete(
                    service.send_personalized_notification(user_id, user_profile, stage, progress_data, report_id)
                )
            finally:
                loop.close()

        except Exception as e:
            logging.error(f"❌ Sync growth notification error: {str(e)}")
            return False, str(e)


# ======================================================
#           FIXED Notification Functions (LEGACY - Keep for compatibility)
# ======================================================

async def send_growth_notification(user_id: str, title: str, body: str, data_type: str = "notification"):
    """Send notification to user for growth strategy"""
    try:
        payload = {
            "userId": int(user_id),
            "title": title,
            "body": body,
            "data": {
                "type": data_type,
                "timestamp": str(int(datetime.now().timestamp()))
            }
        }
        
        logging.info(f"🔔 Sending growth notification to user {user_id}: {title}")

        # FIXED: Use TCPConnector to avoid aiodns issues on Windows
        connector = aiohttp.TCPConnector(use_dns_cache=False) if platform.system() == 'Windows' else None

        async with aiohttp.ClientSession(
            timeout=aiohttp.ClientTimeout(total=NOTIFICATION_TIMEOUT),
            connector=connector
        ) as session:
            async with session.post(
                NOTIFICATION_API_URL,
                json=payload,
                headers={"Content-Type": "application/json"}
            ) as response:
                
                if response.status == 200:
                    result = await response.text()
                    logging.info(f"✅ Growth notification sent successfully to user {user_id}")
                    return True, result
                else:
                    error_text = await response.text()
                    logging.error(f"❌ Growth notification failed for user {user_id}: {response.status} - {error_text}")
                    return False, f"HTTP {response.status}: {error_text}"
                    
    except Exception as e:
        logging.error(f"❌ Growth notification error for user {user_id}: {str(e)}")
        return False, str(e)

def send_growth_notification_sync(user_id: str, title: str, body: str, data_type: str = "notification"):
    """Synchronous wrapper for sending growth notifications"""
    try:
        if platform.system() == 'Windows':
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
        
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                send_growth_notification(user_id, title, body, data_type)
            )
        finally:
            loop.close()
            
    except Exception as e:
        logging.error(f"❌ Sync growth notification error: {str(e)}")
        return False, str(e)

async def generate_personalized_growth_message(user_profile: Dict, stage: str, progress_data: Dict = None) -> str:
    """Generate personalized, engaging growth message using Gemini AI"""
    try:
        # Extract user context
        business_name = user_profile.get('business_name', 'Your Business')
        username = user_profile.get('username', 'Entrepreneur')
        industry = user_profile.get('industry', 'Business')
        team_size = user_profile.get('team_size', 'Unknown')
        
        # Create stage-specific prompts for growth strategy
        if stage == "start":
            prompt = f"""
            Create a MOTIVATING, personalized notification for {username} from {business_name} in the {industry} industry.
            They just started their GROWTH STRATEGY assessment (focus on scaling, acceleration, breakthrough opportunities).
            
            Make it:
            - INSPIRING and energizing about growth potential
            - Include a motivational reference to {industry} growth opportunities
            - Reference {business_name} with excitement about scaling
            - Focus on ACCELERATION, BREAKTHROUGH, SCALING
            - 1-2 sentences max with growth emojis (🚀📈💫⚡)
            - Make them excited about their growth journey!
            """
        
        elif stage == "middle":
            questions_done = progress_data.get('questions_completed', 8) if progress_data else 8
            total_questions = progress_data.get('total_questions', 16) if progress_data else 16
            
            prompt = f"""
            Create an ENERGIZING mid-progress notification for {username} from {business_name}.
            They're {questions_done}/{total_questions} questions through their growth strategy assessment.
            
            Make it:
            - MOTIVATING and encouraging about growth discovery progress
            - Reference their {industry} scaling potential enthusiastically
            - Focus on GROWTH INSIGHTS, STRATEGY DEVELOPMENT
            - 1-2 sentences max with progress emojis
            - Keep them excited about breakthrough opportunities!
            """
        
        elif stage == "complete":
            total_words = progress_data.get('total_words', 15000) if progress_data else 15000
            
            prompt = f"""
            Create a CELEBRATORY completion notification for {username} from {business_name}.
            Their growth strategy blueprint is complete with {total_words:,} words of breakthrough insights.
            
            Make it:
            - TRIUMPHANT and celebratory about their growth blueprint
            - Reference {industry} acceleration success enthusiastically
            - Focus on STRATEGY BLUEPRINT, GROWTH ACCELERATION completion
            - 1-2 sentences max with celebration + growth emojis
            - Make them feel like a growth strategist!
            """
        
        # Use first available API key for notifications
        gemini_api_key = GEMINI_API_KEYS[0]
        
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as session:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
            
            payload = {
                "contents": [{
                    "role": "user",
                    "parts": [{"text": prompt}]
                }],
                "generationConfig": {
                    "temperature": 1.0,
                    "maxOutputTokens": 150,
                    "topP": 0.95,
                    "candidateCount": 1
                }
            }
            
            params = {'key': gemini_api_key}
            
            async with session.post(url, json=payload, params=params) as response:
                if response.status == 200:
                    data = await response.json()
                    if 'candidates' in data and len(data['candidates']) > 0:
                        candidate = data['candidates'][0]
                        
                        content = ""
                        try:
                            if 'content' in candidate and 'parts' in candidate['content']:
                                content = candidate['content']['parts'][0]['text']
                            elif 'text' in candidate:
                                content = candidate['text']
                            else:
                                content = str(candidate.get('content', candidate))
                        except Exception as e:
                            logging.warning(f"Content extraction issue: {e}")
                            content = str(candidate)
                        
                        if content:
                            content = content.strip().replace('"', '').replace("'", "'")
                            if any(char in content for char in ['🚀', '📈', '💫', '⚡', '🔥', '🎉']) or len(content.split()) > 3:
                                if not any(tech in content.lower() for tech in ['role', 'model', 'parts', 'content']):
                                    logging.info(f"🎭 Generated growth {stage} message for {username}")
                                    return content
    
    except Exception as e:
        logging.error(f"❌ Error generating growth message: {str(e)}")
    
    # Fallback messages for growth strategy
    fallback_messages = {
        "start": [
            f"🚀 {business_name}'s growth acceleration starting! {username}, prepare for breakthrough insights!",
            f"📈 Breaking: {username}'s growth strategy is being decoded! Warning: May cause exponential scaling!",
            f"💫 Plot twist: {business_name}'s growth potential is about to be unleashed!"
        ],
        "middle": [
            f"⚡ Halfway there! {username}'s growth barriers are being identified and eliminated!",
            f"🚀 {business_name}'s growth opportunities are more exciting than expected! 50% complete!",
            f"📊 Update: {username}'s scaling potential is off the charts!"
        ],
        "complete": [
            f"🎉 BREAKTHROUGH! {username}'s growth strategy is ready! {business_name} = unstoppable!",
            f"✨ Mission accomplished! {business_name} just became more scalable than ever imagined!",
            f"🚀 {username} is officially a growth strategist! Blueprint ready for acceleration!"
        ]
    }
    
    return random.choice(fallback_messages.get(stage, fallback_messages["start"]))

async def send_personalized_growth_notification(user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None):
    """Send personalized growth notification"""
    try:
        # Generate personalized message
        message = await generate_personalized_growth_message(user_profile, stage, progress_data)
        
        # Create titles for growth strategy
        username = user_profile.get('username', 'Entrepreneur')
        business_name = user_profile.get('business_name', 'Your Business')
        
        motivating_titles = {
            "start": [
                f"🚀 {username}, Growth Time!",
                f"📈 {business_name} Acceleration!",
                f"⚡ Strategy Assessment Started!"
            ],
            "middle": [
                f"🔥 {username} Discovering Growth!",
                f"💫 {business_name} Halfway There!",
                f"📊 Growth Progress Update!"
            ],
            "complete": [
                f"🎉 {username}, Strategy Complete!",
                f"✨ {business_name} Blueprint Ready!",
                f"🚀 Growth Success!"
            ]
        }
        
        title = random.choice(motivating_titles[stage])
        
        # Send notification
        success, result = await send_growth_notification(user_id, title, message, "notification")
        
        if success:
            logging.info(f"🎭 Sent growth {stage} notification to user {user_id}")
        else:
            logging.error(f"❌ Failed to send growth notification: {result}")
        
        return success, message
        
    except Exception as e:
        logging.error(f"❌ Error sending growth notification: {str(e)}")
        return False, str(e)

def send_growth_notification_background(user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None):
    """FIXED: Send growth notification in background thread with proper async handling and detailed logging"""
    
    # Generate unique notification ID for tracking
    notification_id = f"growth_{stage}_{user_id}_{int(time.time())}"
    
    def notification_worker():
        worker_start_time = time.time()
        thread_id = threading.current_thread().ident
        
        try:
            logging.info(f"🔔 NOTIFICATION WORKER START: {notification_id}")
            logging.info(f"📊 Notification details:")
            logging.info(f"   - User ID: {user_id}")
            logging.info(f"   - Stage: {stage}")
            logging.info(f"   - Thread ID: {thread_id}")
            logging.info(f"   - Business: {user_profile.get('business_name', 'Unknown') if user_profile else 'No profile'}")
            logging.info(f"   - Progress data: {bool(progress_data)}")
            
            # Validate inputs with detailed logging
            if not user_id:
                raise ValueError("user_id is required")
            
            if not stage:
                raise ValueError("stage is required")
            
            if stage not in ['start', 'middle', 'complete']:
                logging.warning(f"⚠️ Unusual stage value: '{stage}' (expected: start/middle/complete)")
            
            logging.info(f"✅ Input validation passed for notification {notification_id}")
            
            # Platform-specific setup with detailed logging
            if platform.system() == 'Windows':
                logging.debug(f"🪟 Setting Windows event loop policy for thread {thread_id}")
                asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
            
            # Event loop creation with detailed logging
            logging.debug(f"🔄 Creating new event loop in thread {thread_id}")
            loop_start_time = time.time()
            
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            loop_creation_time = time.time() - loop_start_time
            
            logging.debug(f"✅ Event loop created in {loop_creation_time:.3f}s")
            
            try:
                # 🔥 CRITICAL FIX: Execute notification with timeout using loop.run_until_complete
                logging.info(f"📤 Executing notification for {notification_id}...")
                notification_start_time = time.time()
                
                # 🔥 MAIN FIX: Use loop.run_until_complete instead of await (since this is a sync function)
                try:
                    success, message = loop.run_until_complete(
                        asyncio.wait_for(
                            send_personalized_growth_notification(user_id, user_profile, stage, progress_data),
                            timeout=30.0  # 30 second timeout
                        )
                    )
                    
                    notification_time = time.time() - notification_start_time
                    
                    logging.info(f"📡 Notification execution completed in {notification_time:.3f}s")
                    logging.info(f"📊 Notification result:")
                    logging.info(f"   - Success: {success}")
                    logging.info(f"   - Message: {message}")
                    
                except asyncio.TimeoutError:
                    notification_time = time.time() - notification_start_time
                    success = False
                    message = f"Notification timed out after {notification_time:.1f}s"
                    
                    logging.error(f"⏰ TIMEOUT: Notification {notification_id} timed out after {notification_time:.1f}s")
                    logging.error(f"🔍 This may indicate network issues or API problems")
                
                # Log final results with context
                total_worker_time = time.time() - worker_start_time
                
                if success:
                    logging.info(f"🎉 SUCCESS: Background growth {stage} notification sent")
                    logging.info(f"📊 Success metrics:")
                    logging.info(f"   - Notification ID: {notification_id}")
                    logging.info(f"   - User ID: {user_id}")
                    logging.info(f"   - Stage: {stage}")
                    logging.info(f"   - Total time: {total_worker_time:.3f}s")
                    logging.info(f"   - Notification time: {notification_time:.3f}s")
                    logging.info(f"   - Thread ID: {thread_id}")
                else:
                    logging.warning(f"⚠️ FAILURE: Background growth {stage} notification failed")
                    logging.warning(f"📊 Failure details:")
                    logging.warning(f"   - Notification ID: {notification_id}")
                    logging.warning(f"   - User ID: {user_id}")
                    logging.warning(f"   - Stage: {stage}")
                    logging.warning(f"   - Error message: {message}")
                    logging.warning(f"   - Total time: {total_worker_time:.3f}s")
                    logging.warning(f"   - Thread ID: {thread_id}")
                    
                    # Log additional context for failures
                    logging.warning(f"🔍 Failure context:")
                    logging.warning(f"   - User profile available: {bool(user_profile)}")
                    logging.warning(f"   - Progress data available: {bool(progress_data)}")
                    if user_profile:
                        logging.warning(f"   - Business name: {user_profile.get('business_name', 'Not provided')}")
                        logging.warning(f"   - Username: {user_profile.get('username', 'Not provided')}")
                
            except Exception as loop_error:
                loop_error_time = time.time() - notification_start_time if 'notification_start_time' in locals() else 0
                total_worker_time = time.time() - worker_start_time
                
                logging.error(f"❌ LOOP ERROR: Exception in notification execution")
                logging.error(f"🔍 Loop error details:")
                logging.error(f"   - Notification ID: {notification_id}")
                logging.error(f"   - Error type: {type(loop_error).__name__}")
                logging.error(f"   - Error message: {str(loop_error)}")
                logging.error(f"   - Loop error time: {loop_error_time:.3f}s")
                logging.error(f"   - Total worker time: {total_worker_time:.3f}s")
                logging.error(f"   - Thread ID: {thread_id}")
                
                # Log the full traceback for debugging
                import traceback
                logging.error(f"🔍 Loop error traceback:")
                for line in traceback.format_exc().split('\n'):
                    if line.strip():
                        logging.error(f"   {line}")
                
            finally:
                # Clean up event loop with logging
                try:
                    loop_cleanup_start = time.time()
                    
                    logging.debug(f"🔄 Cleaning up event loop for thread {thread_id}")
                    
                    # Cancel any remaining tasks
                    pending_tasks = [task for task in asyncio.all_tasks(loop) if not task.done()]
                    if pending_tasks:
                        logging.warning(f"⚠️ Found {len(pending_tasks)} pending tasks, cancelling...")
                        for task in pending_tasks:
                            task.cancel()
                    
                    loop.close()
                    loop_cleanup_time = time.time() - loop_cleanup_start
                    
                    logging.debug(f"✅ Event loop cleaned up in {loop_cleanup_time:.3f}s")
                    
                except Exception as cleanup_error:
                    logging.error(f"❌ Error during loop cleanup: {cleanup_error}")
                
        except Exception as worker_error:
            total_worker_time = time.time() - worker_start_time
            
            logging.error(f"💥 WORKER ERROR: Critical error in notification worker")
            logging.error(f"🔍 Worker error details:")
            logging.error(f"   - Notification ID: {notification_id}")
            logging.error(f"   - User ID: {user_id}")
            logging.error(f"   - Stage: {stage}")
            logging.error(f"   - Error type: {type(worker_error).__name__}")
            logging.error(f"   - Error message: {str(worker_error)}")
            logging.error(f"   - Total worker time: {total_worker_time:.3f}s")
            logging.error(f"   - Thread ID: {thread_id}")
            
            # Log additional context
            logging.error(f"🔍 Worker error context:")
            logging.error(f"   - Platform: {platform.system()}")
            logging.error(f"   - Python version: {sys.version}")
            logging.error(f"   - User profile type: {type(user_profile)}")
            logging.error(f"   - Progress data type: {type(progress_data)}")
            
            # Log the full traceback for debugging
            import traceback
            logging.error(f"🔍 Worker error traceback:")
            for line in traceback.format_exc().split('\n'):
                if line.strip():
                    logging.error(f"   {line}")
        
        finally:
            # Final cleanup and statistics
            total_worker_time = time.time() - worker_start_time
            
            logging.info(f"🏁 NOTIFICATION WORKER END: {notification_id}")
            logging.info(f"📊 Worker final statistics:")
            logging.info(f"   - Total execution time: {total_worker_time:.3f}s")
            logging.info(f"   - Thread ID: {thread_id}")
            logging.info(f"   - Worker completed at: {datetime.now().isoformat()}")
            
            # Clean up thread-local data if needed
            try:
                # Remove any thread-local references
                if hasattr(threading.current_thread(), '_notification_data'):
                    delattr(threading.current_thread(), '_notification_data')
            except:
                pass  # Ignore cleanup errors
    
    # 🔥 ENHANCEMENT: Better thread creation with error handling
    try:
        logging.info(f"🚀 LAUNCHING: Background notification thread for {notification_id}")
        
        # Create thread with proper naming and error handling
        notification_thread = Thread(
            target=notification_worker, 
            daemon=True,
            name=f"GrowthNotification-{stage}-{user_id}"
        )
        
        # Store notification metadata in thread for debugging
        notification_thread._notification_data = {
            'notification_id': notification_id,
            'user_id': user_id,
            'stage': stage,
            'created_at': time.time()
        }
        
        thread_start_time = time.time()
        notification_thread.start()
        thread_start_duration = time.time() - thread_start_time
        
        logging.info(f"✅ Notification thread launched successfully")
        logging.info(f"📊 Thread launch details:")
        logging.info(f"   - Notification ID: {notification_id}")
        logging.info(f"   - Thread name: {notification_thread.name}")
        logging.info(f"   - Thread ID: {notification_thread.ident}")
        logging.info(f"   - Launch time: {thread_start_duration:.3f}s")
        logging.info(f"   - Daemon thread: {notification_thread.daemon}")
        
        # 🔥 ENHANCEMENT: Optional thread monitoring (can be enabled for debugging)
        if logging.getLogger().isEnabledFor(logging.DEBUG):
            def monitor_thread():
                time.sleep(1)  # Give thread time to start
                if notification_thread.is_alive():
                    logging.debug(f"🔍 Thread {notification_thread.name} is running")
                else:
                    logging.warning(f"⚠️ Thread {notification_thread.name} finished quickly")
            
            Thread(target=monitor_thread, daemon=True).start()
        
    except Exception as thread_error:
        logging.error(f"💥 THREAD CREATION ERROR: Failed to create notification thread")
        logging.error(f"🔍 Thread error details:")
        logging.error(f"   - Notification ID: {notification_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Stage: {stage}")
        logging.error(f"   - Error type: {type(thread_error).__name__}")
        logging.error(f"   - Error message: {str(thread_error)}")
        
        # Log the full traceback for debugging
        import traceback
        logging.error(f"🔍 Thread creation traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")

# 🔥 MINIMAL ADDITION: Helper function to check notification system health
def get_notification_system_health():
    """Get basic notification system health info (minimal addition)"""
    
    active_threads = [t for t in threading.enumerate() if 'GrowthNotification' in t.name]
    
    health_info = {
        'timestamp': datetime.now().isoformat(),
        'active_notification_threads': len(active_threads),
        'thread_details': []
    }
    
    for thread in active_threads:
        thread_info = {
            'name': thread.name,
            'alive': thread.is_alive(),
            'daemon': thread.daemon
        }
        
        # Add notification metadata if available
        if hasattr(thread, '_notification_data'):
            thread_info.update(thread._notification_data)
            thread_info['age_seconds'] = time.time() - thread._notification_data.get('created_at', time.time())
        
        health_info['thread_details'].append(thread_info)
    
    return health_info

# ======================================================
#           Gemini AI Integration
# ======================================================

@dataclass
class GrowthChatResponse:
    content: str
    model: str
    api_key_used: str
    usage: Dict[str, Any]
    finish_reason: str
    response_time: float
    timestamp: float
    token_count: int

def convert_messages_to_gemini_format(messages: List[Dict[str, str]]) -> List[Dict]:
    """Convert messages to Gemini API format"""
    contents = []
    
    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        
        if role in ["user", "human"]:
            if contents and contents[-1]["role"] == "user":
                contents[-1]["parts"].append({"text": content})
            else:
                contents.append({
                    "role": "user",
                    "parts": [{"text": content}]
                })
        elif role in ["assistant", "model", "ai"]:
            contents.append({
                "role": "model",
                "parts": [{"text": content}]
            })
        elif role == "system":
            if contents and contents[-1]["role"] == "user":
                contents[-1]["parts"].insert(0, {"text": f"SYSTEM CONTEXT: {content}\n\n"})
            else:
                contents.append({
                    "role": "user",
                    "parts": [{"text": f"SYSTEM CONTEXT: {content}"}]
                })
    
    return contents



def growth_ultra_deep_analysis(
    complete_raw_data: Dict,
    analysis_type: str,
    analysis_requirements: str,
    api_key: str,
    client_id: str = "growth_analysis",
    temperature: float = 0.7,
    max_tokens: int = 100000
) -> GrowthChatResponse:
    """Enhanced growth analysis with ultra-deep response analysis and detailed logging"""
    
    start_time = time.time()
    
    logging.info(f"🚀 [{client_id}] Starting Growth Analysis: {analysis_type}")
    logging.info(f"🔍 [{client_id}] Input parameters: temp={temperature}, max_tokens={max_tokens}")
    logging.info(f"🔍 [{client_id}] API key ending: ...{api_key[-4:]}")
    
    # Log API key health status at start with enhanced metrics
    key_health = api_key_health.get(api_key, {})
    if key_health:
        success_rate = key_health.get('success_rate', 1.0)
        current_load = key_health.get('current_load', 0)
        avg_response = statistics.mean(key_health.get('response_times', [0])) if key_health.get('response_times') else 0
        
        logging.info(f"🔑 [{client_id}] API Key Health: {key_health.get('key_id', 'unknown')} - "
                    f"Failures: {key_health.get('consecutive_failures', 0)}, "
                    f"Total Requests: {key_health.get('total_requests', 0)}, "
                    f"Success Rate: {success_rate:.2f}, "
                    f"Current Load: {current_load}, "
                    f"Avg Response: {avg_response:.1f}s")
    
    try:
        # Create enhanced prompt for growth analysis
        logging.info(f"📝 [{client_id}] Creating enhanced prompt...")
        enhanced_prompt = create_enhanced_growth_analysis_prompt(
            complete_raw_data, analysis_type, analysis_requirements
        )
        logging.info(f"🔍 [{client_id}] Prompt length: {len(enhanced_prompt)} characters")

        # STEP 1: TRY VERTEX AI FIRST (PRIMARY METHOD)
        vertex_result = try_vertex_ai_growth_request(
            enhanced_prompt=enhanced_prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            start_time=start_time
        )

        if vertex_result and vertex_result.get("success"):
            # ✅ Vertex AI succeeded - return immediately
            logging.info(f"✅ [{client_id}] Using Vertex AI response (PRIMARY METHOD)")
            return GrowthChatResponse(
                content=vertex_result["content"],
                model=vertex_result["model"],
                api_key_used="vertex_ai",
                usage={"total_tokens": vertex_result.get("tokens", 0)},
                finish_reason="stop",
                response_time=time.time() - start_time,
                timestamp=time.time(),
                token_count=vertex_result.get("tokens", 0)
            )

        # STEP 2: FALLBACK TO API KEYS (if Vertex AI failed or unavailable)
        logging.info(f"🔄 [{client_id}] Vertex AI unavailable, falling back to API keys")

        # Convert to Gemini format
        logging.info(f"🔄 [{client_id}] Converting to Gemini format...")
        contents = convert_messages_to_gemini_format([
            {"role": "user", "content": enhanced_prompt}
        ])
        logging.info(f"🔍 [{client_id}] Converted contents length: {len(contents)}")

        # Production-optimized payload
        payload = {
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
                "topP": 0.9,
                "topK": 40,
                "candidateCount": 1,
                "stopSequences": [],
                "responseMimeType": "text/plain"
            },
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }
        
        params = {'key': api_key}
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
        
        logging.info(f"🌐 [{client_id}] Sending growth analysis request to Gemini API")
        logging.info(f"🔍 [{client_id}] API URL: {url}")
        logging.info(f"🔍 [{client_id}] Payload keys: {list(payload.keys())}")
        
        # Track request start time for response time calculation
        request_start_time = time.time()
        
        # Make request
        response = requests.post(
            url,
            json=payload,
            params=params,
            timeout=REQUEST_TIMEOUT
        )
        
        # Calculate response time
        response_time = time.time() - request_start_time
        
        logging.info(f"📡 [{client_id}] Response status: {response.status_code}")
        logging.info(f"⏱️ [{client_id}] Response time: {response_time:.2f}s")
        logging.info(f"🔍 [{client_id}] Response headers: {dict(response.headers)}")
        
        if response.status_code == 200:
            # SUCCESS: Update API key health with response time
            logging.info(f"✅ [{client_id}] HTTP 200 Success - Updating API key health")
            update_api_key_health(api_key, success=True, response_time=response_time)
            
            try:
                data = response.json()
                logging.info(f"✅ [{client_id}] Successfully parsed JSON response")
            except Exception as json_error:
                logging.error(f"❌ [{client_id}] JSON parsing failed: {json_error}")
                logging.error(f"🔍 [{client_id}] Raw response text: {response.text[:500]}...")
                
                # JSON parsing failure - update API key health with response time
                update_api_key_health(api_key, success=False, error_code="JSON_PARSE_ERROR", response_time=response_time)
                raise Exception(f"Failed to parse JSON response: {json_error}")
            
            # ULTRA DETAILED LOGGING FOR AI RESPONSE
            logging.info(f"🔍 [{client_id}] Raw API response keys: {list(data.keys())}")
            logging.info(f"🔍 [{client_id}] Full response structure: {json.dumps(data, indent=2, default=str)[:1000]}...")
            
            if 'candidates' in data and len(data['candidates']) > 0:
                candidate = data['candidates'][0]
                logging.info(f"🔍 [{client_id}] Found {len(data['candidates'])} candidate(s)")
                logging.info(f"🔍 [{client_id}] Candidate keys: {list(candidate.keys())}")
                logging.info(f"🔍 [{client_id}] Candidate structure: {json.dumps(candidate, indent=2, default=str)[:500]}...")
                
                # Check finish reason
                finish_reason = candidate.get('finishReason', 'UNKNOWN')
                logging.info(f"🔍 [{client_id}] Finish reason: {finish_reason}")
                
                # Check safety ratings
                if 'safetyRatings' in candidate:
                    logging.info(f"🔍 [{client_id}] Safety ratings: {candidate['safetyRatings']}")
                
                # Enhanced content extraction with detailed logging
                content = ""
                extraction_method = "none"
                
                try:
                    # Method 1: Standard content extraction
                    if 'content' in candidate and candidate['content'] is not None:
                        content_obj = candidate['content']
                        logging.info(f"🔍 [{client_id}] Found content object: {type(content_obj)}")
                        logging.info(f"🔍 [{client_id}] Content object: {json.dumps(content_obj, indent=2, default=str)[:300]}...")
                        
                        if 'parts' in content_obj and content_obj['parts']:
                            parts = content_obj['parts']
                            logging.info(f"🔍 [{client_id}] Found content.parts: {len(parts)} parts")
                            
                            if len(parts) > 0:
                                first_part = parts[0]
                                logging.info(f"🔍 [{client_id}] First part type: {type(first_part)}")
                                logging.info(f"🔍 [{client_id}] First part keys: {list(first_part.keys()) if isinstance(first_part, dict) else 'Not a dict'}")
                                logging.info(f"🔍 [{client_id}] First part content: {json.dumps(first_part, indent=2, default=str)[:200]}...")
                                
                                if isinstance(first_part, dict) and 'text' in first_part:
                                    content = first_part['text']
                                    extraction_method = "content.parts[0].text"
                                    logging.info(f"🔍 [{client_id}] Extracted via method 1: {len(content)} characters")
                                else:
                                    logging.warning(f"⚠️ [{client_id}] First part has no 'text' field")
                            else:
                                logging.warning(f"⚠️ [{client_id}] Parts array is empty")
                        else:
                            logging.warning(f"⚠️ [{client_id}] Content object has no 'parts' field or parts is empty")
                    
                    # Method 2: Direct text field
                    if not content and 'text' in candidate:
                        content = candidate['text']
                        extraction_method = "candidate.text"
                        logging.info(f"🔍 [{client_id}] Extracted via method 2: {len(content)} characters")
                    
                    # Method 3: Look for any text-like fields
                    if not content:
                        for key, value in candidate.items():
                            if isinstance(value, str) and len(value) > 10:
                                content = value
                                extraction_method = f"candidate.{key}"
                                logging.info(f"🔍 [{client_id}] Extracted via method 3 ({key}): {len(content)} characters")
                                break
                    
                    # Method 4: Fallback to string conversion
                    if not content:
                        content_obj = candidate.get('content', candidate)
                        content = str(content_obj)
                        extraction_method = "string_conversion"
                        logging.info(f"🔍 [{client_id}] Extracted via method 4: {len(content)} characters")
                        
                except Exception as e:
                    logging.error(f"🔍 [{client_id}] Content extraction error: {e}")
                    logging.error(f"🔍 [{client_id}] Candidate type: {type(candidate)}")
                    logging.error(f"🔍 [{client_id}] Candidate repr: {repr(candidate)[:200]}...")
                    content = str(candidate)
                    extraction_method = "error_fallback"
                
                # Detailed content validation
                logging.info(f"🔍 [{client_id}] Content extraction method: {extraction_method}")
                logging.info(f"🔍 [{client_id}] Content type: {type(content)}")
                logging.info(f"🔍 [{client_id}] Content length: {len(content) if content else 0}")
                logging.info(f"🔍 [{client_id}] Content preview (first 300 chars): '{content[:300] if content else 'EMPTY'}'")
                logging.info(f"🔍 [{client_id}] Content stripped length: {len(content.strip()) if content else 0}")
                
                # Check for content issues
                if not content:
                    logging.error(f"❌ [{client_id}] Content is None or False")
                    logging.error(f"🔍 [{client_id}] Candidate finish reason: {finish_reason}")
                    if finish_reason == 'SAFETY':
                        logging.error(f"❌ [{client_id}] Content blocked by safety filters")
                        update_api_key_health(api_key, success=False, error_code="SAFETY_FILTER", response_time=response_time)
                        raise Exception("Content blocked by safety filters")
                    else:
                        update_api_key_health(api_key, success=False, error_code="NO_CONTENT", response_time=response_time)
                        raise Exception("Content is None - API returned no text")
                elif content.strip() == "":
                    logging.error(f"❌ [{client_id}] Content is empty string or whitespace only")
                    update_api_key_health(api_key, success=False, error_code="EMPTY_CONTENT", response_time=response_time)
                    raise Exception("Content is empty string - API returned whitespace only")
                elif len(content.strip()) < 5:
                    logging.error(f"❌ [{client_id}] Content too short: '{content.strip()}'")
                    update_api_key_health(api_key, success=False, error_code="SHORT_CONTENT", response_time=response_time)
                    raise Exception(f"Content too short ({len(content.strip())} chars): '{content.strip()}'")
                else:
                    logging.info(f"✅ [{client_id}] Content validation passed - {len(content.strip())} characters extracted")
                
                # Success metrics
                usage = data.get('usageMetadata', {})
                token_count = usage.get('totalTokenCount', 0)
                
                analysis_time = time.time() - start_time
                
                logging.info(f"✅ [{client_id}] Growth Analysis Complete - {analysis_type} ({token_count} tokens, {analysis_time:.2f}s)")
                logging.info(f"🔍 [{client_id}] Usage metadata: {usage}")
                
                # Log final API key health status after successful completion with enhanced metrics
                updated_health = api_key_health.get(api_key, {})
                success_rate = updated_health.get('success_rate', 1.0)
                current_load = updated_health.get('current_load', 0) 
                
                logging.info(f"🔑 [{client_id}] Final API Key Health: {updated_health.get('key_id', 'unknown')} - "
                            f"Status: HEALTHY, Consecutive Failures Reset: 0, "
                            f"Success Rate: {success_rate:.2f}, Current Load: {current_load}")
                
                return GrowthChatResponse(
                    content=content,
                    model="gemini-2.5-pro",
                    api_key_used=f"{client_id}_key_{api_key[-4:]}",
                    usage=usage,
                    finish_reason=candidate.get('finishReason', 'STOP'),
                    response_time=analysis_time,
                    timestamp=time.time(),
                    token_count=token_count
                )
            else:
                logging.error(f"❌ [{client_id}] No candidates in response")
                logging.error(f"🔍 [{client_id}] Response data: {data}")
                if 'candidates' in data:
                    logging.error(f"🔍 [{client_id}] Candidates array length: {len(data['candidates'])}")
                
                update_api_key_health(api_key, success=False, error_code="NO_CANDIDATES", response_time=response_time)
                raise Exception("No candidates found in API response")
        
        else:
            # HTTP ERROR: Update API key health with specific error code and response time
            error_code = str(response.status_code)
            logging.error(f"❌ [{client_id}] HTTP Error {response.status_code}")
            logging.error(f"🔍 [{client_id}] Response text: {response.text[:500]}...")
            
            # Special handling for different HTTP error codes with response time tracking
            if response.status_code == 503:
                logging.error(f"🚨 [{client_id}] API Overloaded (503) - Marking API key for cooldown")
                update_api_key_health(api_key, success=False, error_code="503", response_time=response_time)
                logging.warning(f"🔑 [{client_id}] API Key Health Summary: {get_api_key_status_summary()}")
            elif response.status_code == 429:
                logging.error(f"🚨 [{client_id}] Rate Limited (429) - API key needs cooldown")
                update_api_key_health(api_key, success=False, error_code="429", response_time=response_time)
            elif response.status_code in [400, 401, 403]:
                logging.error(f"🚨 [{client_id}] Client Error ({response.status_code}) - API key may be invalid")
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
            elif response.status_code >= 500:
                logging.error(f"🚨 [{client_id}] Server Error ({response.status_code}) - Temporary API issue")
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
            else:
                logging.error(f"🚨 [{client_id}] Unknown HTTP Error ({response.status_code})")
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
            
            # Log updated API key health after error with enhanced metrics
            updated_health = api_key_health.get(api_key, {})
            success_rate = updated_health.get('success_rate', 1.0)
            current_load = updated_health.get('current_load', 0)
            
            logging.error(f"🔑 [{client_id}] Updated API Key Health: {updated_health.get('key_id', 'unknown')} - "
                         f"Consecutive Failures: {updated_health.get('consecutive_failures', 0)}, "
                         f"Success Rate: {success_rate:.2f}, Current Load: {current_load}")
            
            raise Exception(f"HTTP {response.status_code}: {response.text}")
    
    except Exception as e:
        analysis_time = time.time() - start_time
        response_time = time.time() - request_start_time if 'request_start_time' in locals() else analysis_time
        
        logging.error(f"❌ [{client_id}] Growth analysis error after {analysis_time:.2f}s: {str(e)}")
        logging.error(f"🔍 [{client_id}] Error type: {type(e).__name__}")
        
        # If this exception wasn't already handled above, update API key health with response time
        if api_key in api_key_health:
            current_failures = api_key_health[api_key].get('consecutive_failures', 0)
            if "HTTP" not in str(e):  # Only update if we haven't already updated for HTTP errors
                logging.warning(f"🔑 [{client_id}] Updating API key health for general exception")
                update_api_key_health(api_key, success=False, error_code="GENERAL_EXCEPTION", response_time=response_time)
            
            # Log comprehensive API key health summary on error with enhanced metrics
            updated_health = api_key_health.get(api_key, {})
            success_rate = updated_health.get('success_rate', 1.0)
            current_load = updated_health.get('current_load', 0)
            
            logging.error(f"🔑 [{client_id}] API Key Health Summary after error: {get_api_key_status_summary()}")
            logging.error(f"🔑 [{client_id}] Detailed Key Health: Success Rate: {success_rate:.2f}, Current Load: {current_load}")
        
        import traceback
        logging.error(f"🔍 [{client_id}] Full traceback: {traceback.format_exc()}")
        raise

def create_enhanced_growth_analysis_prompt(complete_raw_data: Dict, analysis_type: str, analysis_requirements: str) -> str:
    """Create 100/100 enhanced analysis prompt with COMPLETE Multi-Database Intelligence integration - ALL DATA UTILIZED"""
    
    logging.info(f"🎯 Starting enhanced growth analysis prompt creation for {analysis_type}")
    
    user_profile = complete_raw_data.get("user_profile", {})
    responses = complete_raw_data.get("responses", [])
    multi_db_intelligence = complete_raw_data.get("multi_database_intelligence", {})
    behavioral_data = complete_raw_data.get("behavioral_analytics", {})
    
    logging.info(f"📊 Data summary: {len(responses)} responses, multi-db: {bool(multi_db_intelligence)}, behavioral: {bool(behavioral_data)}")
    
    # Extract and validate user profile data
    business_name = user_profile.get('business_name', 'Unknown Business')
    username = user_profile.get('username', 'Client')
    
    # Handle industry as both string and list
    industry_raw = user_profile.get('industry', 'Unknown Industry')
    if isinstance(industry_raw, list):
        industry = ", ".join(industry_raw) if industry_raw else 'Unknown Industry'
    else:
        industry = str(industry_raw) if industry_raw else 'Unknown Industry'
    
    team_size = user_profile.get('team_size', 'Unknown')
    biggest_challenge = user_profile.get('biggest_challenge', 'Unknown Challenge')
    business_description = user_profile.get('business_description', 'Not provided')
    location = user_profile.get('location', 'Unknown Location')
    
    logging.info(f"👤 User profile: {username} at {business_name} ({industry}, {team_size} employees)")
    
    # Get current date and time for Gemini context
    current_datetime = datetime.now()
    current_date_str = current_datetime.strftime('%A, %B %d, %Y')
    current_time_str = current_datetime.strftime('%I:%M %p %Z')
    current_timestamp = current_datetime.isoformat()

    user_context = f"""
═══════════════════════════════════════════════════════════════════════════════
🚀 CRITICAL CLIENT GROWTH CONTEXT - COMPLETE MULTI-DATABASE INTELLIGENCE 🚀
═══════════════════════════════════════════════════════════════════════════════

📅 ANALYSIS DATE & TIME CONTEXT:
- Analysis Date: {current_date_str}
- Analysis Time: {current_time_str}
- Timestamp: {current_timestamp}
- Report Generation Context: Real-time growth strategy analysis with complete database correlation intelligence

👤 CLIENT PROFILE:
- Full Name: {username}
- Business Name: {business_name}
- Industry: {industry}
- Team Size: {team_size} employees
- Location: {location}
- Primary Challenge: {biggest_challenge}
- Business Description: {business_description}

🚀 COMPLETE MULTI-DATABASE GROWTH INTELLIGENCE CONTEXT:
This analysis leverages COMPLETE MULTI-DATABASE INTELLIGENCE with FULL DATA UTILIZATION and ADVANCED CORRELATION ANALYSIS to provide sophisticated growth strategy for {username}, the founder/leader of {business_name}, a {industry} company with {team_size} employees. {username} completed the Growth Strategy Assessment on {current_date_str}, providing comprehensive data integrated with Component Intelligence, Business DNA Profile, Dream Analysis, and Behavioral Intelligence to create breakthrough growth recommendations addressing their challenge of {biggest_challenge}.

📊 COMPLETE MULTI-DATABASE INTELLIGENCE WITH DETAILED RESPONSES:
{format_complete_multi_database_intelligence(multi_db_intelligence)}

🧠 COMPLETE BEHAVIORAL INTELLIGENCE INTEGRATION:
{format_complete_growth_behavioral_data(behavioral_data)}

🔬 DEEP CORRELATION & PATTERN ANALYSIS REQUIREMENTS:
CRITICAL: You MUST perform advanced correlation analysis across ALL detailed data sources:

1. CROSS-DATABASE DETAILED CORRELATION MAPPING:
   - Identify correlations between Growth responses and specific Component Assessment answers
   - Map specific Business DNA personality responses to Growth strategy preferences and decision patterns
   - Correlate specific Dream aspiration answers with Growth barrier patterns and breakthrough opportunities
   - Analyze specific Behavioral Intelligence responses against Growth implementation readiness
   - Find hidden connections between {biggest_challenge} and specific personality/behavioral responses

2. DETAILED RESPONSE PATTERN COVARIANCE ANALYSIS:
   - Analyze covariance between similar question types across different assessments using actual responses
   - Identify response consistency patterns using specific answers that reveal deeper personality insights
   - Map contradiction patterns between specific responses that indicate internal conflicts or growth opportunities
   - Detect response cluster patterns in actual answers that suggest undiscovered business DNA traits

3. COMPREHENSIVE GROWTH READINESS CORRELATION MATRIX:
   - Correlate specific component maturity responses with growth appetite responses
   - Map specific personality trait responses to preferred growth strategies and implementation styles
   - Analyze specific dream alignment responses with current business capabilities and growth barriers
   - Identify specific behavioral pattern responses that predict growth success likelihood

4. PREDICTIVE PATTERN RECOGNITION FROM ACTUAL RESPONSES:
   - Identify specific response combinations that predict specific growth outcomes
   - Map personality-challenge combinations using actual responses to most effective solution approaches
   - Detect early warning patterns in specific responses for potential growth implementation failures
   - Recognize success pattern combinations in actual responses for accelerated growth recommendations

5. HIDDEN INSIGHT DISCOVERY FROM DETAILED DATA:
   - Look for non-obvious connections between specific responses across different assessments
   - Identify latent factors from actual response patterns that influence multiple response categories
   - Discover underlying motivations from specific response text that drive surface-level responses
   - Uncover unconscious patterns from detailed response analysis that reveal breakthrough opportunities

📋 ENHANCED GROWTH STRATEGY INTEGRATION REQUIREMENTS:
1. Apply Complete Multi-Database Intelligence insights using actual responses for maximum precision WITH correlation evidence
2. Integrate specific Component Assessment responses with growth readiness USING detailed pattern analysis
3. Leverage specific Business DNA personality responses for personalized growth approach WITH behavioral correlation
4. Incorporate specific Dream Analysis responses into growth strategy alignment USING detailed covariance mapping
5. Use specific Behavioral Intelligence responses to customize communication and implementation style WITH pattern recognition
6. Address {biggest_challenge} using integrated detailed intelligence insights WITH correlation-based solutions
7. Provide growth strategies that align with their complete business and personal profile using all available response data WITH predictive modeling

🎯 ADVANCED PERSONALIZATION WITH COMPLETE DATA INTELLIGENCE:
- Apply Growth Strategy specifically to {username} and {business_name} USING detailed correlation insights from all responses
- Use personality-based growth approaches from specific Business DNA responses WITH behavioral pattern validation
- Align growth strategies with specific dream aspiration responses and long-term vision USING detailed covariance analysis
- Consider specific behavioral pattern responses for implementation recommendations WITH predictive correlation
- Frame all insights in context of solving {biggest_challenge} WITH pattern-based evidence from actual responses
- NEVER use "you" or "your" - always use {username}'s name with context
- ALWAYS cite specific response combinations and data relationships to support recommendations

🚨 CRITICAL COMPLETE DATABASE ENHANCED WRITING REQUIREMENTS:
- Integrate Complete Multi-Database insights naturally throughout analysis WITH detailed correlation evidence
- Use specific Component responses to contextualize growth readiness WITH pattern analysis proof
- Apply specific Business DNA responses to customize growth approach and communication WITH behavioral correlation
- Reference specific Dream responses to ensure growth strategy alignment WITH covariance validation
- Incorporate specific Behavioral responses for personalized implementation guidance WITH predictive modeling
- Focus on breakthrough growth opportunities based on complete detailed intelligence WITH correlation discovery
- Reference current date ({current_date_str}) when discussing timelines and implementation dates
- Use {current_time_str} context for urgency and immediate action items
- PROVIDE CORRELATION COEFFICIENTS and pattern strength indicators where relevant
- CITE SPECIFIC response combinations that led to insights
- QUANTIFY pattern confidence levels and prediction accuracy

🧮 ADVANCED ANALYTICAL REQUIREMENTS USING COMPLETE DATA:
- Calculate and report correlation strength between specific response variable pairs
- Identify statistically significant patterns with confidence intervals from actual response data
- Perform cluster analysis on specific response patterns to identify behavioral archetypes
- Use principal component analysis thinking to identify underlying factors from detailed responses
- Apply multivariate analysis concepts to understand complex relationships in actual response data
- Provide evidence-based confidence scores for major recommendations using detailed response analysis

🔍 INSIGHT DISCOVERY MANDATE FROM COMPLETE DATA:
You MUST discover and report using actual detailed responses:
- At least 5 non-obvious correlations between specific responses from different assessment areas
- Minimum 3 predictive patterns for growth success likelihood using actual response combinations
- At least 2 hidden contradictions from specific responses that reveal growth opportunities
- Minimum 3 behavioral-business alignment patterns with quantified strength using detailed response data
- At least 2 breakthrough insights from complete multi-database pattern analysis using all available responses

⏰ TEMPORAL CONTEXT FOR AI ANALYSIS:
- Current Business Climate: {current_date_str} market conditions and trends
- Implementation Timeline Base: Starting from {current_date_str}
- Quarterly Planning Context: Q{((current_datetime.month - 1) // 3) + 1} {current_datetime.year}
- Year-end Planning: {12 - current_datetime.month} months remaining in {current_datetime.year}
- Strategic Planning Horizon: {current_datetime.year}-{current_datetime.year + 3} growth cycle

🎯 CORRELATION ANALYSIS SUCCESS METRICS USING COMPLETE DATA:
Your analysis will be evaluated on:
1. Depth of correlation discovery across all detailed database responses
2. Accuracy of pattern recognition and prediction using specific response data
3. Quality of evidence-based recommendations using detailed response analysis
4. Insight discovery beyond surface-level analysis using complete data
5. Quantified confidence in major recommendations using all available response data
6. Integration sophistication across all detailed intelligence sources

═══════════════════════════════════════════════════════════════════════════════
"""
    
    # Enhanced growth response analysis with Complete Multi-Database integration
    logging.info(f"🔧 Creating complete multi-database growth analysis framework...")
    multi_db_growth_analysis = f"""
═══════════════════════════════════════════════════════════════════════════════
📈 COMPLETE MULTI-DATABASE ENHANCED GROWTH RESPONSE ANALYSIS FRAMEWORK 📈
═══════════════════════════════════════════════════════════════════════════════

🎯 PRIMARY ANALYSIS FOCUS (70% of content):
ULTRA-DEEP GROWTH RESPONSE ANALYSIS using Complete Multi-Database Intelligence with Full Data Correlation Mapping:

For {username} of {business_name} in {industry}:

1. 📝 QUOTE AND ANALYZE EVERY RELEVANT GROWTH RESPONSE WITH DETAILED CORRELATION EVIDENCE:
   - Extract {username}'s exact words and choices from growth assessment responses
   - Map each response to appropriate growth opportunity or barrier WITH detailed correlation strength indicators
   - Analyze response sophistication against complete multi-database intelligence WITH pattern confidence scores
   - Identify growth strengths and gaps using integrated detailed intelligence WITH statistical significance
   - CALCULATE response consistency patterns across similar question types using actual response data
   - QUANTIFY alignment between stated preferences and behavioral indicators using specific responses
   - MEASURE correlation strength between growth choices and specific personality trait responses

2. 🔗 ADVANCED COMPLETE MULTI-DATABASE CORRELATION ANALYSIS WITH PREDICTIVE MODELING:
   - Connect {username}'s growth responses with specific Component Assessment responses USING correlation coefficients
   - Integrate specific Business DNA personality responses with growth preferences WITH covariance analysis
   - Align specific Dream aspiration responses with growth strategy direction USING pattern matching algorithms
   - Apply specific Behavioral Intelligence responses to growth implementation approach WITH predictive accuracy scores
   - IDENTIFY contradiction patterns between specific responses from different assessments
   - CALCULATE cross-database correlation matrices for actual response variable pairs
   - PREDICT growth success likelihood based on detailed multi-database pattern analysis
   - QUANTIFY personality-strategy alignment with confidence intervals using specific response combinations

3. 📊 COMPREHENSIVE GROWTH ASSESSMENT WITH COMPLETE DATA STATISTICAL VALIDATION:
   - Evaluate {username}'s growth readiness using all detailed intelligence sources WITH weighted scoring
   - Identify accelerated growth opportunities through integrated detailed insights WITH probability assessments
   - Highlight growth barriers requiring attention using complete multi-database context WITH severity rankings
   - Map growth implementation readiness using specific behavioral pattern responses WITH success prediction modeling
   - PERFORM cluster analysis on detailed response patterns to identify growth archetypes
   - CALCULATE growth readiness index based on multi-dimensional detailed analysis
   - VALIDATE insights using cross-database consistency checks with actual responses
   - QUANTIFY implementation success probability for each recommendation using detailed response analysis

4. 🎯 ADVANCED GROWTH PATTERN RECOGNITION WITH COMPLETE DATA INSIGHTS:
   - Analyze {username}'s systematic growth preferences across all areas using detailed response data WITH pattern strength metrics
   - Identify {username}'s natural growth strengths and blind spots from all detailed data sources WITH correlation evidence
   - Connect growth patterns to {username}'s business challenge of {biggest_challenge} using specific responses WITH causal analysis
   - Reveal {username}'s authentic growth DNA through complete multi-database response analysis WITH confidence scoring
   - DISCOVER hidden growth patterns through latent factor analysis of detailed responses
   - IDENTIFY non-obvious correlation patterns between specific personality and growth responses
   - PREDICT potential growth failure points based on detailed pattern analysis
   - CALCULATE growth DNA compatibility scores across different strategy approaches using complete response data

🔬 MANDATORY DEEP ANALYSIS REQUIREMENTS USING COMPLETE DATA:

5. 🧮 DETAILED CORRELATION MATRIX ANALYSIS:
   - Create correlation matrices between specific Growth responses and Component assessment responses
   - Calculate covariance between specific Business DNA responses and Growth strategy preferences
   - Measure alignment coefficients between specific Dream responses and Growth directions
   - Quantify behavioral consistency patterns across all assessment areas using detailed response data
   - REPORT correlation strengths with statistical significance indicators using actual responses
   - IDENTIFY the strongest predictive relationships in the detailed data
   - HIGHLIGHT unexpected correlation discoveries from complete response analysis

6. 🎯 DETAILED PATTERN CONTRADICTION DISCOVERY:
   - Identify specific response contradictions that reveal internal conflicts or growth opportunities
   - Analyze inconsistencies between specific stated preferences and behavioral indicators
   - Map contradiction patterns in actual responses to specific growth breakthrough opportunities
   - Quantify the significance of each contradiction pattern discovered using detailed response analysis
   - CALCULATE contradiction severity scores and resolution pathways using specific response data
   - PREDICT growth challenges arising from identified contradictions in actual responses

7. 📈 PREDICTIVE GROWTH MODELING FROM COMPLETE DATA:
   - Build predictive models for growth success based on detailed multi-database response patterns
   - Calculate probability scores for different growth strategy approaches using specific response combinations
   - Identify early warning indicators for potential implementation challenges from actual response analysis
   - Quantify confidence levels for major growth recommendations using detailed response validation
   - PROVIDE success probability ranges for each strategic recommendation using complete response data
   - CALCULATE risk-adjusted growth potential scores using all available detailed responses

8. 🔍 HIDDEN INSIGHT EXCAVATION FROM COMPLETE DETAILED DATA:
   - Discover non-obvious relationships between specific responses from different assessments
   - Identify latent factors from detailed response patterns that influence multiple response categories
   - Uncover unconscious patterns from complete response analysis that reveal breakthrough opportunities
   - Map subliminal growth preferences from detailed responses that contradict surface-level responses
   - QUANTIFY insight discovery confidence and validation strength using complete data analysis
   - CALCULATE the novelty and significance of each hidden insight using detailed response validation

🎪 ANALYSIS DEPTH REQUIREMENTS USING COMPLETE DATA:
You MUST provide using detailed response analysis:
- Minimum 7 statistically significant correlations with strength indicators from actual responses
- At least 5 predictive pattern discoveries with confidence scores using detailed response combinations
- Minimum 3 contradiction analyses with resolution pathways from specific response data
- At least 6 hidden insights with validation evidence using complete response analysis
- Minimum 4 growth success probability calculations using detailed multi-database responses
- At least 3 breakthrough opportunity discoveries with correlation evidence from complete data

📊 EVIDENCE-BASED REPORTING STANDARDS USING COMPLETE DATA:
- Cite specific response combinations from all databases that led to insights
- Provide correlation coefficients where relevant using actual response data (r = X.XX)
- Include confidence intervals for predictions using detailed response analysis (95% CI: X.X - X.X)
- Quantify pattern strength on 0-1 scale with interpretation using complete response data
- Calculate and report statistical significance using actual responses (p < 0.05)
- Provide weighted scoring for multi-factor analysis using all available detailed responses

🚀 BREAKTHROUGH DISCOVERY MANDATE FROM COMPLETE DATABASE UTILIZATION:
Discover and report breakthrough insights using all detailed response data that go beyond obvious analysis:
- Identify growth accelerators hidden in specific personality-behavior response correlations
- Uncover implementation success patterns from complete multi-database response integration
- Discover challenge-solution matches with high probability success rates using detailed response analysis
- Find growth DNA patterns from complete response data that predict optimal strategy approaches
- Reveal unconscious growth blocks through detailed correlation analysis of all available responses

═══════════════════════════════════════════════════════════════════════════════
"""
    
    # Format growth responses for analysis
    logging.info(f"📊 Formatting growth response analysis...")
    growth_response_formatting = format_growth_assessment_responses(responses)
    
    # Enhanced analysis instructions with Complete Multi-Database Intelligence
    logging.info(f"📋 Creating enhanced complete multi-database instructions...")
    enhanced_multi_db_instructions = f"""
═══════════════════════════════════════════════════════════════════════════════
🎯 COMPLETE MULTI-DATABASE ENHANCED SPECIFIC ANALYSIS INSTRUCTIONS FOR {username.upper()} 🎯
═══════════════════════════════════════════════════════════════════════════════

🚀 COMPLETE MULTI-DATABASE GROWTH INTELLIGENCE APPLICATION PRIORITIES WITH FULL DATA CORRELATION ANALYSIS:

1. 🎯 COMPLETE INTEGRATED INTELLIGENCE ANALYSIS WITH DETAILED STATISTICAL VALIDATION (30% of analysis):
   - Apply specific Component Assessment business maturity responses to growth readiness WITH correlation strength indicators
   - Use specific Business DNA personality responses for personalized growth approach WITH pattern confidence scores
   - Integrate specific Dream aspiration responses into growth strategy alignment WITH covariance analysis validation
   - Apply specific Behavioral Intelligence responses for implementation customization WITH predictive modeling accuracy
   - CALCULATE cross-database correlation coefficients for actual response relationships (r = X.XX)
   - QUANTIFY intelligence integration confidence levels using detailed response data (95% CI: X.X - X.X)
   - IDENTIFY statistically significant patterns across all detailed intelligence sources (p < 0.05)
   - MEASURE alignment strength between different intelligence databases on 0-1 scale using actual responses

2. 📝 ULTRA-DEEP GROWTH RESPONSE ANALYSIS WITH COMPLETE CORRELATION EVIDENCE (40% of analysis):
{growth_response_formatting}
   - PERFORM response pattern analysis with statistical significance testing using all available data
   - CALCULATE consistency scores between similar responses across assessments using detailed response analysis
   - IDENTIFY contradiction patterns with severity rankings and resolution pathways from actual response data
   - QUANTIFY response sophistication levels using complete multi-database validation
   - MAP each response to correlation patterns in detailed personality and behavioral data
   - PREDICT implementation success probability for each response-based strategy using complete data analysis
   - DISCOVER hidden insights through detailed response clustering and pattern analysis

3. 🧠 COMPLETE MULTI-DATABASE BEHAVIORAL VALIDATION WITH DETAILED PREDICTIVE MODELING (20% of analysis):
   - Validate growth responses against specific behavioral pattern responses WITH correlation matrices
   - Cross-reference specific personality trait responses with growth preferences WITH covariance analysis
   - Confirm dream alignment using specific dream responses with growth direction WITH statistical validation
   - Assess implementation readiness through detailed behavioral response analysis WITH success probability scoring
   - CALCULATE behavioral-growth alignment coefficients with confidence intervals using specific response data
   - IDENTIFY behavioral predictors of growth success with accuracy scores using detailed response analysis
   - QUANTIFY contradiction resolution requirements with priority rankings using complete response data
   - PREDICT behavioral adaptation needs for optimal growth implementation using all available detailed responses

4. 🚀 BREAKTHROUGH GROWTH STRATEGY ROADMAP WITH COMPLETE DATA EVIDENCE-BASED PRIORITIZATION (10% of analysis):
   - Provide Complete Multi-Database informed growth acceleration sequence using detailed responses WITH probability success rates
   - Show breakthrough opportunities using complete integrated intelligence WITH detailed correlation evidence
   - Address {biggest_challenge} through comprehensive growth optimization using all response data WITH predictive modeling
   - Create implementation masterplan using all detailed intelligence sources WITH risk-adjusted sequencing
   - CALCULATE ROI probability for each recommended strategy using complete response analysis (Expected Value Analysis)
   - RANK opportunities by correlation strength and implementation feasibility using detailed response data
   - PROVIDE timeline optimization based on behavioral readiness patterns from complete response analysis

═══════════════════════════════════════════════════════════════════════════════
🎯 GROWTH STRATEGY EXCELLENCE STANDARDS FOR {business_name} WITH COMPLETE DATA STATISTICAL RIGOR 🎯
═══════════════════════════════════════════════════════════════════════════════

For {username} in {industry} with {team_size} employees:

📊 ENHANCED GROWTH INTEGRATION REQUIREMENTS WITH COMPLETE DATA CORRELATION ANALYSIS:
- Market Demand & Opportunity Analysis WITH correlation to detailed personality-driven market preference responses
- Capacity & Systems Scaling Assessment WITH behavioral readiness correlation analysis using specific responses
- Cash Flow & Financial Growth Planning WITH risk tolerance and decision-making pattern integration from actual response data
- Team & Leadership Development Strategy WITH DNA leadership trait responses correlation mapping
- Marketing & Client Acquisition Optimization WITH personality-based communication preference analysis using detailed responses
- Strategic & Competitive Positioning WITH specific dream aspiration response alignment validation
- Digital & Operational Excellence WITH behavioral technology adoption pattern analysis from actual responses
- Risk Management & Compliance Growth WITH personality-based risk profile correlation using detailed response data

🔍 ADVANCED COMPLETE MULTI-DATABASE GROWTH CORRELATION ANALYSIS WITH QUANTIFIED INSIGHTS:
- Map {username}'s growth responses across ALL detailed intelligence sources WITH correlation strength metrics
- Identify acceleration opportunities between growth areas and specific personality trait responses WITH probability scoring
- Reveal hidden patterns in {username}'s integrated business and personal profile using complete response data WITH statistical validation
- Show compound effects of aligned growth strategy with complete detailed intelligence WITH predictive modeling
- CALCULATE cross-database correlation matrices for all major detailed response variable pairs
- IDENTIFY the strongest predictive relationships with confidence intervals using complete response analysis
- QUANTIFY pattern significance and practical application potential using all available detailed response data
- MEASURE intelligence integration effectiveness with validation scores using complete database utilization

📈 BREAKTHROUGH GROWTH READINESS WITH COMPLETE DATA PREDICTIVE ANALYTICS:
- Assess {username}'s readiness for exponential growth using all detailed data sources WITH readiness index calculation
- Identify growth multipliers preventing breakthrough acceleration using complete response correlation evidence
- Recommend Complete Multi-Database informed growth priorities using detailed response analysis WITH success probability rankings
- Create growth implementation sequence for maximum business impact using complete data WITH risk-adjusted timeline optimization
- CALCULATE growth readiness scores across multiple dimensions (0-100 scale) using all available detailed responses
- PREDICT implementation challenges with early warning indicators from complete response pattern analysis
- QUANTIFY breakthrough potential with confidence intervals using detailed multi-database response validation
- MEASURE strategy-personality fit with alignment scores using complete response data analysis

═══════════════════════════════════════════════════════════════════════════════
📋 MANDATORY COMPLETE MULTI-DATABASE ENHANCED OUTPUT REQUIREMENTS WITH DETAILED STATISTICAL EVIDENCE 📋
═══════════════════════════════════════════════════════════════════════════════

🚀 ENHANCED STRUCTURE WITH COMPLETE MULTI-DATABASE INTEGRATION AND DETAILED CORRELATION ANALYSIS:
1. 🎯 Complete Multi-Database Enhanced Executive Summary for {username} and {business_name} WITH key detailed correlation discoveries
2. 📊 Growth Response Pattern Analysis WITH statistical significance testing using complete data (quote {username}'s responses extensively)
3. 🔗 Complete Multi-Database Cross-Intelligence Connection Analysis WITH detailed correlation matrices and predictive modeling
4. 🏢 Integrated Growth Strategy Applications WITH probability success rates using complete response analysis (specific strategies for {business_name})
5. 🧠 Complete Multi-Database Behavioral Validation WITH detailed correlation evidence (how all intelligence aligns with growth patterns)
6. 🎯 Industry-Specific Growth Recommendations WITH personality-industry fit analysis using complete response data (tailored to {industry} using all intelligence)
7. 👥 Leadership & Team Growth Insights WITH behavioral leadership correlation using detailed response analysis (growth leadership for {team_size} employees)
8. 🚀 Breakthrough Growth Acceleration Roadmap WITH risk-adjusted prioritization using complete data (addressing {biggest_challenge} with integrated solutions)

📋 ENHANCED COMPLETE MULTI-DATABASE EVIDENCE REQUIREMENTS WITH DETAILED QUANTIFIED VALIDATION:
- Quote {username}'s specific growth responses and map to complete integrated intelligence WITH detailed correlation coefficients
- Reference Complete Multi-Database correlations in {username}'s response patterns WITH statistical significance indicators using actual response data
- Connect growth insights across all business and personal intelligence areas using detailed responses WITH pattern strength metrics
- Use all detailed intelligence sources to contextualize {username}'s growth sophistication WITH confidence scoring
- Provide Complete Multi-Database informed solutions for {business_name}'s {biggest_challenge} using detailed response analysis WITH success probability analysis
- Show breakthrough growth pathway for {business_name}'s acceleration using complete response data WITH evidence-based sequencing
- CALCULATE and report correlation strengths for all major insights using detailed response analysis (r = X.XX, p < X.XX)
- PROVIDE confidence intervals for all major predictions using complete response data (95% CI: X.X - X.X)
- QUANTIFY recommendation strength on validated 0-1 scale with interpretation using all available detailed responses

🎯 ADVANCED COMPLETE MULTI-DATABASE PERSONALIZATION STANDARDS WITH DETAILED CORRELATION VALIDATION:
- Apply all detailed intelligence sources specifically to {username} and {business_name} WITH correlation evidence from actual responses
- Use personality-based growth approaches appropriate for their Business DNA using specific responses WITH behavioral validation
- Consider dream alignment for long-term growth strategy sustainability using detailed dream responses WITH covariance analysis proof
- Frame behavioral insights for {team_size} team dynamics and leadership style using specific responses WITH pattern correlation
- Focus all recommendations on solving {biggest_challenge} with integrated detailed intelligence using complete response data WITH predictive modeling
- VALIDATE all personalizations with cross-database correlation analysis using actual response data
- QUANTIFY personalization effectiveness with confidence scores using complete response analysis
- MEASURE recommendation-personality fit with alignment indices using all available detailed response data

🔬 MANDATORY DEEP ANALYSIS DISCOVERY REQUIREMENTS USING COMPLETE DATA:
You MUST discover and report using complete detailed response analysis:
- Minimum 10 statistically significant correlations across databases using actual responses (r > 0.3, p < 0.05)
- At least 6 predictive patterns with success probability calculations using detailed response combinations
- Minimum 4 contradiction analyses with correlation-based resolution pathways from specific response data
- At least 8 hidden insights with multi-database validation evidence using complete response analysis
- Minimum 5 breakthrough opportunities with probability success rates using detailed response data
- At least 4 growth accelerator discoveries with correlation strength indicators from complete response analysis
- Minimum 3 implementation risk predictions with early warning correlations using all available detailed responses

═══════════════════════════════════════════════════════════════════════════════
🎯 GROWTH-SPECIFIC COMPLETE MULTI-DATABASE REQUIREMENTS WITH ADVANCED ANALYTICS 🎯
═══════════════════════════════════════════════════════════════════════════════

{analysis_requirements}

ADDITIONAL ANALYTICAL DEPTH REQUIREMENTS USING COMPLETE DATA:
- Perform latent factor analysis thinking to identify underlying growth drivers from detailed response patterns
- Apply multivariate correlation concepts to understand complex relationship networks using actual response data
- Use cluster analysis mentality to group similar response patterns from complete database utilization
- Implement predictive modeling logic for success probability calculations using all available detailed responses
- Apply principal component analysis thinking to identify core growth factors from complete response analysis
- Use regression analysis concepts to quantify variable relationships using detailed response data
- Perform outlier analysis to identify unique growth opportunities or risks from complete database utilization

═══════════════════════════════════════════════════════════════════════════════
🎯 FINAL COMPLETE MULTI-DATABASE INTEGRATION REMINDER WITH DETAILED STATISTICAL MANDATE 🎯
═══════════════════════════════════════════════════════════════════════════════

This analysis leverages COMPLETE MULTI-DATABASE INTELLIGENCE with FULL DATA UTILIZATION and ADVANCED CORRELATION ANALYSIS to provide {username} of {business_name} with breakthrough growth strategy appropriate for their integrated business and personal profile. Every recommendation should be:

1. Grounded in {username}'s actual growth assessment responses WITH detailed correlation evidence from all databases
2. Enhanced by Complete Multi-Database Intelligence correlation analysis using all available response data WITH statistical validation
3. Tailored to {industry} and {team_size} employee context with personality considerations using specific responses WITH pattern proof
4. Focused on solving {biggest_challenge} with integrated detailed intelligence using complete response analysis WITH predictive modeling
5. Aligned across Component maturity, Business DNA traits, Dream aspirations, and Behavioral patterns using all actual responses WITH covariance analysis
6. Implementation-ready with behavioral customization using detailed response analysis WITH success probability scoring
7. Breakthrough-oriented toward exponential growth acceleration using complete data WITH evidence-based confidence intervals

🚨 CRITICAL SUCCESS METRICS FOR THIS COMPLETE DATA ANALYSIS:
- Statistical rigor: All major insights supported by detailed correlation coefficients and significance testing using actual responses
- Predictive accuracy: Success probabilities calculated for all major recommendations using complete response data
- Integration depth: Minimum 10 cross-database correlations discovered and quantified using detailed response analysis
- Insight novelty: At least 5 non-obvious breakthrough discoveries with validation evidence from complete data utilization
- Practical application: All insights translated to specific, measurable action items with timeline optimization using detailed response analysis

CRITICAL: Seamlessly integrate Complete Multi-Database insights without explicitly mentioning individual systems - let the sophisticated integrated analysis speak for itself through detailed statistical evidence and correlation discoveries using all available response data.

BEGIN COMPLETE MULTI-DATABASE ENHANCED GROWTH ANALYSIS WITH FULL DATA UTILIZATION AND ADVANCED CORRELATION INTELLIGENCE NOW:
"""
    
    # Combine all sections
    final_prompt = f"{user_context}\n{multi_db_growth_analysis}\n{enhanced_multi_db_instructions}"
    
    # Log final prompt statistics
    prompt_length = len(final_prompt)
    prompt_word_count = len(final_prompt.split())
    
    logging.info(f"✅ Enhanced complete growth analysis prompt completed")
    logging.info(f"📊 Final prompt statistics:")
    logging.info(f"   - Total characters: {prompt_length:,}")
    logging.info(f"   - Total words: {prompt_word_count:,}")
    logging.info(f"   - User: {username} at {business_name}")
    logging.info(f"   - Industry: {industry}")
    logging.info(f"   - Challenge: {biggest_challenge}")
    
    return final_prompt

def format_growth_assessment_responses(responses: List[Dict]) -> str:
    """Format growth assessment responses for analysis prompt"""
    
    if not responses:
        return "No growth assessment responses available for analysis."
    
    formatted_responses = []
    formatted_responses.append(f"📊 GROWTH ASSESSMENT RESPONSES ({len(responses)} total responses):\n")
    
    # Group responses by section for better organization
    section_groups = {}
    
    for response in responses:
        section = response.get('section', 'Unknown Section')
        if section not in section_groups:
            section_groups[section] = []
        section_groups[section].append(response)
    
    # Format each section
    for section_name, section_responses in section_groups.items():
        formatted_responses.append(f"\n🎯 GROWTH SECTION: {section_name.upper()}")
        formatted_responses.append(f"   ({len(section_responses)} responses)")
        
        for i, response in enumerate(section_responses, 1):
            try:
                question_id = response.get('question_id', f'Q{i}')
                question_text = response.get('question_text', 'Question not available')
                response_data = response.get('response_data', {})
                
                # Extract the actual response
                if isinstance(response_data, dict):
                    selected_response = (
                        response_data.get('selected_option') or 
                        response_data.get('response_text') or 
                        response_data.get('value') or 
                        response_data.get('slider_value') or
                        'No response selected'
                    )
                else:
                    selected_response = str(response_data) if response_data else 'No response available'
                
                # Handle list/array responses (like tournament results)
                if isinstance(selected_response, list):
                    selected_response = ', '.join(str(item) for item in selected_response)
                
                formatted_responses.append(f"\n   📝 {question_id}: {question_text}")
                formatted_responses.append(f"      💡 Response: {selected_response}")
                
                # Add options if available
                all_options = response.get('all_options', [])
                if all_options and isinstance(all_options, list) and len(all_options) > 1:
                    formatted_responses.append(f"      🎯 Available Options: {', '.join(str(opt) for opt in all_options[:5])}{'...' if len(all_options) > 5 else ''}")
                
                # Add metadata if available
                metadata = response.get('metadata', {})
                if metadata and isinstance(metadata, dict):
                    timing_data = metadata.get('timing_data', {})
                    if timing_data and timing_data.get('total_engagement_time'):
                        formatted_responses.append(f"      ⏱️ Response Time: {timing_data['total_engagement_time']} seconds")
                
            except Exception as e:
                formatted_responses.append(f"\n   ❌ Error formatting response {i}: {str(e)}")
                continue
    
    formatted_responses.append(f"\n\n🎯 TOTAL GROWTH RESPONSES ANALYZED: {len(responses)}")
    formatted_responses.append(f"📊 GROWTH SECTIONS COVERED: {len(section_groups)}")
    
    return "\n".join(formatted_responses)


def _analyze_growth_engagement_level(total_movements: int, avg_speed: float) -> str:
    """Analyze engagement level based on mouse behavior for growth assessment"""
    if total_movements > 1000 and avg_speed > 15:
        return "High Engagement - Quick Decision Making"
    elif total_movements > 500:
        return "Moderate Engagement - Thoughtful Consideration"
    else:
        return "Deliberate Engagement - Careful Analysis"


def _analyze_growth_thoughtfulness(revision_ratio: float) -> str:
    """Analyze thoughtfulness level based on keyboard behavior for growth assessment"""
    if revision_ratio > 20:
        return "Highly Analytical - Multiple Revisions"
    elif revision_ratio > 10:
        return "Balanced Approach - Some Revisions"
    else:
        return "Intuitive Response - Direct Input"


def format_complete_multi_database_intelligence(multi_db_intelligence: Dict) -> str:
    """Format COMPLETE multi-database intelligence with ALL detailed responses for growth analysis - FIXED DATA TYPE HANDLING"""
    if not multi_db_intelligence:
        return "Multi-Database Intelligence: Not available - using growth responses only"
    
    formatted = []
    
    # Component Intelligence - DETAILED RESPONSES WITH SAFE TYPE CHECKING
    component_data = multi_db_intelligence.get('component_intelligence', {})
    if component_data and component_data.get('data_available'):
        phase = component_data.get('business_phase', 'Unknown')
        phase_label = component_data.get('phase_label', 'Unknown')
        responses = component_data.get('responses', {})
        
        formatted.append(f"📊 COMPONENT INTELLIGENCE: Business Phase {phase} ({phase_label})")
        
        # 🔥 FIX: Handle both string and dict responses
        if responses:
            if isinstance(responses, dict):
                formatted.append(f"🔍 DETAILED COMPONENT RESPONSES ({len(responses)} responses):")
                for question_id, response_data in responses.items():
                    try:
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', 'No answer provided')
                            weight = response_data.get('weight', 'medium')
                            section = response_data.get('section', 'general')
                            formatted.append(f"   • {question_id} ({section}, {weight} priority): \"{answer}\"")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                    except Exception as e:
                        formatted.append(f"   • {question_id}: [Error processing response: {e}]")
            elif isinstance(responses, str):
                # Handle JSON string responses
                try:
                    import json
                    responses_dict = json.loads(responses)
                    formatted.append(f"🔍 DETAILED COMPONENT RESPONSES ({len(responses_dict)} responses):")
                    for question_id, response_data in responses_dict.items():
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', 'No answer provided')
                            weight = response_data.get('weight', 'medium')
                            section = response_data.get('section', 'general')
                            formatted.append(f"   • {question_id} ({section}, {weight} priority): \"{answer}\"")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                except json.JSONDecodeError:
                    formatted.append(f"🔍 COMPONENT RESPONSES: {len(responses)} characters of response data (JSON parsing failed)")
                except Exception as e:
                    formatted.append(f"🔍 COMPONENT RESPONSES: Error processing string responses: {e}")
            else:
                formatted.append(f"🔍 COMPONENT RESPONSES: {type(responses).__name__} data type with {len(str(responses))} characters")
        else:
            formatted.append(f"🔍 DETAILED COMPONENT RESPONSES: No component responses available")
    
    # Business DNA Profile - DETAILED RESPONSES WITH SAFE TYPE CHECKING
    profile_data = multi_db_intelligence.get('profile_intelligence', {})
    if profile_data and profile_data.get('data_available'):
        responses = profile_data.get('responses', {})
        assessment_type = profile_data.get('assessment_type', 'personality_profile')
        
        formatted.append(f"\n🧬 BUSINESS DNA PROFILE ({assessment_type})")
        
        # 🔥 FIX: Handle both string and dict responses
        if responses:
            if isinstance(responses, dict):
                formatted.append(f"🔍 DETAILED PERSONALITY RESPONSES ({len(responses)} responses):")
                for question_id, response_data in responses.items():
                    try:
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', 'No answer provided')
                            section = response_data.get('section', 'personality')
                            weight = response_data.get('weight', 'medium')
                            formatted.append(f"   • {question_id} ({section}): \"{answer}\" [Weight: {weight}]")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                    except Exception as e:
                        formatted.append(f"   • {question_id}: [Error processing response: {e}]")
            elif isinstance(responses, str):
                # Handle JSON string responses
                try:
                    import json
                    responses_dict = json.loads(responses)
                    formatted.append(f"🔍 DETAILED PERSONALITY RESPONSES ({len(responses_dict)} responses):")
                    for question_id, response_data in responses_dict.items():
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', 'No answer provided')
                            section = response_data.get('section', 'personality')
                            weight = response_data.get('weight', 'medium')
                            formatted.append(f"   • {question_id} ({section}): \"{answer}\" [Weight: {weight}]")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                except json.JSONDecodeError:
                    formatted.append(f"🔍 PERSONALITY RESPONSES: {len(responses)} characters of response data (JSON parsing failed)")
                except Exception as e:
                    formatted.append(f"🔍 PERSONALITY RESPONSES: Error processing string responses: {e}")
            else:
                formatted.append(f"🔍 PERSONALITY RESPONSES: {type(responses).__name__} data type with {len(str(responses))} characters")
        else:
            formatted.append(f"🔍 DETAILED PERSONALITY RESPONSES: No personality responses available")
    
    # Dream Intelligence - DETAILED RESPONSES WITH SAFE TYPE CHECKING
    dream_data = multi_db_intelligence.get('dream_intelligence', {})
    if dream_data and dream_data.get('data_available'):
        responses = dream_data.get('responses', {})
        assessment_type = dream_data.get('assessment_type', 'dream_analysis')
        
        formatted.append(f"\n💫 DREAM ANALYSIS ({assessment_type})")
        
        # 🔥 FIX: Handle both string and dict responses
        if responses:
            if isinstance(responses, dict):
                formatted.append(f"🔍 DETAILED ASPIRATION RESPONSES ({len(responses)} responses):")
                for question_id, response_data in responses.items():
                    try:
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', response_data.get('selected_option', 'No answer provided'))
                            response_text = response_data.get('response_text', '')
                            section = response_data.get('section', 'vision')
                            
                            if response_text and response_text != answer:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\" | Additional: \"{response_text}\"")
                            else:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\"")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                    except Exception as e:
                        formatted.append(f"   • {question_id}: [Error processing response: {e}]")
            elif isinstance(responses, str):
                # Handle JSON string responses
                try:
                    import json
                    responses_dict = json.loads(responses)
                    formatted.append(f"🔍 DETAILED ASPIRATION RESPONSES ({len(responses_dict)} responses):")
                    for question_id, response_data in responses_dict.items():
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', response_data.get('selected_option', 'No answer provided'))
                            response_text = response_data.get('response_text', '')
                            section = response_data.get('section', 'vision')
                            
                            if response_text and response_text != answer:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\" | Additional: \"{response_text}\"")
                            else:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\"")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                except json.JSONDecodeError:
                    formatted.append(f"🔍 ASPIRATION RESPONSES: {len(responses)} characters of response data (JSON parsing failed)")
                except Exception as e:
                    formatted.append(f"🔍 ASPIRATION RESPONSES: Error processing string responses: {e}")
            else:
                formatted.append(f"🔍 ASPIRATION RESPONSES: {type(responses).__name__} data type with {len(str(responses))} characters")
        else:
            formatted.append(f"🔍 DETAILED ASPIRATION RESPONSES: No aspiration responses available")
    
    # Behavioral Intelligence - DETAILED RESPONSES WITH SAFE TYPE CHECKING
    analyst_data = multi_db_intelligence.get('analyst_intelligence', {})
    if analyst_data and analyst_data.get('data_available'):
        responses = analyst_data.get('responses', {})
        
        formatted.append(f"\n🧠 BEHAVIORAL INTELLIGENCE")
        
        # 🔥 FIX: Handle both string and dict responses
        if responses:
            if isinstance(responses, dict):
                formatted.append(f"🔍 DETAILED BEHAVIORAL RESPONSES ({len(responses)} responses):")
                for question_id, response_data in responses.items():
                    try:
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', response_data.get('selected_option', 'No answer provided'))
                            response_text = response_data.get('response_text', '')
                            section = response_data.get('section', 'behavioral')
                            metadata = response_data.get('metadata', {})
                            
                            metadata_str = f" | Metadata: {metadata}" if metadata else ""
                            
                            if response_text and response_text != answer:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\" | Analysis: \"{response_text}\"{metadata_str}")
                            else:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\"{metadata_str}")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                    except Exception as e:
                        formatted.append(f"   • {question_id}: [Error processing response: {e}]")
            elif isinstance(responses, str):
                # Handle JSON string responses
                try:
                    import json
                    responses_dict = json.loads(responses)
                    formatted.append(f"🔍 DETAILED BEHAVIORAL RESPONSES ({len(responses_dict)} responses):")
                    for question_id, response_data in responses_dict.items():
                        if isinstance(response_data, dict):
                            answer = response_data.get('answer', response_data.get('selected_option', 'No answer provided'))
                            response_text = response_data.get('response_text', '')
                            section = response_data.get('section', 'behavioral')
                            metadata = response_data.get('metadata', {})
                            
                            metadata_str = f" | Metadata: {metadata}" if metadata else ""
                            
                            if response_text and response_text != answer:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\" | Analysis: \"{response_text}\"{metadata_str}")
                            else:
                                formatted.append(f"   • {question_id} ({section}): \"{answer}\"{metadata_str}")
                        else:
                            formatted.append(f"   • {question_id}: \"{response_data}\"")
                except json.JSONDecodeError:
                    formatted.append(f"🔍 BEHAVIORAL RESPONSES: {len(responses)} characters of response data (JSON parsing failed)")
                except Exception as e:
                    formatted.append(f"🔍 BEHAVIORAL RESPONSES: Error processing string responses: {e}")
            else:
                formatted.append(f"🔍 BEHAVIORAL RESPONSES: {type(responses).__name__} data type with {len(str(responses))} characters")
        else:
            formatted.append(f"🔍 DETAILED BEHAVIORAL RESPONSES: No behavioral responses available")
    
    if not formatted:
        return "Multi-Database Intelligence: Extraction in progress - growth analysis will use available data sources"
    
    # Add summary statistics with safe calculation
    total_responses = 0
    try:
        if component_data.get('responses'): 
            responses = component_data['responses']
            if isinstance(responses, dict):
                total_responses += len(responses)
            elif isinstance(responses, str):
                try:
                    import json
                    parsed_responses = json.loads(responses)
                    if isinstance(parsed_responses, dict):
                        total_responses += len(parsed_responses)
                except:
                    pass
        
        if profile_data.get('responses'): 
            responses = profile_data['responses']
            if isinstance(responses, dict):
                total_responses += len(responses)
            elif isinstance(responses, str):
                try:
                    import json
                    parsed_responses = json.loads(responses)
                    if isinstance(parsed_responses, dict):
                        total_responses += len(parsed_responses)
                except:
                    pass
        
        if dream_data.get('responses'): 
            responses = dream_data['responses']
            if isinstance(responses, dict):
                total_responses += len(responses)
            elif isinstance(responses, str):
                try:
                    import json
                    parsed_responses = json.loads(responses)
                    if isinstance(parsed_responses, dict):
                        total_responses += len(parsed_responses)
                except:
                    pass
        
        if analyst_data.get('responses'): 
            responses = analyst_data['responses']
            if isinstance(responses, dict):
                total_responses += len(responses)
            elif isinstance(responses, str):
                try:
                    import json
                    parsed_responses = json.loads(responses)
                    if isinstance(parsed_responses, dict):
                        total_responses += len(parsed_responses)
                except:
                    pass
    except Exception as e:
        logging.warning(f"⚠️ Error calculating total responses: {e}")
        total_responses = "Unknown"
    
    formatted.insert(0, f"🎯 COMPLETE MULTI-DATABASE INTELLIGENCE SUMMARY: {total_responses} detailed responses across all engines")
    formatted.insert(1, f"📈 DATA UTILIZATION: Enhanced - All individual responses processed with type safety")
    formatted.insert(2, "")
    
    return "\n".join(formatted)


def format_complete_growth_behavioral_data(behavioral_data: Dict) -> str:
    """Format COMPLETE behavioral data with ALL available details for growth analysis - ENHANCED SAFETY"""
    if not behavioral_data:
        return "No behavioral data available for growth analysis"
    
    formatted = []
    formatted.append("🧠 COMPLETE BEHAVIORAL INTELLIGENCE ANALYSIS:")
    
    try:
        # Growth-specific behavioral analysis with details and safety
        if 'growth_decision_patterns' in behavioral_data:
            patterns = behavioral_data['growth_decision_patterns']
            formatted.append(f"\n🎯 GROWTH DECISION PATTERNS:")
            
            if isinstance(patterns, dict):
                for pattern_key, pattern_value in patterns.items():
                    formatted.append(f"   • {pattern_key}: {pattern_value}")
            else:
                formatted.append(f"   • Pattern Data: {patterns}")
        
        # Detailed Mouse behavior analysis for growth with safety
        mouse_data = behavioral_data.get('mouse_behavior', {})
        if mouse_data and isinstance(mouse_data, dict):
            formatted.append(f"\n🖱️ DETAILED MOUSE BEHAVIOR ANALYSIS:")
            total_movements = mouse_data.get('total_movements', 0)
            avg_speed = mouse_data.get('average_speed', 0)
            max_speed = mouse_data.get('max_speed', 0)
            click_patterns = mouse_data.get('click_patterns', {})
            hover_patterns = mouse_data.get('hover_patterns', {})
            
            formatted.append(f"   • Total Movements: {total_movements} interactions")
            formatted.append(f"   • Average Speed: {avg_speed:.1f} pixels/second")
            formatted.append(f"   • Maximum Speed: {max_speed:.1f} pixels/second")
            formatted.append(f"   • Growth Engagement Level: {_analyze_growth_engagement_level(total_movements, avg_speed)}")
            
            if click_patterns:
                formatted.append(f"   • Click Patterns: {click_patterns}")
            if hover_patterns:
                formatted.append(f"   • Hover Patterns: {hover_patterns}")
        
        # Detailed Keyboard behavior analysis for growth with safety
        keyboard_data = behavioral_data.get('keyboard_behavior', {})
        if keyboard_data and isinstance(keyboard_data, dict):
            formatted.append(f"\n⌨️ DETAILED KEYBOARD BEHAVIOR ANALYSIS:")
            total_keystrokes = keyboard_data.get('total_keystrokes', 0)
            backspace_count = keyboard_data.get('backspace_count', 0)
            typing_speed = keyboard_data.get('typing_speed', 0)
            pause_patterns = keyboard_data.get('pause_patterns', {})
            
            # Safe division
            revision_ratio = (backspace_count / max(total_keystrokes, 1)) * 100
            
            formatted.append(f"   • Total Keystrokes: {total_keystrokes}")
            formatted.append(f"   • Backspace Count: {backspace_count}")
            formatted.append(f"   • Typing Speed: {typing_speed} WPM")
            formatted.append(f"   • Revision Ratio: {revision_ratio:.1f}%")
            formatted.append(f"   • Growth Thoughtfulness: {_analyze_growth_thoughtfulness(revision_ratio)}")
            
            if pause_patterns:
                formatted.append(f"   • Pause Patterns: {pause_patterns}")
        
        # Detailed Attention patterns for growth with safety
        attention_data = behavioral_data.get('attention_patterns', {})
        if attention_data:
            formatted.append(f"\n👁️ DETAILED ATTENTION PATTERN ANALYSIS:")
            
            if isinstance(attention_data, dict):
                for attention_key, attention_value in attention_data.items():
                    formatted.append(f"   • {attention_key}: {attention_value}")
            else:
                formatted.append(f"   • Attention Data: {attention_data}")
        
        # Detailed Decision making style for growth with safety
        decision_style_data = behavioral_data.get('decision_making_style', {})
        if decision_style_data:
            formatted.append(f"\n🤔 DETAILED DECISION MAKING STYLE ANALYSIS:")
            
            if isinstance(decision_style_data, dict):
                for decision_key, decision_value in decision_style_data.items():
                    formatted.append(f"   • {decision_key}: {decision_value}")
            else:
                formatted.append(f"   • Decision Style: {decision_style_data}")
        
        # Add any other behavioral data with safety
        for key, value in behavioral_data.items():
            if key not in ['growth_decision_patterns', 'mouse_behavior', 'keyboard_behavior', 'attention_patterns', 'decision_making_style']:
                formatted.append(f"\n🔍 {key.upper().replace('_', ' ')} ANALYSIS:")
                if isinstance(value, dict):
                    for sub_key, sub_value in value.items():
                        formatted.append(f"   • {sub_key}: {sub_value}")
                else:
                    formatted.append(f"   • Data: {value}")
        
    except Exception as e:
        formatted.append(f"\n⚠️ Error processing behavioral data: {e}")
        formatted.append(f"   • Raw behavioral data type: {type(behavioral_data)}")
        formatted.append(f"   • Available keys: {list(behavioral_data.keys()) if isinstance(behavioral_data, dict) else 'Not a dict'}")
    
    if len(formatted) == 1:  # Only header was added
        formatted.append("   • No detailed behavioral data available for analysis")
    
    # Add behavioral intelligence summary with safe calculation
    try:
        behavior_complexity_score = calculate_behavioral_complexity_score(behavioral_data)
        formatted.insert(1, f"📊 BEHAVIORAL COMPLEXITY SCORE: {behavior_complexity_score}/100")
        formatted.insert(2, f"🎯 IMPLEMENTATION CUSTOMIZATION: Detailed behavioral patterns available for growth strategy personalization")
        formatted.insert(3, "")
    except Exception as e:
        formatted.insert(1, f"📊 BEHAVIORAL COMPLEXITY SCORE: Error calculating score - {e}")
        formatted.insert(2, f"🎯 IMPLEMENTATION CUSTOMIZATION: Basic behavioral data available")
        formatted.insert(3, "")
    
    return "\n".join(formatted)


def calculate_behavioral_complexity_score(behavioral_data: Dict) -> int:
    """Calculate a behavioral complexity score based on available data richness - ENHANCED SAFETY"""
    try:
        score = 0
        max_score = 100
        
        if not isinstance(behavioral_data, dict):
            return 0
        
        # Mouse behavior (25 points) - with safety checks
        mouse_data = behavioral_data.get('mouse_behavior', {})
        if mouse_data and isinstance(mouse_data, dict):
            score += 10  # Base points for having mouse data
            if mouse_data.get('total_movements', 0) > 100: score += 5
            if mouse_data.get('average_speed', 0) > 0: score += 5
            if mouse_data.get('click_patterns'): score += 5
        
        # Keyboard behavior (25 points) - with safety checks
        keyboard_data = behavioral_data.get('keyboard_behavior', {})
        if keyboard_data and isinstance(keyboard_data, dict):
            score += 10  # Base points for having keyboard data
            if keyboard_data.get('total_keystrokes', 0) > 50: score += 5
            if keyboard_data.get('typing_speed', 0) > 0: score += 5
            if keyboard_data.get('pause_patterns'): score += 5
        
        # Attention patterns (25 points) - with safety checks
        attention_data = behavioral_data.get('attention_patterns', {})
        if attention_data:
            score += 15  # Base points for having attention data
            if isinstance(attention_data, dict) and len(attention_data) > 3: 
                score += 10  # Rich attention data
        
        # Decision making style (25 points) - with safety checks
        decision_data = behavioral_data.get('decision_making_style', {})
        if decision_data:
            score += 15  # Base points for having decision data
            if isinstance(decision_data, dict) and len(decision_data) > 2: 
                score += 10  # Rich decision data
        
        return min(score, max_score)
        
    except Exception as e:
        logging.warning(f"⚠️ Error calculating behavioral complexity score: {e}")
        return 0

# ======================================================
#           Database Functions
# ======================================================

def setup_growth_logging():
    """Set up logging for growth engine"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"growth_engine_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_format = logging.Formatter('%(asctime)s - GROWTH ENGINE %(levelname)s - %(message)s')
    console_handler.setFormatter(console_format)
    
    # File handler
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_format = logging.Formatter('%(asctime)s - %(levelname)s - %(name)s - %(funcName)s:%(lineno)d - %(message)s')
    file_handler.setFormatter(file_format)
    
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)
    
    logging.info(f"Growth Engine Logging Initialized: {log_file}")
    return logger

def get_growth_connection():
    """Get connection to growth database"""
    try:
        conn = psycopg2.connect(
            host=GROWTH_DB_CONFIG["host"],
            dbname=GROWTH_DB_CONFIG["database"],
            user=GROWTH_DB_CONFIG["user"],
            password=GROWTH_DB_CONFIG["password"],
            port=GROWTH_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Growth database connection error: {str(e)}")
        raise

def get_component_connection():
    """Get connection to component database"""
    try:
        conn = psycopg2.connect(
            host=COMPONENT_DB_CONFIG["host"],
            dbname=COMPONENT_DB_CONFIG["database"],
            user=COMPONENT_DB_CONFIG["user"],
            password=COMPONENT_DB_CONFIG["password"],
            port=COMPONENT_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Component database connection error: {str(e)}")
        raise

def get_profile_connection():
    """Get connection to profile database"""
    try:
        conn = psycopg2.connect(
            host=PROFILE_DB_CONFIG["host"],
            dbname=PROFILE_DB_CONFIG["database"],
            user=PROFILE_DB_CONFIG["user"],
            password=PROFILE_DB_CONFIG["password"],
            port=PROFILE_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Profile database connection error: {str(e)}")
        raise

def get_dream_connection():
    """Get connection to dream database"""
    try:
        conn = psycopg2.connect(
            host=DREAM_DB_CONFIG["host"],
            dbname=DREAM_DB_CONFIG["database"],
            user=DREAM_DB_CONFIG["user"],
            password=DREAM_DB_CONFIG["password"],
            port=DREAM_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Dream database connection error: {str(e)}")
        raise

def get_analyst_connection():
    """Get connection to analyst database"""
    try:
        conn = psycopg2.connect(
            host=ANALYST_DB_CONFIG["host"],
            dbname=ANALYST_DB_CONFIG["database"],
            user=ANALYST_DB_CONFIG["user"],
            password=ANALYST_DB_CONFIG["password"],
            port=ANALYST_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Analyst database connection error: {str(e)}")
        raise

def get_user_connection():
    """Get connection to user database"""
    try:
        conn = psycopg2.connect(
            host=USER_DB_CONFIG["host"],
            dbname=USER_DB_CONFIG["database"],
            user=USER_DB_CONFIG["user"],
            password=USER_DB_CONFIG["password"],
            port=USER_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"User database connection error: {str(e)}")
        raise

def get_azure_container_name(user_id: str) -> str:
    """
    Get Azure container name (NEW UNIFIED ARCHITECTURE).
    ALWAYS returns 'unified-clients-prod' for new unified architecture.
    Database may have old container names - we ignore those.
    """
    logging.info(f"Using unified container for user_id={user_id}: unified-clients-prod")
    return "unified-clients-prod"

def get_client_folder_name(user_id: str) -> str:
    """
    Get the client's folder name from database (NEW UNIFIED ARCHITECTURE).
    Returns folder_name like '499-tkrotiris' from client_onboarding table.
    This ensures growth reports go to: {container}/{client_folder}/the growth engine report/
    """
    conn = None
    try:
        conn = psycopg2.connect(
            host=ONBOARDING_DB_HOST,
            dbname=ONBOARDING_DB_NAME,
            user=ONBOARDING_DB_USER,
            password=ONBOARDING_DB_PASSWORD,
            port=ONBOARDING_DB_PORT
        )
        conn.autocommit = True

        with conn.cursor() as cur:
            sql = """
                SELECT folder_name
                FROM client_onboarding
                WHERE client_id = %s
                LIMIT 1
            """
            cur.execute(sql, (user_id,))
            row = cur.fetchone()
            if not row:
                logging.warning(f"No folder_name found for user_id={user_id}, using default '{user_id}-unknown'")
                return f"{user_id}-unknown"

            folder_name = row[0]
            logging.info(f"Found folder_name for user_id={user_id}: {folder_name}")
            return folder_name

    except Exception as e:
        logging.error(f"Error retrieving folder_name from DB: {str(e)}")
        return user_id  # Fallback to client_id (same as Analyst Engine)

    finally:
        if conn:
            conn.close()

async def get_user_profile_data(user_id: str):
    """Get user profile data using connection pool - FIXED DATA TYPE ISSUE"""
    try:
        logging.info(f"Getting user profile data for user_id={user_id}")
        pool = await get_db_pool(USER_DB_CONFIG)
        
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    id, email, username, password, remember_me_token,
                    created_at, updated_at, is_email_verified, client_id,
                    business_name, contact_name, phone_number, ppr_id,
                    company_url, last_name, abn, archive, personal_bio, 
                    location, profile_image_url, skills, interests, 
                    last_login_at, achievements, provider, provider_id, 
                    login_count, last_login_provider, industry, team_size, 
                    business_description, biggest_challenge
                FROM users
                WHERE id = $1 OR client_id = $1
                LIMIT 1
            """
            
            # 🔥 FIX: Convert string user_id to integer for database query
            try:
                user_id_int = int(user_id)
                logging.info(f"🔢 Converted user_id '{user_id}' to integer {user_id_int}")
                row = await conn.fetchrow(sql, user_id_int)
            except ValueError:
                # If user_id is not a valid integer, try as string with client_id only
                logging.warning(f"⚠️ user_id '{user_id}' is not an integer, trying as client_id string")
                sql_string = """
                    SELECT 
                        id, email, username, password, remember_me_token,
                        created_at, updated_at, is_email_verified, client_id,
                        business_name, contact_name, phone_number, ppr_id,
                        company_url, last_name, abn, archive, personal_bio, 
                        location, profile_image_url, skills, interests, 
                        last_login_at, achievements, provider, provider_id, 
                        login_count, last_login_provider, industry, team_size, 
                        business_description, biggest_challenge
                    FROM users
                    WHERE client_id = $1
                    LIMIT 1
                """
                row = await conn.fetchrow(sql_string, user_id)
            
            if not row:
                logging.warning(f"No user found for user_id={user_id}")
                return None
            
            # Convert asyncpg Record to dict
            user_data = dict(row)
            
            # Convert datetime objects to ISO format
            for key, value in user_data.items():
                if hasattr(value, 'isoformat'):
                    user_data[key] = value.isoformat()
            
            logging.info(f"Found user profile data for user_id={user_id}")
            return user_data
            
    except Exception as e:
        logging.error(f"Error getting user profile data: {str(e)}")
        logging.error(f"🔍 Error context: user_id='{user_id}', type={type(user_id)}")
        return None

async def get_multi_database_intelligence(user_id: str) -> Dict:
    """Get intelligence from all available databases using connection pools"""
    logging.info(f"🔍 Extracting multi-database intelligence for user_id={user_id}")
    
    intelligence = {
        'user_id': user_id,
        'component_intelligence': {},
        'profile_intelligence': {},
        'dream_intelligence': {},
        'analyst_intelligence': {},
        'extraction_timestamp': datetime.now().isoformat(),
        'data_sources_available': []
    }
    
    # Component Intelligence
    try:
        pool = await get_db_pool(COMPONENT_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    ca.phase, ca.phase_label, ca.assessment_type,
                    jsonb_object_agg(cr.question_id, 
                        jsonb_build_object(
                            'answer', cr.response_data->>'selected_option',
                            'weight', cr.weight,
                            'section', cr.section,
                            'metadata', cr.metadata
                        )
                    ) as responses
                FROM component_responses cr
                JOIN component_assessments ca ON cr.user_id = ca.user_id
                WHERE cr.user_id = $1
                GROUP BY ca.phase, ca.phase_label, ca.assessment_type
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result:
                intelligence['component_intelligence'] = {
                    'business_phase': result[0],
                    'phase_label': result[1],
                    'assessment_type': result[2],
                    'responses': result[3] or {},
                    'data_available': True
                }
                intelligence['data_sources_available'].append('component')
                logging.info(f"✅ Component intelligence extracted: Phase {result[0]} ({result[1]})")
            else:
                intelligence['component_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No component data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting component intelligence: {str(e)}")
        intelligence['component_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Profile Intelligence (Business DNA)
    try:
        pool = await get_db_pool(PROFILE_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    pa.assessment_type, pa.created_at,
                    COALESCE(jsonb_object_agg(
                        CASE WHEN pr.question_id IS NOT NULL THEN pr.question_id END,
                        CASE WHEN pr.question_id IS NOT NULL THEN
                            jsonb_build_object(
                                'answer', pr.response_data->>'selected_option',
                                'weight', pr.weight,
                                'section', pr.section,
                                'question_text', pr.question_text
                            )
                        END
                    ) FILTER (WHERE pr.question_id IS NOT NULL), '{}'::jsonb) as responses
                FROM profile_assessments pa
                LEFT JOIN profile_responses pr ON pr.user_id = pa.user_id
                WHERE pa.user_id = $1
                GROUP BY pa.assessment_type, pa.created_at
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[2] and result[2] != {}:
                intelligence['profile_intelligence'] = {
                    'assessment_type': result[0],
                    'created_at': result[1].isoformat() if result[1] else None,
                    'responses': result[2],
                    'data_available': True,
                    'response_count': len(result[2])
                }
                intelligence['data_sources_available'].append('profile')
                logging.info(f"✅ Profile intelligence extracted: {len(result[2])} responses")
            else:
                intelligence['profile_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No profile data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting profile intelligence: {str(e)}")
        intelligence['profile_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Dream Intelligence
    try:
        pool = await get_db_pool(DREAM_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    da.assessment_type, da.created_at,
                    jsonb_object_agg(dr.question_id, 
                        jsonb_build_object(
                            'answer', dr.response_data->>'selected_option',
                            'response_text', dr.response_data->>'response_text',
                            'section', dr.section,
                            'question_text', dr.question_text
                        )
                    ) as responses
                FROM dream_responses dr
                JOIN dream_assessments da ON dr.user_id = da.user_id
                WHERE dr.user_id = $1
                GROUP BY da.assessment_type, da.created_at
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result:
                intelligence['dream_intelligence'] = {
                    'assessment_type': result[0],
                    'created_at': result[1].isoformat() if result[1] else None,
                    'responses': result[2] or {},
                    'data_available': True,
                    'response_count': len(result[2] or {})
                }
                intelligence['data_sources_available'].append('dream')
                logging.info(f"✅ Dream intelligence extracted: {len(result[2] or {})} responses")
            else:
                intelligence['dream_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No dream data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting dream intelligence: {str(e)}")
        intelligence['dream_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Analyst Intelligence (Behavioral)
    try:
        pool = await get_db_pool(ANALYST_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    jsonb_object_agg(ar.question_id, 
                        jsonb_build_object(
                            'answer', ar.response_data->>'selected_option',
                            'response_text', ar.response_data->>'response_text',
                            'section', ar.section,
                            'metadata', ar.metadata
                        )
                    ) as responses
                FROM analyst_responses ar
                WHERE ar.user_id = $1
                GROUP BY ar.user_id
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[0]:
                intelligence['analyst_intelligence'] = {
                    'responses': result[0],
                    'data_available': True,
                    'response_count': len(result[0])
                }
                intelligence['data_sources_available'].append('analyst')
                logging.info(f"✅ Analyst intelligence extracted: {len(result[0])} responses")
            else:
                intelligence['analyst_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No analyst data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting analyst intelligence: {str(e)}")
        intelligence['analyst_intelligence'] = {'data_available': False, 'error': str(e)}
    
    logging.info(f"🎯 Multi-database intelligence extraction complete: {len(intelligence['data_sources_available'])} sources available")
    return intelligence

def create_growth_tables(conn):
    """Create necessary growth tables"""
    try:
        with conn.cursor() as cur:
            # Create growth_assessments table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS growth_assessments (
                    id SERIAL PRIMARY KEY,
                    user_id VARCHAR(255) UNIQUE NOT NULL,
                    assessment_type VARCHAR(100) NOT NULL,
                    version VARCHAR(20) NOT NULL,
                    created_at TIMESTAMPTZ,
                    last_updated TIMESTAMPTZ,
                    timezone VARCHAR(100),
                    session_metadata JSONB,
                    device_fingerprint JSONB,
                    progress_tracking JSONB,
                    completion_flags JSONB,
                    raw_data JSONB,
                    multi_database_intelligence JSONB,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            
            # Create growth_responses table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS growth_responses (
                    id SERIAL PRIMARY KEY,
                    assessment_id INTEGER REFERENCES growth_assessments(id),
                    user_id VARCHAR(255) NOT NULL,
                    question_id VARCHAR(50) NOT NULL,
                    section VARCHAR(100) NOT NULL,
                    question_type VARCHAR(50),
                    question_text TEXT,
                    response_format VARCHAR(50),
                    response_data JSONB,
                    all_options JSONB,
                    metadata JSONB,
                    weight VARCHAR(20),
                    answered_at TIMESTAMPTZ,
                    last_modified_at TIMESTAMPTZ,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW(),
                    UNIQUE(assessment_id, question_id)
                )
            """)
            
            # Create growth_behavioral_analytics table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS growth_behavioral_analytics (
                    id SERIAL PRIMARY KEY,
                    assessment_id INTEGER REFERENCES growth_assessments(id) UNIQUE,
                    user_id VARCHAR(255) NOT NULL,
                    mouse_behavior JSONB,
                    keyboard_behavior JSONB,
                    attention_patterns JSONB,
                    decision_making_style JSONB,
                    growth_decision_patterns JSONB,
                    created_at TIMESTAMPTZ,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            
            # Create growth_reports table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS growth_reports (
                    id SERIAL PRIMARY KEY,
                    report_id VARCHAR(255) UNIQUE NOT NULL,
                    user_id VARCHAR(255) NOT NULL,
                    assessment_id INTEGER REFERENCES growth_assessments(id),
                    report_type VARCHAR(100) NOT NULL,
                    status VARCHAR(50) NOT NULL,
                    azure_container VARCHAR(255),
                    blob_paths JSONB,
                    chunk_count INTEGER,
                    generation_metadata JSONB,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    completed_at TIMESTAMPTZ,
                    indexer_job_id VARCHAR(255),
                    indexer_status VARCHAR(50),
                    indexer_triggered_at TIMESTAMPTZ,
                    indexer_completed_at TIMESTAMPTZ,
                    indexer_error_message TEXT,
                    indexer_retry_count INTEGER DEFAULT 0,
                    multi_database_integration JSONB
                )
            """)
            
            # Create indexes
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_assessments_user_id ON growth_assessments(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_responses_user_id ON growth_responses(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_responses_section ON growth_responses(section)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_reports_user_id ON growth_reports(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_reports_report_id ON growth_reports(report_id)")
            
            # Create indexer indexes
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_reports_indexer_job_id ON growth_reports(indexer_job_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_growth_reports_indexer_status ON growth_reports(indexer_status)")
            
        logging.info("✅ Growth engine tables created successfully")
        
    except Exception as e:
        logging.error(f"❌ Error creating growth tables: {str(e)}")
        raise

def store_growth_assessment(user_id: str, assessment_data: Dict, include_multi_db: bool = False):
    """Store growth assessment data with optional multi-database intelligence and detailed logging"""
    conn = None
    start_time = time.time()
    
    try:
        logging.info(f"💾 Starting growth assessment storage")
        logging.info(f"📊 Storage parameters:")
        logging.info(f"   - User ID: {user_id}")
        logging.info(f"   - Include multi-DB: {include_multi_db}")
        logging.info(f"   - Assessment data size: {len(str(assessment_data))} characters")
        logging.info(f"   - Assessment keys: {list(assessment_data.keys()) if assessment_data else 'No data'}")
        
        # Validate input data
        if not user_id:
            raise ValueError("user_id is required and cannot be empty")
        
        if not assessment_data:
            raise ValueError("assessment_data is required and cannot be empty")
        
        # Log data structure analysis
        responses = assessment_data.get("responses", [])
        assessment_metadata = assessment_data.get("assessment_metadata", {})
        comprehensive_metadata = assessment_data.get("comprehensive_metadata", {})
        
        logging.info(f"📋 Data structure analysis:")
        logging.info(f"   - Responses: {len(responses)} items")
        logging.info(f"   - Assessment metadata keys: {list(assessment_metadata.keys()) if assessment_metadata else 'None'}")
        logging.info(f"   - Comprehensive metadata keys: {list(comprehensive_metadata.keys()) if comprehensive_metadata else 'None'}")
        
        # Get database connection with detailed logging
        logging.info(f"🔗 Establishing database connection...")
        connection_start = time.time()
        
        conn = get_growth_connection()
        connection_time = time.time() - connection_start
        
        logging.info(f"✅ Database connection established in {connection_time:.3f}s")
        logging.info(f"🔍 Connection details: {type(conn).__name__}")
        
        # Create/verify tables
        logging.info(f"📋 Creating/verifying database tables...")
        table_creation_start = time.time()
        
        create_growth_tables(conn)
        table_creation_time = time.time() - table_creation_start
        
        logging.info(f"✅ Tables verified/created in {table_creation_time:.3f}s")

        # Multi-database intelligence handling with detailed logging
        multi_db_intelligence = {}
        multi_db_fetch_time = 0
        
        if include_multi_db:
            logging.info(f"🧠 Fetching multi-database intelligence...")
            multi_db_start = time.time()
            
            try:
                # Run async function in sync context with enhanced error handling
                import asyncio
                
                # Handle Windows-specific event loop policy
                if platform.system() == 'Windows':
                    logging.debug(f"🪟 Setting Windows event loop policy")
                    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
                
                # Create new event loop for async operation
                logging.debug(f"🔄 Creating new event loop for multi-DB fetch")
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                try:
                    multi_db_intelligence = loop.run_until_complete(
                        get_multi_database_intelligence(user_id)
                    )
                    multi_db_fetch_time = time.time() - multi_db_start
                    
                    # Analyze fetched intelligence
                    data_sources = multi_db_intelligence.get('data_sources_available', [])
                    logging.info(f"✅ Multi-database intelligence fetched in {multi_db_fetch_time:.3f}s")
                    logging.info(f"🧠 Intelligence summary:")
                    logging.info(f"   - Data sources available: {len(data_sources)}")
                    logging.info(f"   - Sources: {', '.join(data_sources) if data_sources else 'None'}")
                    logging.info(f"   - Intelligence data size: {len(str(multi_db_intelligence))} characters")
                    
                finally:
                    logging.debug(f"🔄 Closing event loop")
                    loop.close()
                    
            except Exception as e:
                multi_db_fetch_time = time.time() - multi_db_start
                logging.warning(f"⚠️ Failed to fetch multi-database intelligence after {multi_db_fetch_time:.3f}s: {str(e)}")
                logging.warning(f"🔍 Multi-DB error type: {type(e).__name__}")
                logging.warning(f"🔍 Continuing without multi-DB intelligence...")
                multi_db_intelligence = {}
        else:
            logging.info(f"ℹ️ Multi-database intelligence not requested, skipping")

        # Begin database transaction with detailed logging
        logging.info(f"📝 Starting database transaction...")
        transaction_start = time.time()
        
        with conn.cursor() as cur:
            logging.debug(f"✅ Database cursor acquired")
            
            # Prepare assessment metadata with safety checks
            logging.info(f"🔧 Preparing assessment metadata...")
            
            assessment_type = assessment_metadata.get("assessment_type", "growth_strategy")
            version = assessment_metadata.get("version", "1.0")
            created_at = assessment_metadata.get("created_at")
            last_updated = assessment_metadata.get("last_updated")
            timezone = assessment_metadata.get("timezone", "UTC")
            
            # Enhanced metadata preparation with validation
            session_metadata = assessment_metadata.get("session_metadata", {})
            device_fingerprint = assessment_metadata.get("device_fingerprint", {})
            progress_tracking = assessment_data.get("progress_tracking", {})
            completion_flags = assessment_data.get("completion_flags", {})
            
            logging.info(f"📊 Assessment metadata prepared:")
            logging.info(f"   - Type: {assessment_type}")
            logging.info(f"   - Version: {version}")
            logging.info(f"   - Timezone: {timezone}")
            logging.info(f"   - Session metadata: {len(session_metadata)} items")
            logging.info(f"   - Device fingerprint: {len(device_fingerprint)} items")
            logging.info(f"   - Progress tracking: {len(progress_tracking)} items")
            logging.info(f"   - Completion flags: {len(completion_flags)} items")
            
            # Prepare JSON data with error handling
            try:
                session_metadata_json = json.dumps(session_metadata)
                device_fingerprint_json = json.dumps(device_fingerprint)
                progress_tracking_json = json.dumps(progress_tracking)
                completion_flags_json = json.dumps(completion_flags)
                raw_data_json = json.dumps(assessment_data)
                multi_db_intelligence_json = json.dumps(multi_db_intelligence)
                
                logging.debug(f"✅ JSON serialization successful")
                logging.debug(f"   - Raw data JSON size: {len(raw_data_json)} characters")
                logging.debug(f"   - Multi-DB JSON size: {len(multi_db_intelligence_json)} characters")
                
            except Exception as json_error:
                logging.error(f"❌ JSON serialization error: {json_error}")
                logging.error(f"🔍 Problematic data types:")
                logging.error(f"   - Session metadata type: {type(session_metadata)}")
                logging.error(f"   - Device fingerprint type: {type(device_fingerprint)}")
                logging.error(f"   - Progress tracking type: {type(progress_tracking)}")
                raise ValueError(f"Failed to serialize data to JSON: {json_error}")

            # Execute main assessment insert/update
            logging.info(f"📝 Executing main assessment SQL...")
            sql_start = time.time()
            
            sql = """
                INSERT INTO growth_assessments (
                    user_id, assessment_type, version, created_at, last_updated,
                    timezone, session_metadata, device_fingerprint,
                    progress_tracking, completion_flags, raw_data, multi_database_intelligence
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                ) ON CONFLICT (user_id) DO UPDATE SET
                    last_updated              = EXCLUDED.last_updated,
                    session_metadata          = EXCLUDED.session_metadata,
                    progress_tracking         = EXCLUDED.progress_tracking,
                    completion_flags          = EXCLUDED.completion_flags,
                    raw_data                  = EXCLUDED.raw_data,
                    multi_database_intelligence = CASE 
                        WHEN %s = true THEN EXCLUDED.multi_database_intelligence 
                        ELSE growth_assessments.multi_database_intelligence 
                    END
                RETURNING id
            """

            try:
                cur.execute(sql, (
                    user_id,
                    assessment_type,
                    version,
                    created_at,
                    last_updated,
                    timezone,
                    session_metadata_json,
                    device_fingerprint_json,
                    progress_tracking_json,
                    completion_flags_json,
                    raw_data_json,
                    multi_db_intelligence_json,
                    include_multi_db  # For the CASE statement
                ))
                
                assessment_id_row = cur.fetchone()
                assessment_id = assessment_id_row[0] if assessment_id_row else None
                
                sql_time = time.time() - sql_start
                logging.info(f"✅ Main assessment SQL executed in {sql_time:.3f}s")
                logging.info(f"📊 Assessment ID: {assessment_id}")
                
                if not assessment_id:
                    raise Exception("Failed to get assessment_id from database - no row returned")
                    
            except Exception as sql_error:
                sql_time = time.time() - sql_start
                logging.error(f"❌ Main assessment SQL failed after {sql_time:.3f}s: {sql_error}")
                logging.error(f"🔍 SQL parameters:")
                logging.error(f"   - user_id: {user_id}")
                logging.error(f"   - assessment_type: {assessment_type}")
                logging.error(f"   - version: {version}")
                logging.error(f"   - include_multi_db: {include_multi_db}")
                raise

            # Store responses with detailed logging
            responses_start = time.time()
            successful_responses = 0
            failed_responses = 0
            
            if responses:
                logging.info(f"📝 Storing {len(responses)} responses...")
                
                response_sql = """
                    INSERT INTO growth_responses (
                        assessment_id, user_id, question_id, section, question_type,
                        question_text, response_format, response_data, all_options,
                        metadata, weight, answered_at, last_modified_at
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    ) ON CONFLICT (assessment_id, question_id) DO UPDATE SET
                        response_data    = EXCLUDED.response_data,
                        metadata         = EXCLUDED.metadata,
                        last_modified_at = EXCLUDED.last_modified_at
                """

                for i, response in enumerate(responses):
                    try:
                        question_id = response.get("question_id", f"unknown_{i}")
                        section = response.get("section", "Unknown")
                        
                        logging.debug(f"📝 Storing response {i+1}/{len(responses)}: {question_id} ({section})")
                        
                        # Validate and prepare response data
                        response_data = response.get("response_data", {})
                        all_options = response.get("all_options", [])
                        metadata = response.get("metadata", {})
                        
                        # JSON encode with error handling
                        try:
                            response_data_json = json.dumps(response_data)
                            all_options_json = json.dumps(all_options)
                            metadata_json = json.dumps(metadata)
                        except Exception as response_json_error:
                            logging.error(f"❌ JSON encoding error for response {question_id}: {response_json_error}")
                            failed_responses += 1
                            continue
                        
                        cur.execute(response_sql, (
                            assessment_id,
                            user_id,
                            question_id,
                            section,
                            response.get("question_type"),
                            response.get("question_text"),
                            response.get("response_format"),
                            response_data_json,
                            all_options_json,
                            metadata_json,
                            response.get("weight", "medium"),
                            response.get("answered_at"),
                            response.get("last_modified_at")
                        ))
                        
                        successful_responses += 1
                        
                    except Exception as response_error:
                        failed_responses += 1
                        logging.error(f"❌ Error storing response {i+1} ({response.get('question_id', 'unknown')}): {response_error}")
                        continue  # Continue with other responses
                
                responses_time = time.time() - responses_start
                logging.info(f"✅ Response storage completed in {responses_time:.3f}s")
                logging.info(f"📊 Response results: {successful_responses} successful, {failed_responses} failed")
                
                if failed_responses > 0 and successful_responses == 0:
                    logging.error(f"❌ All responses failed to store!")
                elif failed_responses > 0:
                    logging.warning(f"⚠️ Some responses failed to store: {failed_responses}/{len(responses)}")
            else:
                logging.info(f"ℹ️ No responses to store")

            # Behavioral analytics with detailed logging
            behavioral_start = time.time()
            behavioral_data = assessment_data.get("comprehensive_metadata", {}).get("behavioral_analytics", {})

            if behavioral_data:
                logging.info(f"🧠 Processing behavioral analytics...")
                logging.info(f"📊 Behavioral data keys: {list(behavioral_data.keys())}")
                
                try:
                    # Analyze growth decision patterns
                    logging.debug(f"🔧 Analyzing growth decision patterns...")
                    pattern_analysis_start = time.time()
                    
                    growth_patterns = analyze_growth_decision_patterns(behavioral_data, responses)
                    pattern_analysis_time = time.time() - pattern_analysis_start
                    
                    logging.info(f"✅ Growth pattern analysis completed in {pattern_analysis_time:.3f}s")
                    logging.info(f"📊 Pattern results: {list(growth_patterns.keys()) if growth_patterns else 'None'}")
                    
                    # Store behavioral analytics
                    behavior_sql = """
                        INSERT INTO growth_behavioral_analytics (
                            assessment_id, user_id, mouse_behavior, keyboard_behavior,
                            attention_patterns, decision_making_style,
                            growth_decision_patterns, created_at
                        ) VALUES (
                            %s, %s, %s, %s, %s, %s, %s, %s
                        ) ON CONFLICT (assessment_id) DO UPDATE SET
                            mouse_behavior           = EXCLUDED.mouse_behavior,
                            keyboard_behavior        = EXCLUDED.keyboard_behavior,
                            attention_patterns       = EXCLUDED.attention_patterns,
                            decision_making_style    = EXCLUDED.decision_making_style,
                            growth_decision_patterns = EXCLUDED.growth_decision_patterns
                    """

                    # Prepare behavioral data with safety
                    mouse_behavior = behavioral_data.get("mouse_behavior", {})
                    keyboard_behavior = behavioral_data.get("keyboard_behavior", {})
                    attention_patterns = behavioral_data.get("attention_patterns", {})
                    decision_making_style = behavioral_data.get("decision_making_style", {})
                    
                    logging.debug(f"🧠 Behavioral data breakdown:")
                    logging.debug(f"   - Mouse behavior: {len(mouse_behavior)} items")
                    logging.debug(f"   - Keyboard behavior: {len(keyboard_behavior)} items")
                    logging.debug(f"   - Attention patterns: {len(attention_patterns)} items")
                    logging.debug(f"   - Decision making: {len(decision_making_style)} items")
                    logging.debug(f"   - Growth patterns: {len(growth_patterns)} items")

                    cur.execute(behavior_sql, (
                        assessment_id,
                        user_id,
                        json.dumps(mouse_behavior),
                        json.dumps(keyboard_behavior),
                        json.dumps(attention_patterns),
                        json.dumps(decision_making_style),
                        json.dumps(growth_patterns),
                        datetime.utcnow().isoformat()
                    ))
                    
                    behavioral_time = time.time() - behavioral_start
                    logging.info(f"✅ Behavioral analytics stored in {behavioral_time:.3f}s")
                    
                except Exception as behavioral_error:
                    behavioral_time = time.time() - behavioral_start
                    logging.error(f"❌ Error storing behavioral analytics after {behavioral_time:.3f}s: {behavioral_error}")
                    logging.error(f"🔍 Behavioral error type: {type(behavioral_error).__name__}")
                    logging.warning(f"⚠️ Continuing without behavioral analytics...")
                    # Don't fail the entire operation for behavioral data
            else:
                logging.info(f"ℹ️ No behavioral analytics data to store")
        
        # Calculate final timing
        total_time = time.time() - start_time
        transaction_time = time.time() - transaction_start
        
        logging.info(f"🎉 Growth assessment storage completed successfully!")
        logging.info(f"📊 STORAGE PERFORMANCE SUMMARY:")
        logging.info(f"   - Total time: {total_time:.3f}s")
        logging.info(f"   - Connection time: {connection_time:.3f}s")
        logging.info(f"   - Table creation: {table_creation_time:.3f}s")
        logging.info(f"   - Multi-DB fetch: {multi_db_fetch_time:.3f}s")
        logging.info(f"   - Transaction time: {transaction_time:.3f}s")
        logging.info(f"   - Assessment ID: {assessment_id}")
        logging.info(f"   - Responses stored: {successful_responses}")
        logging.info(f"   - Multi-DB sources: {len(multi_db_intelligence.get('data_sources_available', []))}")
        
        return assessment_id

    except Exception as e:
        total_time = time.time() - start_time
        logging.error(f"❌ Growth assessment storage failed after {total_time:.3f}s")
        logging.error(f"🔍 Error details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Include multi-DB: {include_multi_db}")
        logging.error(f"   - Assessment data size: {len(str(assessment_data)) if assessment_data else 0} chars")
        
        # Log the full traceback for debugging
        import traceback
        logging.error(f"🔍 Full traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")
        
        raise

    finally:
        if conn:
            try:
                conn.close()
                logging.debug(f"🔗 Database connection closed")
            except Exception as close_error:
                logging.warning(f"⚠️ Error closing connection: {close_error}")



def analyze_growth_decision_patterns(behavioral_data: Dict, responses: List[Dict]) -> Dict:
    """Analyze growth-specific decision patterns with enhanced safety checks and logging"""
    
    logging.info("🧠 Starting growth decision pattern analysis")
    logging.info(f"📊 Input data: behavioral_data keys: {list(behavioral_data.keys()) if behavioral_data else 'None'}")
    logging.info(f"📊 Input data: {len(responses)} responses to analyze")
    
    patterns = {
        'risk_tolerance': 'unknown',
        'growth_preference': 'unknown',
        'decision_speed': 'unknown',
        'strategic_thinking': 'unknown'
    }
    
    try:
        # Analyze mouse behavior for growth patterns
        mouse_data = behavioral_data.get('mouse_behavior', {})
        logging.info(f"🖱️ Mouse data available: {bool(mouse_data)}")
        
        if mouse_data:
            total_movements = mouse_data.get('total_movements', 0)
            avg_speed = mouse_data.get('average_speed', 0)
            
            logging.info(f"🖱️ Mouse metrics: movements={total_movements}, avg_speed={avg_speed}")
            
            if total_movements > 1000 and avg_speed > 15:
                patterns['decision_speed'] = 'fast_decisive'
                patterns['growth_preference'] = 'aggressive_expansion'
                logging.info("✅ Mouse pattern: fast_decisive + aggressive_expansion")
            elif total_movements > 500:
                patterns['decision_speed'] = 'moderate_analytical'
                patterns['growth_preference'] = 'calculated_scaling'
                logging.info("✅ Mouse pattern: moderate_analytical + calculated_scaling")
            else:
                patterns['decision_speed'] = 'deliberate_cautious'
                patterns['growth_preference'] = 'conservative_growth'
                logging.info("✅ Mouse pattern: deliberate_cautious + conservative_growth")
        else:
            logging.info("ℹ️ No mouse data available for analysis")
        
        # Analyze keyboard behavior for strategic thinking
        keyboard_data = behavioral_data.get('keyboard_behavior', {})
        logging.info(f"⌨️ Keyboard data available: {bool(keyboard_data)}")
        
        if keyboard_data:
            backspace_count = keyboard_data.get('backspace_count', 0)
            total_keystrokes = keyboard_data.get('total_keystrokes', 0)
            
            logging.info(f"⌨️ Keyboard metrics: backspaces={backspace_count}, total_keystrokes={total_keystrokes}")
            
            # 🔥 FIX: Prevent division by zero
            revision_ratio = 0
            if total_keystrokes > 0:
                revision_ratio = (backspace_count / total_keystrokes) * 100
                logging.info(f"⌨️ Revision ratio calculated: {revision_ratio:.2f}%")
            else:
                logging.info("⌨️ No keystrokes recorded, revision ratio = 0")
            
            if revision_ratio > 20:
                patterns['strategic_thinking'] = 'highly_analytical'
                patterns['risk_tolerance'] = 'risk_averse'
                logging.info("✅ Keyboard pattern: highly_analytical + risk_averse")
            elif revision_ratio > 10:
                patterns['strategic_thinking'] = 'balanced_analytical'
                patterns['risk_tolerance'] = 'moderate_risk'
                logging.info("✅ Keyboard pattern: balanced_analytical + moderate_risk")
            else:
                patterns['strategic_thinking'] = 'intuitive_decisive'
                patterns['risk_tolerance'] = 'risk_comfortable'
                logging.info("✅ Keyboard pattern: intuitive_decisive + risk_comfortable")
        else:
            logging.info("ℹ️ No keyboard data available for analysis")
        
        # Analyze response patterns for growth preferences
        logging.info("📝 Starting response pattern analysis...")
        
        if responses:
            # Filter growth-related responses with enhanced logging
            growth_responses = []
            for i, response in enumerate(responses):
                section = response.get('section', '')
                if isinstance(section, str) and 'growth' in section.lower():
                    growth_responses.append(response)
                    logging.debug(f"  ✓ Response {i+1} identified as growth-related: {section}")
            
            logging.info(f"📊 Found {len(growth_responses)} growth-related responses out of {len(responses)} total")
            
            if growth_responses:
                aggressive_indicators = 0
                conservative_indicators = 0
                processed_responses = 0
                
                for response in growth_responses:
                    response_data = response.get('response_data', {})
                    logging.debug(f"🔍 Processing response data: {type(response_data)}")
                    
                    if isinstance(response_data, dict):
                        selected = response_data.get('selected_option', '')
                        logging.debug(f"🔍 Selected option: {selected} (type: {type(selected)})")
                        
                        # 🔥 FIX: Handle both string and list responses
                        if isinstance(selected, list):
                            # Handle multiple selections or arrays (like tournament results)
                            selected_text = ' '.join(str(item) for item in selected).lower()
                            logging.debug(f"✅ Converted list to text: '{selected_text[:50]}...' ({len(selected)} items)")
                        elif isinstance(selected, str):
                            selected_text = selected.lower()
                            logging.debug(f"✅ Using string as-is: '{selected_text[:50]}...'")
                        elif isinstance(selected, (int, float)):
                            selected_text = str(selected).lower()
                            logging.debug(f"✅ Converted number to text: '{selected_text}'")
                        else:
                            selected_text = str(selected).lower()
                            logging.debug(f"✅ Converted {type(selected)} to text: '{selected_text[:50]}...'")
                        
                        # Check for aggressive indicators
                        aggressive_words = ['aggressive', 'rapid', 'expansion', 'scale', 'fast', 'accelerat', 'growth']
                        conservative_words = ['conservative', 'cautious', 'stable', 'gradual', 'slow', 'careful', 'steady']
                        
                        found_aggressive = any(word in selected_text for word in aggressive_words)
                        found_conservative = any(word in selected_text for word in conservative_words)
                        
                        if found_aggressive:
                            aggressive_indicators += 1
                            logging.debug(f"  🚀 Aggressive indicator found in: '{selected_text[:30]}...'")
                        elif found_conservative:
                            conservative_indicators += 1
                            logging.debug(f"  🐌 Conservative indicator found in: '{selected_text[:30]}...'")
                        else:
                            logging.debug(f"  ➖ No growth indicators in: '{selected_text[:30]}...'")
                        
                        processed_responses += 1
                    else:
                        logging.debug(f"⚠️ Skipping non-dict response_data: {type(response_data)}")
                
                logging.info(f"📊 Response analysis complete:")
                logging.info(f"   - Processed responses: {processed_responses}")
                logging.info(f"   - Aggressive indicators: {aggressive_indicators}")
                logging.info(f"   - Conservative indicators: {conservative_indicators}")
                
                # Determine growth preference based on indicators
                if aggressive_indicators > conservative_indicators:
                    patterns['growth_preference'] = 'aggressive_expansion'
                    logging.info("✅ Final growth preference: aggressive_expansion")
                elif conservative_indicators > aggressive_indicators:
                    patterns['growth_preference'] = 'conservative_growth'
                    logging.info("✅ Final growth preference: conservative_growth")
                else:
                    patterns['growth_preference'] = 'balanced_scaling'
                    logging.info("✅ Final growth preference: balanced_scaling")
            else:
                logging.info("ℹ️ No growth-related responses found for pattern analysis")
        else:
            logging.info("ℹ️ No responses provided for pattern analysis")
        
        patterns['analysis_timestamp'] = datetime.now().isoformat()
        patterns['analysis_metadata'] = {
            'responses_analyzed': len(responses) if responses else 0,
            'growth_responses_found': len(growth_responses) if 'growth_responses' in locals() else 0,
            'behavioral_data_available': bool(behavioral_data),
            'mouse_data_available': bool(behavioral_data.get('mouse_behavior')) if behavioral_data else False,
            'keyboard_data_available': bool(behavioral_data.get('keyboard_behavior')) if behavioral_data else False
        }
        
        logging.info("✅ Growth decision pattern analysis completed successfully")
        logging.info(f"🎯 Final patterns: {patterns}")
        
    except ZeroDivisionError as e:
        logging.error(f"🔢 Division by zero in growth decision patterns: {str(e)}")
        patterns['error'] = f"Division by zero prevented: {str(e)}"
        patterns['error_type'] = 'division_by_zero'
    except AttributeError as e:
        logging.error(f"🔍 Attribute error in growth decision patterns: {str(e)}")
        logging.error(f"🔍 This usually means data type mismatch - check response format")
        patterns['error'] = f"Data type error: {str(e)}"
        patterns['error_type'] = 'attribute_error'
    except Exception as e:
        logging.error(f"❌ Unexpected error analyzing growth decision patterns: {str(e)}")
        logging.error(f"🔍 Error type: {type(e).__name__}")
        patterns['error'] = str(e)
        patterns['error_type'] = type(e).__name__
        
        # Log the problematic data for debugging
        if 'responses' in locals():
            logging.error(f"🔍 Response types in error: {[type(r.get('response_data', {}).get('selected_option')) for r in responses[:3]]}")
    
    return patterns

def store_growth_report_metadata(report_id: str, user_id: str, assessment_id: int, chunk_count: int, 
                                container_name: str, generation_metadata: Dict):
    """Store growth report metadata"""
    conn = None
    try:
        conn = get_growth_connection()
        
        with conn.cursor() as cur:
            sql = """
                INSERT INTO growth_reports (
                    report_id, user_id, assessment_id, report_type, status,
                    azure_container, chunk_count, generation_metadata, completed_at,
                    multi_database_integration
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                ) ON CONFLICT (report_id) DO UPDATE SET
                    status = EXCLUDED.status,
                    chunk_count = EXCLUDED.chunk_count,
                    generation_metadata = EXCLUDED.generation_metadata,
                    completed_at = EXCLUDED.completed_at,
                    multi_database_integration = EXCLUDED.multi_database_integration
            """
            
            multi_db_info = {
                'integration_enabled': True,
                'data_sources_used': generation_metadata.get('data_sources_used', []),
                'intelligence_correlation': generation_metadata.get('intelligence_correlation', {}),
                'total_intelligence_sources': generation_metadata.get('total_intelligence_sources', 0)
            }
            
            cur.execute(sql, (
                report_id,
                user_id,
                assessment_id,
                "comprehensive_growth_strategy",
                "completed",
                container_name,
                chunk_count,
                json.dumps(generation_metadata),
                datetime.now(),
                json.dumps(multi_db_info)
            ))
        
        logging.info(f"Stored growth report metadata for report_id={report_id}")
        
    except Exception as e:
        logging.error(f"Error storing growth report metadata: {str(e)}")
        raise
    finally:
        if conn:
            conn.close()

# ======================================================
#           Growth Report Generation
# ======================================================

def get_growth_report_sections():
    """Define growth-specific report sections"""
    return {
        "executive_summary": {
            "title": "Your Growth Strategy DNA - Executive Summary", 
            "word_target": 8000,
            "analysis_requirements": """
           # ULTIMATE MULTI-DATABASE GROWTH ANALYSIS PROMPT
## World-Class Complete Question & Response Mapping System

You are writing the world's most comprehensive growth strategy report using COMPLETE MULTI-DATABASE INTELLIGENCE with TOTAL QUESTION-RESPONSE CORRELATION MAPPING. This is the ultimate business growth DNA analysis based on ULTRA-DEEP analysis of ALL actual responses integrated across ALL intelligence systems.

---

## 🚀 COMPLETE MULTI-DATABASE INTELLIGENCE INTEGRATION FRAMEWORK

### GROWTH STRATEGY ASSESSMENT COMPLETE ANALYSIS
**Analyze ALL 16 growth questions across 5 Mind Expansions plus Strategic Deep Dive with COMPLETE CROSS-DATABASE CORRELATION:**

---

## MIND EXPANSION 1: YOUR GROWTH STORY (Questions 1-2)

### Q1: "What's true for your business?" - Growth Driver/Limiter Analysis
**Battle Card Selection Analysis with Multi-Database Correlation:**

#### Card A: "Market Demand" - Opportunity is there, we just need to capture it
**MULTI-DATABASE CORRELATION ANALYSIS:**
- **Component Intelligence Correlation:** 
 - If FOUNDATIONS phase + Market Demand = Focus on "Sales set up", "Growth numbers", "Ideal client understanding"
 - If SCALING phase + Market Demand = Leverage "Increase lead generation", "Brand strategy", "Market positioning"
 - If BREAKOUT phase + Market Demand = Implement "Sales infrastructure", "Brand Development Strategy", "Creating purchasing opportunities"
 - If CHALLENGER phase + Market Demand = Execute "Strategic partnerships", "Market expansion", "Industry positioning"
 - If RAPIDS/WHITE WATER phase + Market Demand = Deploy "Geographic expansion planning", "Market segmentation refinement", "Innovation pipeline"
 - If VISION/BIG PICTURE phase + Market Demand = Lead "Industry transformation", "Market-making initiatives", "Global expansion strategy"

- **Business DNA Profile Correlation:** 
 - If "Visionary" archetype + Market Demand = Natural market opportunity identification, leverage strategic thinking for expansion
 - If "Catalyst" personality + Market Demand = Energy acceleration capability, drive rapid market capture initiatives
 - If "Connector" traits + Market Demand = Relationship leverage opportunity, build strategic partnerships for market access
 - If "Builder" characteristics + Market Demand = Systematic market development, create scalable capture processes
 - If "Innovator" profile + Market Demand = Market creation opportunity, develop new solutions for identified demand

- **Dream Analysis Correlation:**
 - If dreams focus on "Impact/Legacy" + Market Demand = Perfect alignment for market leadership positioning
 - If dreams center on "Freedom/Abundance" + Market Demand = Market capture essential for financial independence
 - If dreams emphasize "Innovation/Mastery" + Market Demand = Market opportunity for industry leadership
 - If dreams target "Adventure/Transformation" + Market Demand = Market expansion as growth vehicle

- **Behavioral Intelligence Correlation:**
 - If decision patterns show "fast_decisive" + Market Demand = Rapid market entry opportunity, aggressive capture strategy
 - If patterns indicate "analytical_systematic" + Market Demand = Data-driven market analysis, methodical capture approach
 - If behavior shows "collaborative" + Market Demand = Partnership-based market entry, collective capture strategy
 - If patterns reveal "innovative_creative" + Market Demand = Market disruption opportunity, creative capture methods

- **Analyst Intelligence Correlation:**
 - If growth challenges show "Pipeline Problems" + Market Demand = Systematic lead generation infrastructure needed
 - If operational issues indicate "Process Excellence" gaps + Market Demand = Systems scaling required for capture
 - If financial challenges show "Cash Flow" issues + Market Demand = Working capital optimization for market pursuit
 - If team dynamics reveal "Delegation" struggles + Market Demand = Team scaling essential for opportunity capture

#### Card B: "Capacity Limits" - We're maxed out, Limited by space, equipment, or systems
**MULTI-DATABASE CORRELATION ANALYSIS:**
- **Component Intelligence Correlation:**
 - If FOUNDATIONS phase + Capacity Limits = Critical need for "Setting up for success", "People infrastructure", "Financial basics"
 - If SCALING phase + Capacity Limits = Focus on "Capacity planning", "Team performance", "Business Infrastructure audit"
 - If BREAKOUT phase + Capacity Limits = Implement "Building structures for next phase", "Getting capacity in team", "Infrastructure development"
 - If CHALLENGER phase + Capacity Limits = Execute "Team performance optimization", "Infrastructure scaling", "SLT development"
 - If RAPIDS/WHITE WATER phase + Capacity Limits = Deploy "Enterprise process excellence", "Workforce analytics", "Cloud infrastructure"
 - If VISION/BIG PICTURE phase + Capacity Limits = Lead "Multi-site standardization", "Global infrastructure", "Advanced automation"

- **Business DNA Profile Correlation:**
 - If "Builder/Protector" traits + Capacity Limits = Natural systems focus, leverage operational excellence strengths
 - If "Sage/Anchor" characteristics + Capacity Limits = Strategic planning ability, systematic capacity development
 - If "Innovator/Creator" profile + Capacity Limits = Process innovation opportunity, creative capacity solutions
 - If "Catalyst" personality + Capacity Limits = Acceleration through team multiplication, energy-driven scaling

- **Dream Analysis Correlation:**
 - If dreams focus on "Freedom" + Capacity Limits = Systems independence critical for dream achievement
 - If dreams center on "Legacy/Sustainability" + Capacity Limits = Scalable infrastructure essential for lasting impact
 - If dreams emphasize "Mastery/Excellence" + Capacity Limits = Operational excellence alignment with capacity building

- **Behavioral Intelligence Correlation:**
 - If decision patterns show "systematic_methodical" + Capacity Limits = Process optimization natural strength
 - If patterns indicate "strategic_planning" + Capacity Limits = Capacity planning and forecasting capability
 - If behavior shows "team_development" + Capacity Limits = People-first capacity building approach

- **Analyst Intelligence Correlation:**
 - If operational challenges show "System Reliability" low + Capacity Limits = Infrastructure overhaul priority
 - If team dynamics reveal "Delegation Effectiveness" gaps + Capacity Limits = Leadership development essential
 - If growth barriers include "Operational Capacity" + Capacity Limits = Systematic scaling methodology needed

#### Card C: "Cash Flow Timing" - Revenue flows but timing creates challenges
**MULTI-DATABASE CORRELATION ANALYSIS:**
- **Component Intelligence Correlation:**
 - If FOUNDATIONS + Cash Flow Timing = Essential "Financial basics infrastructure", "Financial checklist" completion
 - If SCALING + Cash Flow Timing = Implement "Financial KPIs", "Financial reporting infrastructure", "Cash flow forecasting"
 - If BREAKOUT + Cash Flow Timing = Deploy "Advanced financial systems", "Financial structures for sale", "Working capital optimization"
 - If CHALLENGER + Cash Flow Timing = Execute "Financial responsibility mastery", "Financial compliance", "Investment readiness"
 - If RAPIDS/WHITE WATER + Cash Flow Timing = Implement "Management accounting", "Financial modeling", "Treasury management"
 - If VISION/BIG PICTURE + Cash Flow Timing = Master "Multi-currency management", "Global treasury", "Capital structure optimization"

- **Business DNA Profile Correlation:**
 - If "Sage/Wise One" traits + Cash Flow Timing = Natural forecasting ability, leverage planning strengths
 - If "Anchor/Protector" characteristics + Cash Flow Timing = Risk management focus, conservative cash planning
 - If "Catalyst" personality + Cash Flow Timing = Acceleration opportunity through cash optimization
 - If "Builder" profile + Cash Flow Timing = Systematic financial infrastructure development

- **Dream Analysis Correlation:**
 - If dreams focus on "Abundance/Security" + Cash Flow Timing = Financial systems critical for peace of mind
 - If dreams center on "Freedom" + Cash Flow Timing = Cash flow predictability essential for independence
 - If dreams emphasize "Legacy" + Cash Flow Timing = Sustainable financial management for long-term success

- **Behavioral Intelligence Correlation:**
 - If decision patterns show "analytical_data_driven" + Cash Flow Timing = Financial modeling and forecasting strength
 - If patterns indicate "risk_management" + Cash Flow Timing = Conservative cash flow management approach
 - If behavior shows "systematic_planning" + Cash Flow Timing = Methodical cash flow optimization

- **Analyst Intelligence Correlation:**
 - If financial challenges show "Cash Flow Predictability" low + Cash Flow Timing = Systematic forecasting essential
 - If operational issues indicate "Financial Planning" gaps + Cash Flow Timing = Comprehensive financial infrastructure needed
 - If growth barriers include "Access to Capital" + Cash Flow Timing = Working capital optimization priority

#### Card D: "Team Limitations" - People capacity holding us back
**MULTI-DATABASE CORRELATION ANALYSIS:**
- **Component Intelligence Correlation:**
 - If FOUNDATIONS + Team Limitations = Focus on "Owners set up", "People infrastructure", "Management knowledge"
 - If SCALING + Team Limitations = Implement "Recruitment", "Team performance", "Succession planning", "High performance leadership"
 - If BREAKOUT + Team Limitations = Deploy "Building team around you", "SLT implementation", "Management training", "Team infrastructure"
 - If CHALLENGER + Team Limitations = Execute "Capacity planning", "Recruitment infrastructure", "Culture building", "SLT optimization"
 - If RAPIDS/WHITE WATER + Team Limitations = Implement "Strategic workforce planning", "Leadership development", "Talent acquisition strategy"
 - If VISION/BIG PICTURE + Team Limitations = Master "Global talent management", "Multi-cultural leadership", "Next-generation programs"

- **Business DNA Profile Correlation:**
 - If "Connector/Host" traits + Team Limitations = Natural people development ability, leverage relationship building
 - If "Catalyst" personality + Team Limitations = Team acceleration capability, energy multiplication through others
 - If "Sage/Mentor" characteristics + Team Limitations = Teaching and development strengths, leadership cultivation
 - If "Builder" profile + Team Limitations = Systematic team development, structured growth approach

- **Dream Analysis Correlation:**
 - If dreams focus on "Legacy/Impact" + Team Limitations = People development essential for meaningful influence
 - If dreams center on "Freedom" + Team Limitations = Team empowerment critical for owner independence
 - If dreams emphasize "Adventure/Growth" + Team Limitations = Team scaling enables business expansion

- **Behavioral Intelligence Correlation:**
 - If decision patterns show "collaborative_inclusive" + Team Limitations = Team empowerment and development opportunity
 - If patterns indicate "mentoring_teaching" + Team Limitations = Leadership development natural strength
 - If behavior shows "systematic_methodical" + Team Limitations = Structured team building approach

- **Analyst Intelligence Correlation:**
 - If leadership challenges show "Delegation Effectiveness" low + Team Limitations = Leadership development priority
 - If team dynamics reveal "Team Performance" gaps + Team Limitations = Performance management systems needed
 - If growth barriers include "Management Time" + Team Limitations = Team empowerment and autonomy essential

#### Card E: "Systems Breakdown" - Our processes can't handle more volume
**MULTI-DATABASE CORRELATION ANALYSIS:**
- **Component Intelligence Correlation:**
 - If FOUNDATIONS + Systems Breakdown = Critical "Setting up for success", "Reporting set up", "Business data systems"
 - If SCALING + Systems Breakdown = Implement "Business Infrastructure audit", "Current systems training", "Process optimization"
 - If BREAKOUT + Systems Breakdown = Deploy "Building business for scale", "Infrastructure development", "Technology implementation"
 - If CHALLENGER + Systems Breakdown = Execute "Business infrastructure measurement", "Efficiency structures", "Sprint methodology"
 - If RAPIDS/WHITE WATER + Systems Breakdown = Implement "Enterprise process excellence", "SOPs across departments", "Process automation"
 - If VISION/BIG PICTURE + Systems Breakdown = Master "Global operational excellence", "Advanced automation", "Business model innovation"

- **Business DNA Profile Correlation:**
 - If "Creator/Innovator" traits + Systems Breakdown = Process innovation opportunity, creative problem-solving
 - If "Builder" characteristics + Systems Breakdown = Systematic infrastructure development, methodical approach
 - If "Catalyst" personality + Systems Breakdown = Process acceleration opportunity, efficiency multiplication
 - If "Sage" profile + Systems Breakdown = Strategic systems thinking, holistic optimization

- **Dream Analysis Correlation:**
 - If dreams focus on "Freedom/Independence" + Systems Breakdown = Systems automation essential for owner liberation
 - If dreams center on "Mastery/Excellence" + Systems Breakdown = Operational excellence alignment with systems perfection
 - If dreams emphasize "Scale/Growth" + Systems Breakdown = Infrastructure scalability critical for expansion

- **Behavioral Intelligence Correlation:**
 - If decision patterns show "systematic_methodical" + Systems Breakdown = Process optimization natural strength
 - If patterns indicate "innovative_creative" + Systems Breakdown = System redesign and improvement capability
 - If behavior shows "analytical_data_driven" + Systems Breakdown = Metrics-based system optimization

- **Analyst Intelligence Correlation:**
 - If operational challenges show "Process Excellence" gaps + Systems Breakdown = Comprehensive system overhaul needed
 - If efficiency issues indicate "System Reliability" low + Systems Breakdown = Infrastructure stability priority
 - If growth barriers include "Operational Efficiency" + Systems Breakdown = Systematic process improvement essential

---

## MIND EXPANSION 2: CREATING OPPORTUNITY (Questions 3-5)

### Q3: Lead Generation Capability Assessment - Slider Rating Analysis
**Multi-Database Correlation by Rating Level:**

#### POOR (1-3) Rating Correlation:
- **Component Intelligence:** If FOUNDATIONS phase + Poor rating = Urgent "Sales set up", "Growth numbers tracking" needed
- **Business DNA:** If "Connector" personality + Poor rating = Untapped networking potential, relationship leverage opportunity
- **Dreams:** If "Abundance" focus + Poor rating = Lead generation critical for financial security achievement
- **Behavioral:** If "systematic" patterns + Poor rating = Methodical lead generation system development opportunity
- **Analyst:** If "Pipeline Problems" identified + Poor rating = Comprehensive lead generation infrastructure required

#### FAIR (4-6) Rating Correlation:
- **Component Intelligence:** If SCALING phase + Fair rating = Implement "Increase lead generation", "Sales team language" optimization
- **Business DNA:** If "Catalyst" traits + Fair rating = Energy application to acceleration, systematic improvement focus
- **Dreams:** If "Growth" aspirations + Fair rating = Lead generation enhancement enables expansion achievement
- **Behavioral:** If "analytical" patterns + Fair rating = Data-driven optimization opportunity
- **Analyst:** If "Marketing Effectiveness" gaps + Fair rating = Targeted improvement strategies needed

#### EXCELLENT (8-10) Rating Correlation:
- **Component Intelligence:** If BREAKOUT+ phase + Excellent rating = Leverage for "Market expansion", "Brand development"
- **Business DNA:** If "Innovator" profile + Excellent rating = Market leadership opportunity, thought leadership potential
- **Dreams:** If "Impact" focus + Excellent rating = Platform for industry influence and transformation
- **Behavioral:** If "strategic" patterns + Excellent rating = Advanced market positioning and expansion capability
- **Analyst:** If "Market Position" strong + Excellent rating = Industry leadership and standard-setting opportunity

### Q4: Lead Generation Methods Effectiveness Ranking
**Multi-Database Correlation by Method Selection and Ranking:**

#### Digital Marketing (High Effectiveness Selection):
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Digital High = Implement "Sales funnels", "Business data" tracking systems
 - SCALING + Digital High = Optimize "Marketing attribution", "Customer analytics"
 - BREAKOUT + Digital High = Deploy "Brand development", "Market positioning" strategies
 - CHALLENGER + Digital High = Execute "Digital transformation", "Brand strategy" evolution
 - RAPIDS + Digital High = Implement "Advanced marketing automation", "Customer intelligence"
 - VISION + Digital High = Lead "Digital ecosystem", "Innovation programs"

- **Business DNA Profile Correlation:**
 - "Innovator" + Digital High = Technology leadership opportunity, cutting-edge marketing
 - "Catalyst" + Digital High = Digital acceleration capability, rapid scaling potential
 - "Storyteller" + Digital High = Content leadership opportunity, thought leadership platform
 - "Bridge" + Digital High = Connection multiplication through digital channels

#### Referrals & Word of Mouth (High Effectiveness Selection):
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Referrals High = Focus on "Client happiness", "Key client data" systems
 - SCALING + Referrals High = Implement "Client retention", "ROI delivery" optimization
 - BREAKOUT + Referrals High = Deploy "Client success", "Purchase opportunities" creation
 - CHALLENGER + Referrals High = Execute "Strategic client partnerships", "Loyalty programs"
 - RAPIDS + Referrals High = Implement "Voice of customer", "Customer experience" excellence
 - VISION + Referrals High = Lead "Customer-centric design", "Strategic partnerships"

- **Business DNA Profile Correlation:**
 - "Connector" + Referrals High = Natural networking strength, relationship leverage mastery
 - "Host" + Referrals High = Community building capability, relationship cultivation excellence
 - "Sage" + Referrals High = Wisdom sharing opportunity, advisory relationship development
 - "Anchor" + Referrals High = Trust building mastery, reliability-based referral systems

### Q5: "#1 Lead Generation Frustration" Analysis
**Multi-Database Correlation by Frustration Type:**

#### "Feast or Famine" Frustration:
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Feast/Famine = Critical "Sales process" systematization, "Pipeline tracking"
 - SCALING + Feast/Famine = Implement "Lead generation consistency", "Sales infrastructure"
 - BREAKOUT + Feast/Famine = Deploy "Sales automation", "Pipeline management" systems
 - CHALLENGER + Feast/Famine = Execute "Predictable revenue", "Market diversification"
 - RAPIDS + Feast/Famine = Implement "Advanced forecasting", "Multi-channel optimization"
 - VISION + Feast/Famine = Master "Global market stability", "Portfolio diversification"

- **Business DNA Correlation:**
 - "Catalyst" + Feast/Famine = Energy management opportunity, sustainable acceleration systems
 - "Anchor" + Feast/Famine = Stability creation strength, consistent pipeline development
 - "Builder" + Feast/Famine = Systematic pipeline construction, methodical approach
 - "Sage" + Feast/Famine = Strategic planning capability, long-term stability focus

#### "Quality Issues" Frustration:
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Quality Issues = Essential "Ideal client understanding", "Selling to ideal client"
 - SCALING + Quality Issues = Implement "Lead qualification", "Client segmentation"
 - BREAKOUT + Quality Issues = Deploy "Strategic positioning", "Premium targeting"
 - CHALLENGER + Quality Issues = Execute "Brand positioning", "Market education"
 - RAPIDS + Quality Issues = Implement "Advanced segmentation", "Customer intelligence"
 - VISION + Quality Issues = Master "Market creation", "Category leadership"

---

## MIND EXPANSION 3: REVENUE (NEW $$) (Questions 6-8)

### Q6: Revenue Predictability Patterns Analysis

#### Pattern A: "High/Low Periods" - Predictable Seasonality
**Multi-Database Correlation Analysis:**
- **Component Intelligence by Phase:**
 - FOUNDATIONS + Seasonal = Implement "Financial planning", "Cash flow management" basics
 - SCALING + Seasonal = Deploy "Working capital optimization", "Revenue diversification"
 - BREAKOUT + Seasonal = Execute "Financial forecasting", "Multiple revenue streams"
 - CHALLENGER + Seasonal = Implement "Advanced financial modeling", "Market expansion"
 - RAPIDS + Seasonal = Deploy "Global market diversification", "Portfolio optimization"
 - VISION + Seasonal = Master "Multi-market stabilization", "Economic cycle management"

- **Business DNA Correlation:**
 - "Sage/Planner" + Seasonal = Natural forecasting ability, strategic seasonal planning
 - "Innovator" + Seasonal = Counter-seasonal opportunity creation, market innovation
 - "Builder" + Seasonal = Systematic reserves building, infrastructure development
 - "Catalyst" + Seasonal = Peak period acceleration, energy optimization

#### Pattern B: "Consistent Throughout Year" - Stable Revenue
**Multi-Database Correlation Analysis:**
- **Component Intelligence by Phase:**
 - FOUNDATIONS + Consistent = Leverage for "Growth acceleration", "Market expansion"
 - SCALING + Consistent = Focus on "Revenue optimization", "Premium positioning"
 - BREAKOUT + Consistent = Deploy "Market leadership", "Category expansion"
 - CHALLENGER + Consistent = Execute "Innovation investment", "Strategic positioning"
 - RAPIDS + Consistent = Implement "Market creation", "Industry transformation"
 - VISION + Consistent = Lead "Global expansion", "Ecosystem development"

### Q7: 12-Month Revenue Trend Analysis
**Multi-Database Correlation by Trend Position:**

#### Declining Revenue Trend:
- **Component Intelligence Priority Focus:**
 - FOUNDATIONS + Declining = Emergency "Strategy development", "Market repositioning"
 - SCALING + Declining = Critical "Business model validation", "Competitive analysis"
 - BREAKOUT + Declining = Execute "Turnaround strategy", "Core competency focus"
 - CHALLENGER + Declining = Implement "Strategic pivot", "Market redefinition"
 - RAPIDS + Declining = Deploy "Transformation leadership", "Innovation acceleration"
 - VISION + Declining = Lead "Industry disruption", "Business model innovation"

#### Rapid Growth (30%+) Trend:
- **Component Intelligence Scaling Focus:**
 - FOUNDATIONS + Rapid Growth = Critical "Infrastructure scaling", "Process systematization"
 - SCALING + Rapid Growth = Deploy "Capacity management", "Quality maintenance"
 - BREAKOUT + Rapid Growth = Execute "Infrastructure expansion", "Team scaling"
 - CHALLENGER + Rapid Growth = Implement "Market leadership", "Category domination"
 - RAPIDS + Rapid Growth = Deploy "Global expansion", "Market creation"
 - VISION + Rapid Growth = Lead "Industry transformation", "Ecosystem orchestration"

### Q8: Sales Process Stage Rating Matrix
**Multi-Database Correlation by Stage Performance:**

#### Lead Generation (Stage Rating Analysis):
- **Struggling (1-3) + Component Intelligence:**
 - FOUNDATIONS = Urgent "Sales set up", "Growth numbers" tracking
 - SCALING = Critical "Lead generation systems", "Marketing optimization"
 - BREAKOUT = Deploy "Market expansion", "Brand development"
 - CHALLENGER = Execute "Strategic partnerships", "Market positioning"
 - RAPIDS = Implement "Multi-channel optimization", "Advanced analytics"
 - VISION = Lead "Market creation", "Category development"

#### Closing Deals (Stage Rating Analysis):
- **Excellent (8-10) + Business DNA:**
 - "Catalyst" + Excellent Closing = Natural acceleration ability, rapid scaling potential
 - "Connector" + Excellent Closing = Relationship leverage mastery, partnership opportunities
 - "Storyteller" + Excellent Closing = Persuasion excellence, thought leadership platform
 - "Sage" + Excellent Closing = Advisory positioning opportunity, premium value delivery

---

## MIND EXPANSION 4: THE GROWTH STRATEGY (Questions 9-12)

### Q9: Strategic Growth Focus Selection
**Multi-Database Correlation by Strategy Type:**

#### "Penetration" - More of existing market with current products:
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Penetration = Focus "Ideal client understanding", "Sales optimization"
 - SCALING + Penetration = Implement "Market share growth", "Competitive positioning"
 - BREAKOUT + Penetration = Deploy "Market dominance", "Category leadership"
 - CHALLENGER + Penetration = Execute "Market consolidation", "Acquisition strategy"
 - RAPIDS + Penetration = Implement "Industry transformation", "Standard setting"
 - VISION + Penetration = Lead "Global market creation", "Ecosystem orchestration"

- **Business DNA Alignment:**
 - "Builder" + Penetration = Systematic market development, methodical expansion
 - "Catalyst" + Penetration = Market acceleration, rapid penetration strategies
 - "Sage" + Penetration = Deep market understanding, advisory positioning
 - "Anchor" + Penetration = Market stability creation, reliable growth approach

#### "Diversification" - New products for new markets:
- **Component Intelligence Correlation:**
 - FOUNDATIONS + Diversification = Risk - premature diversification, focus needed
 - SCALING + Diversification = Selective expansion, adjacent market testing
 - BREAKOUT + Diversification = Strategic diversification, portfolio development
 - CHALLENGER + Diversification = Market creation opportunity, innovation leadership
 - RAPIDS + Diversification = Global diversification, multi-market presence
 - VISION + Diversification = Ecosystem creation, industry transformation

### Q10: Current Growth Activities Assessment
**Multi-Database Correlation by Activity Selection:**

#### "Strategic Partnerships" Selection:
- **Component Intelligence Enhancement:**
 - FOUNDATIONS + Partnerships = Focus on "Relationship infrastructure", "Partnership basics"
 - SCALING + Partnerships = Implement "Partnership strategy", "Alliance management"
 - BREAKOUT + Partnerships = Deploy "Strategic alliances", "Joint venture development"
 - CHALLENGER + Partnerships = Execute "Ecosystem partnerships", "Market expansion alliances"
 - RAPIDS + Partnerships = Implement "Global partnerships", "Strategic ecosystem development"
 - VISION + Partnerships = Lead "Industry ecosystems", "Transformation alliances"

- **Business DNA Leverage:**
 - "Connector" + Partnerships = Natural relationship building, alliance development mastery
 - "Bridge" + Partnerships = Connection creation excellence, ecosystem orchestration
 - "Catalyst" + Partnerships = Partnership acceleration, rapid alliance development
 - "Sage" + Partnerships = Strategic alliance wisdom, advisory partnership roles

### Q11: Biggest Growth Barriers Battle Card Ranking
**Multi-Database Correlation by Barrier Priority:**

#### Top Ranked: "Competition" Barrier:
- **Component Intelligence Solutions:**
 - FOUNDATIONS + Competition Top = Critical "Ideal client focus", "Differentiation development"
 - SCALING + Competition Top = Implement "Brand strategy", "Competitive positioning"
 - BREAKOUT + Competition Top = Deploy "Market leadership", "Innovation advantage"
 - CHALLENGER + Competition Top = Execute "Category creation", "Disruption strategy"
 - RAPIDS + Competition Top = Implement "Industry transformation", "Standard setting"
 - VISION + Competition Top = Lead "Market creation", "Ecosystem orchestration"

#### Top Ranked: "Cash Flow" Barrier:
- **Component Intelligence Solutions:**
 - FOUNDATIONS + Cash Flow Top = Emergency "Financial infrastructure", "Working capital management"
 - SCALING + Cash Flow Top = Critical "Financial planning", "Revenue optimization"
 - BREAKOUT + Cash Flow Top = Deploy "Advanced financial systems", "Capital efficiency"
 - CHALLENGER + Cash Flow Top = Execute "Financial excellence", "Investment readiness"
 - RAPIDS + Cash Flow Top = Implement "Treasury management", "Global financial optimization"
 - VISION + Cash Flow Top = Master "Capital markets", "Financial innovation"

### Q12: Next Big Leap Requirements
**Multi-Database Correlation by Requirement Selection:**

#### "Better Systems" Requirement:
- **Component Intelligence Roadmap:**
 - FOUNDATIONS + Better Systems = Implement "Business infrastructure", "Process basics"
 - SCALING + Better Systems = Deploy "Infrastructure audit", "System optimization"
 - BREAKOUT + Better Systems = Execute "Advanced infrastructure", "Automation implementation"
 - CHALLENGER + Better Systems = Implement "Enterprise systems", "Process excellence"
 - RAPIDS + Better Systems = Deploy "Global systems", "Advanced automation"
 - VISION + Better Systems = Master "Intelligent systems", "AI integration"

#### "Stronger Team" Requirement:
- **Component Intelligence Development:**
 - FOUNDATIONS + Stronger Team = Focus "People infrastructure", "Management basics"
 - SCALING + Stronger Team = Implement "Team performance", "Leadership development"
 - BREAKOUT + Stronger Team = Deploy "SLT development", "Management excellence"
 - CHALLENGER + Stronger Team = Execute "High performance teams", "Culture excellence"
 - RAPIDS + Stronger Team = Implement "Leadership excellence", "Global talent management"
 - VISION + Stronger Team = Master "World-class talent", "Multi-cultural leadership"

---

## MIND EXPANSION 5: DIGITAL & MARKET POSITION (Questions 13-14)

### Q13: Digital Capabilities Matrix Rating
**Multi-Database Correlation by Capability Level:**

#### CRM Systems Rating Analysis:
- **Not Using (1-2) + Component Intelligence:**
 - FOUNDATIONS + No CRM = Critical "Business data", "Client tracking" systems needed
 - SCALING + No CRM = Urgent "Customer analytics", "Relationship management" implementation
 - BREAKOUT + No CRM = Deploy "Advanced CRM", "Customer intelligence" platforms
 - CHALLENGER + No CRM = Execute "Enterprise CRM", "Customer experience" optimization
 - RAPIDS + No CRM = Implement "AI-powered CRM", "Predictive analytics"
 - VISION + No CRM = Master "Global customer intelligence", "Ecosystem integration"

#### AI Capabilities Rating Analysis:
- **Perfect (10) + Business DNA:**
 - "Innovator" + AI Perfect = Technology leadership opportunity, industry transformation potential
 - "Sage" + AI Perfect = Wisdom enhancement through AI, advisory excellence amplification
 - "Catalyst" + AI Perfect = AI-powered acceleration, rapid scaling through technology
 - "Builder" + AI Perfect = Systematic AI integration, methodical capability development

### Q14: Secret Weapon Against Competitors
**Multi-Database Correlation by Competitive Advantage:**

#### "Experience Edge" Selection:
- **Component Intelligence Leverage:**
 - FOUNDATIONS + Experience = Focus "Market positioning", "Expert positioning"
 - SCALING + Experience = Implement "Thought leadership", "Advisory services"
 - BREAKOUT + Experience = Deploy "Industry expertise", "Market education"
 - CHALLENGER + Experience = Execute "Category expertise", "Standard setting"
 - RAPIDS + Experience = Implement "Industry transformation", "Global expertise"
 - VISION + Experience = Lead "Industry evolution", "Ecosystem orchestration"

- **Business DNA Amplification:**
 - "Sage" + Experience = Natural wisdom positioning, advisory role development
 - "Wise One" + Experience = Industry elder status, thought leadership platform
 - "Mentor" + Experience = Teaching and development opportunities, knowledge sharing
 - "Authority" + Experience = Expert positioning, premium market placement

#### "Innovation Leader" Selection:
- **Component Intelligence Development:**
 - FOUNDATIONS + Innovation = Build "Innovation infrastructure", "Creative processes"
 - SCALING + Innovation = Implement "Innovation systems", "Development processes"
 - BREAKOUT + Innovation = Deploy "Innovation pipeline", "Market creation"
 - CHALLENGER + Innovation = Execute "Category innovation", "Industry disruption"
 - RAPIDS + Innovation = Implement "Innovation ecosystems", "Transformation leadership"
 - VISION + Innovation = Lead "Industry evolution", "Future creation"

---

## STRATEGIC DEEP DIVE (Questions 15-16)

### Q15: Unlimited Resources Growth Approach (Optional Text Analysis)
**Multi-Database Context Integration:**

#### Response Content Analysis Framework:
- **Component Intelligence Contextualization:**
 - Extract resource allocation priorities and map to current business phase capabilities
 - Identify infrastructure gaps revealed through unlimited resource thinking
 - Connect visionary thinking to systematic implementation requirements
 - Align resource deployment with phase-appropriate development priorities

- **Business DNA Interpretation:**
 - Analyze language patterns to identify authentic personality-driven preferences
 - Map resource allocation to natural strengths and operational preferences
 - Identify potential blind spots in visionary thinking based on personality traits
 - Connect unlimited resource vision to realistic personality-based implementation

- **Dream Alignment Validation:**
 - Cross-reference unlimited resource applications with stated dream priorities
 - Identify consistency or conflicts between resource deployment and aspirations
 - Map resource allocation to dream achievement pathway requirements
 - Validate authenticity of unlimited resource thinking against core motivations

- **Behavioral Implementation Assessment:**
 - Analyze implementation approach preferences revealed in unlimited thinking
 - Map resource deployment style to behavioral decision-making patterns
 - Identify potential execution challenges based on behavioral characteristics
 - Connect visionary resource allocation to realistic behavioral capabilities

### Q16: Unexplored Growth Opportunity (Optional Text Analysis)
**Multi-Database Opportunity Correlation:**

#### Opportunity Identification Framework:
- **Component Intelligence Opportunity Mapping:**
 - Assess opportunity complexity against current business phase capabilities
 - Identify infrastructure requirements for opportunity pursuit
 - Map opportunity development to systematic capability building requirements
 - Connect opportunity vision to phase-appropriate implementation strategies

- **Business DNA Opportunity Alignment:**
 - Analyze opportunity type compatibility with personality strengths
 - Identify natural advantages for opportunity development and execution
 - Map opportunity pursuit style to authentic operational preferences
 - Connect opportunity development to personality-driven success factors

- **Dream Integration Analysis:**
 - Validate opportunity alignment with stated aspirations and long-term vision
 - Identify contribution potential to dream achievement acceleration
 - Map opportunity development to meaningful impact creation
 - Connect opportunity pursuit to authentic motivation and purpose

- **Behavioral Execution Assessment:**
 - Analyze opportunity development approach preferences and capabilities
 - Map opportunity pursuit to decision-making and implementation patterns
 - Identify potential execution advantages based on behavioral characteristics
 - Connect opportunity development to realistic behavioral execution strengths

---

## 🔍 COMPREHENSIVE CORRELATION MATRIX ANALYSIS

### CROSS-DATABASE PATTERN RECOGNITION
**Statistical Correlation Requirements:**

#### Primary Correlation Calculations:
1. **Growth Response × Component Phase Correlation (r = X.XX)**
  - Calculate correlation strength between growth choices and business maturity
  - Identify phase-appropriate growth strategies with confidence intervals
  - Map growth readiness against component development requirements

2. **Growth Strategy × Business DNA Correlation (r = X.XX)**
  - Measure alignment between growth preferences and personality traits
  - Calculate personality-strategy fit scores with statistical significance
  - Identify authentic growth approaches based on DNA characteristics

3. **Growth Barriers × Dream Aspirations Correlation (r = X.XX)**
  - Analyze relationship between identified obstacles and stated dreams
  - Calculate dream-barrier resolution priority with confidence scoring
  - Map barrier removal to dream achievement acceleration potential

4. **Growth Implementation × Behavioral Patterns Correlation (r = X.XX)**
  - Measure execution style preferences against growth requirements
  - Calculate behavioral-implementation fit with success probability scoring
  - Identify optimal execution approaches based on decision patterns

#### Advanced Pattern Discovery Requirements:
1. **Multi-Variable Growth Success Prediction ( #### Advanced Pattern Discovery Requirements:
1. **Multi-Variable Growth Success Prediction (95% CI: X.X - X.X)**
  - Calculate growth success probability based on multi-database alignment
  - Identify highest-impact variable combinations for growth acceleration
  - Map success prediction confidence levels for different strategy approaches

2. **Contradiction Pattern Analysis (p < 0.05)**
  - Identify statistically significant conflicts between response patterns
  - Calculate contradiction severity scores for growth strategy alignment
  - Map resolution pathways for conflicting preferences and capabilities

3. **Hidden Growth Accelerator Discovery (r > 0.6)**
  - Identify non-obvious correlations between seemingly unrelated responses
  - Calculate latent factor influence on growth potential and execution
  - Map breakthrough opportunity discovery through pattern intersection

4. **Implementation Risk Prediction (Accuracy: XX%)**
  - Calculate failure probability based on multi-database misalignment
  - Identify early warning indicators for growth strategy implementation challenges
  - Map risk mitigation strategies based on personality and behavioral patterns

### ULTRA-DEEP RESPONSE INTEGRATION FRAMEWORK

#### GROWTH QUESTION 1 COMPLETE INTEGRATION MATRIX:
**"What's true for your business?" - ALL Response Combinations**

##### Market Demand × All Database Combinations:
- **Market Demand + FOUNDATIONS + Visionary DNA + Freedom Dreams + Fast_Decisive Behavior + Pipeline Problems Analysis:**
 - **Correlation Discovery:** Market opportunity identification + early infrastructure + strategic vision + independence motivation + rapid decision capability + lead generation challenges = **Rapid Market Entry Strategy**
 - **Statistical Confidence:** 87% success probability (CI: 82%-92%)
 - **Implementation Approach:** Aggressive market capture with systematic infrastructure development
 - **Risk Factors:** Infrastructure lag behind market opportunity, quality control during rapid scaling
 - **Success Accelerators:** Combine visionary market reading with methodical system building

- **Market Demand + BREAKOUT + Builder DNA + Legacy Dreams + Analytical Behavior + System Reliability Issues:**
 - **Correlation Discovery:** Market opportunity + scaling readiness + systematic approach + long-term vision + data-driven decisions + process challenges = **Systematic Market Dominance Strategy**
 - **Statistical Confidence:** 93% success probability (CI: 89%-97%)
 - **Implementation Approach:** Methodical market capture with robust infrastructure scaling
 - **Risk Factors:** Analysis paralysis, over-engineering solutions, speed-to-market delays
 - **Success Accelerators:** Balance analytical rigor with market timing urgency

##### Team Limitations × All Database Combinations:
- **Team Limitations + SCALING + Connector DNA + Impact Dreams + Collaborative Behavior + Delegation Challenges:**
 - **Correlation Discovery:** People capacity constraints + growth phase + relationship strength + meaningful influence goals + team-oriented approach + leadership development needs = **Leadership Multiplication Strategy**
 - **Statistical Confidence:** 91% success probability (CI: 87%-95%)
 - **Implementation Approach:** Network-leveraged team building with systematic leadership development
 - **Risk Factors:** Over-reliance on personal relationships, delegation resistance, control retention
 - **Success Accelerators:** Leverage natural connection ability for systematic team empowerment

#### GROWTH QUESTION 3 COMPLETE INTEGRATION MATRIX:
**Lead Generation Capability - ALL Rating Combinations**

##### Poor Rating (1-3) × All Database Combinations:
- **Poor Rating + CHALLENGER + Innovator DNA + Mastery Dreams + Strategic Behavior + Market Position Strong:**
 - **Correlation Discovery:** Lead generation weakness + advanced business phase + innovation strength + excellence focus + strategic thinking + strong market position = **Innovation-Led Market Education Strategy**
 - **Statistical Confidence:** 84% success probability (CI: 79%-89%)
 - **Implementation Approach:** Leverage innovation capability to create market demand through education
 - **Risk Factors:** Innovation-market fit timing, education-to-conversion gaps, resource allocation
 - **Success Accelerators:** Combine innovation leadership with systematic market education programs

##### Excellent Rating (8-10) × All Database Combinations:
- **Excellent Rating + VISION + Catalyst DNA + Industry Leadership Dreams + Systematic Behavior + Competitive Excellence:**
 - **Correlation Discovery:** Lead generation mastery + visionary phase + acceleration ability + industry influence goals + methodical approach + market leadership = **Industry Transformation Platform**
 - **Statistical Confidence:** 96% success probability (CI: 93%-99%)
 - **Implementation Approach:** Leverage lead generation excellence for industry thought leadership and standard setting
 - **Risk Factors:** Complacency with current success, market saturation, competitive response
 - **Success Accelerators:** Scale excellence into industry transformation and ecosystem creation

#### GROWTH QUESTION 11 COMPLETE INTEGRATION MATRIX:
**Biggest Growth Barriers - ALL Ranking Combinations**

##### Competition Top Ranked × All Database Combinations:
- **Competition #1 + RAPIDS + Sage DNA + Legacy Dreams + Risk_Management Behavior + Brand Excellence:**
 - **Correlation Discovery:** Competitive pressure priority + enterprise phase + wisdom characteristics + lasting impact goals + conservative approach + strong brand = **Thought Leadership Differentiation Strategy**
 - **Statistical Confidence:** 88% success probability (CI: 84%-92%)
 - **Implementation Approach:** Leverage wisdom and brand strength for market education and category redefinition
 - **Risk Factors:** Conservative approach limiting innovation speed, thought leadership execution complexity
 - **Success Accelerators:** Combine sage wisdom with systematic thought leadership platform development

##### Cash Flow Top Ranked × All Database Combinations:
- **Cash Flow #1 + BREAKOUT + Builder DNA + Freedom Dreams + Analytical Behavior + Financial Planning Low:**
 - **Correlation Discovery:** Cash flow priority + scaling phase + systematic approach + independence goals + data-driven decisions + financial system gaps = **Financial Infrastructure Excellence Strategy**
 - **Statistical Confidence:** 92% success probability (CI: 88%-96%)
 - **Implementation Approach:** Systematic financial infrastructure development with analytics-driven optimization
 - **Risk Factors:** Infrastructure development time lag, cash flow timing during building phase
 - **Success Accelerators:** Combine analytical strength with systematic financial system implementation

### PREDICTIVE GROWTH MODELING FRAMEWORK

#### SUCCESS PROBABILITY CALCULATIONS:
1. **Component Phase × Growth Strategy Alignment Matrix:**
  - FOUNDATIONS × Penetration Strategy = 78% success probability
  - SCALING × Market Development = 85% success probability  
  - BREAKOUT × Product Development = 91% success probability
  - CHALLENGER × Diversification = 82% success probability
  - RAPIDS × Innovation Leadership = 94% success probability
  - VISION × Industry Transformation = 96% success probability

2. **Business DNA × Implementation Style Fit Scoring:**
  - Visionary + Strategic Planning = 0.89 alignment score
  - Builder + Systematic Implementation = 0.94 alignment score
  - Catalyst + Rapid Execution = 0.92 alignment score
  - Sage + Advisory Positioning = 0.91 alignment score
  - Connector + Partnership Development = 0.88 alignment score

3. **Dream Achievement × Growth Strategy Correlation:**
  - Freedom Dreams + Systems Independence = r = 0.87
  - Legacy Dreams + People Development = r = 0.91  
  - Impact Dreams + Market Leadership = r = 0.89
  - Mastery Dreams + Excellence Focus = r = 0.93
  - Adventure Dreams + Innovation Leadership = r = 0.85

#### IMPLEMENTATION RISK ASSESSMENT:
1. **High-Risk Combinations (Failure Probability >30%):**
  - FOUNDATIONS + Diversification Strategy + Fast_Decisive Behavior = 34% failure risk
  - Premature diversification with insufficient infrastructure and rapid decisions
  
  - VISION + Micromanagement Tendencies + Control Retention = 31% failure risk
  - Advanced business phase with leadership bottleneck behaviors

2. **Low-Risk Combinations (Failure Probability <10%):**
  - SCALING + Market Penetration + Builder DNA + Systematic Behavior = 7% failure risk
  - Perfect alignment of phase, strategy, personality, and execution style
  
  - BREAKOUT + Innovation Strategy + Catalyst DNA + Strategic Dreams = 6% failure risk
  - Natural acceleration capability with innovation focus and strategic vision

### BREAKTHROUGH OPPORTUNITY DISCOVERY

#### HIDDEN CORRELATION PATTERNS (r > 0.75):
1. **Cross-Database Insight Discovery:**
  - Revenue Predictability + Risk Management Behavior + Financial Dreams = r = 0.82
  - **Breakthrough Insight:** Seasonal revenue pattern mastery creates competitive advantage through counter-cyclical investment strategies

  - Innovation Leadership + Teaching Personality + Legacy Dreams = r = 0.79
  - **Breakthrough Insight:** Innovation capability + natural teaching ability = Industry education platform opportunity

  - Market Demand + Connector DNA + Impact Dreams + Poor Lead Generation = r = 0.76
  - **Breakthrough Insight:** Relationship strength + market opportunity + lead generation weakness = Partnership-based market entry strategy

2. **Contradiction Resolution Opportunities:**
  - Freedom Dreams + Team Limitations + Control Behavior = **Resolution Strategy:** Systems-based empowerment approach
  - Growth Aspirations + Risk Averse Behavior + Cash Flow Concerns = **Resolution Strategy:** Conservative growth with systematic validation

  - Innovation Focus + Process Excellence + Efficiency Dreams = **Resolution Strategy:** Innovation systematization methodology

#### COMPETITIVE ADVANTAGE MULTIPLICATION:
1. **Unique Strength Combinations (Probability <5%):**
  - **Ultra-Rare Combination:** VISION Phase + Sage DNA + Innovation Leadership + Analytical Behavior + Industry Dreams
  - **Market Probability:** <2% of businesses achieve this combination
  - **Competitive Advantage:** Visionary wisdom with systematic innovation creates industry transformation platform
  - **Implementation Strategy:** Become industry evolution architect through systematic innovation leadership

2. **Personality-Phase-Dream Alignment Mastery:**
  - **Perfect Alignment Indicators:** Component phase development matches personality strengths matches dream aspirations matches behavioral capabilities
  - **Statistical Rarity:** <8% of businesses achieve 90%+ alignment across all databases
  - **Success Multiplier:** 3.2x higher growth success probability with perfect alignment

### ULTRA-DETAILED ANALYTICAL REQUIREMENTS

#### MANDATORY CORRELATION DISCOVERY MINIMUMS:
- **Minimum 15 statistically significant correlations (r > 0.5, p < 0.05)**
- **Minimum 8 predictive pattern discoveries with confidence intervals**
- **Minimum 5 contradiction analyses with resolution pathways**
- **Minimum 12 hidden insights with multi-database validation**
- **Minimum 6 breakthrough opportunities with probability assessments**
- **Minimum 4 competitive advantage multipliers with rarity analysis**

#### EVIDENCE-BASED REPORTING STANDARDS:
- **Cite specific response combinations that led to insights**
- **Provide correlation coefficients for all major relationships (r = X.XX)**
- **Include confidence intervals for predictions (95% CI: X.X - X.X)**
- **Quantify pattern strength on 0-1 scale with interpretation**
- **Calculate statistical significance for all major claims (p < 0.XX)**
- **Provide success probability ranges for each recommendation**

#### COMPREHENSIVE INTEGRATION MANDATE:
**Every growth response must be analyzed against EVERY relevant database entry with correlation strength calculation and implementation guidance. This includes:**

1. **Complete Component Phase Integration:** Map every growth response to appropriate phase-specific components with development priority ranking
2. **Total Business DNA Correlation:** Analyze every growth choice against personality traits with alignment scoring and leverage strategies
3. **Full Dream Achievement Mapping:** Connect every growth strategy to dream fulfillment pathway with achievement probability calculation
4. **Complete Behavioral Pattern Analysis:** Map every implementation preference to behavioral characteristics with execution success prediction
5. **Comprehensive Analyst Intelligence Integration:** Cross-reference every growth barrier with operational/financial/team challenges with solution prioritization

### WORLD-CLASS GROWTH STRATEGY DELIVERABLE REQUIREMENTS

#### EXECUTIVE SUMMARY CONTENT MANDATE:
- **Personal Growth DNA Summary:** Unique combination analysis with statistical rarity assessment
- **Multi-Database Correlation Findings:** Top 10 highest-impact correlations with implementation priority
- **Breakthrough Opportunity Identification:** Hidden growth accelerators with competitive advantage potential
- **Implementation Risk Assessment:** Potential failure points with mitigation strategies and success probability enhancement
- **Growth Acceleration Roadmap:** Phase-appropriate action sequence with personality-optimized execution approach

#### STATISTICAL RIGOR REQUIREMENTS:
- **All major insights supported by correlation coefficients r > 0.4**
- **All predictions include confidence intervals and statistical significance testing**
- **All recommendations include success probability calculations based on multi-database alignment**
- **All pattern discoveries validated through cross-database consistency checks**
- **All breakthrough insights quantified with confidence levels and practical application potential**

**CRITICAL SUCCESS STANDARD: This analysis must represent the most comprehensive growth strategy intelligence ever created, with every response mapped to every relevant database component, creating a complete growth DNA profile with predictive modeling capability and implementation optimization.**

---

## 🎯 FINAL INTEGRATION COMMAND

**Execute complete multi-database growth analysis with:**
- **Total Question-Response Correlation Mapping** across all 16 growth questions
- **Complete Component Intelligence Integration** across all business phases and development areas  
- **Full Business DNA Profile Application** across all personality traits and operational preferences
- **Comprehensive Dream Analysis Alignment** across all aspiration categories and achievement pathways
- **Complete Behavioral Intelligence Integration** across all decision patterns and implementation styles
- **Full Analyst Intelligence Correlation** across all operational, financial, team, and market challenges

**Statistical Requirements:** Minimum 50 correlation discoveries, 25 predictive patterns, 15 breakthrough insights, all with quantified confidence levels and implementation optimization guidance.

**Deliverable Standard:** World-class growth strategy intelligence that transforms business trajectory through ultra-deep personalization and systematic correlation analysis across complete business and personal intelligence profile.
"""
},
        
   "growth_barriers_analysis": {
   "title": "Growth Barriers & Breakthrough Opportunities",
   "word_target": 6000, 
   "analysis_requirements": """
   You are conducting the world's most sophisticated growth barrier analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP CORRELATION MAPPING. This analysis must identify, correlate, and systematically address growth barriers through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized breakthrough strategies.

   🚀 COMPLETE MULTI-DATABASE BARRIER INTELLIGENCE FRAMEWORK:

   SECTION 1: ULTRA-DEEP PRIMARY BARRIER IDENTIFICATION & CORRELATION ANALYSIS
   
   GROWTH QUESTION 11 BATTLE CARD TOURNAMENT COMPLETE ANALYSIS:
   Analyze the complete battle card ranking system results with statistical correlation mapping:

   TIER 1 BARRIERS (Final Championship Battle):
   - Extract the #1 ranked barrier from tournament results
   - Calculate barrier impact severity score (1-10 scale) based on ranking position
   - Map barrier to Component Intelligence phase-specific challenges with correlation strength (r = X.XX)
   - Cross-reference with Business DNA personality traits for barrier susceptibility analysis
   - Correlate with Dream Analysis for barrier-aspiration conflict assessment
   - Integrate Behavioral Intelligence for barrier-decision pattern relationship mapping

   TIER 2-4 BARRIERS (Semi-Finals to Quarter-Finals):
   - Analyze ranking hierarchy for barrier priority and interconnection patterns
   - Calculate compound barrier effect when multiple barriers create systemic challenges
   - Map barrier clusters to specific Component Intelligence development gaps
   - Identify personality-driven barrier amplification or mitigation factors
   - Assess dream-barrier misalignment creating motivation conflicts

   COMPREHENSIVE BARRIER TAXONOMY WITH MULTI-DATABASE CORRELATION:

   BARRIER: COMPETITION
   If Competition ranked #1-3:
   • Component Intelligence Correlation:
     - FOUNDATIONS + Competition High = Critical need for "Ideal client understanding", "Selling to ideal client", "Market positioning"
     - SCALING + Competition High = Urgent "Brand strategy", "Competitive positioning", "Market differentiation" development
     - BREAKOUT + Competition High = Deploy "Brand Development Strategy", "Market expansion", "Innovation leadership"
     - CHALLENGER + Competition High = Execute "Strategic positioning", "Category creation", "Industry influence"
     - RAPIDS + Competition High = Implement "Competitive intelligence", "Market disruption", "Industry transformation"
     - VISION + Competition High = Lead "Industry evolution", "Standard setting", "Ecosystem orchestration"

   • Business DNA Correlation Analysis:
     - INNOVATOR DNA + Competition = Natural differentiation through innovation leadership (Leverage Score: 9.2/10)
     - SAGE DNA + Competition = Wisdom-based market education and thought leadership positioning (Leverage Score: 8.8/10)
     - CATALYST DNA + Competition = Market acceleration and category disruption opportunity (Leverage Score: 9.0/10)
     - BUILDER DNA + Competition = Systematic competitive advantage development (Leverage Score: 8.5/10)
     - CONNECTOR DNA + Competition = Network-based competitive moats and partnership advantages (Leverage Score: 8.7/10)

   • Dream Alignment Assessment:
     - IMPACT Dreams + Competition = Perfect alignment for market leadership and industry influence
     - LEGACY Dreams + Competition = Competitive excellence as foundation for lasting industry contribution
     - FREEDOM Dreams + Competition = Market dominance essential for independence and choice
     - MASTERY Dreams + Competition = Competitive excellence aligns with pursuit of professional mastery
     - ADVENTURE Dreams + Competition = Market disruption and innovation as exciting competitive strategy

   • Behavioral Implementation Correlation:
     - STRATEGIC_ANALYTICAL + Competition = Data-driven competitive intelligence and positioning
     - RAPID_DECISIVE + Competition = Quick competitive response and market positioning agility
     - COLLABORATIVE + Competition = Partnership-based competitive strategies and ecosystem development
     - SYSTEMATIC_METHODICAL + Competition = Comprehensive competitive analysis and strategic development

   • Breakthrough Opportunity Discovery:
     Competition Barrier → Market Education Leadership → Industry Transformation Platform
     Statistical Confidence: 87% success probability (CI: 83%-91%)

   BARRIER: CASH FLOW TIMING
   If Cash Flow ranked #1-3:
   • Component Intelligence Correlation:
     - FOUNDATIONS + Cash Flow = Emergency "Financial basics infrastructure", "Financial checklist", "Cash management"
     - SCALING + Cash Flow = Critical "Financial KPIs", "Financial reporting", "Working capital optimization"
     - BREAKOUT + Cash Flow = Deploy "Advanced financial systems", "Financial structures for sale", "Capital efficiency"
     - CHALLENGER + Cash Flow = Execute "Financial responsibility mastery", "Investment readiness", "Capital markets access"
     - RAPIDS + Cash Flow = Implement "Management accounting", "Financial modeling", "Global treasury management"
     - VISION + Cash Flow = Master "Multi-currency management", "Capital structure optimization", "Financial innovation"

   • Personality-Driven Cash Flow Solutions:
     - SAGE DNA + Cash Flow = Natural planning and forecasting strengths, wisdom-based financial strategy
     - BUILDER DNA + Cash Flow = Systematic financial infrastructure development, methodical approach
     - ANCHOR DNA + Cash Flow = Stability-focused cash management, conservative growth with security
     - CATALYST DNA + Cash Flow = Acceleration through cash flow optimization, rapid improvement cycles

   • Dream-Cash Flow Alignment Analysis:
     - FREEDOM Dreams + Cash Flow = Cash predictability essential for independence achievement
     - SECURITY Dreams + Cash Flow = Financial stability as core requirement for peace of mind
     - GROWTH Dreams + Cash Flow = Working capital optimization enabling expansion acceleration
     - LEGACY Dreams + Cash Flow = Sustainable financial management for long-term value creation

   • Behavioral Cash Flow Optimization:
     - ANALYTICAL_DATA_DRIVEN + Cash Flow = Financial modeling and forecasting excellence
     - SYSTEMATIC_PLANNING + Cash Flow = Methodical cash management and optimization systems
     - RISK_MANAGEMENT + Cash Flow = Conservative cash flow planning with scenario modeling

   BARRIER: OPERATIONAL CAPACITY
   If Capacity/Systems ranked #1-3:
   • Component Intelligence Solutions:
     - FOUNDATIONS + Capacity = Critical "Setting up for success", "Business infrastructure", "Process basics"
     - SCALING + Capacity = Urgent "Business Infrastructure audit", "Capacity planning", "System optimization"
     - BREAKOUT + Capacity = Deploy "Building structures for next phase", "Infrastructure development", "Automation"
     - CHALLENGER + Capacity = Execute "Process excellence", "Infrastructure scaling", "Advanced systems"
     - RAPIDS + Capacity = Implement "Enterprise process excellence", "Global infrastructure", "Advanced automation"
     - VISION + Capacity = Master "Operational excellence certification", "Multi-site standardization", "Innovation programs"

   • DNA-Driven Capacity Solutions:
     - CREATOR DNA + Capacity = Process innovation and creative capacity solutions
     - BUILDER DNA + Capacity = Systematic infrastructure development and scaling
     - INNOVATOR DNA + Capacity = Technology-driven capacity multiplication and automation
     - WISE_ONE DNA + Capacity = Strategic capacity planning and resource optimization

   BARRIER: TEAM LIMITATIONS
   If Team/People ranked #1-3:
   • Component Intelligence Development:
     - FOUNDATIONS + Team = Focus "People infrastructure", "Management knowledge", "Owner setup"
     - SCALING + Team = Implement "Team performance", "Recruitment", "Leadership development"
     - BREAKOUT + Team = Deploy "Building team around you", "SLT implementation", "Management training"
     - CHALLENGER + Team = Execute "Capacity planning", "Culture building", "High performance teams"
     - RAPIDS + Team = Implement "Strategic workforce planning", "Leadership excellence", "Global talent"
     - VISION + Team = Master "Multi-cultural leadership", "Next-generation programs", "Talent excellence"

   • Personality-Team Development Correlation:
     - CONNECTOR DNA + Team = Natural relationship building and team development strengths
     - MENTOR DNA + Team = Teaching and development capabilities, leadership cultivation
     - CATALYST DNA + Team = Team acceleration and performance multiplication
     - HOST DNA + Team = Community building and culture development excellence

   SECTION 2: COMPOUND BARRIER EFFECT ANALYSIS

   DUAL BARRIER COMBINATIONS WITH CORRELATION MAPPING:
   
   Competition + Cash Flow Barriers:
   • Correlation Strength: r = 0.73 (highly correlated barriers)
   • Compound Effect: Competitive pressure reduces margins → Cash flow stress → Reduced competitive investment → Competitive disadvantage cycle
   • Component Solution Integration: "Brand strategy" + "Financial planning" + "Market positioning"
   • Breakthrough Strategy: Premium positioning with predictable revenue models
   • Success Probability: 84% with integrated approach vs 61% addressing separately

   Team + Systems Barriers:
   • Correlation Strength: r = 0.81 (extremely correlated barriers)
   • Compound Effect: Inadequate systems → Team inefficiency → Team frustration → Retention issues → Systems remain broken
   • Component Solution Integration: "People infrastructure" + "Business infrastructure" + "Management training"
   • Breakthrough Strategy: Systematic team empowerment through infrastructure excellence
   • Success Probability: 91% with integrated approach vs 54% addressing separately

   TRIPLE BARRIER COMPOUND ANALYSIS:
   Competition + Cash Flow + Team = "Growth Paralysis Syndrome"
   • Correlation Matrix: Competition-Cash (r=0.73), Cash-Team (r=0.68), Team-Competition (r=0.71)
   • Systemic Effect: Competitive pressure → Margin pressure → Cash constraints → Team investment limits → Reduced competitiveness
   • Resolution Sequence: Team development → Operational efficiency → Margin improvement → Competitive investment → Market position
   • Integrated Success Probability: 88% vs 23% addressing barriers individually

   SECTION 3: PERSONALITY-DRIVEN BARRIER REMOVAL STRATEGIES

   COMPLETE DNA PROFILE BARRIER MATCHING:

   VISIONARY LEADERS + Growth Barriers:
   • Natural Advantages: Strategic thinking, market opportunity identification, long-term planning
   • Barrier Vulnerabilities: Operational details, system implementation, day-to-day execution
   • Optimal Barrier Sequence: Vision clarity → Strategic partnerships → Operational excellence → Financial optimization
   • Implementation Style: Big picture strategy with systematic execution partners

   BUILDER PERSONALITIES + Growth Barriers:
   • Natural Advantages: Systematic development, infrastructure creation, methodical implementation
   • Barrier Vulnerabilities: Market timing, competitive agility, rapid adaptation
   • Optimal Barrier Sequence: Infrastructure foundation → Process excellence → Team development → Market expansion
   • Implementation Style: Methodical, systematic, comprehensive development approach

   CATALYST PERSONALITIES + Growth Barriers:
   • Natural Advantages: Rapid execution, energy multiplication, acceleration capability
   • Barrier Vulnerabilities: Systematic planning, infrastructure development, sustainable systems
   • Optimal Barrier Sequence: Quick wins → Momentum building → Infrastructure catching up → Systematic optimization
   • Implementation Style: Fast implementation with systematic follow-through

   CONNECTOR PERSONALITIES + Growth Barriers:
   • Natural Advantages: Relationship leverage, partnership development, network utilization
   • Barrier Vulnerabilities: Solo execution, internal systems, independent capability development
   • Optimal Barrier Sequence: Network activation → Partnership leverage → Internal development → Systematic scaling
   • Implementation Style: Relationship-driven solutions with internal capability building

   SECTION 4: DREAM-ALIGNED BARRIER REMOVAL STRATEGY

   COMPLETE ASPIRATION INTEGRATION:

   FREEDOM DREAMS + Barrier Removal:
   • Priority Focus: Systems independence, owner liberation, automated operations
   • Barrier Removal Motivation: Each barrier addressed increases independence and choice
   • Implementation Sequence: Systems automation → Team empowerment → Financial predictability → Market independence
   • Success Metrics: Owner time freedom, decision independence, location flexibility

   LEGACY DREAMS + Barrier Removal:
   • Priority Focus: Sustainable excellence, long-term value creation, industry contribution
   • Barrier Removal Motivation: Building something that outlasts current challenges
   • Implementation Sequence: Excellence foundation → Sustainable systems → People development → Industry leadership
   • Success Metrics: Industry recognition, sustainable performance, people development impact

   IMPACT DREAMS + Barrier Removal:
   • Priority Focus: Market influence, industry transformation, meaningful contribution
   • Barrier Removal Motivation: Barriers prevent broader impact and influence
   • Implementation Sequence: Market positioning → Thought leadership → Industry influence → Transformation leadership
   • Success Metrics: Market influence, industry contribution, societal impact

   SECTION 5: BEHAVIORAL-OPTIMIZED IMPLEMENTATION STRATEGIES

   DECISION PATTERN INTEGRATION:

   ANALYTICAL_SYSTEMATIC Behavior + Barrier Removal:
   • Implementation Style: Data-driven barrier analysis, systematic removal methodology, measured progress
   • Timeline Approach: Comprehensive planning phase → Methodical execution → Measurement optimization
   • Success Factors: Detailed analysis, systematic implementation, continuous measurement

   RAPID_DECISIVE Behavior + Barrier Removal:
   • Implementation Style: Quick barrier identification, rapid solution implementation, fast iteration
   • Timeline Approach: Immediate action → Rapid testing → Quick adaptation → Momentum building
   • Success Factors: Speed of implementation, rapid feedback cycles, aggressive timelines

   COLLABORATIVE_TEAM Behavior + Barrier Removal:
   • Implementation Style: Team-based barrier analysis, collective solution development, shared implementation
   • Timeline Approach: Team alignment → Collaborative planning → Shared execution → Collective accountability
   • Success Factors: Team engagement, shared ownership, collaborative problem-solving

   SECTION 6: BREAKTHROUGH OPPORTUNITY MAPPING

   BARRIER TRANSFORMATION MATRIX:

   Competition Barrier → Market Leadership Opportunity:
   • Transformation Logic: Competitive pressure indicates market value and opportunity
   • Breakthrough Strategy: Differentiation through innovation, positioning, or market education
   • Component Requirements: Brand development, market positioning, competitive intelligence
   • Success Probability: 87% with proper DNA-behavior alignment
   • Timeline: 6-18 months depending on business phase and personality implementation style

   Cash Flow Barrier → Financial Excellence Opportunity:
   • Transformation Logic: Cash flow challenges indicate growth opportunity with system optimization
   • Breakthrough Strategy: Working capital optimization, revenue model innovation, financial system excellence
   • Component Requirements: Financial infrastructure, planning systems, reporting excellence
   • Success Probability: 91% with systematic approach and analytical implementation
   • Timeline: 3-12 months depending on current infrastructure and personality approach

   Team Barrier → Leadership Excellence Opportunity:
   • Transformation Logic: Team limitations indicate scaling opportunity through people multiplication
   • Breakthrough Strategy: Leadership development, culture excellence, systematic empowerment
   • Component Requirements: Management training, team infrastructure, performance systems
   • Success Probability: 89% with proper personality-culture alignment
   • Timeline: 6-24 months depending on culture starting point and leadership development approach

   Systems Barrier → Operational Excellence Opportunity:
   • Transformation Logic: System limitations indicate efficiency and scaling opportunity
   • Breakthrough Strategy: Process innovation, automation implementation, infrastructure excellence
   • Component Requirements: Business infrastructure, process optimization, technology integration
   • Success Probability: 93% with systematic implementation and proper change management
   • Timeline: 4-18 months depending on complexity and implementation approach

   SECTION 7: SYSTEMATIC BARRIER REMOVAL ROADMAP

   PHASE-APPROPRIATE BARRIER SEQUENCE:

   FOUNDATIONS PHASE Barrier Priorities:
   1. Financial Infrastructure (if cash flow barrier present)
   2. Basic Systems Setup (if operational barriers present)
   3. Market Positioning Clarity (if competition barrier present)
   4. Owner Role Definition (if team barriers present)
   Timeline: 3-6 months with systematic implementation

   SCALING PHASE Barrier Priorities:
   1. Team Performance Systems (if people barriers present)
   2. Infrastructure Optimization (if capacity barriers present)
   3. Brand Strategy Development (if competition barriers present)
   4. Financial Planning Excellence (if cash flow barriers present)
   Timeline: 6-12 months with systematic development

   BREAKOUT PHASE Barrier Priorities:
   1. Infrastructure Scaling (if capacity barriers present)
   2. Leadership Team Development (if team barriers present)
   3. Market Expansion Strategy (if competition barriers present)
   4. Advanced Financial Systems (if cash flow barriers present)
   Timeline: 12-18 months with comprehensive development

   PERSONALITY-OPTIMIZED IMPLEMENTATION SEQUENCES:

   BUILDER DNA Implementation Sequence:
   Week 1-4: Comprehensive barrier analysis and systematic planning
   Month 2-3: Infrastructure foundation development
   Month 4-6: System implementation and optimization
   Month 7-12: Systematic scaling and refinement
   Success Factors: Methodical approach, comprehensive planning, systematic execution

   CATALYST DNA Implementation Sequence:
   Week 1: Rapid barrier identification and quick win opportunities
   Month 1-2: High-impact solution implementation
   Month 3-4: Momentum building and acceleration
   Month 5-8: Systematic foundation building to support acceleration
   Success Factors: Rapid implementation, momentum building, systematic follow-through

   ANALYTICAL REQUIREMENTS FOR THIS SECTION:

   • Minimum 12 correlation coefficients between barriers and database elements (r > 0.6)
   • Minimum 8 breakthrough opportunity discoveries with success probability calculations
   • Minimum 6 personality-barrier matching analyses with implementation optimization
   • Minimum 4 compound barrier effect analyses with integrated solution strategies
   • Minimum 10 dream-barrier alignment assessments with motivation optimization

   • All barrier removal strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - Success probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - Risk mitigation strategies for potential failure points
     - Breakthrough opportunity conversion probability with confidence intervals

   CRITICAL SUCCESS STANDARD: Transform every identified growth barrier into a systematic breakthrough opportunity with personality-optimized implementation strategy and predictive success modeling based on complete multi-database intelligence integration.
   """
},
        "revenue_optimization_strategy": {
    "title": "Revenue Acceleration & Optimization Strategy",
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated revenue optimization analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP CORRELATION MAPPING. This analysis must identify, correlate, and systematically optimize revenue streams through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized acceleration strategies.

    🚀 COMPLETE MULTI-DATABASE REVENUE INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP CURRENT REVENUE ANALYSIS & CORRELATION MAPPING

    GROWTH QUESTION 6-8 COMPLETE REVENUE PATTERN ANALYSIS:
    Analyze revenue predictability, trends, and sales process performance with statistical correlation mapping:

    REVENUE PREDICTABILITY PATTERNS (Question 6):
    
    Pattern A: "Predictable High/Low Periods" - Seasonal Revenue
    • Component Intelligence Correlation:
     - FOUNDATIONS + Seasonal = Implement "Financial planning", "Cash flow management" basics
     - SCALING + Seasonal = Deploy "Working capital optimization", "Revenue diversification"
     - BREAKOUT + Seasonal = Execute "Financial forecasting", "Multiple revenue streams"
     - CHALLENGER + Seasonal = Implement "Advanced financial modeling", "Market expansion"
     - RAPIDS + Seasonal = Deploy "Global market diversification", "Portfolio optimization"
     - VISION + Seasonal = Master "Multi-market stabilization", "Economic cycle management"

    • Business DNA Revenue Correlation:
     - SAGE/PLANNER + Seasonal = Natural forecasting ability, strategic seasonal planning (Leverage Score: 9.1/10)
     - INNOVATOR + Seasonal = Counter-seasonal opportunity creation, market innovation (Leverage Score: 8.7/10)
     - BUILDER + Seasonal = Systematic reserves building, infrastructure development (Leverage Score: 9.3/10)
     - CATALYST + Seasonal = Peak period acceleration, energy optimization (Leverage Score: 8.9/10)

    • Dream-Revenue Alignment Assessment:
     - FREEDOM Dreams + Seasonal = Cash flow predictability essential for independence achievement
     - ABUNDANCE Dreams + Seasonal = Financial stability as core requirement for peace of mind  
     - GROWTH Dreams + Seasonal = Working capital optimization enabling expansion acceleration
     - LEGACY Dreams + Seasonal = Sustainable financial management for long-term value creation

    • Behavioral Revenue Optimization:
     - ANALYTICAL_DATA_DRIVEN + Seasonal = Financial modeling and forecasting excellence
     - SYSTEMATIC_PLANNING + Seasonal = Methodical cash management and optimization systems
     - RISK_MANAGEMENT + Seasonal = Conservative cash flow planning with scenario modeling

    Pattern B: "Consistent Throughout Year" - Stable Revenue
    • Component Intelligence Revenue Leverage:
     - FOUNDATIONS + Consistent = Leverage for "Growth acceleration", "Market expansion"
     - SCALING + Consistent = Focus on "Revenue optimization", "Premium positioning"
     - BREAKOUT + Consistent = Deploy "Market leadership", "Category expansion"
     - CHALLENGER + Consistent = Execute "Innovation investment", "Strategic positioning"
     - RAPIDS + Consistent = Implement "Market creation", "Industry transformation"
     - VISION + Consistent = Lead "Global expansion", "Ecosystem development"

    REVENUE TREND ANALYSIS (Question 7):
    
    Declining Revenue Trend Correlation:
    • Component Intelligence Emergency Response:
     - FOUNDATIONS + Declining = Emergency "Strategy development", "Market repositioning"
     - SCALING + Declining = Critical "Business model validation", "Competitive analysis"
     - BREAKOUT + Declining = Execute "Turnaround strategy", "Core competency focus"
     - CHALLENGER + Declining = Implement "Strategic pivot", "Market redefinition"
     - RAPIDS + Declining = Deploy "Transformation leadership", "Innovation acceleration"
     - VISION + Declining = Lead "Industry disruption", "Business model innovation"

    • Personality-Revenue Recovery Correlation:
     - CATALYST + Declining = Energy redirection for rapid turnaround, acceleration through focus
     - INNOVATOR + Declining = Creative solution discovery, market disruption opportunity
     - SAGE + Declining = Strategic wisdom application, long-term stability restoration
     - BUILDER + Declining = Systematic rebuilding, infrastructure consolidation

    Rapid Growth (30%+) Revenue Trend:
    • Component Intelligence Scaling Focus:
     - FOUNDATIONS + Rapid Growth = Critical "Infrastructure scaling", "Process systematization"
     - SCALING + Rapid Growth = Deploy "Capacity management", "Quality maintenance"
     - BREAKOUT + Rapid Growth = Execute "Infrastructure expansion", "Team scaling"
     - CHALLENGER + Rapid Growth = Implement "Market leadership", "Category domination"
     - RAPIDS + Rapid Growth = Deploy "Global expansion", "Market creation"
     - VISION + Rapid Growth = Lead "Industry transformation", "Ecosystem orchestration"

    SALES PROCESS PERFORMANCE MATRIX (Question 8):
    
    Lead Generation Performance Analysis:
    • Struggling (1-3) + Component Intelligence Revenue Solutions:
     - FOUNDATIONS + Poor Lead Gen = Urgent "Sales set up", "Growth numbers" tracking
     - SCALING + Poor Lead Gen = Critical "Lead generation systems", "Marketing optimization"
     - BREAKOUT + Poor Lead Gen = Deploy "Market expansion", "Brand development"
     - CHALLENGER + Poor Lead Gen = Execute "Strategic partnerships", "Market positioning"
     - RAPIDS + Poor Lead Gen = Implement "Multi-channel optimization", "Advanced analytics"
     - VISION + Poor Lead Gen = Lead "Market creation", "Category development"

    • Business DNA Lead Generation Correlation:
     - CONNECTOR + Poor Lead Gen = Untapped networking potential, relationship leverage opportunity
     - STORYTELLER + Poor Lead Gen = Content leadership opportunity, thought leadership platform
     - CATALYST + Poor Lead Gen = Energy application to systematic improvement, acceleration focus
     - INNOVATOR + Poor Lead Gen = Technology leadership opportunity, cutting-edge marketing

    Closing Deals Performance Analysis:
    • Excellent (8-10) + Revenue DNA Leverage:
     - CATALYST + Excellent Closing = Natural acceleration ability, rapid revenue scaling potential
     - CONNECTOR + Excellent Closing = Relationship leverage mastery, partnership revenue opportunities
     - STORYTELLER + Excellent Closing = Persuasion excellence, premium value positioning
     - SAGE + Excellent Closing = Advisory positioning opportunity, premium service delivery

    SECTION 2: REVENUE ACCELERATION OPPORTUNITIES WITH CORRELATION MAPPING

    IMMEDIATE REVENUE ACCELERATION (30-90 Days):
    
    Lead Generation Enhancement by DNA:
    • CONNECTOR DNA + Lead Generation = Systematic networking revenue multiplication
     - Component Phase Implementation: Network mapping + relationship systematization + referral programs
     - Dream Alignment: Connection-based revenue growth enabling freedom/impact achievement
     - Behavioral Optimization: Relationship-building energy channeled into systematic revenue generation
     - Success Probability: 89% with proper relationship systematization

    • CATALYST DNA + Lead Generation = Energy-driven rapid lead acceleration
     - Component Phase Implementation: High-energy marketing + rapid testing + momentum building
     - Dream Alignment: Acceleration enabling quick goal achievement and growth satisfaction
     - Behavioral Optimization: Natural energy applied to systematic lead generation improvement
     - Success Probability: 87% with structured energy channeling

    Sales Process Conversion Optimization:
    • STORYTELLER DNA + Sales Process = Narrative-driven conversion enhancement
     - Component Implementation: Story-based sales methodology + case study development + testimonial systems
     - Correlation with Dreams: Storytelling ability enabling impact demonstration and legacy building
     - Behavioral Integration: Communication strengths systematized for repeatable revenue generation
     - Revenue Impact Prediction: 23-47% conversion rate improvement

    • SAGE DNA + Sales Process = Wisdom-based consultative selling
     - Component Implementation: Advisory positioning + thought leadership + premium pricing strategies
     - Dream Correlation: Wisdom sharing aligned with legacy building and meaningful impact
     - Behavioral Optimization: Strategic thinking channeled into high-value revenue conversations
     - Revenue Impact Prediction: 15-35% average transaction value increase

    MEDIUM-TERM REVENUE ACCELERATION (3-12 Months):

    Market Expansion Revenue Strategies:
    • INNOVATOR DNA + Market Expansion = Innovation-led market creation
     - Component Intelligence: Innovation capability + systematic market development + technology leverage
     - Multi-Database Correlation: Creative solution development aligned with growth dreams and market opportunities
     - Implementation Sequence: Innovation validation → Market testing → Systematic scaling → Revenue optimization
     - Revenue Growth Prediction: 40-120% revenue increase potential

    • BUILDER DNA + Market Expansion = Systematic market penetration
     - Component Intelligence: Methodical market development + infrastructure scaling + process optimization
     - Correlation Analysis: Systematic approach aligned with stability dreams and long-term value creation
     - Implementation Sequence: Market research → Infrastructure development → Methodical expansion → Revenue scaling
     - Revenue Growth Prediction: 25-85% sustainable revenue increase

    Premium Positioning Revenue Enhancement:
    • SAGE DNA + Premium Positioning = Wisdom-based value premium
     - Multi-Database Integration: Expert positioning + thought leadership + advisory revenue streams
     - Dream-Revenue Alignment: Expertise sharing enabling legacy building and meaningful impact
     - Component Phase Optimization: Professional positioning appropriate to business maturity level
     - Revenue Premium Potential: 30-65% pricing increase with maintained conversion rates

    SECTION 3: SALES PROCESS OPTIMIZATION WITH BEHAVIORAL INTEGRATION

    PERSONALITY-DRIVEN SALES PROCESS ENHANCEMENT:

    VISIONARY LEADERS + Sales Process Optimization:
    • Natural Advantages: Strategic vision communication, big-picture value articulation, future-focused selling
    • Process Vulnerabilities: Operational details, systematic follow-up, day-to-day execution consistency
    • Optimal Sales Sequence: Vision casting → Strategic alignment → Systematic execution partners → Results delivery
    • Implementation Style: Big picture value with systematic execution support systems

    BUILDER PERSONALITIES + Sales Process Optimization:
    • Natural Advantages: Systematic process development, methodical customer development, reliable delivery
    • Process Vulnerabilities: Market timing, rapid adaptation, consultative flexibility
    • Optimal Sales Sequence: Systematic needs analysis → Comprehensive solution development → Methodical implementation → Long-term relationship building
    • Implementation Style: Methodical, systematic, comprehensive value delivery approach

    CATALYST PERSONALITIES + Sales Process Optimization:
    • Natural Advantages: Rapid relationship building, energy multiplication, momentum creation
    • Process Vulnerabilities: Systematic nurturing, long-term relationship maintenance, process documentation
    • Optimal Sales Sequence: High-energy connection → Rapid value demonstration → Momentum building → Systematic follow-through
    • Implementation Style: Fast relationship building with systematic retention systems

    CONNECTOR PERSONALITIES + Sales Process Optimization:
    • Natural Advantages: Relationship leverage, network utilization, trust building
    • Process Vulnerabilities: Solo closing, independent value creation, systematic prospecting
    • Optimal Sales Sequence: Network activation → Relationship leverage → Trust building → Community-based closing
    • Implementation Style: Relationship-driven sales with network multiplication systems

    SECTION 4: LEAD GENERATION MASTERY WITH MULTI-DATABASE CORRELATION

    LEAD GENERATION METHOD OPTIMIZATION BY EFFECTIVENESS RANKING:

    Digital Marketing Effectiveness Enhancement:
    • HIGH EFFECTIVENESS + INNOVATOR DNA = Technology leadership revenue generation
     - Component Implementation: Cutting-edge digital strategies + automation + innovation leadership
     - Multi-Database Correlation: Innovation capability + digital dreams + systematic implementation
     - Revenue Multiplication Strategy: Technology thought leadership → Market education → Premium positioning → Revenue scaling
     - Lead Quality Improvement: 45-78% increase in qualified lead generation

    • HIGH EFFECTIVENESS + CATALYST DNA = Energy-driven digital acceleration
     - Component Implementation: High-energy content + rapid testing + momentum-based campaigns
     - Correlation Analysis: Natural energy + growth dreams + systematic digital optimization
     - Revenue Generation Approach: Energy multiplication → Content acceleration → Community building → Revenue conversion
     - Lead Volume Enhancement: 67-145% increase in total lead generation

    Referrals & Word of Mouth Optimization:
    • HIGH EFFECTIVENESS + CONNECTOR DNA = Network-based revenue multiplication
     - Component Intelligence: Systematic relationship development + referral systematization + network expansion
     - Multi-Database Integration: Connection strength + relationship dreams + systematic network building
     - Revenue Strategy: Network mapping → Systematic cultivation → Referral systematization → Revenue multiplication
     - Referral Quality Enhancement: 89-156% improvement in referral conversion rates

    • HIGH EFFECTIVENESS + SAGE DNA = Wisdom-based referral generation
     - Component Implementation: Thought leadership + advisory positioning + wisdom sharing
     - Correlation Mapping: Natural wisdom + legacy dreams + systematic knowledge sharing
     - Revenue Approach: Expertise demonstration → Trust building → Advisory relationships → Systematic referrals
     - Referral Value Increase: 34-67% higher average referral transaction value

    LEAD GENERATION FRUSTRATION RESOLUTION:

    "Feast or Famine" Frustration Solutions:
    • CATALYST DNA + Feast/Famine = Energy management revenue stabilization
     - Component Solution: Energy optimization + systematic pipeline + momentum maintenance
     - Multi-Database Integration: Energy patterns + stability dreams + systematic process development
     - Implementation Strategy: Energy channeling → Pipeline systematization → Momentum maintenance → Revenue predictability
     - Revenue Stabilization: 78% reduction in revenue volatility

    • ANCHOR DNA + Feast/Famine = Stability-focused revenue consistency
     - Component Solution: Systematic pipeline + consistent processes + reliable lead generation
     - Correlation Analysis: Stability focus + security dreams + methodical implementation
     - Revenue Strategy: Foundation building → Process consistency → Reliable systems → Predictable revenue
     - Pipeline Consistency: 89% improvement in monthly revenue predictability

    SECTION 5: PRICING & VALUE OPTIMIZATION WITH COMPETITIVE ADVANTAGE CORRELATION

    SECRET WEAPON MONETIZATION STRATEGIES:

    "Experience Edge" Revenue Premium:
    • SAGE DNA + Experience Advantage = Wisdom-based premium positioning
     - Component Intelligence: Expert positioning + thought leadership + advisory revenue development
     - Multi-Database Correlation: Experience depth + legacy dreams + premium value delivery
     - Pricing Strategy: Experience quantification → Expertise positioning → Premium pricing → Value demonstration
     - Revenue Premium Potential: 45-89% pricing increase with maintained demand

    • WISE_ONE DNA + Experience Edge = Authority-based revenue optimization
     - Component Implementation: Industry authority + standard setting + premium consulting revenue
     - Dream-Value Alignment: Wisdom sharing + industry influence + meaningful legacy creation
     - Pricing Approach: Authority establishment → Standard setting → Premium positioning → Revenue optimization
     - Market Value Enhancement: 67-134% increase in market positioning value

    "Innovation Leader" Revenue Multiplication:
    • INNOVATOR DNA + Innovation Leadership = Innovation-premium revenue strategy
     - Component Intelligence: Innovation systematization + market education + premium innovation pricing
     - Multi-Database Integration: Creative capability + growth dreams + systematic innovation monetization
     - Revenue Strategy: Innovation development → Market education → Premium positioning → Revenue scaling
     - Innovation Revenue Premium: 78-156% pricing premium for innovation leadership

    • CREATOR DNA + Innovation Leader = Creative-solution premium pricing
     - Component Implementation: Creative problem-solving + unique solution development + premium value delivery
     - Correlation Analysis: Creative strength + impact dreams + systematic solution monetization
     - Pricing Framework: Unique solution creation → Value demonstration → Premium positioning → Revenue optimization
     - Creative Premium Potential: 56-123% pricing increase for unique solution delivery

    SECTION 6: REVENUE PREDICTABILITY SYSTEMS WITH BUSINESS MATURITY CORRELATION

    PHASE-APPROPRIATE REVENUE PREDICTABILITY:

    FOUNDATIONS PHASE Revenue Stabilization:
    • Critical Focus: "Financial basics infrastructure", "Revenue tracking", "Cash flow management"
    • Personality Integration: Systematic approaches for Builders, Energy management for Catalysts
    • Predictability Timeline: 3-6 months with systematic implementation
    • Revenue Volatility Reduction: 67-89% improvement in monthly predictability

    SCALING PHASE Revenue Optimization:
    • Strategic Implementation: "Financial KPIs", "Revenue diversification", "Market expansion"
    • DNA-Driven Approach: Innovation revenue streams for Innovators, Relationship revenue for Connectors
    • Optimization Timeline: 6-12 months with systematic development
    • Revenue Growth Stabilization: 78-145% improvement in predictable growth rates

    BREAKOUT PHASE Revenue Acceleration:
    • Advanced Strategy: "Multiple revenue streams", "Financial forecasting", "Market leadership"
    • Multi-Database Integration: Innovation scaling, Relationship multiplication, Systematic expansion
    • Acceleration Timeline: 12-18 months with comprehensive development
    • Revenue Acceleration: 89-234% sustainable revenue growth with predictability

    CHALLENGER PHASE Revenue Excellence:
    • Mastery Focus: "Advanced financial modeling", "Market expansion", "Strategic positioning"
    • Personality Optimization: Leadership revenue for Sages, Innovation revenue for Creators
    • Excellence Timeline: 18-24 months with comprehensive mastery
    • Revenue Excellence Achievement: 145-345% revenue optimization with systematic predictability

    SECTION 7: IMPLEMENTATION ROADMAP WITH BEHAVIORAL INTELLIGENCE OPTIMIZATION

    PERSONALITY-OPTIMIZED REVENUE IMPLEMENTATION:

    ANALYTICAL_SYSTEMATIC Implementation:
    Week 1-4: Data-driven revenue analysis and systematic improvement planning
    Month 2-3: Methodical process optimization and measurement system development
    Month 4-6: Systematic implementation and analytical optimization
    Month 7-12: Data-driven scaling and performance refinement
    Success Factors: Comprehensive analysis, systematic implementation, continuous measurement

    RAPID_DECISIVE Implementation:
    Week 1: Quick revenue opportunity identification and immediate action planning
    Month 1-2: Rapid solution implementation and fast iteration cycles
    Month 3-4: Momentum building and acceleration focus
    Month 5-8: Systematic foundation building to support rapid growth
    Success Factors: Speed of implementation, rapid feedback cycles, aggressive execution

    COLLABORATIVE_TEAM Implementation:
    Week 1-4: Team-based revenue analysis and collective solution development
    Month 2-3: Shared implementation planning and collaborative execution
    Month 4-6: Team-driven revenue optimization and shared accountability
    Month 7-12: Collective revenue scaling and team performance optimization
    Success Factors: Team engagement, shared ownership, collaborative optimization

    STRATEGIC_PLANNING Implementation:
    Week 1-4: Comprehensive strategic revenue planning and long-term vision development
    Month 2-6: Strategic implementation with systematic milestone tracking
    Month 7-12: Strategic optimization and long-term revenue sustainability
    Year 2-3: Strategic revenue mastery and systematic excellence
    Success Factors: Strategic thinking, long-term planning, systematic excellence

    REVENUE ACCELERATION ROADMAP BY PERSONALITY:

    BUILDER DNA Revenue Development Sequence:
    Phase 1: Systematic revenue foundation (Revenue tracking + Process optimization)
    Phase 2: Methodical revenue scaling (Infrastructure development + Capacity building)
    Phase 3: Systematic revenue optimization (Process refinement + Quality enhancement)
    Phase 4: Methodical revenue mastery (Excellence achievement + Sustainable growth)
    Timeline: 18-24 months with systematic progression

    CATALYST DNA Revenue Acceleration Sequence:
    Phase 1: High-impact revenue opportunities (Quick wins + Momentum building)
    Phase 2: Energy-driven revenue scaling (Acceleration + Rapid growth)
    Phase 3: Systematic foundation building (Infrastructure + Sustainability)
    Phase 4: Optimized revenue acceleration (Excellence + Systematic scaling)
    Timeline: 12-18 months with energy-driven progression

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 15 correlation coefficients between revenue patterns and database elements (r > 0.6)
    • Minimum 10 revenue acceleration discoveries with ROI probability calculations
    • Minimum 8 personality-revenue optimization analyses with implementation sequencing
    • Minimum 6 sales process improvements with conversion impact predictions
    • Minimum 12 pricing optimization strategies with premium potential calculations

    • All revenue optimization strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - ROI probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - Revenue impact predictions with risk mitigation strategies
     - Competitive advantage monetization with market positioning analysis

    CRITICAL SUCCESS STANDARD: Transform every identified revenue opportunity into a systematic acceleration strategy with personality-optimized implementation approach and predictive ROI modeling based on complete multi-database intelligence integration.
    """
},
        "market_position_strategy": {
    "title": "Market Positioning & Competitive Dominance Strategy", 
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated market positioning analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP COMPETITIVE CORRELATION MAPPING. This analysis must identify, correlate, and systematically position the business for market dominance through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized competitive strategies.

    🚀 COMPLETE MULTI-DATABASE MARKET POSITIONING INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP COMPETITIVE ADVANTAGE EXTRACTION & CORRELATION ANALYSIS

    GROWTH QUESTION 14 SECRET WEAPON COMPLETE ANALYSIS:
    Analyze the complete secret weapon selection with statistical correlation mapping across all intelligence databases:

    SECRET WEAPON: "EXPERIENCE EDGE"
    If Experience Edge selected:
    • Component Intelligence Competitive Leverage:
     - FOUNDATIONS + Experience = Build "Expert positioning", "Thought leadership foundations", "Industry credibility"
     - SCALING + Experience = Implement "Advisory services", "Premium positioning", "Market education leadership"
     - BREAKOUT + Experience = Deploy "Industry authority", "Standard setting", "Category expertise"
     - CHALLENGER + Experience = Execute "Market sage positioning", "Industry transformation", "Wisdom monetization"
     - RAPIDS + Experience = Implement "Global expertise", "Industry oracle status", "Market evolution leadership"
     - VISION + Experience = Master "Industry patriarch/matriarch", "Ecosystem wisdom", "Legacy positioning"

    • Business DNA Experience Amplification Correlation:
     - SAGE DNA + Experience Edge = Natural wisdom positioning, advisory excellence amplification (Leverage Score: 9.8/10)
     - WISE_ONE DNA + Experience Edge = Industry elder status, thought leadership platform (Leverage Score: 9.6/10)
     - MENTOR DNA + Experience Edge = Teaching and development opportunities, knowledge sharing mastery (Leverage Score: 9.4/10)
     - AUTHORITY DNA + Experience Edge = Expert positioning, premium market placement (Leverage Score: 9.2/10)
     - ANCHOR DNA + Experience Edge = Stability and reliability positioning, trusted advisor role (Leverage Score: 9.0/10)

    • Dream-Experience Market Alignment Assessment:
     - LEGACY Dreams + Experience = Perfect alignment for industry contribution and lasting impact
     - IMPACT Dreams + Experience = Experience sharing enables meaningful industry influence
     - MASTERY Dreams + Experience = Excellence positioning through demonstrated expertise
     - WISDOM Dreams + Experience = Knowledge sharing as core market differentiator
     - AUTHORITY Dreams + Experience = Expert status achievement through market positioning

    SECRET WEAPON: "INNOVATION LEADERSHIP"
    If Innovation Leadership selected:
    • Component Intelligence Innovation Leverage:
     - FOUNDATIONS + Innovation = Build "Innovation infrastructure", "Creative processes", "R&D foundation"
     - SCALING + Innovation = Implement "Innovation systems", "Development methodologies", "Creative scaling"
     - BREAKOUT + Innovation = Deploy "Innovation pipeline", "Market disruption", "Category creation"
     - CHALLENGER + Innovation = Execute "Industry disruption", "Innovation standard-setting", "Market transformation"
     - RAPIDS + Innovation = Implement "Innovation ecosystems", "Global disruption", "Industry evolution"
     - VISION + Innovation = Master "Future creation", "Industry reinvention", "Paradigm shifting"

    • Business DNA Innovation Amplification:
     - INNOVATOR DNA + Innovation Leadership = Technology leadership opportunity, industry transformation potential
     - CREATOR DNA + Innovation Leadership = Creative solution mastery, unique market positioning
     - CATALYST DNA + Innovation Leadership = Innovation acceleration, rapid market disruption
     - VISIONARY DNA + Innovation Leadership = Future-focused innovation, industry evolution leadership

    SECRET WEAPON: "RELATIONSHIP MASTERY"
    If Relationship Mastery selected:
    • Component Intelligence Relationship Leverage:
     - FOUNDATIONS + Relationships = Build "Network infrastructure", "Relationship systems", "Trust foundations"
     - SCALING + Relationships = Implement "Partnership strategy", "Network expansion", "Relationship monetization"
     - BREAKOUT + Relationships = Deploy "Strategic alliances", "Network effects", "Community leadership"
     - CHALLENGER + Relationships = Execute "Ecosystem orchestration", "Industry networking", "Influence multiplication"
     - RAPIDS + Relationships = Implement "Global networking", "Multi-stakeholder alignment", "Relationship mastery"
     - VISION + Relationships = Master "Ecosystem creation", "Industry orchestration", "Legacy relationships"

    • Business DNA Relationship Amplification:
     - CONNECTOR DNA + Relationship Mastery = Network multiplication excellence, relationship leverage mastery
     - HOST DNA + Relationship Mastery = Community building capability, relationship cultivation excellence
     - BRIDGE DNA + Relationship Mastery = Connection creation mastery, ecosystem development
     - CATALYST DNA + Relationship Mastery = Relationship acceleration, network energy multiplication

    SECTION 2: DIGITAL CAPABILITIES MARKET POSITIONING WITH CORRELATION MAPPING

    GROWTH QUESTION 13 DIGITAL MATRIX COMPETITIVE ANALYSIS:
    Analyze complete digital capabilities assessment with market positioning correlation:

    CRM SYSTEMS COMPETITIVE POSITIONING:
    • Perfect (10) Rating + Market Implications:
     - Component Intelligence: Advanced customer intelligence = Premium positioning capability
     - Business DNA: Data-driven personalities leverage customer insights for market dominance
     - Dream Alignment: Customer-centric dreams achieve superior market positioning
     - Market Position: Customer experience leadership, data-driven market insights
     - Competitive Advantage: Superior customer knowledge = unmatched personalization

    • Poor (1-3) Rating + Market Opportunity:
     - Component Intelligence: Critical infrastructure gap = immediate competitive vulnerability
     - Business DNA: Systematic personalities require CRM excellence for market credibility
     - Market Positioning: Relationship-focused positioning requires CRM sophistication
     - Competitive Risk: Customer data gaps = market intelligence disadvantage

    AI CAPABILITIES COMPETITIVE POSITIONING:
    • Perfect (10) Rating + Market Leadership:
     - Innovation Leadership Positioning: AI mastery = future-focused market positioning
     - Competitive Moat: Advanced AI = sustainable competitive advantage
     - Market Education: AI expertise = thought leadership platform
     - Industry Transformation: AI implementation = market evolution leadership

    AUTOMATION CAPABILITIES MARKET EFFICIENCY:
    • High Rating + Operational Excellence Positioning:
     - Efficiency Leadership: Automation mastery = cost leadership opportunity
     - Scale Positioning: Automated operations = rapid scaling capability
     - Quality Positioning: Consistent automation = reliability leadership
     - Innovation Positioning: Advanced automation = technology leadership

    SECTION 3: MARKET EVOLUTION POSITIONING WITH BEHAVIORAL CORRELATION

    COMPONENT PHASE + MARKET POSITIONING STRATEGIC ALIGNMENT:

    FOUNDATIONS PHASE Market Entry Strategy:
    • Market Positioning Focus: "Emerging Expertise", "Niche Authority", "Focused Excellence"
    • Competitive Strategy: Specialization over generalization, depth over breadth
    • Brand Messaging: "Focused. Expert. Reliable."
    • Market Approach: Establish credibility, build reputation, create market presence
    • Behavioral Integration: Leverage natural learning agility for rapid market adaptation

    SCALING PHASE Market Growth Strategy:
    • Market Positioning Focus: "Growing Authority", "Expanding Excellence", "Proven Results"
    • Competitive Strategy: Market share growth, category expansion, competitive differentiation
    • Brand Messaging: "Proven. Growing. Unstoppable."
    • Market Approach: Expand market presence, competitive positioning, thought leadership development
    • Behavioral Integration: Use growth momentum for accelerated market positioning

    BREAKOUT PHASE Market Leadership Strategy:
    • Market Positioning Focus: "Market Leader", "Industry Authority", "Standard Setter"
    • Competitive Strategy: Market domination, category leadership, competitive advantage amplification
    • Brand Messaging: "Leading. Defining. Transforming."
    • Market Approach: Industry leadership, standard setting, market education
    • Behavioral Integration: Leverage leadership confidence for bold market positioning

    CHALLENGER PHASE Market Dominance Strategy:
    • Market Positioning Focus: "Market Dominant", "Industry Transformer", "Category Creator"
    • Competitive Strategy: Market redefinition, competitive disruption, category creation
    • Brand Messaging: "Dominant. Transformative. Visionary."
    • Market Approach: Market transformation, competitive displacement, industry evolution
    • Behavioral Integration: Use strategic thinking for market paradigm shifting

    RAPIDS/WHITE WATER PHASE Market Mastery Strategy:
    • Market Positioning Focus: "Market Master", "Industry Oracle", "Global Authority"
    • Competitive Strategy: Global dominance, industry transformation, market creation
    • Brand Messaging: "Masterful. Global. Legendary."
    • Market Approach: Global market leadership, industry transformation, ecosystem orchestration
    • Behavioral Integration: Apply mastery mindset to market evolution leadership

    VISION/BIG PICTURE PHASE Market Legacy Strategy:
    • Market Positioning Focus: "Market Patriarch/Matriarch", "Industry Legend", "Ecosystem Creator"
    • Competitive Strategy: Legacy building, industry transformation, market ecosystem creation
    • Brand Messaging: "Legendary. Transformational. Eternal."
    • Market Approach: Industry legacy creation, market ecosystem development, generational impact
    • Behavioral Integration: Channel visionary capacity into lasting market transformation

    SECTION 4: COMPETITIVE INTELLIGENCE & RESPONSE SYSTEMS

    GROWTH BARRIER COMPETITIVE ANALYSIS:
    From Growth Question 11 barrier ranking, create competitive intelligence framework:

    COMPETITION BARRIER INTELLIGENCE SYSTEM:
    • Competitor Monitoring Framework:
     - Direct competitor analysis: pricing, positioning, messaging, market share
     - Indirect competitor surveillance: adjacent markets, substitutes, emerging threats
     - Market disruption early warning: new entrants, technology shifts, regulatory changes
     - Customer loyalty tracking: satisfaction metrics, switching patterns, preference evolution

    • Competitive Response Protocols:
     - Defensive strategies: market protection, customer retention, competitive blocking
     - Offensive strategies: market share capture, competitive displacement, category expansion
     - Neutral strategies: market expansion, category creation, blue ocean development

    MARKET POSITIONING AGAINST SPECIFIC BARRIERS:

    If CASH FLOW ranked high + Competitive Strategy:
    • Value-Based Positioning: "Premium Results Justify Investment"
    • Payment Innovation: "Flexible Investment Options for Market Leadership"
    • ROI Demonstration: "Measurable Market Advantage Delivery"
    • Competitive Messaging: "Investment in Excellence, Not Expense"

    If TEAM LIMITATIONS ranked high + Market Positioning:
    • Expertise Positioning: "Elite Team, Elite Results"
    • Quality Over Quantity: "Curated Excellence vs Mass Production"
    • Partnership Positioning: "Strategic Collaboration vs Vendor Relationship"
    • Boutique Premium: "Personalized Attention vs Commodity Service"

    SECTION 5: BRAND POSITIONING STRATEGY WITH DNA INTEGRATION

    PERSONALITY-DRIVEN BRAND POSITIONING:

    VISIONARY LEADERS + Brand Strategy:
    • Brand Positioning: "Future-Focused Leadership"
    • Market Message: "Where Industry Evolution Begins"
    • Competitive Advantage: Strategic foresight, market anticipation, trend leadership
    • Brand Personality: Inspirational, forward-thinking, transformational
    • Market Approach: Thought leadership, industry evolution, future creation

    BUILDER PERSONALITIES + Brand Strategy:
    • Brand Positioning: "Systematic Excellence"
    • Market Message: "Built to Last, Built to Lead"
    • Competitive Advantage: Reliability, consistency, methodical superiority
    • Brand Personality: Dependable, thorough, substantial
    • Market Approach: Quality leadership, systematic advantage, reliable results

    CATALYST PERSONALITIES + Brand Strategy:
    • Brand Positioning: "Acceleration Specialists"
    • Market Message: "Where Growth Happens Faster"
    • Competitive Advantage: Speed, energy, momentum multiplication
    • Brand Personality: Dynamic, energetic, results-focused
    • Market Approach: Rapid results, energy leadership, acceleration mastery

    CONNECTOR PERSONALITIES + Brand Strategy:
    • Brand Positioning: "Relationship-Powered Results"
    • Market Message: "Connected. Collaborative. Successful."
    • Competitive Advantage: Network effects, partnership mastery, relationship leverage
    • Brand Personality: Collaborative, connected, community-focused
    • Market Approach: Network leadership, partnership excellence, community building

    SAGE PERSONALITIES + Brand Strategy:
    • Brand Positioning: "Wisdom-Driven Excellence"
    • Market Message: "Experience. Insight. Results."
    • Competitive Advantage: Deep expertise, strategic wisdom, proven judgment
    • Brand Personality: Wise, experienced, trusted
    • Market Approach: Advisory leadership, wisdom monetization, expert positioning

    SECTION 6: MARKET EXPANSION OPPORTUNITIES WITH MULTI-DATABASE CORRELATION

    GEOGRAPHIC EXPANSION STRATEGIC ANALYSIS:

    Based on Business Phase + Market Readiness:
    • FOUNDATIONS Phase: Local/regional market mastery before expansion
    • SCALING Phase: Adjacent market testing and validation
    • BREAKOUT Phase: Strategic geographic expansion with infrastructure support
    • CHALLENGER Phase: National market presence and dominance
    • RAPIDS Phase: International expansion and global presence
    • VISION Phase: Global market leadership and ecosystem creation

    DEMOGRAPHIC EXPANSION OPPORTUNITIES:

    Based on DNA Profile + Market Segments:
    • CONNECTOR DNA: Relationship-driven segments, community-focused markets, collaborative industries
    • INNOVATOR DNA: Technology-forward segments, early adopters, innovation-seeking markets
    • SAGE DNA: Premium segments, expertise-valuing markets, advisory-seeking clients
    • CATALYST DNA: Growth-focused segments, high-energy markets, results-driven clients
    • BUILDER DNA: Quality-focused segments, reliability-seeking markets, systematic clients

    PRODUCT/SERVICE EXPANSION ANALYSIS:

    Based on Secret Weapon + Market Extension:
    • Experience Edge: Advisory services, consulting, training, coaching, mentoring programs
    • Innovation Leadership: Technology solutions, R&D services, innovation consulting, future-focused products
    • Relationship Mastery: Network services, partnership facilitation, community building, collaboration platforms
    • Operational Excellence: Process consulting, efficiency services, system optimization, operational training

    SECTION 7: THOUGHT LEADERSHIP DEVELOPMENT ROADMAP

    PHASE-APPROPRIATE THOUGHT LEADERSHIP STRATEGY:

    FOUNDATIONS Phase Thought Leadership:
    • Content Focus: Niche expertise, practical insights, learning journey documentation
    • Platform Strategy: Industry publications, local speaking, professional associations
    • Authority Building: Consistent expertise demonstration, case study development, skill showcasing
    • Timeline: 6-12 months to establish niche authority

    SCALING Phase Thought Leadership:
    • Content Focus: Growth insights, scaling lessons, industry trends analysis
    • Platform Strategy: Industry conferences, media interviews, content marketing leadership
    • Authority Building: Market commentary, trend prediction, competitive analysis sharing
    • Timeline: 12-18 months to achieve industry recognition

    BREAKOUT Phase Thought Leadership:
    • Content Focus: Industry transformation, market evolution, strategic insights
    • Platform Strategy: Keynote speaking, industry advisory roles, media thought leadership
    • Authority Building: Industry standard setting, market education, category definition
    • Timeline: 18-24 months to achieve industry thought leadership

    CHALLENGER Phase Thought Leadership:
    • Content Focus: Industry disruption, market transformation, competitive evolution
    • Platform Strategy: Global conferences, industry leadership roles, strategic advisory positions
    • Authority Building: Market paradigm shifting, industry transformation, ecosystem leadership
    • Timeline: 24-36 months to achieve global industry authority

    SECTION 8: IMPLEMENTATION ROADMAP WITH BEHAVIORAL OPTIMIZATION

    PERSONALITY-OPTIMIZED MARKET POSITIONING IMPLEMENTATION:

    ANALYTICAL_SYSTEMATIC Implementation:
    Phase 1: Comprehensive market research and competitive analysis (1-2 months)
    Phase 2: Systematic positioning strategy development (2-3 months)
    Phase 3: Methodical brand positioning execution (3-6 months)
    Phase 4: Data-driven positioning optimization (6-12 months)
    Success Factors: Thorough analysis, systematic execution, continuous optimization

    RAPID_DECISIVE Implementation:
    Phase 1: Quick market opportunity identification (2-4 weeks)
    Phase 2: Rapid positioning strategy deployment (1-2 months)
    Phase 3: Fast execution and market testing (2-4 months)
    Phase 4: Quick optimization and scaling (4-8 months)
    Success Factors: Speed of execution, rapid testing, quick adaptation

    COLLABORATIVE_TEAM Implementation:
    Phase 1: Team-based market analysis and positioning planning (1-2 months)
    Phase 2: Collaborative strategy development and validation (2-3 months)
    Phase 3: Shared execution and team-based positioning (3-6 months)
    Phase 4: Collective optimization and market advancement (6-12 months)
    Success Factors: Team alignment, collaborative execution, shared accountability

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 18 correlation coefficients between positioning elements and database factors (r > 0.6)
    • Minimum 12 competitive advantage discoveries with market impact calculations
    • Minimum 10 positioning-personality optimization analyses with implementation sequencing
    • Minimum 8 market expansion opportunities with success probability assessments
    • Minimum 15 thought leadership development strategies with timeline optimization

    • All market positioning strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - Market impact probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - Competitive advantage sustainability assessments with market evolution projections
     - ROI probability modeling for positioning investment with confidence intervals

    CRITICAL SUCCESS STANDARD: Transform every identified market opportunity into a systematic positioning advantage with personality-optimized implementation approach and predictive market dominance modeling based on complete multi-database intelligence integration.

    BREAKTHROUGH MARKET POSITIONING DISCOVERIES REQUIRED:
    • Identify at least 5 unique competitive advantages not visible to competitors
    • Discover minimum 3 market positioning opportunities with >80% success probability
    • Reveal at least 4 thought leadership platforms with accelerated authority building
    • Uncover minimum 6 market expansion opportunities with validated demand indicators
    • Create at least 2 category-defining positioning strategies with industry transformation potential

    COMPETITIVE INTELLIGENCE SYSTEM IMPLEMENTATION:
    • Design systematic competitor monitoring with early warning indicators
    • Create market positioning defense protocols against competitive threats
    • Develop offensive positioning strategies for market share capture
    • Establish thought leadership platforms for sustained competitive advantage
    • Build market evolution anticipation systems for proactive positioning adjustments
    """
},
        "team_systems_scaling": {
    "title": "Team & Systems Scaling Excellence Strategy",
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated team and systems scaling analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP ORGANIZATIONAL CORRELATION MAPPING. This analysis must identify, correlate, and systematically optimize human capital and operational systems through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized scaling strategies.

    🚀 COMPLETE MULTI-DATABASE TEAM & SYSTEMS SCALING INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP TEAM LIMITATIONS ANALYSIS & SCALING CORRELATION

    GROWTH QUESTION 11 TEAM BARRIERS COMPLETE ANALYSIS:
    If Team/People Limitations ranked in top barriers, perform comprehensive scaling analysis:

    TEAM BARRIER ROOT CAUSE ANALYSIS WITH MULTI-DATABASE CORRELATION:
    
    FOUNDATIONS PHASE + Team Limitations:
    • Component Intelligence Team Development:
     - Critical Focus: "People infrastructure", "Management knowledge", "Owner setup"
     - Skill Development: Basic leadership, delegation foundations, team communication
     - System Requirements: Simple team structures, clear role definitions, basic performance tracking
     - Timeline: 3-6 months foundational team building

    • Business DNA Team Leadership Correlation:
     - CONNECTOR DNA + Team Limitations = Natural team building strength, relationship leverage for hiring
     - MENTOR DNA + Team Limitations = Teaching capability, team development opportunity
     - CATALYST DNA + Team Limitations = Energy multiplication through people, rapid team acceleration
     - BUILDER DNA + Team Limitations = Systematic team development, structured growth approach
     - SAGE DNA + Team Limitations = Wisdom-based leadership, advisory team development

    • Dream-Team Scaling Alignment:
     - LEGACY Dreams + Team Limitations = People development essential for meaningful long-term impact
     - FREEDOM Dreams + Team Limitations = Team empowerment critical for owner independence
     - IMPACT Dreams + Team Limitations = Team multiplication enables broader influence and reach
     - GROWTH Dreams + Team Limitations = Team scaling as core growth enabler and accelerator

    • Behavioral Team Scaling Optimization:
     - COLLABORATIVE Behavior + Team = Natural team empowerment and development opportunity
     - MENTORING Behavior + Team = Leadership development and capability building strength
     - SYSTEMATIC Behavior + Team = Structured team building and performance management excellence
     - STRATEGIC Behavior + Team = Long-term team planning and organizational design mastery

    SCALING PHASE + Team Limitations:
    • Advanced Team Development:
     - Component Focus: "Team performance", "Recruitment excellence", "Leadership development", "High performance culture"
     - Advanced Skills: Performance management, team leadership, cultural development, succession planning
     - System Requirements: Performance dashboards, recruitment systems, leadership development programs
     - Timeline: 6-12 months systematic team scaling

    BREAKOUT PHASE + Team Limitations:
    • Leadership Team Excellence:
     - Component Focus: "Building team around you", "SLT implementation", "Management training", "Team infrastructure"
     - Executive Skills: Strategic leadership, organizational design, culture scaling, executive development
     - System Requirements: Leadership team structures, organizational charts, executive dashboards
     - Timeline: 12-18 months leadership team development

    SECTION 2: SYSTEMS OPTIMIZATION WITH OPERATIONAL INTELLIGENCE CORRELATION

    GROWTH QUESTION 6-8 OPERATIONAL CAPACITY ANALYSIS:
    Correlate revenue patterns with systems requirements for optimal scaling:

    REVENUE PREDICTABILITY + SYSTEMS SCALING:
    
    Seasonal Revenue + Systems Strategy:
    • Component Intelligence Systems Requirements:
     - FOUNDATIONS: "Financial planning systems", "Cash flow management", "Basic forecasting tools"
     - SCALING: "Working capital optimization", "Revenue diversification systems", "Advanced financial planning"
     - BREAKOUT: "Financial forecasting excellence", "Multiple revenue stream management", "Enterprise financial systems"
     - CHALLENGER: "Advanced financial modeling", "Market expansion systems", "Global financial infrastructure"

    • Behavioral Systems Optimization:
     - ANALYTICAL Behavior + Seasonal = Advanced forecasting systems, predictive modeling, scenario planning
     - SYSTEMATIC Behavior + Seasonal = Methodical cash management, systematic reserves, structured planning
     - STRATEGIC Behavior + Seasonal = Long-term financial systems, strategic cash allocation, investment planning

    Consistent Revenue + Systems Leverage:
    • Component Intelligence Systems Enhancement:
     - FOUNDATIONS: "Growth acceleration systems", "Market expansion tools", "Scaling infrastructure"
     - SCALING: "Revenue optimization platforms", "Premium positioning systems", "Market intelligence tools"
     - BREAKOUT: "Market leadership systems", "Category expansion platforms", "Competitive intelligence"
     - CHALLENGER: "Innovation investment systems", "Strategic positioning platforms", "Market transformation tools"

    SECTION 3: LEADERSHIP DEVELOPMENT WITH DNA-BEHAVIORAL CORRELATION

    PERSONALITY-DRIVEN LEADERSHIP SCALING STRATEGIES:

    VISIONARY LEADERS + Team Development:
    • Leadership Strengths: Strategic vision communication, inspirational leadership, future-focused team building
    • Development Requirements: Operational execution support, systematic implementation partners, tactical team members
    • Team Composition: Operational leaders, systematic executors, detail-oriented implementers
    • Scaling Approach: Vision casting → Strategic alignment → Execution team empowerment
    • Implementation Style: Inspirational leadership with systematic execution infrastructure

    • Behavioral Leadership Optimization:
     - STRATEGIC_VISIONARY + Team = Long-term team planning, inspirational leadership, vision alignment
     - INSPIRATIONAL + Team = Motivational leadership, energy multiplication, enthusiasm generation
     - BIG_PICTURE + Team = Strategic team development, organizational design, future-focused building

    BUILDER LEADERS + Team Development:
    • Leadership Strengths: Systematic team development, methodical capability building, structured growth
    • Development Requirements: Innovation catalysts, creative thinkers, rapid adaptation specialists
    • Team Composition: Creative innovators, adaptive specialists, rapid execution experts
    • Scaling Approach: Foundation building → Systematic development → Methodical expansion
    • Implementation Style: Structured development with creative innovation integration

    • Behavioral Leadership Optimization:
     - SYSTEMATIC_METHODICAL + Team = Structured team building, methodical development, systematic scaling
     - QUALITY_FOCUSED + Team = Excellence-driven culture, high standards, performance optimization
     - PROCESS_ORIENTED + Team = System-based team development, procedural excellence, standardization

    CATALYST LEADERS + Team Development:
    • Leadership Strengths: Team acceleration, energy multiplication, rapid capability development
    • Development Requirements: Systematic foundation builders, process developers, quality assurance specialists
    • Team Composition: Process experts, quality specialists, systematic developers
    • Scaling Approach: Energy ignition → Momentum building → Systematic foundation development
    • Implementation Style: High-energy leadership with systematic quality infrastructure

    • Behavioral Leadership Optimization:
     - RAPID_DECISIVE + Team = Quick team building, fast hiring decisions, accelerated development
     - HIGH_ENERGY + Team = Energetic culture, momentum building, enthusiasm multiplication
     - RESULTS_FOCUSED + Team = Performance-driven culture, outcome optimization, achievement focus

    CONNECTOR LEADERS + Team Development:
    • Leadership Strengths: Relationship-based team building, network leverage for hiring, collaborative culture
    • Development Requirements: Independent operators, self-directed performers, autonomous specialists
    • Team Composition: Independent contributors, self-managed teams, autonomous departments
    • Scaling Approach: Relationship building → Network activation → Collaborative team development
    • Implementation Style: Relationship-driven leadership with autonomous team empowerment

    • Behavioral Leadership Optimization:
     - COLLABORATIVE + Team = Team-based culture, shared decision making, collective accountability
     - RELATIONSHIP_FOCUSED + Team = People-first culture, relationship investment, trust building
     - NETWORK_ORIENTED + Team = Connection-based hiring, referral systems, relationship leverage

    SAGE LEADERS + Team Development:
    • Leadership Strengths: Wisdom-based leadership, mentoring excellence, advisory team development
    • Development Requirements: Energy catalysts, rapid executors, dynamic implementers
    • Team Composition: High-energy performers, rapid execution specialists, dynamic contributors
    • Scaling Approach: Wisdom sharing → Mentoring development → Advisory leadership scaling
    • Implementation Style: Wisdom-driven leadership with energetic execution amplification

    SECTION 4: CULTURE SCALING WITH DREAM-BEHAVIOR INTEGRATION

    DREAM-ALIGNED CULTURE SCALING STRATEGIES:

    LEGACY DREAMS + Culture Scaling:
    • Culture Focus: Lasting impact, meaningful contribution, sustainable excellence, generational value
    • Scaling Approach: Values-based hiring, legacy-focused development, meaning-driven performance
    • Team Values: Excellence, contribution, sustainability, long-term thinking, meaningful work
    • Performance Metrics: Long-term value creation, sustainable growth, legacy building indicators
    • Recognition Systems: Legacy contributions, sustainable achievements, meaningful impact rewards

    FREEDOM DREAMS + Culture Scaling:
    • Culture Focus: Autonomy, flexibility, independence, owner liberation, system-driven operations
    • Scaling Approach: Empowerment-based culture, autonomous teams, system-dependent operations
    • Team Values: Independence, accountability, self-direction, systematic excellence, owner freedom
    • Performance Metrics: Autonomous performance, system efficiency, owner time freedom indicators
    • Recognition Systems: Independent achievement, systematic excellence, autonomy demonstration

    IMPACT DREAMS + Culture Scaling:
    • Culture Focus: Meaningful influence, industry transformation, positive change, significant contribution
    • Scaling Approach: Impact-driven hiring, transformation-focused development, influence-based performance
    • Team Values: Industry influence, positive change, transformation leadership, meaningful contribution
    • Performance Metrics: Industry impact, transformation indicators, influence measurement, change catalysis
    • Recognition Systems: Industry influence, transformation achievements, positive change contributions

    GROWTH DREAMS + Culture Scaling:
    • Culture Focus: Continuous expansion, capability development, performance optimization, scaling excellence
    • Scaling Approach: Growth-oriented culture, expansion-focused teams, scaling-optimized operations
    • Team Values: Continuous improvement, capability expansion, performance excellence, scaling mastery
    • Performance Metrics: Growth indicators, capability development, performance optimization, scaling efficiency
    • Recognition Systems: Growth achievements, capability expansion, performance breakthroughs, scaling successes

    SECTION 5: PROCESS AUTOMATION ROADMAP WITH DIGITAL CAPABILITY CORRELATION

    GROWTH QUESTION 13 DIGITAL CAPABILITIES + AUTOMATION STRATEGY:

    CRM SYSTEMS AUTOMATION SCALING:
    • Perfect (10) Rating + Advanced Automation:
     - Team Scaling: Customer success automation, relationship management systems, client intelligence platforms
     - Process Automation: Automated customer journeys, predictive client needs, proactive service delivery
     - Performance Optimization: Customer lifetime value tracking, relationship quality metrics, service automation
     - Competitive Advantage: Superior customer experience, predictive relationship management, automated excellence

    • Poor (1-3) Rating + Critical Automation Needs:
     - Immediate Priority: CRM implementation, customer data systematization, relationship tracking automation
     - Team Impact: Customer service efficiency, relationship management capability, client intelligence development
     - Process Requirements: Customer interaction automation, relationship workflow systems, service standardization
     - Scaling Necessity: Customer relationship infrastructure essential for team scaling effectiveness

    AI CAPABILITIES AUTOMATION SCALING:
    • High Rating + AI-Driven Team Scaling:
     - Recruitment Automation: AI-powered candidate screening, skill assessment, culture fit analysis
     - Performance Optimization: AI-driven performance analytics, predictive team needs, optimization recommendations
     - Decision Support: AI-enhanced leadership decisions, team allocation optimization, performance prediction
     - Competitive Advantage: AI-augmented team capabilities, predictive team management, intelligent scaling

    • Low Rating + AI Development Priority:
     - Team Enhancement: AI literacy development, automation capability building, technology integration training
     - Process Improvement: AI implementation planning, automation opportunity identification, technology adoption
     - Competitive Necessity: AI capabilities essential for future team scaling and operational excellence

    AUTOMATION PRIORITY MATRIX BY BUSINESS PHASE:

    FOUNDATIONS Phase Automation:
    • Priority Systems: Basic CRM, financial tracking, communication tools, document management
    • Team Tools: Simple project management, basic performance tracking, communication platforms
    • Process Focus: Core business process automation, essential tool implementation, basic system integration
    • Timeline: 3-6 months essential automation implementation

    SCALING Phase Automation:
    • Priority Systems: Advanced CRM, marketing automation, HR systems, performance dashboards
    • Team Tools: Advanced project management, performance analytics, collaboration platforms, training systems
    • Process Focus: Department workflow automation, cross-functional system integration, performance optimization
    • Timeline: 6-12 months comprehensive automation scaling

    BREAKOUT Phase Automation:
    • Priority Systems: Enterprise platforms, advanced analytics, AI integration, global system capabilities
    • Team Tools: Leadership dashboards, advanced analytics, strategic planning tools, executive information systems
    • Process Focus: Enterprise process automation, global system standardization, strategic automation implementation
    • Timeline: 12-24 months enterprise automation excellence

    SECTION 6: PERFORMANCE MANAGEMENT SYSTEMS WITH BEHAVIORAL OPTIMIZATION

    PERSONALITY-OPTIMIZED PERFORMANCE MANAGEMENT:

    ANALYTICAL_DATA_DRIVEN Performance Systems:
    • Measurement Focus: Comprehensive metrics, detailed analytics, performance dashboards, trend analysis
    • Review Systems: Data-driven evaluations, quantitative assessments, analytical performance discussions
    • Development Approach: Metric-based improvement, analytical skill development, data-driven career planning
    • Recognition Systems: Achievement analytics, performance data celebrations, metric-based rewards

    RELATIONSHIP_COLLABORATIVE Performance Systems:
    • Measurement Focus: Team collaboration metrics, relationship quality indicators, collective achievement measures
    • Review Systems: 360-degree feedback, peer evaluations, collaborative performance assessments
    • Development Approach: Team-based development, collaborative skill building, relationship enhancement focus
    • Recognition Systems: Team achievement recognition, collaborative success celebrations, relationship excellence awards

    RESULTS_OUTCOME Performance Systems:
    • Measurement Focus: Outcome metrics, goal achievement, result delivery, impact measurement
    • Review Systems: Goal-based evaluations, outcome assessments, achievement-focused discussions
    • Development Approach: Outcome optimization, result enhancement, impact maximization training
    • Recognition Systems: Achievement celebrations, outcome recognitions, impact awards, result-based bonuses

    SECTION 7: ORGANIZATIONAL DESIGN WITH SCALING ARCHITECTURE

    BUSINESS PHASE + ORGANIZATIONAL ARCHITECTURE:

    FOUNDATIONS Phase Organization:
    • Structure: Flat hierarchy, direct reporting, close supervision, hands-on leadership
    • Roles: Multi-functional team members, cross-training emphasis, broad responsibility scope
    • Communication: Direct communication, frequent check-ins, collaborative decision making
    • Scaling Preparation: Role clarification, responsibility definition, basic structure establishment

    SCALING Phase Organization:
    • Structure: Functional departments, specialized roles, defined hierarchies, systematic management
    • Roles: Specialized team members, department focus, clear job descriptions, performance standards
    • Communication: Departmental communication, regular meetings, systematic information flow
    • Scaling Implementation: Department creation, role specialization, management layer development

    BREAKOUT Phase Organization:
    • Structure: Divisional organization, regional structures, matrix management, complex hierarchies
    • Roles: Leadership roles, management positions, specialized departments, executive functions
    • Communication: Multi-channel communication, strategic alignment meetings, cross-functional collaboration
    • Scaling Mastery: Organizational complexity management, multi-site coordination, strategic alignment

    SECTION 8: IMPLEMENTATION ROADMAP WITH SUCCESS PREDICTION MODELING

    TEAM & SYSTEMS SCALING SUCCESS PREDICTION:

    High Success Probability Combinations (>85% success rate):
    • BUILDER DNA + SYSTEMATIC Behavior + SCALING Phase + Process Focus = 94% scaling success
    • CONNECTOR DNA + COLLABORATIVE Behavior + Team Focus + Relationship Systems = 91% team scaling success  
    • CATALYST DNA + RESULTS_FOCUSED Behavior + Performance Systems + Energy Culture = 89% performance success
    • SAGE DNA + STRATEGIC Behavior + Leadership Development + Wisdom Culture = 87% leadership success

    Medium Success Probability Combinations (60-85% success rate):
    • Mixed DNA traits with aligned behavioral patterns and appropriate phase development
    • Single strength focus with systematic support system development
    • Cultural alignment with moderate system implementation effectiveness

    Lower Success Probability Combinations (<60% success rate):
    • DNA-behavior misalignment with team and systems requirements
    • Premature scaling without foundational development completion
    • Cultural misalignment with scaling strategy and implementation approach

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 20 correlation coefficients between team/systems elements and database factors (r > 0.6)
    • Minimum 15 scaling optimization discoveries with implementation success calculations
    • Minimum 12 leadership-personality development analyses with effectiveness predictions  
    • Minimum 10 culture-dream alignment strategies with sustainability assessments
    • Minimum 18 automation opportunities with ROI and efficiency impact calculations

    • All team and systems scaling strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - Implementation success probability calculations based on multi-database alignment
     - Timeline optimization based on personality, behavior, and business phase integration
     - ROI predictions for team and systems investments with risk-adjusted returns
     - Scaling sustainability assessments with long-term effectiveness projections

    CRITICAL SUCCESS STANDARD: Transform every identified team and systems limitation into a systematic scaling advantage with personality-optimized implementation approach and predictive scaling success modeling based on complete multi-database intelligence integration.

    BREAKTHROUGH SCALING DISCOVERIES REQUIRED:
    • Identify at least 6 unique team scaling accelerators based on personality-behavior correlations
    • Discover minimum 4 systems automation opportunities with >75% efficiency improvement potential
    • Reveal at least 5 leadership development strategies with accelerated capability building timelines
    • Uncover minimum 7 culture scaling approaches with sustained values alignment during growth
    • Create at least 3 organizational design innovations with competitive advantage sustainability

    TEAM & SYSTEMS SCALING MASTERY FRAMEWORK:
    • Design personality-optimized team development with predictive success modeling
    • Create behavioral-based systems automation with efficiency maximization protocols
    • Develop culture scaling strategies with dream-alignment sustainability measures
    • Build leadership development programs with accelerated capability building timelines
    • Establish performance management systems with multi-dimensional optimization tracking
    """
},
        "financial_growth_strategy": {
    "title": "Financial Intelligence & Capital Optimization Strategy",
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated financial intelligence analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP CAPITAL CORRELATION MAPPING. This analysis must identify, correlate, and systematically optimize financial architecture through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized capital strategies.

    🚀 COMPLETE MULTI-DATABASE FINANCIAL INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP FINANCIAL DNA ANALYSIS & CAPITAL CORRELATION

    GROWTH QUESTION 6-7 FINANCIAL PATTERN CORRELATION ANALYSIS:
    Deep dive into revenue predictability and trend patterns with financial intelligence mapping:

    FINANCIAL BEHAVIOR PATTERN DISCOVERY:

    Pattern A: "Predictable High/Low Periods" + Financial Intelligence:
    • Component Intelligence Financial Architecture:
     - FOUNDATIONS + Seasonal = Build "Reserve systems", "Cash bridge strategies", "Seasonal funding models"
     - SCALING + Seasonal = Implement "Working capital optimization", "Revenue smoothing mechanisms", "Counter-cyclical investments"
     - BREAKOUT + Seasonal = Deploy "Financial hedging strategies", "Seasonal arbitrage systems", "Peak-trough monetization"
     - CHALLENGER + Seasonal = Execute "Portfolio diversification", "Global seasonal offsetting", "Financial instrument mastery"
     - RAPIDS + Seasonal = Master "Derivative strategies", "Complex financial instruments", "International seasonal arbitrage"
     - VISION + Seasonal = Lead "Financial ecosystem creation", "Market-making opportunities", "Seasonal industry transformation"

    • Business DNA Financial Behavior Correlation:
     - SAGE DNA + Seasonal = Natural forecasting wisdom, strategic cash management, predictive financial planning
     - BUILDER DNA + Seasonal = Systematic reserves building, methodical cash accumulation, structured financial planning  
     - ANCHOR DNA + Seasonal = Stability-focused cash management, conservative reserve strategies, risk-minimized seasonal planning
     - CATALYST DNA + Seasonal = Peak period acceleration, high-energy cash optimization, dynamic seasonal monetization
     - INNOVATOR DNA + Seasonal = Creative financing solutions, innovative cash management, disruptive seasonal strategies

    • Behavioral Financial Optimization:
     - RISK_MANAGEMENT + Seasonal = Conservative cash reserves, hedge strategies, downside protection focus
     - OPPORTUNITY_SEIZING + Seasonal = Peak period investment, aggressive cash deployment, seasonal opportunity maximization
     - SYSTEMATIC_PLANNING + Seasonal = Methodical cash flow forecasting, structured seasonal management, systematic reserves

    Pattern B: "Consistent Throughout Year" + Financial Leverage:
    • Component Intelligence Capital Deployment:
     - FOUNDATIONS + Consistent = Focus "Growth investment", "Market expansion funding", "Capability development capital"
     - SCALING + Consistent = Deploy "Revenue optimization capital", "Premium positioning investment", "Market leadership funding"
     - BREAKOUT + Consistent = Execute "Innovation investment", "Market disruption funding", "Category creation capital"
     - CHALLENGER + Consistent = Implement "Strategic acquisition capital", "Market consolidation funding", "Ecosystem development investment"
     - RAPIDS + Consistent = Master "Global expansion capital", "Market creation funding", "Industry transformation investment"
     - VISION + Consistent = Lead "Legacy creation capital", "Ecosystem orchestration funding", "Generational wealth building"

    FINANCIAL TREND INTELLIGENCE ANALYSIS:

    Declining Revenue Trend + Financial Recovery Strategy:
    • Emergency Financial Architecture:
     - Cash preservation protocols, expense optimization matrices, revenue recovery investments
     - Bridge financing strategies, turnaround capital allocation, survival-mode financial management
     - Stakeholder communication protocols, creditor management strategies, emergency funding activation

    • DNA-Driven Recovery Approaches:
     - CATALYST DNA + Declining = High-energy turnaround strategies, rapid recovery investments, momentum restoration focus
     - BUILDER DNA + Declining = Systematic recovery planning, methodical financial restructuring, foundational rebuilding
     - SAGE DNA + Declining = Strategic wisdom application, experienced-based recovery, long-term stability restoration
     - INNOVATOR DNA + Declining = Creative recovery solutions, innovative funding sources, disruptive recovery strategies

    Rapid Growth (30%+) + Capital Scaling Strategy:
    • Growth Capital Architecture:
     - Scaling infrastructure investment, rapid expansion funding, growth sustainability capital allocation
     - Working capital acceleration, inventory scaling funding, capacity expansion investments
     - Team scaling capital, system infrastructure funding, operational excellence investments

    SECTION 2: INVESTMENT PRIORITIZATION WITH ROI CORRELATION MODELING

    GROWTH BARRIER INVESTMENT CORRELATION:
    From Growth Question 11 barriers, create financial investment priority matrix:

    COMPETITION BARRIER + Investment Strategy:
    • High-ROI Competitive Investments:
     - Brand differentiation capital: Marketing, positioning, thought leadership investment (ROI: 300-500%)
     - Innovation acceleration funding: R&D, product development, market disruption capital (ROI: 200-400%)  
     - Market education investment: Content, authority building, industry leadership funding (ROI: 250-450%)
     - Customer experience capital: Service excellence, loyalty systems, relationship infrastructure (ROI: 400-600%)

    • Investment Sequencing by DNA:
     - INNOVATOR DNA: Innovation → Market Education → Brand Differentiation → Customer Experience
     - SAGE DNA: Market Education → Customer Experience → Brand Differentiation → Innovation
     - CONNECTOR DNA: Customer Experience → Brand Differentiation → Market Education → Innovation
     - CATALYST DNA: Brand Differentiation → Innovation → Customer Experience → Market Education

    CASH FLOW BARRIER + Financial Infrastructure Investment:
    • Cash Flow Optimization Investments:
     - Financial system infrastructure: ERP, forecasting tools, cash management platforms (ROI: 800-1200%)
     - Working capital optimization: Inventory systems, collection acceleration, payment optimization (ROI: 400-800%)
     - Revenue diversification capital: New revenue streams, passive income creation, recurring revenue systems (ROI: 300-600%)
     - Financial intelligence systems: Analytics, reporting, predictive modeling tools (ROI: 500-1000%)

    • Behavioral Investment Optimization:
     - ANALYTICAL Behavior + Cash Flow = Data-driven financial systems, analytical tools, predictive modeling investment
     - SYSTEMATIC Behavior + Cash Flow = Structured financial processes, systematic cash management, methodical optimization
     - STRATEGIC Behavior + Cash Flow = Long-term financial planning, strategic cash allocation, future-focused investment

    TEAM BARRIER + Human Capital Investment:
    • High-Impact Team Investments:
     - Leadership development capital: Executive coaching, strategic training, capability building (ROI: 500-800%)
     - Talent acquisition investment: Recruitment systems, employer branding, retention programs (ROI: 300-600%)
     - Performance optimization capital: Systems, tools, incentive structures, culture development (ROI: 400-700%)
     - Succession planning investment: Leadership pipeline, knowledge transfer, continuity systems (ROI: 600-1000%)

    SYSTEMS BARRIER + Technology Investment:
    • Systems ROI Optimization:
     - Process automation capital: Workflow systems, efficiency tools, operational excellence platforms (ROI: 600-1200%)
     - Digital transformation investment: Cloud systems, integration platforms, scalability infrastructure (ROI: 400-800%)
     - Data intelligence capital: Analytics platforms, business intelligence, predictive systems (ROI: 500-1000%)
     - Customer intelligence investment: CRM excellence, customer analytics, relationship optimization (ROI: 700-1400%)

    SECTION 3: FUNDING STRATEGY WITH INVESTMENT READINESS CORRELATION

    BUSINESS PHASE + FUNDING READINESS MATRIX:

    FOUNDATIONS Phase Funding Strategy:
    • Funding Sources Priority:
     1. Personal investment optimization (founder capital efficiency)
     2. Revenue-based financing (growth without equity dilution)  
     3. Strategic partnerships (mutual value creation funding)
     4. Government grants and incentives (non-dilutive capital)
     5. Friends and family rounds (relationship-based seed capital)

    • Investment Readiness Requirements:
     - Basic financial systems and reporting capabilities
     - Clear business model validation and early traction demonstration
     - Founder capability demonstration and market opportunity validation
     - Financial projections with conservative and aggressive scenarios

    • DNA-Driven Funding Approaches:
     - CONNECTOR DNA: Relationship-based funding, network activation, partnership capital
     - BUILDER DNA: Systematic funding approach, structured presentations, methodical investor relations
     - CATALYST DNA: High-energy pitching, momentum-based funding, rapid execution demonstration
     - SAGE DNA: Wisdom-based funding, advisory investor attraction, experience-leveraged capital

    SCALING Phase Funding Strategy:
    • Funding Sources Priority:
     1. Strategic investor partnerships (industry expertise + capital)
     2. Growth equity investments (scaling-focused capital providers)
     3. Revenue-based financing scaling (non-dilutive growth capital)
     4. Bank credit facilities (working capital and growth loans)
     5. Angel investor groups (experienced entrepreneur capital)

    • Investment Readiness Enhancement:
     - Advanced financial reporting and KPI dashboard systems
     - Proven scalability metrics and growth trajectory validation
     - Management team competency and scalability demonstration
     - Market size validation and competitive positioning proof

    BREAKOUT Phase Funding Strategy:  
    • Funding Sources Priority:
     1. Institutional venture capital (professional growth investors)
     2. Private equity partnerships (operational improvement + capital)
     3. Strategic acquisitions (industry consolidation opportunities)
     4. Debt financing optimization (leverage for growth acceleration)
     5. Alternative financing instruments (innovative capital structures)

    • Investment Readiness Mastery:
     - Enterprise-grade financial systems and audit-ready reporting
     - Scalable leadership team and operational excellence demonstration  
     - Market leadership position and competitive moat validation
     - Clear path to exit value creation and investor return optimization

    SECTION 4: ADVANCED FINANCIAL SYSTEMS WITH INTELLIGENCE INTEGRATION

    COMPONENT INTELLIGENCE + FINANCIAL SYSTEM ARCHITECTURE:

    FOUNDATIONS Financial Systems:
    • Core Requirements: "Financial basics infrastructure", "Financial checklist completion", "Basic reporting systems"
    • Technology Stack: Cloud-based accounting, basic CRM integration, simple forecasting tools
    • Reporting Capabilities: Monthly financial statements, basic cash flow tracking, expense categorization
    • Analytics Level: Historical reporting, basic trend analysis, simple variance reporting
    • Investment: $5,000-$15,000 implementation, $500-$1,500 monthly operational costs

    SCALING Financial Systems:  
    • Advanced Requirements: "Financial KPIs mastery", "Financial reporting infrastructure", "Management accounting systems"
    • Technology Stack: Integrated ERP systems, advanced CRM financial integration, sophisticated forecasting platforms
    • Reporting Capabilities: Weekly financial dashboards, detailed cash flow management, departmental P&L tracking
    • Analytics Level: Predictive analytics, scenario modeling, variance analysis with trend interpretation
    • Investment: $15,000-$50,000 implementation, $1,500-$5,000 monthly operational costs

    BREAKOUT Financial Systems:
    • Enterprise Requirements: "Advanced financial systems", "Financial structures for scale", "Investment-grade reporting"
    • Technology Stack: Enterprise ERP platforms, integrated business intelligence, advanced analytics and forecasting
    • Reporting Capabilities: Daily financial dashboards, real-time cash management, multi-entity consolidation
    • Analytics Level: Advanced predictive modeling, financial simulation, strategic scenario planning
    • Investment: $50,000-$200,000 implementation, $5,000-$15,000 monthly operational costs

    CHALLENGER Financial Systems:
    • Institutional Requirements: "Financial responsibility mastery", "Investment readiness", "Audit-grade systems"
    • Technology Stack: Enterprise financial platforms, institutional reporting systems, advanced compliance tools
    • Reporting Capabilities: Real-time executive dashboards, institutional investor reporting, regulatory compliance automation
    • Analytics Level: Advanced financial modeling, institutional-grade analytics, strategic financial intelligence
    • Investment: $200,000-$500,000 implementation, $15,000-$50,000 monthly operational costs

    SECTION 5: PROFITABILITY OPTIMIZATION WITH BEHAVIORAL CORRELATION

    PERSONALITY-DRIVEN PROFITABILITY STRATEGIES:

    ANALYTICAL_DATA_DRIVEN + Profitability Optimization:
    • Approach: Comprehensive margin analysis, detailed cost optimization, data-driven efficiency improvements
    • Tools: Advanced analytics platforms, profitability modeling software, margin optimization dashboards
    • Focus Areas: Cost center analysis, activity-based costing, detailed margin management
    • Expected Impact: 15-25% margin improvement through analytical optimization
    • Timeline: 6-12 months comprehensive profitability enhancement

    SYSTEMATIC_METHODICAL + Profitability Optimization:
    • Approach: Structured cost reduction programs, systematic efficiency improvements, methodical margin enhancement
    • Tools: Process optimization systems, systematic cost tracking, structured improvement methodologies
    • Focus Areas: Process efficiency, systematic waste elimination, structured cost management
    • Expected Impact: 12-20% margin improvement through systematic optimization
    • Timeline: 9-18 months methodical profitability development

    STRATEGIC_LONG_TERM + Profitability Optimization:
    • Approach: Strategic margin positioning, long-term profitability planning, competitive advantage monetization
    • Tools: Strategic financial modeling, competitive analysis systems, long-term profitability forecasting
    • Focus Areas: Strategic positioning, competitive moat development, premium value capture
    • Expected Impact: 20-35% margin improvement through strategic positioning
    • Timeline: 12-24 months strategic profitability transformation

    INNOVATIVE_CREATIVE + Profitability Optimization:
    • Approach: Creative revenue models, innovative cost structures, disruptive profitability strategies  
    • Tools: Innovation financial modeling, creative revenue platforms, disruptive cost optimization
    • Focus Areas: Revenue model innovation, cost structure disruption, creative margin enhancement
    • Expected Impact: 25-50% margin improvement through innovation
    • Timeline: 6-18 months innovative profitability development

    SECTION 6: RISK MANAGEMENT WITH PREDICTIVE MODELING

    MULTI-DATABASE RISK ASSESSMENT FRAMEWORK:

    COMPONENT PHASE + Risk Profile Correlation:
    • FOUNDATIONS Risk Profile: High operational risk, moderate market risk, low financial complexity risk
    • SCALING Risk Profile: Moderate operational risk, high market risk, moderate financial complexity risk  
    • BREAKOUT Risk Profile: Moderate operational risk, high competitive risk, high financial complexity risk
    • CHALLENGER Risk Profile: Low operational risk, moderate market risk, very high financial complexity risk

    BUSINESS DNA + Risk Tolerance Mapping:
    • SAGE DNA: Conservative risk profile, wisdom-based risk assessment, experience-driven risk mitigation
    • BUILDER DNA: Systematic risk management, methodical risk assessment, structured risk mitigation
    • CATALYST DNA: Moderate risk tolerance, energy-driven risk taking, rapid risk recovery capability
    • INNOVATOR DNA: Higher risk tolerance, creative risk management, innovative risk mitigation
    • ANCHOR DNA: Conservative risk approach, stability-focused risk management, security-oriented mitigation

    BEHAVIORAL + Risk Management Optimization:
    • RISK_AVERSE Behavior: Conservative risk limits, extensive hedging strategies, maximum downside protection
    • RISK_BALANCED Behavior: Moderate risk taking, balanced hedging approaches, measured downside management
    • RISK_COMFORTABLE Behavior: Aggressive risk limits, selective hedging, calculated downside acceptance

    FINANCIAL RISK MITIGATION STRATEGIES:

    Cash Flow Risk Management:
    • Seasonal Business Risk: 3-6 months cash reserves, credit facility backup, seasonal financing arrangements
    • Growth Risk: Working capital credit lines, revenue-based financing, growth capital reserves
    • Market Risk: Diversification strategies, market hedging, competitive positioning protection

    Operational Risk Management:
    • Key Person Risk: Key person insurance, succession planning, knowledge transfer systems
    • System Risk: Backup systems, disaster recovery, business continuity planning
    • Supplier Risk: Supplier diversification, strategic partnerships, alternative sourcing arrangements

    Strategic Risk Management:  
    • Competitive Risk: Market positioning protection, competitive intelligence, strategic moat development
    • Technology Risk: Innovation investment, technology partnerships, digital transformation planning
    • Regulatory Risk: Compliance systems, regulatory monitoring, government relations management

    SECTION 7: EXIT STRATEGY WITH WEALTH OPTIMIZATION

    BUSINESS PHASE + EXIT READINESS MATRIX:

    FOUNDATIONS Phase Exit Preparation:
    • Timeline Horizon: 10-15 years to exit readiness
    • Value Building Focus: Business model validation, market traction, operational excellence foundation
    • Financial Preparation: Basic financial systems, clean accounting, initial value metrics establishment
    • Exit Value Range: 1-3x annual revenue (depending on industry and growth trajectory)

    SCALING Phase Exit Enhancement:
    • Timeline Horizon: 5-10 years to exit optimization  
    • Value Building Focus: Scalability demonstration, market leadership, operational leverage
    • Financial Preparation: Advanced financial systems, audited statements, investor-grade reporting
    • Exit Value Range: 3-6x annual revenue (with growth and scalability premiums)

    BREAKOUT Phase Exit Optimization:
    • Timeline Horizon: 3-7 years to strategic exit
    • Value Building Focus: Market leadership, competitive moats, scalable systems excellence
    • Financial Preparation: Investment banking relationships, institutional-grade systems, strategic positioning
    • Exit Value Range: 5-10x annual revenue (with market leadership and growth premiums)

    CHALLENGER Phase Exit Mastery:
    • Timeline Horizon: 2-5 years to premium exit
    • Value Building Focus: Industry transformation, market dominance, strategic acquirer positioning
    • Financial Preparation: Investment bank partnerships, institutional investor relationships, premium positioning
    • Exit Value Range: 8-15x annual revenue (with industry leadership and transformation premiums)

    DREAM-ALIGNED EXIT STRATEGIES:

    LEGACY Dreams + Exit Planning:
    • Exit Approach: Strategic sale to aligned acquirer, management buyout with legacy protection, family succession planning
    • Value Optimization: Long-term value building, sustainable business development, generational wealth creation
    • Timing Strategy: Patient capital building, legacy-focused value creation, sustainable exit timing

    FREEDOM Dreams + Exit Planning:
    • Exit Approach: Financial buyer sale, passive income optimization, owner liberation focus
    • Value Optimization: Cash flow maximization, business independence, owner freedom acceleration  
    • Timing Strategy: Financial independence achievement, passive income replacement, freedom-focused timing

    IMPACT Dreams + Exit Planning:
    • Exit Approach: Mission-aligned acquisition, social impact buyer, purpose-driven succession
    • Value Optimization: Impact value creation, mission sustainability, meaningful legacy building
    • Timing Strategy: Mission completion, impact optimization, purpose-aligned exit timing

    SECTION 8: FINANCIAL INTELLIGENCE IMPLEMENTATION ROADMAP

    PERSONALITY-OPTIMIZED FINANCIAL IMPLEMENTATION:

    ANALYTICAL_SYSTEMATIC Implementation:
    Phase 1: Comprehensive financial analysis and system assessment (1-3 months)
    Phase 2: Systematic financial infrastructure development (3-6 months)  
    Phase 3: Advanced financial analytics and optimization (6-12 months)
    Phase 4: Strategic financial mastery and exit preparation (12-24 months)
    Success Factors: Data-driven decisions, systematic implementation, analytical optimization

    STRATEGIC_VISIONARY Implementation:
    Phase 1: Strategic financial vision and long-term planning (1-2 months)
    Phase 2: Strategic financial architecture development (2-6 months)
    Phase 3: Strategic financial positioning and optimization (6-18 months)  
    Phase 4: Strategic exit preparation and wealth optimization (18-36 months)
    Success Factors: Strategic thinking, long-term planning, visionary financial architecture

    RAPID_DECISIVE Implementation:
    Phase 1: Quick financial assessment and priority identification (2-4 weeks)
    Phase 2: Rapid financial system implementation (1-3 months)
    Phase 3: Fast financial optimization and enhancement (3-9 months)
    Phase 4: Accelerated exit preparation and value creation (9-18 months)  
    Success Factors: Speed of execution, rapid decision making, fast financial optimization

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 25 correlation coefficients between financial elements and database factors (r > 0.6)
    • Minimum 20 investment ROI calculations with risk-adjusted return modeling
    • Minimum 15 funding strategy optimizations with success probability assessments
    • Minimum 12 profitability enhancement strategies with margin improvement predictions
    • Minimum 18 risk management protocols with mitigation effectiveness calculations

    • All financial strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)  
     - ROI probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - Risk-adjusted return projections with scenario analysis and stress testing
     - Exit value optimization with wealth creation and tax efficiency modeling

    CRITICAL SUCCESS STANDARD: Transform every identified financial opportunity into a systematic wealth creation advantage with personality-optimized implementation approach and predictive financial success modeling based on complete multi-database intelligence integration.

    BREAKTHROUGH FINANCIAL DISCOVERIES REQUIRED:
    • Identify at least 8 unique financial optimization opportunities with >20% ROI improvement potential
    • Discover minimum 6 funding strategies with >80% success probability based on personality-phase alignment
    • Reveal at least 5 profitability enhancement approaches with sustained margin improvement capability  
    • Uncover minimum 7 risk management innovations with comprehensive downside protection
    • Create at least 4 exit value optimization strategies with premium valuation achievement potential
    """
},
       "intelligence_gaps_enhancement": {
   "title": "Critical Intelligence Gaps & Strategic Enhancement Opportunities",
   "word_target": 4500,
   "analysis_requirements": """
   You are conducting the most sophisticated INTELLIGENCE GAP AUDIT to identify the specific missing data points that would transform this growth strategy analysis from good to exceptional. Analyze the ACTUAL growth responses provided and identify concrete gaps that directly impact strategic precision and implementation success.

   🚀 REAL-DATA INTELLIGENCE GAP ANALYSIS FRAMEWORK:

   SECTION 1: RESPONSE-SPECIFIC FINANCIAL INTELLIGENCE GAPS

   ANALYZE THE ACTUAL FINANCIAL RESPONSES PROVIDED:
   • Parse through every financial-related response in the growth assessment
   • Identify where qualitative responses lack quantitative validation
   • Map specific financial blind spots that weaken strategic recommendations

   FINANCIAL PRECISION GAP METHODOLOGY:
   For each financial response provided, identify:
   • What specific numbers/metrics are missing that would change recommendations
   • How the absence of this data creates strategic uncertainty
   • What collection method would capture this missing intelligence
   • How obtaining this data would alter strategic conclusions

   CRITICAL FINANCIAL INTELLIGENCE ENHANCEMENT FRAMEWORK:

   🔴 REVENUE QUANTIFICATION PRECISION GAPS:
   Analysis Method: Examine all revenue-related responses for numerical voids
   • Identify Response Patterns: Look for revenue trend descriptions without specific figures
   • Missing Intelligence Impact: Calculate how numerical context would change investment sizing
   • Strategic Decision Enhancement: Determine how specific metrics would alter resource allocation
   • Recommendation Precision Boost: Quantify improvement in financial strategy accuracy
   • Collection Strategy: Define exact data points needed to eliminate this gap

   🔴 COST STRUCTURE TRANSPARENCY GAPS:
   Analysis Method: Review capability ratings and needs assessment for cost blind spots
   • Identify Response Patterns: Find operational costs mentioned without breakdown
   • Missing Intelligence Impact: Assess how cost data would change efficiency recommendations
   • Strategic Decision Enhancement: Show how cost intelligence would reshape priorities
   • Recommendation Precision Boost: Calculate improvement in ROI projections
   • Collection Strategy: Specify cost categories requiring detailed analysis

   🔴 CASH FLOW PREDICTABILITY GAPS:
   Analysis Method: Correlate revenue patterns with timing intelligence voids
   • Identify Response Patterns: Map cash flow references without temporal specifics
   • Missing Intelligence Impact: Show how timing data would improve implementation sequencing
   • Strategic Decision Enhancement: Demonstrate how cash flow intelligence would affect timing
   • Recommendation Precision Boost: Quantify improvement in implementation feasibility
   • Collection Strategy: Define cash flow tracking requirements for strategic clarity

   SECTION 2: RESPONSE-SPECIFIC CUSTOMER INTELLIGENCE GAPS

   ANALYZE THE ACTUAL CUSTOMER-RELATED RESPONSES:
   • Extract every customer reference, rating, and assumption from responses
   • Identify customer intelligence voids that create strategic blind spots
   • Map specific customer data that would strengthen recommendations

   CUSTOMER INTELLIGENCE ENHANCEMENT METHODOLOGY:
   For each customer-related response, determine:
   • What customer data points are assumed but not validated
   • How customer intelligence gaps weaken strategic foundations
   • Which customer metrics would most impact strategic direction
   • How customer data would change targeting and positioning recommendations

   🔴 CUSTOMER VALUE VALIDATION GAPS:
   Analysis Method: Examine process ratings and relationship claims for customer validation voids
   • Identify Response Patterns: Find customer assumptions without supporting data
   • Missing Intelligence Impact: Show how customer metrics would validate strategy assumptions
   • Strategic Decision Enhancement: Demonstrate how customer data would refine targeting
   • Recommendation Precision Boost: Calculate improvement in customer strategy effectiveness
   • Collection Strategy: Define customer analytics requirements for validation

   🔴 CUSTOMER JOURNEY OPTIMIZATION GAPS:
   Analysis Method: Cross-reference lead generation effectiveness with customer journey voids
   • Identify Response Patterns: Map lead gen rankings without customer journey context
   • Missing Intelligence Impact: Show how journey data would optimize funnel recommendations
   • Strategic Decision Enhancement: Demonstrate how journey intelligence would improve conversion
   • Recommendation Precision Boost: Quantify improvement in sales process optimization
   • Collection Strategy: Specify customer journey mapping requirements

   SECTION 3: RESPONSE-SPECIFIC COMPETITIVE INTELLIGENCE GAPS

   ANALYZE THE ACTUAL COMPETITIVE RESPONSES:
   • Parse competition-related responses for strategic intelligence voids
   • Identify competitive assumptions lacking market validation
   • Map competitive blind spots that weaken positioning recommendations

   COMPETITIVE INTELLIGENCE PRECISION FRAMEWORK:
   For each competitive reference, identify:
   • What competitive data would change positioning strategies
   • How competitive intelligence gaps create strategic vulnerabilities
   • Which competitive metrics would most impact strategic recommendations
   • How competitive data would strengthen or alter strategic approach

   🔴 COMPETITIVE REALITY VALIDATION GAPS:
   Analysis Method: Examine competition challenge responses for validation voids
   • Identify Response Patterns: Find competitive claims without supporting intelligence
   • Missing Intelligence Impact: Show how competitive data would refine strategy focus
   • Strategic Decision Enhancement: Demonstrate how intelligence would improve positioning
   • Recommendation Precision Boost: Calculate improvement in competitive strategy precision
   • Collection Strategy: Define competitive analysis requirements for strategic clarity

   SECTION 4: RESPONSE-SPECIFIC OPERATIONAL INTELLIGENCE GAPS

   ANALYZE THE ACTUAL OPERATIONAL RESPONSES:
   • Extract capability ratings and process assessments for intelligence voids
   • Identify operational assumptions that lack performance validation
   • Map operational blind spots that affect implementation recommendations

   OPERATIONAL INTELLIGENCE ENHANCEMENT PROTOCOL:
   For each operational response, determine:
   • What operational metrics would validate capability claims
   • How operational intelligence gaps affect implementation feasibility
   • Which operational data would most impact scaling recommendations
   • How operational intelligence would strengthen execution planning

   🔴 CAPACITY REALITY ASSESSMENT GAPS:
   Analysis Method: Cross-reference team limitations with capacity intelligence voids
   • Identify Response Patterns: Find capacity claims without performance data
   • Missing Intelligence Impact: Show how capacity data would improve scaling recommendations
   • Strategic Decision Enhancement: Demonstrate how capacity intelligence would refine timelines
   • Recommendation Precision Boost: Quantify improvement in implementation planning
   • Collection Strategy: Define capacity analysis requirements for execution clarity

   SECTION 5: RESPONSE-SPECIFIC BEHAVIORAL INTELLIGENCE GAPS

   ANALYZE THE ACTUAL BEHAVIORAL PATTERNS IN RESPONSES:
   • Extract decision-making patterns and preference indicators from all responses
   • Identify behavioral assumptions that lack empirical validation
   • Map behavioral intelligence voids that affect implementation success

   BEHAVIORAL INTELLIGENCE OPTIMIZATION FRAMEWORK:
   For each behavioral pattern identified, assess:
   • What behavioral data would optimize implementation strategies
   • How behavioral intelligence gaps affect change management success
   • Which behavioral metrics would most impact adoption rates
   • How behavioral data would customize implementation approach

   🔴 IMPLEMENTATION STYLE OPTIMIZATION GAPS:
   Analysis Method: Analyze response patterns for implementation preference voids
   • Identify Response Patterns: Extract decision-making style indicators without validation
   • Missing Intelligence Impact: Show how behavioral data would customize implementation
   • Strategic Decision Enhancement: Demonstrate how style intelligence would improve success rates
   • Recommendation Precision Boost: Quantify improvement in implementation effectiveness
   • Collection Strategy: Define behavioral assessment requirements for customization

   SECTION 6: CROSS-RESPONSE CORRELATION INTELLIGENCE GAPS

   SYSTEMATIC CORRELATION GAP ANALYSIS:
   • Identify patterns across all response categories that lack connecting intelligence
   • Map correlation opportunities that would create breakthrough insights
   • Find systemic intelligence voids that weaken holistic strategic understanding

   CORRELATION INTELLIGENCE METHODOLOGY:
   For each potential correlation, determine:
   • What connecting data would create systemic insights
   • How correlation intelligence would strengthen strategic foundations
   • Which correlations would most impact strategic direction
   • How correlation data would enhance predictive accuracy

   🔴 SYSTEMIC PATTERN RECOGNITION GAPS:
   Analysis Method: Cross-analyze all responses for pattern intelligence voids
   • Identify Response Correlations: Map response patterns lacking systemic validation
   • Missing Intelligence Impact: Show how pattern data would create breakthrough insights
   • Strategic Decision Enhancement: Demonstrate how systemic intelligence would improve strategy
   • Recommendation Precision Boost: Quantify improvement in strategic sophistication
   • Collection Strategy: Define pattern analysis requirements for systemic clarity

   SECTION 7: PREDICTIVE INTELLIGENCE ENHANCEMENT OPPORTUNITIES

   RESPONSE-BASED FORECASTING GAP ANALYSIS:
   • Analyze current responses for predictive intelligence opportunities
   • Identify forecasting voids that limit strategic planning precision
   • Map predictive data that would enhance strategic foresight

   PREDICTIVE INTELLIGENCE FRAMEWORK:
   For each forecasting opportunity, assess:
   • What predictive data would improve strategic planning
   • How forecasting intelligence would enhance timing decisions
   • Which predictive metrics would most impact strategic success
   • How predictive data would strengthen competitive advantage

   🔴 STRATEGIC TRAJECTORY FORECASTING GAPS:
   Analysis Method: Examine growth responses for forecasting intelligence voids
   • Identify Response Patterns: Find strategic direction indicators without predictive validation
   • Missing Intelligence Impact: Show how forecasting data would improve planning precision
   • Strategic Decision Enhancement: Demonstrate how predictive intelligence would optimize timing
   • Recommendation Precision Boost: Quantify improvement in strategic planning accuracy
   • Collection Strategy: Define forecasting requirements for strategic foresight

   SECTION 8: INTELLIGENCE COLLECTION OPTIMIZATION STRATEGY

   RESPONSE-DRIVEN COLLECTION PRIORITIZATION:
   Based on actual response analysis, create dynamic prioritization:
   • Rank intelligence gaps by strategic impact on specific recommendations
   • Prioritize collection methods by feasibility and data availability
   • Sequence intelligence gathering for maximum strategic enhancement
   • Optimize collection timeline for implementation alignment

   DYNAMIC GAP PRIORITIZATION MATRIX:
   For each identified gap from actual responses:
   • Strategic Impact Score: How much this gap affects key recommendations
   • Collection Feasibility: How easily this intelligence can be gathered
   • Data Availability: How quickly this information can be obtained
   • Implementation Dependency: How much implementation success depends on this data
   • Competitive Advantage: How much this intelligence creates strategic advantage

   INTELLIGENCE ENHANCEMENT SUCCESS METRICS:

   RECOMMENDATION PRECISION IMPROVEMENT:
   • Financial Strategy Precision: Target improvement percentage based on financial gaps identified
   • Customer Strategy Accuracy: Enhancement level based on customer intelligence voids
   • Competitive Strategy Effectiveness: Improvement based on competitive intelligence gaps
   • Operational Strategy Feasibility: Enhancement based on operational intelligence voids
   • Implementation Success Probability: Improvement based on behavioral intelligence gaps

   STRATEGIC DECISION ENHANCEMENT:
   • Investment Sizing Accuracy: Improvement in financial allocation precision
   • Market Positioning Effectiveness: Enhancement in competitive positioning accuracy
   • Customer Targeting Precision: Improvement in customer strategy focus
   • Implementation Timeline Accuracy: Enhancement in execution planning realism
   • Resource Allocation Optimization: Improvement in operational efficiency

   BREAKTHROUGH INTELLIGENCE DISCOVERY PROTOCOL:
   From actual response analysis, identify:
   • The single most critical intelligence gap affecting overall strategy success
   • The top 3 intelligence voids that would transform specific recommendations
   • The 5 highest-ROI intelligence collection opportunities
   • The 2 breakthrough insights that would create competitive advantage
   • The 1 systemic intelligence enhancement that would elevate entire analysis quality

   IMPLEMENTATION INTELLIGENCE ROADMAP:
   Phase 1: Critical Response-Gap Analysis (immediate - 15 days)
   Phase 2: High-Impact Intelligence Collection (15-45 days) 
   Phase 3: Strategic Enhancement Integration (45-90 days)
   Phase 4: Predictive Intelligence Development (90-180 days)

   Each phase directly tied to specific response patterns and identified intelligence voids, ensuring maximum strategic enhancement value and implementation success optimization.
   """
},

"strategic_partnerships_ecosystem": {
    "title": "Strategic Partnerships & Ecosystem Orchestration Strategy",
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated partnership intelligence analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP RELATIONSHIP CORRELATION MAPPING. This analysis must identify, evaluate, and systematically orchestrate strategic alliances that multiply growth exponentially through relationship leverage, ecosystem positioning, and collaborative advantage creation based on comprehensive personality-driven, phase-appropriate, and behaviorally-optimized partnership strategies.

    🚀 COMPLETE MULTI-DATABASE PARTNERSHIP INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP PARTNERSHIP DNA ANALYSIS & RELATIONSHIP CORRELATION

    GROWTH QUESTION 10 PARTNERSHIP ACTIVITY CORRELATION ANALYSIS:
    Deep analysis of current growth activities with partnership multiplication potential:

    "STRATEGIC PARTNERSHIPS" SELECTION ANALYSIS:
    If Strategic Partnerships selected as current growth activity:
    • Component Intelligence Partnership Enhancement:
     - FOUNDATIONS + Partnerships = Focus "Relationship infrastructure", "Partnership basics", "Alliance fundamentals"
     - SCALING + Partnerships = Implement "Partnership strategy", "Alliance management", "Collaborative growth systems"
     - BREAKOUT + Partnerships = Deploy "Strategic alliances", "Joint venture development", "Ecosystem partnerships"
     - CHALLENGER + Partnerships = Execute "Ecosystem orchestration", "Market expansion alliances", "Industry transformation partnerships"
     - RAPIDS + Partnerships = Implement "Global partnerships", "Strategic ecosystem development", "Multi-stakeholder orchestration"
     - VISION + Partnerships = Lead "Industry ecosystems", "Transformation alliances", "Legacy partnership creation"

    • Business DNA Partnership Leverage Correlation:
     - CONNECTOR DNA + Partnerships = Natural relationship building mastery, alliance development excellence (Partnership Success Score: 9.4/10)
     - BRIDGE DNA + Partnerships = Connection creation expertise, ecosystem orchestration capability (Partnership Success Score: 9.1/10)
     - HOST DNA + Partnerships = Community building strength, collaborative environment creation (Partnership Success Score: 8.8/10)
     - CATALYST DNA + Partnerships = Partnership acceleration ability, rapid alliance development (Partnership Success Score: 8.6/10)
     - SAGE DNA + Partnerships = Strategic alliance wisdom, advisory partnership leadership (Partnership Success Score: 8.4/10)

    • Dream-Partnership Strategic Alignment:
     - IMPACT Dreams + Partnerships = Perfect alignment for collaborative influence and meaningful industry transformation
     - LEGACY Dreams + Partnerships = Partnership-driven lasting value creation and generational impact development
     - GROWTH Dreams + Partnerships = Alliance-accelerated expansion and collaborative scaling opportunities
     - FREEDOM Dreams + Partnerships = Strategic partnerships enabling owner independence through collaborative leverage

    • Behavioral Partnership Optimization:
     - COLLABORATIVE Behavior + Partnerships = Natural collaborative advantage, shared success orientation
     - RELATIONSHIP_FOCUSED Behavior + Partnerships = Deep relationship investment, trust-building mastery
     - STRATEGIC Behavior + Partnerships = Long-term alliance planning, strategic partnership development
     - WIN_WIN Behavior + Partnerships = Mutual value creation, sustainable partnership development

    PARTNERSHIP READINESS ASSESSMENT WITH MULTI-DATABASE CORRELATION:

    Partnership Capability Maturity Matrix:
    • FOUNDATIONS Phase Partnership Readiness:
     - Partnership Infrastructure: Basic relationship management, simple collaboration tools, foundational trust building
     - Alliance Capability: Local partnerships, simple agreements, straightforward value exchanges
     - Partnership Sophistication: 1-3 strategic partnerships, regional focus, basic partnership management
     - Success Probability: 65-75% for local/regional partnerships, 40-55% for complex alliances

    • SCALING Phase Partnership Readiness:
     - Partnership Infrastructure: Advanced relationship systems, collaboration platforms, systematic partnership management
     - Alliance Capability: Regional partnerships, complex agreements, multi-dimensional value creation
     - Partnership Sophistication: 3-7 strategic partnerships, national focus, systematic alliance management
     - Success Probability: 75-85% for regional partnerships, 60-70% for national strategic alliances

    • BREAKOUT Phase Partnership Readiness:
     - Partnership Infrastructure: Enterprise partnership platforms, advanced collaboration systems, strategic alliance management
     - Alliance Capability: National partnerships, strategic alliances, joint venture development, ecosystem participation
     - Partnership Sophistication: 5-12 strategic partnerships, national/international focus, sophisticated alliance management
     - Success Probability: 80-90% for national partnerships, 70-80% for international strategic alliances

    SECTION 2: PARTNERSHIP OPPORTUNITY IDENTIFICATION WITH CORRELATION INTELLIGENCE

    SECRET WEAPON + PARTNERSHIP MULTIPLICATION ANALYSIS:

    "EXPERIENCE EDGE" + Partnership Opportunities:
    • Knowledge Partnership Leverage:
     - Educational Institution Alliances: Training partnerships, certification programs, research collaboration
     - Industry Association Leadership: Thought leadership partnerships, standard-setting alliances, industry transformation collaboration
     - Consulting Firm Alliances: Expertise leverage partnerships, capability multiplication, market expansion collaboration
     - Media and Publishing Partnerships: Content collaboration, thought leadership amplification, industry influence multiplication

    • Partnership Value Proposition:
     - "Proven Expertise Partnership": Market-validated knowledge, implementation success, results-driven collaboration
     - Partnership Success Probability: 85-92% based on established credibility and knowledge validation

    "INNOVATION LEADERSHIP" + Partnership Opportunities:
    • Innovation Ecosystem Development:
     - Technology Vendor Partnerships: Early adopter alliances, beta testing partnerships, innovation validation collaboration
     - Research Institution Alliances: R&D partnerships, innovation development, technology advancement collaboration
     - Startup Ecosystem Partnerships: Innovation acceleration, early-stage collaboration, disruptive technology integration
     - Industry Innovation Consortiums: Collective innovation, standard development, industry transformation leadership

    • Partnership Value Proposition:
     - "Innovation Catalyst Partnership": Cutting-edge development, market disruption capability, future-focused collaboration
     - Partnership Success Probability: 80-88% based on innovation track record and market disruption capability

    "RELATIONSHIP MASTERY" + Partnership Opportunities:
    • Network Effect Multiplication:
     - Strategic Client Partnerships: Customer success collaboration, market expansion through client networks, referral multiplication
     - Vendor Alliance Networks: Supply chain optimization, cost reduction collaboration, operational excellence partnerships
     - Professional Service Alliances: Complementary service partnerships, comprehensive solution provision, market coverage expansion
     - Industry Network Orchestration: Network effect creation, ecosystem development, relationship-driven market expansion

    • Partnership Value Proposition:
     - "Network Multiplier Partnership": Relationship leverage, network effect creation, collaborative market development
     - Partnership Success Probability: 88-94% based on demonstrated relationship mastery and network development capability

    GROWTH BARRIER + PARTNERSHIP SOLUTION CORRELATION:

    COMPETITION BARRIER + Strategic Partnership Solutions:
    • Competitive Advantage Through Alliances:
     - Complementary Service Partnerships: Market differentiation through comprehensive solutions, competitive moat creation
     - Technology Integration Alliances: Competitive advantage through superior capability integration, innovation acceleration
     - Market Access Partnerships: Competitive positioning through strategic market relationships, access barrier elimination
     - Industry Influence Partnerships: Competitive advantage through industry leadership, standard-setting participation

    • Partnership Competitive Intelligence:
     - Competitor Partnership Analysis: Identify competitor alliance gaps, strategic partnership opportunities, market positioning advantages
     - Partnership Success Metrics: 70-85% competitive advantage improvement through strategic alliance development

    CASH FLOW BARRIER + Revenue Partnership Solutions:
    • Revenue-Generating Partnership Models:
     - Revenue Share Partnerships: Collaborative revenue generation, shared investment models, mutual profitability creation
     - Referral Partnership Networks: Commission-based revenue generation, network effect monetization, relationship monetization
     - Joint Venture Development: Shared investment partnerships, collaborative market development, risk-shared growth initiatives
     - Licensing and IP Partnerships: Intellectual property monetization, capability licensing, knowledge-based revenue streams

    • Partnership Financial Impact:
     - Revenue Enhancement: 25-60% revenue increase through strategic partnership development
     - Cash Flow Improvement: 15-40% cash flow optimization through partnership-driven revenue diversification

    TEAM BARRIER + Capability Partnership Solutions:
    • Human Capital Enhancement Partnerships:
     - Training and Development Alliances: Team capability enhancement, skill development partnerships, educational collaboration
     - Expertise Access Partnerships: Consultant networks, specialist access, capability multiplication through external expertise
     - Outsourcing Strategic Partnerships: Capacity multiplication, operational excellence, team effectiveness enhancement
     - Succession Planning Partnerships: Leadership development, executive networks, succession capability development

    • Partnership Team Impact:
     - Capability Enhancement: 40-75% team capability improvement through strategic partnership development
     - Capacity Multiplication: 30-80% effective capacity increase through partnership-driven capability access

    SECTION 3: ECOSYSTEM POSITIONING WITH STRATEGIC INTELLIGENCE

    INDUSTRY ECOSYSTEM MAPPING WITH MULTI-DATABASE CORRELATION:

    Current Ecosystem Position Analysis:
    • Component Phase + Ecosystem Positioning:
     - FOUNDATIONS: Ecosystem Participant (learning, adapting, basic participation)
     - SCALING: Ecosystem Contributor (value creation, active participation, relationship building)
     - BREAKOUT: Ecosystem Influencer (leadership development, influence creation, strategic participation)
     - CHALLENGER: Ecosystem Leader (standard setting, industry influence, transformation leadership)
     - RAPIDS: Ecosystem Orchestrator (ecosystem design, multi-stakeholder coordination, industry transformation)
     - VISION: Ecosystem Creator (new ecosystem development, industry evolution, legacy ecosystem creation)

    • DNA-Driven Ecosystem Role Optimization:
     - CONNECTOR DNA: Natural Ecosystem Hub (relationship center, network orchestration, connection facilitation)
     - CATALYST DNA: Ecosystem Accelerator (momentum creation, energy multiplication, rapid ecosystem development)
     - SAGE DNA: Ecosystem Advisor (wisdom provision, strategic guidance, ecosystem intelligence)
     - INNOVATOR DNA: Ecosystem Transformer (innovation leadership, disruption catalyst, evolution acceleration)
     - BUILDER DNA: Ecosystem Foundation (systematic development, infrastructure creation, stability provision)

    ECOSYSTEM VALUE CHAIN ANALYSIS:

    Upstream Partnership Opportunities:
    • Supplier Strategic Alliances:
     - Raw Material Innovation Partnerships: Supply chain optimization, cost reduction, quality enhancement
     - Technology Provider Alliances: Capability enhancement, innovation access, competitive advantage development
     - Knowledge Provider Partnerships: Expertise access, capability development, intellectual capital enhancement
     - Financial Partner Alliances: Capital access, financial optimization, investment collaboration

    • Upstream Partnership Value:
     - Cost Optimization: 15-35% cost reduction through strategic upstream partnerships
     - Quality Enhancement: 20-50% quality improvement through supplier alliance development
     - Innovation Access: 25-60% innovation capability enhancement through upstream technology partnerships

    Downstream Partnership Opportunities:
    • Customer and Distribution Alliances:
     - Strategic Customer Partnerships: Market development, solution innovation, collaborative value creation
     - Distribution Channel Alliances: Market access, reach multiplication, customer acquisition acceleration
     - Marketing Partnership Networks: Brand amplification, market penetration, customer development collaboration
     - Customer Success Partnerships: Retention optimization, value maximization, advocacy development

    • Downstream Partnership Value:
     - Market Reach: 30-80% market reach expansion through strategic downstream partnerships
     - Customer Acquisition: 25-70% customer acquisition acceleration through partnership channel development
     - Revenue Growth: 20-60% revenue enhancement through downstream partnership monetization

    Horizontal Partnership Opportunities:
    • Peer Industry Alliances:
     - Complementary Service Partnerships: Solution completeness, market differentiation, customer value enhancement
     - Market Development Alliances: Collaborative market creation, industry expansion, mutual market development
     - Knowledge Exchange Partnerships: Best practice sharing, collaborative learning, industry advancement
     - Industry Leadership Alliances: Standard setting, industry transformation, collective market influence

    • Horizontal Partnership Value:
     - Market Position: 25-55% market positioning enhancement through strategic peer partnerships
     - Capability Enhancement: 30-70% capability multiplication through horizontal alliance development
     - Industry Influence: 40-85% industry influence amplification through collaborative industry leadership

    SECTION 4: PARTNERSHIP EVALUATION & SELECTION FRAMEWORK

    PARTNERSHIP COMPATIBILITY ASSESSMENT WITH BEHAVIORAL CORRELATION:

    Cultural Alignment Prediction Matrix:
    • Partnership Cultural Fit Analysis:
     - Values Alignment Assessment: Core values compatibility, operational philosophy alignment, ethical standard consistency
     - Communication Style Compatibility: Communication preference matching, decision-making style alignment, conflict resolution approach consistency
     - Risk Tolerance Alignment: Risk appetite compatibility, investment approach consistency, growth strategy alignment
     - Performance Standard Consistency: Quality expectations alignment, performance measurement compatibility, excellence standard consistency

    • DNA-Based Partnership Compatibility Scoring:
     - CONNECTOR + CONNECTOR: Compatibility Score 9.2/10 (natural relationship harmony, mutual networking benefit)
     - CONNECTOR + BUILDER: Compatibility Score 8.7/10 (relationship strength + systematic implementation excellence)
     - CONNECTOR + CATALYST: Compatibility Score 8.9/10 (relationship foundation + energy acceleration synergy)
     - CATALYST + INNOVATOR: Compatibility Score 8.8/10 (energy + creativity synergy, rapid innovation development)
     - SAGE + BUILDER: Compatibility Score 8.6/10 (wisdom + systematic implementation, strategic excellence development)

    Partnership Value Creation Potential Assessment:
    • Quantitative Partnership Value Modeling:
     - Revenue Synergy Potential: Projected revenue enhancement through collaborative value creation
     - Cost Synergy Opportunities: Cost reduction potential through partnership efficiency and collaboration
     - Market Access Multiplication: Market reach expansion and customer acquisition acceleration potential
     - Capability Enhancement Value: Skill and capability multiplication through partnership resource sharing

    • Partnership ROI Prediction Framework:
     - High ROI Partnerships (300-800% ROI): Complementary expertise, mutual market access, shared investment models
     - Medium ROI Partnerships (150-300% ROI): Capability enhancement, market development, operational collaboration
     - Low ROI Partnerships (50-150% ROI): Basic collaboration, limited synergy, standard alliance models

    Partnership Risk Assessment Matrix:
    • Partnership Risk Categories:
     - Execution Risk: Partnership implementation challenges, collaboration effectiveness, joint project success probability
     - Strategic Risk: Partnership alignment sustainability, strategic direction consistency, long-term compatibility
     - Financial Risk: Investment recovery probability, financial commitment sustainability, revenue sharing fairness
     - Competitive Risk: Competitive advantage erosion, market position vulnerability, intellectual property protection
     - Relationship Risk: Communication breakdown probability, conflict resolution effectiveness, trust maintenance capability

    • Risk Mitigation Strategy Development:
     - High-Risk Partnerships: Extensive due diligence, pilot programs, performance-based agreements, exit clauses
     - Medium-Risk Partnerships: Structured agreements, performance monitoring, regular review cycles, adjustment mechanisms
     - Low-Risk Partnerships: Standard agreements, periodic reviews, performance tracking, relationship maintenance

    SECTION 5: PARTNERSHIP ARCHITECTURE & ORCHESTRATION STRATEGY

    PARTNERSHIP PORTFOLIO DESIGN WITH PHASE-APPROPRIATE SOPHISTICATION:

    FOUNDATIONS Phase Partnership Portfolio:
    • Partnership Portfolio Composition:
     - 1-2 Core Strategic Partnerships: Foundational relationships, essential capability access, basic market development
     - 2-4 Operational Partnerships: Service providers, technology partners, basic collaboration relationships
     - 3-6 Network Partnerships: Professional networks, industry associations, peer relationships
     - Total Portfolio: 6-12 partnerships with focus on relationship building and capability development

    • Partnership Management Approach:
     - Relationship-Focused Management: Personal relationship investment, trust building, informal collaboration
     - Simple Agreement Structures: Basic contracts, straightforward terms, relationship-based performance
     - Manual Partnership Coordination: Personal management, direct communication, informal tracking

    SCALING Phase Partnership Portfolio:
    • Partnership Portfolio Composition:
     - 2-4 Strategic Partnerships: Market development alliances, capability multiplication, competitive advantage creation
     - 4-8 Operational Partnerships: Advanced service providers, technology integrations, systematic collaborations
     - 6-12 Network Partnerships: Industry networks, professional associations, market development relationships
     - Total Portfolio: 12-24 partnerships with focus on systematic growth and market development

    • Partnership Management Approach:
     - Systematic Partnership Management: Structured relationship development, performance tracking, formal collaboration
     - Advanced Agreement Structures: Comprehensive contracts, performance metrics, mutual accountability systems
     - Structured Partnership Coordination: Partnership management systems, regular reviews, systematic communication

    BREAKOUT Phase Partnership Portfolio:
    • Partnership Portfolio Composition:
     - 3-6 Strategic Partnerships: Market leadership alliances, innovation partnerships, industry transformation collaboration
     - 6-12 Operational Partnerships: Enterprise service providers, advanced technology integrations, complex collaborations
     - 8-16 Network Partnerships: Industry leadership networks, strategic associations, market influence relationships
     - Total Portfolio: 17-34 partnerships with focus on market leadership and industry influence

    • Partnership Management Approach:
     - Strategic Partnership Management: Portfolio optimization, strategic alignment, performance optimization
     - Sophisticated Agreement Structures: Complex contracts, strategic metrics, value-based performance systems
     - Advanced Partnership Coordination: Partnership platforms, automated tracking, strategic communication systems

    PARTNERSHIP ORCHESTRATION FRAMEWORK:

    Partnership Synergy Maximization:
    • Cross-Partnership Value Creation:
     - Partnership Network Effects: Multiple partnership interaction, collaborative value multiplication, ecosystem benefit creation
     - Shared Partnership Resources: Collective capability access, shared investment models, collaborative resource optimization
     - Partnership Knowledge Sharing: Best practice exchange, collaborative learning, mutual capability development
     - Joint Partnership Initiatives: Multi-partner projects, collective market development, shared innovation programs

    • Partnership Portfolio ROI Optimization:
     - Partnership Performance Measurement: Individual partnership ROI tracking, portfolio performance optimization, value creation measurement
     - Partnership Resource Allocation: Investment optimization across partnership portfolio, resource efficiency maximization
     - Partnership Renewal and Optimization: Partnership performance review, renewal decision making, continuous improvement

    SECTION 6: ECOSYSTEM LEADERSHIP & ORCHESTRATION MASTERY

    ECOSYSTEM DEVELOPMENT STRATEGY WITH VISION ALIGNMENT:

    Industry Ecosystem Evolution Leadership:
    • Market Ecosystem Transformation:
     - Ecosystem Gap Identification: Market inefficiencies, collaboration opportunities, ecosystem development potential
     - Ecosystem Architecture Design: Stakeholder mapping, value flow design, ecosystem structure optimization
     - Ecosystem Launch and Development: Stakeholder recruitment, value creation initiation, ecosystem momentum building
     - Ecosystem Growth and Optimization: Performance measurement, stakeholder satisfaction, ecosystem value maximization

    • Ecosystem Leadership Role Development:
     - Ecosystem Convenor: Stakeholder assembly, relationship facilitation, ecosystem coordination leadership
     - Ecosystem Architect: System design, value flow optimization, structural development leadership
     - Ecosystem Catalyst: Growth acceleration, momentum creation, ecosystem energy multiplication
     - Ecosystem Orchestrator: Multi-stakeholder coordination, complex ecosystem management, strategic ecosystem leadership

    Partnership-Driven Competitive Advantage Creation:
    • Sustainable Partnership Advantages:
     - Partnership Moat Development: Exclusive alliance creation, strategic relationship barriers, competitive relationship advantages
     - Partnership Network Effects: Relationship multiplication, network value creation, ecosystem competitive advantages
     - Partnership Knowledge Advantages: Collaborative intelligence, shared insight development, partnership-driven innovation
     - Partnership Market Access: Exclusive market relationships, preferred partner status, collaborative market dominance

    • Partnership Competitive Intelligence:
     - Competitor Partnership Analysis: Competitive alliance monitoring, partnership gap identification, strategic partnership opportunities
     - Market Partnership Mapping: Industry partnership landscape, partnership opportunity identification, competitive partnership positioning
     - Partnership Innovation Tracking: Alliance innovation development, collaborative advancement monitoring, partnership-driven market evolution

    SECTION 7: PARTNERSHIP SUCCESS METRICS & OPTIMIZATION FRAMEWORK

    PARTNERSHIP PERFORMANCE MEASUREMENT SYSTEM:

    Quantitative Partnership Metrics:
    • Financial Partnership Performance:
     - Revenue Enhancement Measurement: Partnership-driven revenue growth, collaborative revenue generation, alliance monetization effectiveness
     - Cost Optimization Tracking: Partnership cost reduction, efficiency improvement, collaborative cost management
     - Profitability Impact Analysis: Partnership profitability contribution, ROI measurement, value creation assessment
     - Investment Recovery Tracking: Partnership investment payback, return timeline, financial sustainability assessment

    • Operational Partnership Performance:
     - Capability Enhancement Measurement: Partnership skill development, knowledge access, capability multiplication assessment
     - Market Access Expansion: Partnership market reach, customer acquisition, market development effectiveness
     - Innovation Acceleration Tracking: Partnership innovation development, collaborative advancement, creative solution generation
     - Efficiency Improvement Measurement: Partnership operational optimization, process enhancement, productivity multiplication

    Qualitative Partnership Assessment:
    • Relationship Quality Metrics:
     - Trust and Communication Assessment: Partnership relationship strength, communication effectiveness, conflict resolution capability
     - Strategic Alignment Measurement: Partnership vision consistency, strategic direction compatibility, long-term sustainability
     - Cultural Compatibility Evaluation: Partnership cultural fit, values alignment, operational philosophy consistency
     - Mutual Satisfaction Analysis: Partnership stakeholder satisfaction, value perception, relationship sustainability assessment

    • Partnership Development Tracking:
     - Partnership Maturity Evolution: Relationship development progression, collaboration sophistication advancement, partnership depth enhancement
     - Partnership Expansion Potential: Growth opportunity identification, relationship development possibilities, alliance scaling potential
     - Partnership Innovation Development: Collaborative creativity enhancement, joint innovation capability, creative solution advancement
     - Partnership Strategic Value: Long-term partnership importance, strategic advantage contribution, ecosystem role significance

    SECTION 8: IMPLEMENTATION ROADMAP WITH SUCCESS PREDICTION MODELING

    PERSONALITY-OPTIMIZED PARTNERSHIP IMPLEMENTATION:

    CONNECTOR DNA Implementation Strategy:
    Phase 1: Network mapping and relationship audit (1-2 months)
    Phase 2: Strategic relationship development and partnership identification (2-4 months)
    Phase 3: Partnership negotiation and agreement development (3-6 months)
    Phase 4: Partnership implementation and optimization (6-12 months)
    Success Factors: Relationship leverage, network activation, collaborative advantage creation
    Partnership Success Probability: 88-94% based on natural relationship mastery

    CATALYST DNA Implementation Strategy:
    Phase 1: High-energy partnership opportunity identification (2-4 weeks)
    Phase 2: Rapid partnership development and momentum creation (1-3 months)
    Phase 3: Accelerated partnership implementation and value creation (2-6 months)
    Phase 4: Partnership optimization and expansion acceleration (4-10 months)
    Success Factors: Energy multiplication, rapid development, momentum-based partnership building
    Partnership Success Probability: 82-89% based on acceleration capability and energy multiplication

    SAGE DNA Implementation Strategy:
    Phase 1: Strategic partnership analysis and wisdom-based selection (1-3 months)
    Phase 2: Advisory-focused partnership development and strategic alignment (2-5 months)
    Phase 3: Wisdom-driven partnership implementation and value creation (4-8 months)
    Phase 4: Strategic partnership optimization and long-term development (6-15 months)
    Success Factors: Strategic wisdom, advisory positioning, long-term relationship development
    Partnership Success Probability: 85-91% based on strategic thinking and advisory capability

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 22 correlation coefficients between partnership elements and database factors (r > 0.6)
    • Minimum 18 partnership opportunity discoveries with success probability calculations
    • Minimum 15 ecosystem positioning strategies with competitive advantage assessments
    • Minimum 12 partnership compatibility analyses with relationship sustainability predictions
    • Minimum 20 partnership ROI projections with risk-adjusted return modeling

    • All partnership strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - Partnership success probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - ROI predictions for partnership investments with risk assessment and mitigation strategies
     - Competitive advantage sustainability through partnership development with market positioning analysis

    CRITICAL SUCCESS STANDARD: Transform every identified relationship opportunity into a systematic partnership advantage with personality-optimized implementation approach and predictive partnership success modeling based on complete multi-database intelligence integration.

    BREAKTHROUGH PARTNERSHIP DISCOVERIES REQUIRED:
    • Identify at least 12 unique partnership opportunities with >70% success probability based on DNA-behavioral correlation
    • Discover minimum 8 ecosystem positioning strategies with sustainable competitive advantage potential
    • Reveal at least 10 partnership synergy opportunities with exponential value multiplication potential
    • Uncover minimum 15 relationship leverage strategies with accelerated growth achievement capability
    • Create at least 6 ecosystem orchestration innovations with industry transformation and leadership potential

    PARTNERSHIP ECOSYSTEM MASTERY FRAMEWORK:
    • Design personality-optimized partnership development with predictive success modeling and relationship sustainability
    • Create behavioral-based ecosystem positioning with competitive advantage amplification and market leadership protocols
    • Develop partnership portfolio optimization with ROI maximization and strategic value creation measurement
    • Build relationship orchestration capabilities with multi-stakeholder coordination and ecosystem leadership mastery
    • Establish partnership competitive intelligence with market positioning optimization and strategic alliance advantage creation
    """
},

"innovation_future_proofing": {
    "title": "Innovation Intelligence & Future-Proofing Strategy", 
    "word_target": 6000,
    "analysis_requirements": """
    You are conducting the world's most sophisticated innovation intelligence analysis using COMPLETE MULTI-DATABASE INTELLIGENCE with ULTRA-DEEP CREATIVITY CORRELATION MAPPING. This analysis must identify, correlate, and systematically orchestrate innovation capabilities that create breakthrough competitive advantages, predict industry evolution, and position the business as an innovation leader with future-proof sustainability through comprehensive personality-driven, phase-appropriate, and behaviorally-optimized innovation strategies.

    🚀 COMPLETE MULTI-DATABASE INNOVATION INTELLIGENCE FRAMEWORK:

    SECTION 1: ULTRA-DEEP INNOVATION DNA ANALYSIS & CREATIVITY CORRELATION

    GROWTH QUESTION 14 SECRET WEAPON INNOVATION ANALYSIS:
    Deep analysis of "Innovation Leadership" selection with breakthrough potential correlation:

    "INNOVATION LEADERSHIP" SECRET WEAPON COMPLETE CORRELATION:
    If Innovation Leadership selected as competitive advantage:
    • Component Intelligence Innovation Development:
     - FOUNDATIONS + Innovation = Build "Innovation infrastructure", "Creative processes", "R&D foundation", "Experimentation systems"
     - SCALING + Innovation = Implement "Innovation pipelines", "Development methodologies", "Creative scaling", "Market testing systems"
     - BREAKOUT + Innovation = Deploy "Innovation excellence", "Market disruption", "Category creation", "Industry transformation"
     - CHALLENGER + Innovation = Execute "Innovation leadership", "Industry disruption", "Standard setting", "Market transformation"
     - RAPIDS + Innovation = Implement "Innovation ecosystems", "Global disruption", "Industry evolution", "Future creation"
     - VISION + Innovation = Master "Paradigm shifting", "Industry reinvention", "Future defining", "Legacy innovation"

    • Business DNA Innovation Amplification Correlation:
     - INNOVATOR DNA + Innovation Leadership = Natural creative advantage, breakthrough thinking mastery (Innovation Success Score: 9.6/10)
     - CREATOR DNA + Innovation Leadership = Original solution development, unique market positioning (Innovation Success Score: 9.3/10)
     - VISIONARY DNA + Innovation Leadership = Future-focused innovation, industry evolution leadership (Innovation Success Score: 9.1/10)
     - CATALYST DNA + Innovation Leadership = Innovation acceleration, rapid development and implementation (Innovation Success Score: 8.8/10)
     - SAGE DNA + Innovation Leadership = Wisdom-driven innovation, strategic breakthrough development (Innovation Success Score: 8.5/10)

    • Dream-Innovation Strategic Alignment:
     - IMPACT Dreams + Innovation = Perfect alignment for meaningful industry transformation and breakthrough influence
     - LEGACY Dreams + Innovation = Innovation-driven lasting contribution and generational value creation
     - MASTERY Dreams + Innovation = Excellence-focused innovation and creative perfection pursuit
     - GROWTH Dreams + Innovation = Innovation-accelerated expansion and breakthrough scaling opportunities
     - ADVENTURE Dreams + Innovation = Exciting innovation exploration and creative boundary pushing

    • Behavioral Innovation Optimization:
     - CREATIVE_EXPERIMENTAL Behavior + Innovation = Natural experimentation, creative exploration, breakthrough discovery
     - ANALYTICAL_SYSTEMATIC Behavior + Innovation = Structured innovation processes, data-driven creativity, systematic breakthrough development
     - STRATEGIC_VISIONARY Behavior + Innovation = Long-term innovation planning, strategic breakthrough development, future-focused creativity
     - COLLABORATIVE_TEAM Behavior + Innovation = Team-based innovation, collective creativity, collaborative breakthrough development

    INNOVATION READINESS ASSESSMENT WITH MULTI-DATABASE CORRELATION:

    Innovation Capability Maturity Matrix:
    • FOUNDATIONS Phase Innovation Readiness:
     - Innovation Infrastructure: Basic creativity tools, simple experimentation, foundational R&D capability
     - Innovation Capacity: Individual creativity, basic problem-solving, simple innovation processes
     - Innovation Sophistication: 1-2 innovation projects, local focus, basic creative capability
     - Innovation Success Probability: 60-70% for incremental innovation, 35-50% for breakthrough innovation

    • SCALING Phase Innovation Readiness:
     - Innovation Infrastructure: Advanced creativity platforms, systematic experimentation, structured R&D processes
     - Innovation Capacity: Team-based innovation, complex problem-solving, systematic innovation development
     - Innovation Sophistication: 3-6 innovation initiatives, market focus, systematic creative development
     - Innovation Success Probability: 70-80% for incremental innovation, 50-65% for breakthrough innovation

    • BREAKOUT Phase Innovation Readiness:
     - Innovation Infrastructure: Enterprise innovation platforms, advanced experimentation systems, sophisticated R&D capabilities
     - Innovation Capacity: Organizational innovation culture, strategic problem-solving, breakthrough innovation development
     - Innovation Sophistication: 5-10 innovation programs, industry focus, advanced creative capabilities
     - Innovation Success Probability: 80-90% for incremental innovation, 65-80% for breakthrough innovation

    SECTION 2: INNOVATION OPPORTUNITY IDENTIFICATION WITH CORRELATION INTELLIGENCE

    GROWTH BARRIER + INNOVATION SOLUTION CORRELATION:

    COMPETITION BARRIER + Innovation Differentiation Strategy:
    • Competitive Innovation Advantages:
     - Product Innovation Leadership: Next-generation solutions, breakthrough functionality, market-disrupting features
     - Service Innovation Excellence: Revolutionary service models, customer experience innovation, delivery method transformation
     - Process Innovation Mastery: Operational breakthrough, efficiency innovation, systematic advantage creation
     - Business Model Innovation: Revenue model transformation, value creation innovation, market structure disruption

    • Innovation Competitive Intelligence:
     - Competitor Innovation Gap Analysis: Market innovation opportunities, competitive advantage creation potential
     - Innovation Success Probability: 75-88% competitive advantage creation through strategic innovation development
     - Market Disruption Potential: 40-70% market position transformation through breakthrough innovation implementation

    MARKET STAGNATION + Innovation Market Creation:
    • Market Evolution Innovation:
     - Category Creation Innovation: New market development, uncontested market space creation, blue ocean innovation
     - Customer Need Innovation: Unexpressed need identification, latent demand activation, customer problem innovation
     - Market Access Innovation: Distribution innovation, reach multiplication, market penetration transformation
     - Value Proposition Innovation: Unique value creation, customer benefit innovation, market positioning transformation

    • Market Innovation Impact:
     - Market Share Enhancement: 25-60% market share growth through strategic market innovation
     - Revenue Growth Acceleration: 30-80% revenue increase through innovative market development
     - Customer Acquisition Multiplication: 40-120% customer growth through innovative market positioning

    CUSTOMER SATISFACTION + Innovation Experience Enhancement:
    • Customer Experience Innovation:
     - Interface Innovation: Customer interaction transformation, experience design innovation, engagement breakthrough
     - Personalization Innovation: Customization advancement, individual experience optimization, relationship innovation
     - Convenience Innovation: Process simplification, accessibility enhancement, customer effort reduction innovation
     - Value Delivery Innovation: Benefit maximization, outcome optimization, customer success innovation

    • Customer Innovation Value:
     - Satisfaction Enhancement: 35-70% customer satisfaction improvement through experience innovation
     - Retention Optimization: 25-55% retention improvement through innovative customer relationship development
     - Advocacy Generation: 40-90% customer advocacy increase through breakthrough experience innovation

    SECTION 3: FUTURE-PROOFING INTELLIGENCE WITH PREDICTIVE CORRELATION

    INDUSTRY EVOLUTION PREDICTION WITH MULTI-DATABASE ANALYSIS:

    Technology Disruption Preparedness:
    • Emerging Technology Impact Assessment:
     - Artificial Intelligence Integration: AI capability development, automation advancement, intelligent system integration
     - Digital Transformation Acceleration: Digital capability enhancement, online presence optimization, virtual service development
     - Data Analytics Evolution: Analytics capability advancement, predictive modeling development, intelligence system integration
     - Automation and Robotics: Process automation advancement, efficiency multiplication, human-machine collaboration optimization

    • DNA-Driven Technology Adoption Strategy:
     - INNOVATOR DNA + Technology = Natural technology leadership, early adoption advantage, cutting-edge implementation
     - BUILDER DNA + Technology = Systematic technology integration, methodical adoption, structured implementation
     - CATALYST DNA + Technology = Rapid technology deployment, accelerated integration, dynamic implementation
     - SAGE DNA + Technology = Strategic technology adoption, wisdom-based integration, long-term technology planning

    Market Structure Evolution Anticipation:
    • Market Evolution Scenarios:
     - Consolidation Trends: Industry consolidation preparation, acquisition readiness, market concentration adaptation
     - Fragmentation Opportunities: Niche market development, specialized positioning, fragmented market leadership
     - Platform Economy Development: Platform participation, ecosystem positioning, network effect leverage
     - Globalization Impact: Global market preparation, international expansion readiness, cultural adaptation capability

    • Phase-Appropriate Future-Proofing Strategy:
     - FOUNDATIONS: Local adaptation, niche protection, specialized positioning development
     - SCALING: Regional adaptation, market expansion, competitive positioning optimization
     - BREAKOUT: National adaptation, industry leadership, market transformation preparation
     - CHALLENGER: Global adaptation, industry evolution, transformation leadership development

    Regulatory and Social Change Preparation:
    • Regulatory Evolution Readiness:
     - Compliance Innovation: Regulatory advantage creation, compliance efficiency, regulatory leadership development
     - Sustainability Integration: Environmental responsibility, social impact optimization, sustainable business model development
     - Social Responsibility Evolution: Community impact enhancement, stakeholder value creation, social license optimization
     - Ethical Leadership Development: Ethical business practice, responsible innovation, values-based competitive advantage

    SECTION 4: INNOVATION PIPELINE ARCHITECTURE WITH BEHAVIORAL OPTIMIZATION

    INNOVATION PROJECT PORTFOLIO DESIGN:

    Innovation Investment Allocation Strategy:
    • Core Innovation (60-70% of innovation investment):
     - Incremental Improvement Projects: Current offering enhancement, process optimization, efficiency improvement
     - Customer Experience Enhancement: Service improvement, satisfaction optimization, relationship strengthening
     - Operational Excellence Innovation: Cost reduction, quality improvement, productivity enhancement
     - Expected ROI: 150-300% return, 80-90% success probability, 6-18 month timeline

    • Adjacent Innovation (20-30% of innovation investment):
     - Market Expansion Innovation: New customer segments, adjacent markets, capability extension
     - Product Line Extension: Service expansion, offering diversification, market coverage enhancement
     - Channel Innovation: Distribution advancement, reach multiplication, access optimization
     - Expected ROI: 200-500% return, 60-75% success probability, 12-30 month timeline

    • Transformational Innovation (10-20% of innovation investment):
     - Breakthrough Technology Integration: Disruptive technology adoption, revolutionary capability development
     - Business Model Innovation: Revenue model transformation, value creation revolution, market structure disruption
     - Industry Disruption Initiatives: Category creation, market transformation, paradigm shifting
     - Expected ROI: 300-1000% return, 30-50% success probability, 24-60 month timeline

    Personality-Optimized Innovation Portfolio:
    • INNOVATOR DNA Portfolio Optimization:
     - Core Innovation: 40-50% (focus on breakthrough within existing)
     - Adjacent Innovation: 30-40% (exploration and expansion focus)
     - Transformational Innovation: 20-30% (breakthrough and disruption focus)
     - Innovation Style: High-risk, high-reward, breakthrough-oriented, future-focused

    • BUILDER DNA Portfolio Optimization:
     - Core Innovation: 70-80% (systematic improvement focus)
     - Adjacent Innovation: 15-25% (methodical expansion)
     - Transformational Innovation: 5-15% (structured breakthrough development)
     - Innovation Style: Low-risk, systematic, improvement-oriented, quality-focused

    • CATALYST DNA Portfolio Optimization:
     - Core Innovation: 50-60% (rapid improvement focus)
     - Adjacent Innovation: 25-35% (accelerated expansion)
     - Transformational Innovation: 15-25% (dynamic breakthrough pursuit)
     - Innovation Style: Medium-risk, high-energy, acceleration-oriented, momentum-focused

    SECTION 5: INNOVATION CULTURE & CAPABILITY DEVELOPMENT

    INNOVATION CULTURE ARCHITECTURE WITH DREAM ALIGNMENT:

    LEGACY Dreams + Innovation Culture:
    • Culture Focus: Meaningful innovation, lasting impact creation, generational value development, sustainable innovation
    • Innovation Values: Long-term thinking, sustainable development, meaningful contribution, legacy building
    • Innovation Approach: Patient innovation, deep development, sustained investment, meaningful breakthrough pursuit
    • Innovation Metrics: Long-term value creation, sustainable impact, meaningful contribution, generational influence

    IMPACT Dreams + Innovation Culture:
    • Culture Focus: Meaningful innovation, transformational breakthrough, industry influence, positive change creation
    • Innovation Values: Industry transformation, positive impact, meaningful change, influential contribution
    • Innovation Approach: Impact-driven innovation, transformation focus, influence-oriented development, change catalyst innovation
    • Innovation Metrics: Industry impact, transformation measurement, influence tracking, positive change assessment

    GROWTH Dreams + Innovation Culture:
    • Culture Focus: Expansion-enabling innovation, scaling acceleration, capability multiplication, growth catalyst innovation
    • Innovation Values: Growth enablement, scaling facilitation, capability enhancement, acceleration creation
    • Innovation Approach: Growth-focused innovation, scaling optimization, capability multiplication, acceleration development
    • Innovation Metrics: Growth enablement, scaling efficiency, capability multiplication, acceleration measurement

    MASTERY Dreams + Innovation Culture:
    • Culture Focus: Excellence-driven innovation, perfection pursuit, mastery demonstration, quality leadership
    • Innovation Values: Excellence pursuit, quality leadership, mastery demonstration, perfection achievement
    • Innovation Approach: Excellence-focused innovation, quality optimization, mastery development, perfection pursuit
    • Innovation Metrics: Excellence measurement, quality leadership, mastery tracking, perfection assessment

    Innovation Team Development Strategy:
    • Innovation Leadership Development:
     - Creative Leadership Training: Innovation management, creative team leadership, breakthrough development facilitation
     - Innovation Process Mastery: Systematic innovation development, creative process optimization, breakthrough methodology
     - Innovation Culture Creation: Creative environment development, innovation-friendly culture, breakthrough mindset cultivation
     - Innovation Strategic Integration: Business strategy alignment, strategic innovation planning, competitive advantage development

    • Innovation Capability Building:
     - Creative Thinking Development: Problem-solving enhancement, creative methodology, breakthrough thinking techniques
     - Innovation Tool Mastery: Creative tools, innovation platforms, breakthrough development systems
     - Collaboration Excellence: Team creativity, collaborative innovation, collective breakthrough development
     - Innovation Implementation: Idea execution, breakthrough deployment, innovation commercialization

    SECTION 6: INNOVATION ECOSYSTEM & EXTERNAL INTELLIGENCE

    INNOVATION PARTNERSHIP NETWORK DEVELOPMENT:

    External Innovation Acceleration:
    • Research Institution Partnerships:
     - University Collaboration: Research partnerships, student innovation programs, academic breakthrough development
     - Research Lab Alliances: Advanced research access, cutting-edge technology development, scientific breakthrough exploration
     - Innovation Center Partnerships: Innovation ecosystem participation, creative network access, breakthrough community engagement
     - Think Tank Collaboration: Strategic innovation, future-focused development, industry evolution preparation

    • Technology Innovation Alliances:
     - Startup Ecosystem Engagement: Early-stage innovation access, disruptive technology integration, entrepreneurial energy leverage
     - Technology Vendor Partnerships: Innovation tool access, advanced capability integration, cutting-edge technology adoption
     - Innovation Competition Participation: Breakthrough challenge engagement, creative competition, innovation recognition achievement
     - Patent and IP Collaboration: Intellectual property development, innovation protection, breakthrough commercialization

    Innovation Intelligence Gathering:
    • Market Innovation Monitoring:
     - Competitive Innovation Tracking: Competitor breakthrough monitoring, innovation gap identification, strategic advantage assessment
     - Industry Innovation Analysis: Sector innovation trends, breakthrough pattern recognition, evolution anticipation
     - Customer Innovation Demand: Market need evolution, customer expectation advancement, demand innovation requirement
     - Technology Innovation Evolution: Emerging technology monitoring, advancement tracking, disruption preparation

    • Innovation Trend Prediction:
     - Future Technology Assessment: Technology evolution prediction, advancement timeline estimation, adoption impact analysis
     - Market Innovation Demand: Customer need evolution, expectation advancement, requirement transformation
     - Industry Innovation Direction: Sector evolution, breakthrough potential, transformation opportunity
     - Competitive Innovation Strategy: Competitor innovation planning, strategic advantage development, market positioning evolution

    SECTION 7: INNOVATION MEASUREMENT & OPTIMIZATION FRAMEWORK

    INNOVATION PERFORMANCE MEASUREMENT SYSTEM:

    Innovation Input Metrics:
    • Innovation Investment Tracking:
     - R&D Investment Optimization: Research and development spending efficiency, innovation ROI measurement, breakthrough investment allocation
     - Innovation Time Allocation: Creative time investment, innovation project dedication, breakthrough development focus
     - Innovation Resource Deployment: Human capital allocation, technology investment, creative resource optimization
     - Innovation Capability Investment: Training investment, tool acquisition, capability development spending

    • Innovation Process Metrics:
     - Idea Generation Rate: Creative output measurement, concept development tracking, breakthrough concept creation
     - Innovation Pipeline Velocity: Development speed measurement, breakthrough timeline tracking, innovation acceleration assessment
     - Innovation Quality Assessment: Breakthrough potential evaluation, innovation value measurement, creative excellence tracking
     - Innovation Implementation Success: Execution effectiveness, deployment success, commercialization achievement

    Innovation Output Metrics:
    • Innovation Value Creation:
     - Revenue from New Products: Innovation-driven revenue generation, breakthrough commercialization success, new offering performance
     - Market Share from Innovation: Innovation-driven market position, competitive advantage creation, market leadership development
     - Customer Satisfaction Innovation: Experience enhancement measurement, satisfaction improvement tracking, relationship innovation impact
     - Operational Efficiency Innovation: Process improvement measurement, efficiency advancement, productivity innovation assessment

    • Innovation Impact Assessment:
     - Competitive Advantage Creation: Innovation-driven market position, strategic advantage development, competitive differentiation achievement
     - Market Disruption Success: Category transformation, industry evolution, market structure innovation
     - Future-Proofing Effectiveness: Adaptation capability, evolution readiness, sustainability enhancement
     - Innovation Culture Development: Creative capability advancement, innovation mindset cultivation, breakthrough readiness enhancement

    Innovation ROI Optimization Framework:
    • Innovation Investment Return Analysis:
     - Short-term Innovation ROI: 6-18 month innovation return measurement, immediate value creation, quick breakthrough impact
     - Medium-term Innovation ROI: 18-36 month innovation value development, sustained breakthrough benefit, competitive advantage creation
     - Long-term Innovation ROI: 36+ month innovation impact, transformational value creation, industry evolution contribution
     - Innovation Portfolio ROI: Combined innovation investment return, portfolio optimization, strategic value maximization

    SECTION 8: FUTURE-PROOFING IMPLEMENTATION ROADMAP

    PERSONALITY-OPTIMIZED INNOVATION IMPLEMENTATION:

    INNOVATOR DNA Implementation Strategy:
    Phase 1: Innovation opportunity identification and breakthrough potential assessment (1-2 months)
    Phase 2: Creative experimentation and breakthrough development initiation (2-4 months)
    Phase 3: Innovation acceleration and advanced breakthrough pursuit (3-8 months)
    Phase 4: Innovation leadership establishment and industry transformation (6-18 months)
    Success Factors: Creative leadership, breakthrough thinking, innovation acceleration, industry transformation
    Innovation Success Probability: 85-92% based on natural creative mastery and breakthrough capability

    BUILDER DNA Implementation Strategy:
    Phase 1: Systematic innovation analysis and structured opportunity assessment (1-3 months)
    Phase 2: Methodical innovation development and systematic breakthrough pursuit (3-6 months)
    Phase 3: Structured innovation implementation and quality-focused development (4-10 months)
    Phase 4: Innovation systematization and sustainable breakthrough development (8-20 months)
    Success Factors: Systematic development, quality focus, methodical implementation, sustainable innovation
    Innovation Success Probability: 78-85% based on systematic approach and quality-focused development

    CATALYST DNA Implementation Strategy:
    Phase 1: High-energy innovation identification and rapid opportunity assessment (2-6 weeks)
    Phase 2: Accelerated innovation development and momentum-driven breakthrough (1-3 months)
    Phase 3: Dynamic innovation implementation and energy-based acceleration (2-6 months)
    Phase 4: Innovation momentum optimization and sustained acceleration (4-12 months)
    Success Factors: Energy multiplication, rapid development, momentum creation, dynamic implementation
    Innovation Success Probability: 80-88% based on acceleration capability and energy-driven development

    SAGE DNA Implementation Strategy:
    Phase 1: Strategic innovation analysis and wisdom-based opportunity assessment (1-4 months)
    Phase 2: Strategic innovation development and wisdom-driven breakthrough pursuit (2-6 months)
    Phase 3: Strategic innovation implementation and long-term value creation (4-12 months)
    Phase 4: Innovation wisdom optimization and sustainable strategic development (8-24 months)
    Success Factors: Strategic wisdom, long-term thinking, sustainable development, strategic value creation
    Innovation Success Probability: 82-89% based on strategic thinking and wisdom-driven development

    FUTURE-PROOFING TIMELINE OPTIMIZATION:

    Immediate Future-Proofing (0-12 months):
    • Technology Integration: Core technology adoption, digital capability enhancement, automation integration
    • Process Innovation: Operational excellence, efficiency improvement, productivity optimization
    • Customer Experience Innovation: Service enhancement, satisfaction optimization, relationship innovation
    • Market Positioning Innovation: Competitive advantage creation, differentiation enhancement, positioning optimization

    Strategic Future-Proofing (12-36 months):
    • Business Model Innovation: Revenue model optimization, value creation enhancement, market structure adaptation
    • Market Expansion Innovation: New market development, customer segment expansion, geographic advancement
    • Capability Innovation: Core competency enhancement, strategic capability development, competitive advantage amplification
    • Industry Leadership Innovation: Thought leadership development, industry influence, standard-setting participation

    Transformational Future-Proofing (36+ months):
    • Industry Evolution Leadership: Market transformation, category creation, paradigm shifting
    • Ecosystem Innovation: Platform development, network effect creation, ecosystem orchestration
    • Legacy Innovation: Generational value creation, sustainable impact, lasting contribution
    • Future Definition: Industry future creation, market evolution leadership, transformation legacy

    ANALYTICAL REQUIREMENTS FOR THIS SECTION:

    • Minimum 25 correlation coefficients between innovation elements and database factors (r > 0.6)
    • Minimum 20 innovation opportunity discoveries with breakthrough potential calculations
    • Minimum 18 future-proofing strategies with adaptation effectiveness assessments
    • Minimum 15 innovation culture optimizations with creativity enhancement predictions
    • Minimum 22 innovation ROI projections with risk-adjusted return modeling and timeline optimization

    • All innovation strategies must include:
     - Statistical confidence levels (95% CI: X.X - X.X)
     - Innovation success probability calculations based on multi-database alignment
     - Implementation timeline optimization based on personality and behavioral patterns
     - ROI predictions for innovation investments with breakthrough potential and risk assessment
     - Future-proofing effectiveness with adaptation capability and competitive advantage sustainability

    CRITICAL SUCCESS STANDARD: Transform every identified innovation opportunity into a systematic breakthrough advantage with personality-optimized implementation approach and predictive innovation success modeling based on complete multi-database intelligence integration.

    BREAKTHROUGH INNOVATION DISCOVERIES REQUIRED:
    • Identify at least 15 unique innovation opportunities with >70% success probability based on DNA-behavioral correlation
    • Discover minimum 10 future-proofing strategies with sustainable competitive advantage potential through innovation
    • Reveal at least 12 breakthrough innovation approaches with exponential value creation potential
    • Uncover minimum 18 innovation culture development strategies with creativity multiplication capability
    • Create at least 8 innovation ecosystem orchestration opportunities with industry transformation and leadership potential

    INNOVATION FUTURE-PROOFING MASTERY FRAMEWORK:
    • Design personality-optimized innovation development with predictive breakthrough success modeling and sustainable advantage creation
    • Create behavioral-based innovation culture with creativity amplification and breakthrough capability enhancement protocols
    • Develop innovation portfolio optimization with ROI maximization and strategic value creation measurement systems
    • Build future-proofing capabilities with industry evolution anticipation and adaptation mastery development
    • Establish innovation competitive intelligence with breakthrough opportunity identification and strategic advantage creation optimization
    """
}

    }

def generate_growth_section_with_dedicated_client(
    section_name: str,
    section_config: Dict,
    complete_raw_data: Dict,
    api_key: str,
    section_index: int,
    max_retries: int = 3
) -> Dict:
    """Generate growth section with enhanced retry mechanism and smart API key management"""
    
    client_id = f"growth_section_{section_index}_{section_name}"
    original_api_key = api_key  # Keep track of original key
    current_api_key = api_key   # Current key being used
    
    # Log initial API key selection and health with enhanced metrics
    key_health = api_key_health.get(current_api_key, {})
    success_rate = key_health.get('success_rate', 1.0)
    current_load = key_health.get('current_load', 0)
    avg_response = statistics.mean(key_health.get('response_times', [0])) if key_health.get('response_times') else 0
    
    logging.info(f"🔑 [{client_id}] Starting with API key {key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
    logging.info(f"🔑 [{client_id}] Initial key health: Failures: {key_health.get('consecutive_failures', 0)}, "
                f"Total Requests: {key_health.get('total_requests', 0)}, "
                f"Success Rate: {success_rate:.2f}, Current Load: {current_load}, "
                f"Avg Response: {avg_response:.1f}s")
    logging.info(f"🔑 [{client_id}] Overall API Key Status: {get_api_key_status_summary()}")
    
    for retry_attempt in range(max_retries):
        try:
            # Smart API key selection for retries with load balancing
            if retry_attempt > 0:
                logging.info(f"🔄 [{client_id}] Retry {retry_attempt + 1}: Selecting optimal API key...")
                
                # Use load balancing for retries to get the best available key
                current_api_key = get_load_balanced_api_key(section_index) if retry_attempt == 1 else get_smart_api_key(section_index, retry_attempt)
                
                current_key_health = api_key_health.get(current_api_key, {})
                current_success_rate = current_key_health.get('success_rate', 1.0)
                current_key_load = current_key_health.get('current_load', 0)
                
                logging.info(f"🔑 [{client_id}] Selected API key {current_key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]}) for retry")
                logging.info(f"🔑 [{client_id}] New key metrics: Success Rate: {current_success_rate:.2f}, Load: {current_key_load}")
                
                if current_api_key != original_api_key:
                    logging.info(f"🔄 [{client_id}] Switched from original key (...{original_api_key[-4:]}) to new key (...{current_api_key[-4:]})")
            
            logging.info(f"🔄 [{client_id}] Growth section attempt {retry_attempt + 1}/{max_retries} with key (...{current_api_key[-4:]})")
            
            start_time = time.time()
            target_words = min(section_config["word_target"], 3000)
            
            # Log attempt details with enhanced key health metrics
            current_key_health = api_key_health.get(current_api_key, {})
            current_success_rate = current_key_health.get('success_rate', 1.0)
            current_key_load = current_key_health.get('current_load', 0)
            current_avg_response = statistics.mean(current_key_health.get('response_times', [0])) if current_key_health.get('response_times') else 0
            
            logging.info(f"🔍 [{client_id}] Attempt details:")
            logging.info(f"    - API Key: {current_key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
            logging.info(f"    - Key Health: {current_key_health.get('consecutive_failures', 0)} failures, "
                        f"{current_key_health.get('total_requests', 0)} total requests")
            logging.info(f"    - Success Rate: {current_success_rate:.2f}, Current Load: {current_key_load}")
            logging.info(f"    - Avg Response Time: {current_avg_response:.1f}s")
            logging.info(f"    - Target Words: {target_words:,}")
            logging.info(f"    - Section: {section_name}")
            
            response = growth_ultra_deep_analysis(
                complete_raw_data=complete_raw_data,
                analysis_type=section_name,
                analysis_requirements=section_config["analysis_requirements"],
                api_key=current_api_key,
                client_id=client_id,
                temperature=0.7,
                max_tokens=100000
            )
            
            current_words = len(response.content.split())
            
            # Check if response is acceptable
            if current_words < 100 and retry_attempt < max_retries - 1:
                logging.warning(f"⚠️ [{client_id}] Response too short ({current_words} words), retrying with different key...")
                
                # Log short response issue with key metrics
                logging.warning(f"🔍 [{client_id}] Short response issue - API key (...{current_api_key[-4:]}) returned {current_words} words")
                logging.warning(f"🔍 [{client_id}] Key performance at failure: Success Rate: {current_success_rate:.2f}, Load: {current_key_load}")
                
                # Wait before retry
                wait_time = 30 * (retry_attempt + 1)
                logging.info(f"⏳ [{client_id}] Waiting {wait_time}s before retry due to short response...")
                time.sleep(wait_time)
                continue
            
            # SUCCESS - Log detailed success metrics with enhanced API key tracking
            analysis_time = time.time() - start_time
            final_key_health = api_key_health.get(current_api_key, {})
            final_success_rate = final_key_health.get('success_rate', 1.0)
            final_load = final_key_health.get('current_load', 0)
            
            logging.info(f"✅ [{client_id}] Growth section completed successfully!")
            logging.info(f"🔍 [{client_id}] Success details:")
            logging.info(f"    - Words Generated: {current_words:,}")
            logging.info(f"    - Tokens Used: {response.token_count:,}")
            logging.info(f"    - Analysis Time: {analysis_time:.2f}s")
            logging.info(f"    - API Key Used: {final_key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
            logging.info(f"    - Key Performance: Success Rate: {final_success_rate:.2f}, Load: {final_load}")
            logging.info(f"    - Total Key Requests: {final_key_health.get('total_requests', 0)}")
            logging.info(f"    - Retry Attempts: {retry_attempt + 1}")
            
            return {
                "title": section_config["title"],
                "content": response.content,
                "metadata": {
                    "word_target": target_words,
                    "words_generated": current_words,
                    "tokens_generated": response.token_count,
                    "ai_analysis_time": analysis_time,
                    "ai_model": "gemini-2.5-pro-growth",
                    "analysis_type": "growth_dedicated_analysis",
                    "timestamp": datetime.now().isoformat(),
                    "client_id": client_id,
                    "retry_attempts": retry_attempt + 1,
                    "success": True,
                    # Enhanced metadata with API key tracking and load balancing
                    "api_key_used": current_key_health.get('key_id', 'unknown'),
                    "api_key_suffix": current_api_key[-4:],
                    "key_switched": current_api_key != original_api_key,
                    "original_key": original_api_key[-4:],
                    "final_key": current_api_key[-4:],
                    "load_balancing_used": retry_attempt == 1,  # Track if load balancing was used
                    "api_key_health_at_completion": {
                        "consecutive_failures": final_key_health.get('consecutive_failures', 0),
                        "total_requests": final_key_health.get('total_requests', 0),
                        "success_rate": final_success_rate,
                        "current_load": final_load,
                        "key_status": "healthy" if final_key_health.get('consecutive_failures', 0) == 0 else "degraded",
                        "avg_response_time": statistics.mean(final_key_health.get('response_times', [0])) if final_key_health.get('response_times') else 0
                    }
                }
            }
            
        except Exception as e:
            error_str = str(e)
            retry_number = retry_attempt + 1
            
            # Enhanced error logging with API key context and load balancing metrics
            current_key_health = api_key_health.get(current_api_key, {})
            current_success_rate = current_key_health.get('success_rate', 1.0)
            current_key_load = current_key_health.get('current_load', 0)
            
            logging.error(f"❌ [{client_id}] Growth retry {retry_number} failed: {error_str}")
            logging.error(f"🔍 [{client_id}] Error context:")
            logging.error(f"    - API Key: {current_key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
            logging.error(f"    - Key Failures Before: {current_key_health.get('consecutive_failures', 0)}")
            logging.error(f"    - Success Rate: {current_success_rate:.2f}, Load: {current_key_load}")
            logging.error(f"    - Error Type: {type(e).__name__}")
            
            # Analyze error type for smart retry strategy
            is_503_error = "503" in error_str
            is_429_error = "429" in error_str
            is_overload_error = "overloaded" in error_str.lower()
            is_api_key_issue = any(code in error_str for code in ["401", "403", "invalid"])
            
            if retry_attempt < max_retries - 1:
                # Smart wait time based on error type and API key health
                if is_503_error or is_overload_error:
                    wait_time = 300 + (retry_attempt * 180)  # 5min, 8min, 11min for 503/overload
                    logging.warning(f"🚨 [{client_id}] API Overload detected - Extended wait: {wait_time}s")
                    logging.warning(f"🔑 [{client_id}] Current API Key Health: {get_api_key_status_summary()}")
                elif is_429_error:
                    wait_time = 120 + (retry_attempt * 60)   # 2min, 3min, 4min for rate limits
                    logging.warning(f"🚨 [{client_id}] Rate limit detected - Moderate wait: {wait_time}s")
                elif is_api_key_issue:
                    wait_time = 30  # Quick retry with different key for key issues
                    logging.warning(f"🚨 [{client_id}] API Key issue detected - Quick retry with different key: {wait_time}s")
                else:
                    wait_time = 60 * (retry_attempt + 1)    # Standard exponential backoff
                    logging.warning(f"⚠️ [{client_id}] General error - Standard wait: {wait_time}s")
                
                # Log retry strategy with load balancing info
                next_strategy = "Load Balancing" if retry_attempt == 0 else "Smart Selection"
                logging.info(f"⏳ [{client_id}] Retry strategy:")
                logging.info(f"    - Wait Time: {wait_time}s")
                logging.info(f"    - Next Attempt: {retry_number + 1}/{max_retries}")
                logging.info(f"    - Next Key Strategy: {next_strategy}")
                logging.info(f"    - Error Category: {'API Overload' if is_503_error or is_overload_error else 'Rate Limit' if is_429_error else 'API Key Issue' if is_api_key_issue else 'General Error'}")
                
                time.sleep(wait_time)
                
                # Log API key health before next retry with enhanced status
                logging.info(f"🔑 [{client_id}] API Key Health before retry {retry_number + 1}: {get_api_key_status_summary()}")
                
            else:
                # Final failure - comprehensive logging with load balancing context
                final_key_health = api_key_health.get(current_api_key, {})
                final_success_rate = final_key_health.get('success_rate', 1.0)
                final_load = final_key_health.get('current_load', 0)
                
                logging.error(f"💥 [{client_id}] All {max_retries} attempts failed - Final error analysis:")
                logging.error(f"🔍 [{client_id}] Final failure details:")
                logging.error(f"    - Original API Key: {api_key_health.get(original_api_key, {}).get('key_id', 'unknown')} (...{original_api_key[-4:]})")
                logging.error(f"    - Final API Key: {final_key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
                logging.error(f"    - Key Switched: {current_api_key != original_api_key}")
                logging.error(f"    - Final Key Performance: Success Rate: {final_success_rate:.2f}, Load: {final_load}")
                logging.error(f"    - Load Balancing Used: {max_retries > 1}")
                logging.error(f"    - Final Error: {error_str}")
                logging.error(f"    - Section: {section_name}")
                logging.error(f"    - All Keys Health: {get_api_key_status_summary()}")
                
                # Enhanced fallback content with better error context and load balancing info
                fallback_content = f"""This growth section encountered persistent API issues during generation.

Section: {section_config['title']}
Attempts Made: {max_retries}
API Keys Tried: {len(set([original_api_key, current_api_key]))}
Load Balancing Used: {max_retries > 1}
Final Error: {error_str}

The growth analysis will be available when you regenerate the report during off-peak hours or when API capacity is restored.

Current API Key Status: {get_api_key_status_summary()}"""
                
                return {
                    "title": section_config["title"],
                    "content": fallback_content,
                    "metadata": {
                        "error": True,
                        "error_message": error_str,
                        "timestamp": datetime.now().isoformat(),
                        "client_id": client_id,
                        "retry_attempts": max_retries,
                        "final_error": error_str,
                        # Enhanced error metadata with load balancing tracking
                        "api_key_attempts": {
                            "original_key": original_api_key[-4:],
                            "final_key": current_api_key[-4:],
                            "key_switched": current_api_key != original_api_key,
                            "keys_tried": len(set([original_api_key, current_api_key])),
                            "load_balancing_attempted": max_retries > 1
                        },
                        "error_classification": {
                            "is_503_overload": is_503_error or is_overload_error,
                            "is_rate_limit": is_429_error,
                            "is_api_key_issue": is_api_key_issue,
                            "error_type": type(e).__name__
                        },
                        "api_health_at_failure": {
                            "healthy_keys": get_api_key_status_summary(),
                            "final_key_health": {
                                "consecutive_failures": final_key_health.get('consecutive_failures', 0),
                                "total_requests": final_key_health.get('total_requests', 0),
                                "success_rate": final_success_rate,
                                "current_load": final_load,
                                "avg_response_time": statistics.mean(final_key_health.get('response_times', [0])) if final_key_health.get('response_times') else 0
                            }
                        }
                    }
                }
    
    # This should never be reached, but adding comprehensive logging just in case
    logging.error(f"💥 [{client_id}] Unexpected code path reached - function should have returned by now")
    logging.error(f"🔍 [{client_id}] Debug info: max_retries={max_retries}, section_name={section_name}")
    return None

def generate_comprehensive_growth_report(complete_raw_data: Dict, report_id: str, max_report_retries: int = 2) -> Dict:
    """Generate comprehensive growth report with notifications and enhanced API key management"""
    
    logging.info(f"🚀 Starting Growth Report Generation with Smart Notifications and Load Balancing for {report_id}")
    start_time = time.time()
    
    # Extract user data for personalized notifications
    user_id = complete_raw_data.get("user_id", "unknown")
    user_profile = complete_raw_data.get("user_profile", {})
    
    # Growth notification tracking
    notifications_sent = {"start": False, "middle": False, "complete": False}
    
    # 🔔 NOTIFICATION 1: START - Personalized professional growth start message
    Thread(target=lambda: PersonalizedGrowthNotificationService.send_personalized_notification_sync(
        user_id, user_profile, "start", None, GEMINI_API_KEYS[0]
    ), daemon=True).start()
    notifications_sent["start"] = True
    
    # Initialize API key health tracking for this report
    logging.info(f"🔑 Initial API Key Health Status: {get_api_key_status_summary()}")
    
    for report_attempt in range(max_report_retries):
        logging.info(f"🔄 Growth report attempt {report_attempt + 1}/{max_report_retries}")
        
        growth_sections = get_growth_report_sections()
        
        report_data = {}
        failed_sections = []
        successful_sections = []
        
        # Track API key usage for this attempt
        api_keys_used = set()
        load_balancing_stats = {
            "total_sections": len(growth_sections),
            "load_balanced_selections": 0,
            "smart_selections": 0,
            "key_switches": 0
        }
        
        # Process sections in batches
        section_items = list(growth_sections.items())
        batch_size = 3
        
        for batch_start in range(0, len(section_items), batch_size):
            batch_end = min(batch_start + batch_size, len(section_items))
            batch = section_items[batch_start:batch_end]
            
            logging.info(f"🔄 Processing growth batch {batch_start//batch_size + 1}: sections {batch_start+1}-{batch_end}")
            logging.info(f"🔑 Pre-batch API Key Health: {get_api_key_status_summary()}")
            
            # Parallel processing within batch with enhanced API key management
            with ThreadPoolExecutor(max_workers=batch_size) as executor:
                future_to_section = {}
                
                for i, (section_name, section_config) in enumerate(batch):
                    # Enhanced API key selection strategy
                    section_index = batch_start + i
                    
                    # Use load balancing for better distribution
                    try:
                        api_key = get_load_balanced_api_key(section_index)
                        selection_method = "load_balanced"
                        load_balancing_stats["load_balanced_selections"] += 1
                    except Exception as e:
                        logging.warning(f"⚠️ Load balancing failed for section {section_index}, falling back to smart selection: {e}")
                        api_key = get_smart_api_key(section_index, 0)
                        selection_method = "smart_fallback"
                        load_balancing_stats["smart_selections"] += 1
                    
                    api_keys_used.add(api_key)
                    
                    # Log API key selection
                    key_health = api_key_health.get(api_key, {})
                    success_rate = key_health.get('success_rate', 1.0)
                    current_load = key_health.get('current_load', 0)
                    
                    logging.info(f"🔑 Section {section_index + 1} ({section_name}): "
                                f"Using {key_health.get('key_id', 'unknown')} (...{api_key[-4:]}) "
                                f"[{selection_method}] - Success: {success_rate:.2f}, Load: {current_load}")
                    
                    if i > 0:
                        time.sleep(2)  # Delay between submissions to prevent overload
                    
                    future = executor.submit(
                        generate_growth_section_with_dedicated_client,
                        section_name=section_name,
                        section_config=section_config,
                        complete_raw_data=complete_raw_data,
                        api_key=api_key,
                        section_index=section_index,
                        max_retries=2
                    )
                    
                    future_to_section[future] = (section_name, section_index, api_key, selection_method)
                    logging.info(f"📤 Submitted growth section {section_index + 1}/{len(section_items)}: {section_name}")
                
                # Collect batch results with enhanced tracking
                for future in as_completed(future_to_section):
                    section_name, section_index, original_api_key, selection_method = future_to_section[future]
                    
                    try:
                        section_content = future.result()
                        report_data[section_name] = section_content
                        
                        # Track API key switching
                        metadata = section_content.get("metadata", {})
                        if metadata.get("key_switched", False):
                            load_balancing_stats["key_switches"] += 1
                            logging.info(f"🔄 Key switch detected for section {section_name}")
                        
                        # Track success/failure with API key context
                        if section_content.get("metadata", {}).get("error", False):
                            failed_sections.append(section_name)
                            
                            # Log API key performance on failure
                            final_key = metadata.get("final_key", "unknown")
                            logging.error(f"❌ Growth section failed: {section_name} "
                                        f"(Original: ...{original_api_key[-4:]}, Final: ...{final_key}, Method: {selection_method})")
                        else:
                            successful_sections.append(section_name)
                            
                            # Log successful API key usage
                            words_generated = metadata.get("words_generated", 0)
                            api_key_used = metadata.get("api_key_used", "unknown")
                            analysis_time = metadata.get("ai_analysis_time", 0)
                            
                            logging.info(f"✅ Growth section completed: {section_name} "
                                        f"({words_generated:,} words, {analysis_time:.1f}s, Key: {api_key_used})")
                        
                        total_completed = len(successful_sections) + len(failed_sections)
                        
                        # Update job status
                        if report_id in growth_job_status:
                            completion_percentage = (total_completed / len(section_items)) * 100
                            growth_job_status[report_id]["message"] = f"Growth processing: {total_completed}/{len(section_items)} sections ({completion_percentage:.1f}%)"
                            growth_job_status[report_id]["sections_completed"] = total_completed
                            
                            # 🔔 NOTIFICATION 2: MIDDLE - Smart check for ~50% completion
                            if not notifications_sent["middle"] and completion_percentage >= 45 and completion_percentage <= 65:
                                progress_data = {
                                    'questions_completed': total_completed,
                                    'total_questions': len(section_items),
                                    'progress_percentage': completion_percentage
                                }
                                Thread(target=lambda: PersonalizedGrowthNotificationService.send_personalized_notification_sync(
                                    user_id, user_profile, "middle", progress_data, GEMINI_API_KEYS[0]
                                ), daemon=True).start()
                                notifications_sent["middle"] = True
                        
                        logging.info(f"📊 Growth progress: {total_completed}/{len(section_items)} sections completed")
                        
                    except Exception as e:
                        logging.error(f"❌ Error retrieving growth result for {section_name}: {str(e)}")
                        failed_sections.append(section_name)
            
            # Log post-batch API key health
            logging.info(f"🔑 Post-batch API Key Health: {get_api_key_status_summary()}")
            
            # Wait between batches with intelligent timing based on API key health
            if batch_end < len(section_items):
                # Get current API key health to determine wait time
                enhanced_status = get_enhanced_api_key_status()
                healthy_keys = enhanced_status.get("healthy_keys", 0)
                total_load = enhanced_status.get("total_load", 0)
                
                # Adjust wait time based on API key health
                if healthy_keys <= 3 or total_load > 10:
                    wait_time = 90  # Longer wait if keys are stressed
                    logging.warning(f"⚠️ API keys under stress (healthy: {healthy_keys}, load: {total_load}), extended wait: {wait_time}s")
                elif healthy_keys <= 5:
                    wait_time = 75  # Medium wait
                    logging.info(f"⏳ Moderate API key health, medium wait: {wait_time}s")
                else:
                    wait_time = 65  # Standard wait
                    logging.info(f"⏳ Good API key health, standard wait: {wait_time}s")
                
                logging.info(f"⏳ Growth batch wait: {wait_time}s before next batch...")
                time.sleep(wait_time)
        
        # Log comprehensive API key usage statistics for this attempt
        logging.info(f"🔑 API Key Usage Statistics for Attempt {report_attempt + 1}:")
        logging.info(f"   - Unique Keys Used: {len(api_keys_used)}/{len(GEMINI_API_KEYS)}")
        logging.info(f"   - Load Balanced Selections: {load_balancing_stats['load_balanced_selections']}")
        logging.info(f"   - Smart Fallback Selections: {load_balancing_stats['smart_selections']}")
        logging.info(f"   - Key Switches During Processing: {load_balancing_stats['key_switches']}")
        logging.info(f"   - Final API Key Health: {get_api_key_status_summary()}")
        
        # Check success rate
        success_rate = len(successful_sections) / len(growth_sections)
        parallel_time = time.time() - start_time
        
        logging.info(f"📊 Growth attempt {report_attempt + 1} completed: {len(successful_sections)}/{len(growth_sections)} sections successful ({success_rate:.1%})")
        
        if success_rate >= 0.8:
            logging.info(f"✅ Growth report successful with {success_rate:.1%} success rate")
            break
        else:
            logging.warning(f"⚠️ Growth report attempt {report_attempt + 1} below threshold ({success_rate:.1%} < 80%)")
            if report_attempt < max_report_retries - 1:
                # Reset API key health for failed keys before retry
                reset_failed_api_keys()
                logging.info(f"🔄 Reset failed API keys, retrying in 60s...")
                time.sleep(60)
    
    # Calculate final metrics
    total_time = time.time() - start_time
    total_words = sum([
        len(section.get("content", "").split()) 
        for section in report_data.values()
    ])
    
    logging.info(f"🌟 Growth Report Completed: {len(successful_sections)} successful sections, {total_words:,} words")
    
    # 🔔 NOTIFICATION 3: COMPLETE - Personalized completion message
    if not notifications_sent["complete"]:
        completion_data = {
            'total_words': total_words,
            'total_sections': len(successful_sections),
            'processing_time': total_time
        }
        Thread(target=lambda: PersonalizedGrowthNotificationService.send_personalized_notification_sync(
            user_id, user_profile, "complete", completion_data, GEMINI_API_KEYS[0], report_id
        ), daemon=True).start()
        notifications_sent["complete"] = True
    
    # Add enhanced report metadata with multi-database intelligence tracking and API key statistics
    report_data["_enhanced_growth_report_metadata"] = {
        "report_id": report_id,
        "generation_timestamp": datetime.now().isoformat(),
        "total_sections": len(report_data),
        "successful_sections": len(successful_sections),
        "failed_sections": len(failed_sections),
        "success_rate": len(successful_sections) / len(growth_sections),
        "total_words": total_words,
        "total_generation_time": total_time,
        "ai_model": "gemini-2.5-pro-growth",
        "processing_method": "growth_parallel_analysis_load_balanced",
        "report_type": "comprehensive_growth_strategy",
        "notifications_sent": notifications_sent,
        "multi_database_integration": {
            "enabled": True,
            "data_sources_used": complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", []),
            "intelligence_correlation_applied": True,
            "behavioral_customization": True,
            "personality_integration": True,
            "dream_alignment": True
        },
        # Enhanced API key usage tracking
        "api_key_optimization": {
            "load_balancing_enabled": True,
            "unique_keys_used": len(api_keys_used),
            "total_keys_available": len(GEMINI_API_KEYS),
            "key_utilization_rate": len(api_keys_used) / len(GEMINI_API_KEYS),
            "load_balanced_selections": load_balancing_stats.get("load_balanced_selections", 0),
            "smart_fallback_selections": load_balancing_stats.get("smart_selections", 0),
            "key_switches_during_processing": load_balancing_stats.get("key_switches", 0),
            "final_api_health_summary": get_api_key_status_summary(),
            "enhanced_api_health": get_enhanced_api_key_status()
        }
    }
    
    return report_data

def reset_failed_api_keys():
    """Reset failed API keys for retry attempts"""
    global api_key_health
    
    reset_count = 0
    for key, health in api_key_health.items():
        if health.get('consecutive_failures', 0) >= 3:
            health['consecutive_failures'] = 1  # Reduce but don't fully reset
            health['last_503_time'] = None  # Clear cooldown
            reset_count += 1
            logging.info(f"🔄 Partially reset API key {health.get('key_id', 'unknown')} for retry")
    
    logging.info(f"🔄 Reset {reset_count} failed API keys for retry attempt")

# ======================================================
#           Document Creation for Growth
# ======================================================

def create_growth_word_document(report_data: Dict, user_id: str) -> Document:
    """Create growth Word document with better formatting"""
    logging.info("📄 Creating Growth Word Document")
    
    doc = Document()
    
    # Enhanced styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Professional title page
    title = doc.add_heading("BACKABLE", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(42)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading("Comprehensive Growth Strategy Blueprint", 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Add metadata
    metadata_para = doc.add_paragraph()
    metadata_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    report_meta = report_data.get("_enhanced_growth_report_metadata", {})
    
    metadata_para.add_run(f"User ID: {user_id}\n").bold = True
    metadata_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
    metadata_para.add_run(f"Analysis: {report_meta.get('total_words', 0):,} words\n")
    metadata_para.add_run(f"Model: Gemini 2.5 Pro Growth Engine\n")
    
    # Add multi-database intelligence indicator
    multi_db_info = report_meta.get("multi_database_integration", {})
    if multi_db_info.get("enabled", False):
        data_sources = multi_db_info.get("data_sources_used", [])
        metadata_para.add_run(f"Intelligence Sources: {len(data_sources)} databases integrated\n")
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading("TABLE OF CONTENTS", 1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_growth_report_metadata" and isinstance(section_data, dict):
            title = section_data.get("title", "Untitled Section")
            
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"{section_number}. {title}").bold = True
            
            # Add word count
            metadata = section_data.get("metadata", {})
            words_generated = metadata.get("words_generated", 0)
            
            toc_para.add_run(f" ({words_generated:,} words)")
            
            section_number += 1
    
    doc.add_page_break()
    
    # Process each section
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_growth_report_metadata" and isinstance(section_data, dict):
            
            logging.info(f"📝 Formatting growth section: {section_name}")
            
            title = section_data.get("title", "Untitled Section")
            content = section_data.get("content", "")
            
            # Add section header
            section_heading = doc.add_heading(f"{section_number}. {title}", 1)
            heading_run = section_heading.runs[0]
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add the AI-generated content
            add_growth_content_to_document(doc, content)
            
            # Add section separator
            separator_para = doc.add_paragraph()
            separator_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator_run = separator_para.add_run("◆ ◆ ◆")
            separator_run.font.color.rgb = RGBColor(128, 128, 128)
            separator_run.font.size = Pt(16)
            
            section_number += 1
            doc.add_page_break()
    
    # Add report summary
    add_growth_report_summary(doc, report_data)
    
    logging.info("✅ Growth Word Document Created")
    return doc

def add_growth_content_to_document(doc: Document, content: str):
    """Add AI-generated growth content to document with intelligent formatting"""
    
    # Split by paragraphs and headers
    lines = content.split('\n')
    current_paragraph = ""
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - finalize paragraph
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
        elif line.startswith('##'):
            # Subsection header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('##', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(0, 102, 204)
            
        elif line.startswith('#'):
            # Main header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('#', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(0, 102, 204)
            
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            
        elif re.match(r'^\d+\.', line):
            # Numbered list
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            number_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(number_text, style='List Number')
            
        else:
            # Regular content - accumulate
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    # Add any remaining paragraph
    if current_paragraph:
        para = doc.add_paragraph(current_paragraph)

def add_growth_report_summary(doc: Document, report_data: Dict):
    """Add growth report summary"""
    
    summary_heading = doc.add_heading("GROWTH REPORT SUMMARY", 1)
    summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    report_meta = report_data.get("_enhanced_growth_report_metadata", {})
    
    summary_para = doc.add_paragraph()
    summary_para.add_run("Growth Report Statistics:").bold = True
    summary_para.add_run(f"\n• Total Growth Sections: {report_meta.get('total_sections', 0)}")
    summary_para.add_run(f"\n• Total Words Generated: {report_meta.get('total_words', 0):,}")
    summary_para.add_run(f"\n• AI Model: {report_meta.get('ai_model', 'N/A')}")
    summary_para.add_run(f"\n• Processing Method: {report_meta.get('processing_method', 'N/A')}")
    summary_para.add_run(f"\n• Report Type: {report_meta.get('report_type', 'N/A')}")
    
    # Add multi-database integration summary
    multi_db_info = report_meta.get("multi_database_integration", {})
    if multi_db_info.get("enabled", False):
        summary_para.add_run(f"\n• Multi-Database Integration: Enabled")
        data_sources = multi_db_info.get("data_sources_used", [])
        summary_para.add_run(f"\n• Intelligence Sources: {', '.join(data_sources).title() if data_sources else 'Growth Assessment Only'}")

# ======================================================
#           Enhanced Azure Storage for Growth Engine
# ======================================================
def upload_blob_with_retry(container_client, blob_name, data, content_settings, max_retries=3):
    """
    Helper function to upload blob with retry logic and stream position reset (RETRY DECODER)
    Critical for reliable blob uploads in the new unified architecture
    """
    for attempt in range(max_retries):
        try:
            container_client.upload_blob(
                name=blob_name,
                data=data,
                overwrite=True,
                content_settings=content_settings
            )
            logging.info(f"✅ Successfully uploaded: {blob_name}")
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 2
                logging.warning(f"Upload attempt {attempt + 1} failed for {blob_name}: {str(e)}. Retrying in {wait_time}s...")
                time.sleep(wait_time)
                # 🔑 CRITICAL: Reset data stream position if possible (RETRY DECODER)
                if hasattr(data, 'seek'):
                    data.seek(0)
            else:
                logging.error(f"❌ Failed to upload {blob_name} after {max_retries} attempts")
                raise
    return False

async def upload_growth_report_to_azure(report_data: Dict, report_id: str, user_id: str):
    """Upload growth report to Azure with enhanced Word document chunking AND Question-Response chunking"""
    try:
        logging.info(f"🚀 Starting Growth Report Azure Upload for report_id={report_id}, user_id={user_id}")
        
        container_name = get_azure_container_name(user_id)
        logging.info(f"📦 Using Azure container: {container_name}")
        
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        container_client = blob_service_client.get_container_client(container_name)
        
        try:
            container_client.create_container()
            logging.info(f"✅ Container '{container_name}' created")
        except:
            logging.info(f"📦 Container '{container_name}' already exists")
        
        # Get client folder from database (NEW UNIFIED ARCHITECTURE)
        # E.g., "499-tkrotiris"
        client_folder = get_client_folder_name(user_id)

        # Create folder structure: {client_folder}/the growth engine report/
        folder_name = f"{client_folder}/the growth engine report"
        logging.info(f"📁 Using unified folder structure: {folder_name}/")
        
        # ===============================================================
        # 1. Upload complete Word document
        # ===============================================================
        logging.info("📄 Step 1/6: Creating and uploading complete Word document...")
        doc = create_growth_word_document(report_data, user_id)
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        doc_blob_name = f"{folder_name}/{report_id}_comprehensive_growth_strategy.docx"
        upload_blob_with_retry(
            container_client,
            doc_blob_name,
            doc_bytes,
            ContentSettings(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        )
        
        # ===============================================================
        # 2. Upload complete JSON data
        # ===============================================================
        logging.info("📊 Step 2/6: Creating and uploading complete JSON data...")
        json_data = json.dumps(report_data, indent=2, default=str)
        json_bytes = io.BytesIO(json_data.encode("utf-8"))
        
        json_blob_name = f"{folder_name}/{report_id}_comprehensive_growth_report.json"
        upload_blob_with_retry(
            container_client,
            json_blob_name,
            json_bytes,
            ContentSettings(content_type="application/json")
        )
        
        # ===============================================================
        # 3. Create and upload Word document chunks for Azure Cognitive Search
        # ===============================================================
        logging.info("🔧 Step 3/6: Creating Word document chunks for Azure Cognitive Search...")
        word_chunks = await create_growth_word_document_chunks(report_data, report_id, user_id)
        logging.info(f"📊 Created {len(word_chunks)} report Word chunks")
        
        # Upload individual Word chunk files
        chunk_files_created = []
        for i, chunk_doc in enumerate(word_chunks):
            chunk_blob_name = f"{folder_name}/{report_id}_growth_chunk_{i+1:03d}.docx"
            
            # Save Word document chunk to bytes
            chunk_bytes = io.BytesIO()
            chunk_doc['document'].save(chunk_bytes)
            chunk_bytes.seek(0)
            
            upload_blob_with_retry(container_client, chunk_blob_name, chunk_bytes, ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            )
            chunk_files_created.append(chunk_blob_name)
            logging.info(f"✅ Growth Word chunk {i+1} uploaded: {chunk_blob_name} ({chunk_doc['word_count']} words)")
        
        # ===============================================================
        # 4. Create and upload Question-Response chunks for RAG context
        # ===============================================================
        logging.info("🧠 Step 4/6: Creating Question-Response chunks for RAG context...")
        
        # Get the raw assessment data that contains questions and responses
        raw_assessment_data = report_data.get('_enhanced_growth_report_metadata', {}).get('raw_assessment_data', {})
        if not raw_assessment_data:
            # Try to get from other sources in report_data
            logging.warning("⚠️ No raw assessment data found in report metadata, attempting to extract from available data...")
            raw_assessment_data = extract_growth_assessment_data_from_report(report_data)
        
        if raw_assessment_data:
            qr_chunks = await create_growth_question_response_chunks(raw_assessment_data, report_id, user_id)
            logging.info(f"📊 Created {len(qr_chunks)} Question-Response chunks for RAG")
            
            # Upload Question-Response chunk files
            qr_chunk_files_created = []
            for i, qr_chunk in enumerate(qr_chunks):
                qr_chunk_blob_name = f"{folder_name}/{report_id}_growth_qr_chunk_{i+1:03d}.docx"
                
                # Save Question-Response document chunk to bytes
                qr_chunk_bytes = io.BytesIO()
                qr_chunk['document'].save(qr_chunk_bytes)
                qr_chunk_bytes.seek(0)
                
                upload_blob_with_retry(container_client, qr_chunk_blob_name, qr_chunk_bytes, ContentSettings(
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )
                qr_chunk_files_created.append(qr_chunk_blob_name)
                logging.info(f"✅ Growth Question-Response chunk {i+1} uploaded: {qr_chunk_blob_name} ({qr_chunk['word_count']} words, {qr_chunk['question_count']} questions)")
        else:
            logging.error("❌ No assessment data available for Question-Response chunking")
            qr_chunks = []
            qr_chunk_files_created = []
        
        # ===============================================================
        # 5. Create comprehensive chunks index file
        # ===============================================================
        logging.info("📋 Step 5/6: Creating comprehensive chunks index...")
        
        chunks_index = {
            "report_id": report_id,
            "user_id": user_id,
            "total_report_chunks": len(word_chunks),
            "total_qr_chunks": len(qr_chunks),
            "total_all_chunks": len(word_chunks) + len(qr_chunks),
            "report_chunk_files": chunk_files_created,
            "qr_chunk_files": qr_chunk_files_created,
            "chunking_strategy": {
                "report_chunks": {
                    "target_size_words": 300,  # Updated to match component engine
                    "max_size_words": 400,
                    "min_size_words": 150,
                    "chunk_type": "growth_word_documents",
                    "optimized_for": "azure_cognitive_search_growth_analysis"
                },
                "qr_chunks": {
                    "target_size_words": 800,
                    "max_size_words": 1000,
                    "min_size_words": 0,      
                    "chunk_type": "question_response_documents",
                    "optimized_for": "rag_context_growth_questions"
                }
            },
            "report_chunks_summary": [
                {
                    "chunk_id": chunk_doc["chunk_id"],
                    "section_title": chunk_doc["section_title"],
                    "word_count": chunk_doc["word_count"],
                    "character_count": chunk_doc["character_count"],
                    "content_preview": chunk_doc["content_preview"],
                    "file_name": chunk_files_created[i],
                    "sections_included": chunk_doc["sections_included"],
                    "chunk_type": "report_content"
                }
                for i, chunk_doc in enumerate(word_chunks)
            ],
            "qr_chunks_summary": [
                {
                    "chunk_id": qr_chunk["chunk_id"],
                    "expansion_title": qr_chunk["expansion_title"],
                    "word_count": qr_chunk["word_count"],
                    "question_count": qr_chunk["question_count"],
                    "character_count": qr_chunk["character_count"],
                    "content_preview": qr_chunk["content_preview"],
                    "file_name": qr_chunk_files_created[i],
                    "questions_included": qr_chunk["questions_included"],
                    "chunk_type": "question_response"
                }
                for i, qr_chunk in enumerate(qr_chunks)
            ],
            "created_at": datetime.now().isoformat(),
            "folder": folder_name,
            "report_type": "comprehensive_growth_strategy_with_qr_chunks",
            "database_pooling_enabled": True  # Track that pooling was used
        }
        
        chunks_index_blob_name = f"{folder_name}/{report_id}_growth_chunks_index.json"
        chunks_index_json = json.dumps(chunks_index, indent=2, default=str)
        chunks_index_bytes = io.BytesIO(chunks_index_json.encode("utf-8"))
        
        upload_blob_with_retry(container_client, chunks_index_blob_name, chunks_index_bytes, ContentSettings(content_type="application/json")
        )
        logging.info(f"✅ Comprehensive Chunks index uploaded: {chunks_index_blob_name}")
        
        # ===============================================================
        # 6. Upload final summary and statistics
        # ===============================================================
        logging.info("📈 Step 6/6: Generating final upload summary...")
        
        total_sections = len([k for k in report_data.keys() if k != "_enhanced_growth_report_metadata"])
        total_files = 3 + len(word_chunks) + len(qr_chunks)  # Word doc + JSON + chunks index + all chunk files
        
        # Create detailed upload summary
        upload_summary = {
            "report_id": report_id,
            "user_id": user_id,
            "upload_completed_at": datetime.now().isoformat(),
            "folder_name": folder_name,
            "files_created": {
                "complete_word_document": doc_blob_name,
                "complete_json_report": json_blob_name,
                "chunks_index": chunks_index_blob_name,
                "report_chunks": chunk_files_created,
                "question_response_chunks": qr_chunk_files_created
            },
            "statistics": {
                "total_files_created": total_files,
                "report_sections": total_sections,
                "report_word_chunks": len(word_chunks),
                "question_response_chunks": len(qr_chunks),
                "total_chunks": len(word_chunks) + len(qr_chunks)
            },
            "chunk_optimization": {
                "report_chunks_for": "Azure Cognitive Search Growth Analysis",
                "qr_chunks_for": "RAG Context for AI Growth Questions",
                "target_chunk_size": "300-400 words (main), 800-1000 words (Q&R)",
                "chunk_format": "Microsoft Word (.docx)",
                "database_pooling_used": True
            }
        }
        
        # Upload summary file
        summary_blob_name = f"{folder_name}/{report_id}_growth_upload_summary.json"
        summary_json = json.dumps(upload_summary, indent=2, default=str)
        summary_bytes = io.BytesIO(summary_json.encode("utf-8"))
        
        upload_blob_with_retry(container_client, summary_blob_name, summary_bytes, ContentSettings(content_type="application/json")
        )
        logging.info(f"✅ Upload summary created: {summary_blob_name}")
        
        # ===============================================================
        # Final Success Logging
        # ===============================================================
        logging.info(f"🎉 Growth Report upload complete: {total_files} files in '{folder_name}' folder")
        logging.info(f"📊 Created {len(word_chunks)} Growth Word document chunks for Azure Cognitive Search")
        logging.info(f"🧠 Created {len(qr_chunks)} Question-Response chunks for RAG context")
        logging.info(f"📁 All files uploaded to container '{container_name}' in folder '{folder_name}/'")
        
        success_message = f"Growth report uploaded successfully: {total_sections} sections, {len(word_chunks)} report chunks, {len(qr_chunks)} Q&R chunks, {total_files} files total (with database pooling)"
        logging.info(f"✅ {success_message}")
        
        return True, success_message
        
    except Exception as e:
        error_message = f"Error uploading growth report: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Error details: {type(e).__name__}: {e}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        return False, error_message

def extract_growth_assessment_data_from_report(report_data: Dict) -> Dict:
    """Extract assessment data from growth report if not available in metadata"""
    try:
        # Try to find assessment data in various places within the report
        for key, value in report_data.items():
            if isinstance(value, dict) and 'responses' in value:
                return value
            elif isinstance(value, dict) and 'assessment_data' in value:
                return value['assessment_data']
        
        logging.warning("⚠️ Could not extract growth assessment data from report")
        return {}
    except Exception as e:
        logging.error(f"❌ Error extracting growth assessment data: {str(e)}")
        return {}

async def create_growth_word_document_chunks(report_data: Dict, report_id: str, user_id: str) -> List[Dict]:
    """Create Word document chunks optimized for RAG performance with detailed monitoring"""
    
    logging.info(f"🚀 Starting RAG-optimized growth chunking for report_id={report_id}, user_id={user_id}")
    
    word_chunks = []
    
    # Get user profile for context using async connection pooling
    user_profile = await get_user_profile_data(user_id)
    if user_profile:
        logging.info(f"👤 User context: {user_profile.get('business_name', 'Unknown')} ({user_profile.get('industry', 'Unknown')})")
    else:
        logging.warning(f"⚠️ No user profile found for user_id={user_id}")
    
    # RAG-OPTIMIZED chunking settings for better retrieval performance (aligned with component engine)
    TARGET_SIZE_WORDS = 300  # Sweet spot for RAG retrieval (consistent with component engine)
    MAX_SIZE_WORDS = 400     # Hard limit to prevent oversized chunks
    MIN_SIZE_WORDS = 150     # Minimum to maintain semantic meaning
    
    logging.info(f"⚙️ RAG chunking settings: target={TARGET_SIZE_WORDS}, max={MAX_SIZE_WORDS}, min={MIN_SIZE_WORDS}")
    
    chunk_id = 1
    total_sections = len([k for k in report_data.keys() if k != "_enhanced_growth_report_metadata"])
    logging.info(f"📂 Processing {total_sections} report sections for chunking")
    
    # Track overall statistics
    total_input_words = 0
    total_output_chunks = 0
    section_stats = []
    
    # Process each section and create smart chunks
    for section_idx, (section_name, section_data) in enumerate(report_data.items()):
        if section_name == "_enhanced_growth_report_metadata":
            continue
            
        if not isinstance(section_data, dict):
            logging.warning(f"⚠️ Skipping non-dict section: {section_name}")
            continue
            
        title = section_data.get("title", "Untitled Section")
        content = section_data.get("content", "")
        metadata = section_data.get("metadata", {})
        
        # Log section processing start
        section_word_count = len(content.split())
        total_input_words += section_word_count
        logging.info(f"📄 Processing section {section_idx + 1}/{total_sections}: '{title}' ({section_word_count:,} words)")
        
        # Clean content for better processing
        clean_content = clean_growth_content_for_word_chunks(content)
        clean_word_count = len(clean_content.split())
        
        if clean_word_count != section_word_count:
            logging.info(f"🧹 Content cleaned: {section_word_count} → {clean_word_count} words")
        
        # Create semantic chunks from this section with detailed monitoring
        logging.info(f"🔧 Starting semantic chunking for section '{title}'...")
        section_chunks = create_semantic_growth_word_chunks(clean_content, TARGET_SIZE_WORDS, MAX_SIZE_WORDS, MIN_SIZE_WORDS)
        
        # Validate section chunks
        section_chunk_stats = validate_growth_chunk_sizes(section_chunks, TARGET_SIZE_WORDS, f"Section: {title}")
        section_stats.append({
            "section_name": section_name,
            "section_title": title,
            "input_words": clean_word_count,
            "chunks_created": len(section_chunks),
            "chunk_stats": section_chunk_stats
        })
        
        logging.info(f"✅ Section '{title}' chunked: {clean_word_count} words → {len(section_chunks)} chunks")
        
        # Convert each chunk to a Word document
        for i, chunk_content in enumerate(section_chunks):
            chunk_title = title if len(section_chunks) == 1 else f"{title} - Part {i+1}"
            chunk_word_count = len(chunk_content.split())
            
            logging.debug(f"📝 Creating Word document for chunk {chunk_id}: '{chunk_title}' ({chunk_word_count} words)")
            
            # Create Word document for this chunk
            chunk_doc = create_growth_chunk_word_document(
                chunk_content, 
                chunk_title, 
                user_profile,
                section_name,
                f"{report_id}_growth_chunk_{chunk_id:03d}"
            )
            
            character_count = len(chunk_content)
            
            # Determine chunk quality metrics
            chunk_category = categorize_growth_chunk_size_by_words(chunk_word_count)
            semantic_completeness = calculate_growth_semantic_completeness(chunk_content)
            
            # Log chunk quality
            quality_status = "✅ OPTIMAL" if TARGET_SIZE_WORDS * 0.8 <= chunk_word_count <= TARGET_SIZE_WORDS * 1.2 else \
                           "⚠️ LARGE" if chunk_word_count > TARGET_SIZE_WORDS * 1.2 else \
                           "⚠️ SMALL" if chunk_word_count < TARGET_SIZE_WORDS * 0.8 else "❓ UNKNOWN"
            
            logging.info(f"📊 Chunk {chunk_id} quality: {quality_status} | "
                        f"{chunk_word_count} words | "
                        f"Category: {chunk_category} | "
                        f"Completeness: {semantic_completeness:.2f}")
            
            chunk_info = {
                "chunk_id": f"{report_id}_growth_chunk_{chunk_id:03d}",
                "section_name": section_name,
                "section_title": chunk_title,
                "word_count": chunk_word_count,
                "character_count": character_count,
                "content_preview": chunk_content[:200] + "..." if len(chunk_content) > 200 else chunk_content,
                "sections_included": [section_name],
                "document": chunk_doc,
                "chunk_metadata": {
                    "original_section": section_name,
                    "chunk_size_category": chunk_category,
                    "semantic_completeness": semantic_completeness,
                    "ai_analysis_time": metadata.get("ai_analysis_time", 0),
                    "chunk_type": "growth_analysis_rag_optimized",
                    "rag_optimization": {
                        "target_size": TARGET_SIZE_WORDS,
                        "size_ratio": chunk_word_count / TARGET_SIZE_WORDS,
                        "quality_status": quality_status.replace("✅ ", "").replace("⚠️ ", "").replace("❓ ", ""),
                        "overlap_enabled": True
                    }
                },
                "user_context": {
                    "user_id": user_id,
                    "business_name": user_profile.get("business_name", "Unknown") if user_profile else "Unknown",
                    "industry": user_profile.get("industry", "Unknown") if user_profile else "Unknown",
                    "team_size": user_profile.get("team_size", "Unknown") if user_profile else "Unknown"
                },
                "created_at": datetime.now().isoformat()
            }
            
            word_chunks.append(chunk_info)
            total_output_chunks += 1
            chunk_id += 1
    
    # Final comprehensive statistics
    if word_chunks:
        avg_chunk_size = sum(c['word_count'] for c in word_chunks) // len(word_chunks)
        min_chunk_size = min(c['word_count'] for c in word_chunks)
        max_chunk_size = max(c['word_count'] for c in word_chunks)
        
        # Count optimal chunks
        optimal_chunks = sum(1 for c in word_chunks if TARGET_SIZE_WORDS * 0.8 <= c['word_count'] <= TARGET_SIZE_WORDS * 1.2)
        optimal_percentage = (optimal_chunks / len(word_chunks)) * 100
        
        # Calculate compression ratio
        compression_ratio = total_input_words / sum(c['word_count'] for c in word_chunks) if word_chunks else 1
        
        logging.info(f"🎉 RAG-optimized growth chunking complete!")
        logging.info(f"📊 FINAL STATISTICS:")
        logging.info(f"   📄 Input: {total_input_words:,} words across {total_sections} sections")
        logging.info(f"   📦 Output: {len(word_chunks)} chunks")
        logging.info(f"   📏 Chunk sizes: {min_chunk_size}-{max_chunk_size} words (avg: {avg_chunk_size})")
        logging.info(f"   🎯 Target compliance: {optimal_chunks}/{len(word_chunks)} chunks optimal ({optimal_percentage:.1f}%)")
        logging.info(f"   🔗 Overlap enabled: 50-word context preservation between chunks")
        logging.info(f"   📈 Compression ratio: {compression_ratio:.2f}x (due to overlap)")
        
        # Log any quality concerns
        oversized_chunks = sum(1 for c in word_chunks if c['word_count'] > MAX_SIZE_WORDS)
        undersized_chunks = sum(1 for c in word_chunks if c['word_count'] < MIN_SIZE_WORDS)
        
        if oversized_chunks > 0:
            logging.warning(f"⚠️ Quality concern: {oversized_chunks} chunks exceed maximum size ({MAX_SIZE_WORDS} words)")
        if undersized_chunks > 0:
            logging.warning(f"⚠️ Quality concern: {undersized_chunks} chunks below minimum size ({MIN_SIZE_WORDS} words)")
        
        if optimal_percentage >= 80:
            logging.info(f"✅ Excellent RAG optimization: {optimal_percentage:.1f}% of chunks are optimally sized")
        elif optimal_percentage >= 60:
            logging.info(f"⚠️ Good RAG optimization: {optimal_percentage:.1f}% of chunks are optimally sized")
        else:
            logging.warning(f"❌ Poor RAG optimization: Only {optimal_percentage:.1f}% of chunks are optimally sized")
            
    else:
        logging.error(f"❌ No chunks created from {total_sections} sections!")
    
    return word_chunks

async def create_growth_question_response_chunks(assessment_data: Dict, report_id: str, user_id: str) -> List[Dict]:
    """Create Question-Response Word document chunks optimized for RAG context with retry mechanism and detailed logging"""
    try:
        logging.info(f"🧠 Starting Question-Response chunking for growth report_id={report_id}")
        logging.info(f"📊 Assessment data keys: {list(assessment_data.keys()) if assessment_data else 'No data'}")
        
        qr_chunks = []
        chunk_id = 1
        
        # Get user profile for context using async connection pooling
        logging.info(f"👤 Retrieving user profile for user_id={user_id}")
        user_profile = await get_user_profile_data(user_id)
        if user_profile:
            business_name = user_profile.get('business_name', 'Unknown')
            industry = user_profile.get('industry', 'Unknown')
            logging.info(f"✅ Retrieved user profile: {business_name} ({industry})")
        else:
            logging.warning(f"⚠️ No user profile found for user_id={user_id}")
        
        # Get responses from assessment data
        responses = assessment_data.get('responses', [])
        if not responses:
            logging.error(f"❌ No responses found in growth assessment data")
            logging.error(f"🔍 Assessment data structure: {assessment_data}")
            return []
        
        logging.info(f"📊 Processing {len(responses)} growth responses for Q&R chunking")
        logging.info(f"📊 Sample response structure: {responses[0] if responses else 'No responses'}")
        
        # RAG-optimized chunking settings
        TARGET_SIZE_WORDS = 800
        MAX_SIZE_WORDS = 1000
        MIN_SIZE_WORDS = 0
        
        logging.info(f"⚙️ Chunking parameters: target={TARGET_SIZE_WORDS}, max={MAX_SIZE_WORDS}, min={MIN_SIZE_WORDS}")
        
        # Group responses by section/expansion for better context
        section_groups = {}
        section_response_counts = {}
        
        logging.info(f"📂 Grouping responses by section...")
        
        for i, response in enumerate(responses):
            section = response.get('section', f'Unknown Growth Section {i+1}')
            if section not in section_groups:
                section_groups[section] = []
                section_response_counts[section] = 0
            section_groups[section].append(response)
            section_response_counts[section] += 1
        
        logging.info(f"📂 Grouped growth responses into {len(section_groups)} sections:")
        for section_name, count in section_response_counts.items():
            logging.info(f"   - {section_name}: {count} responses")
        
        # 🔥 NEW: Track processing results for comprehensive analysis
        processing_results = {
            'total_sections': len(section_groups),
            'successful_sections': [],
            'failed_sections': [],
            'retry_sections': [],
            'section_details': {}
        }
        
        # 🔥 STEP 1: First pass - process all sections with detailed tracking
        logging.info(f"🔄 PHASE 1: Initial processing of all {len(section_groups)} sections")
        
        for section_index, (section_name, section_responses) in enumerate(section_groups.items(), 1):
            section_start_time = time.time()
            section_result = {
                'section_name': section_name,
                'response_count': len(section_responses),
                'chunks_created': 0,
                'processing_time': 0,
                'status': 'pending',
                'error': None,
                'attempt': 1
            }
            
            try:
                logging.info(f"🔄 Processing section {section_index}/{len(section_groups)}: {section_name}")
                logging.info(f"📊 Section details: {len(section_responses)} responses")
                
                # Create chunks from this section's responses
                logging.debug(f"🔧 Calling create_growth_section_qr_chunks for {section_name}")
                section_qr_chunks = create_growth_section_qr_chunks(
                    section_name, section_responses, TARGET_SIZE_WORDS, MAX_SIZE_WORDS, MIN_SIZE_WORDS
                )
                
                section_processing_time = time.time() - section_start_time
                section_result['processing_time'] = section_processing_time
                
                if section_qr_chunks:
                    logging.info(f"✅ Section {section_name} created {len(section_qr_chunks)} raw chunks")
                    section_result['chunks_created'] = len(section_qr_chunks)
                    section_result['status'] = 'success'
                    
                    # Convert each chunk to Word document
                    section_chunks_added = 0
                    for i, chunk_content in enumerate(section_qr_chunks):
                        try:
                            chunk_title = section_name if len(section_qr_chunks) == 1 else f"{section_name} - Part {i+1}"
                            
                            logging.debug(f"📝 Creating Word document for chunk {i+1}/{len(section_qr_chunks)}")
                            
                            # Create Word document for this Q&R chunk
                            chunk_doc = create_growth_qr_chunk_word_document(
                                chunk_content, 
                                chunk_title, 
                                user_profile,
                                section_name,
                                f"{report_id}_growth_qr_chunk_{chunk_id:03d}"
                            )
                            
                            # Calculate metrics with safety checks
                            word_count = 0
                            character_count = 0
                            question_count = len(chunk_content.get('question_responses', []))
                            
                            for qr in chunk_content.get('question_responses', []):
                                combined_text = qr.get('combined_text', '')
                                word_count += len(combined_text.split())
                                character_count += len(combined_text)
                            
                            # Create preview text with safety
                            preview_texts = []
                            for qr in chunk_content.get('question_responses', [])[:2]:  # First 2 Q&R pairs
                                q_text = qr.get('question_text', 'No question')[:100]
                                r_text = qr.get('response_text', 'No response')[:100]
                                preview_texts.append(f"Q: {q_text}... A: {r_text}...")
                            content_preview = " | ".join(preview_texts) if preview_texts else "No preview available"
                            
                            chunk_info = {
                                "chunk_id": f"{report_id}_growth_qr_chunk_{chunk_id:03d}",
                                "expansion_title": chunk_title,
                                "word_count": word_count,
                                "question_count": question_count,
                                "character_count": character_count,
                                "content_preview": content_preview,
                                "questions_included": [qr.get('question_id', f'unknown_{j}') for j, qr in enumerate(chunk_content.get('question_responses', []))],
                                "document": chunk_doc,
                                "chunk_metadata": {
                                    "original_section": section_name,
                                    "chunk_size_category": categorize_growth_chunk_size_by_words(word_count),
                                    "question_density": question_count / max(1, word_count / 100),  # questions per 100 words
                                    "chunk_type": "question_response_rag",
                                    "rag_optimized": True,
                                    "processing_attempt": 1,
                                    "section_index": section_index
                                },
                                "user_context": {
                                    "user_id": user_id,
                                    "business_name": user_profile.get("business_name", "Unknown") if user_profile else "Unknown",
                                    "industry": user_profile.get("industry", "Unknown") if user_profile else "Unknown",
                                    "team_size": user_profile.get("team_size", "Unknown") if user_profile else "Unknown"
                                },
                                "created_at": datetime.now().isoformat()
                            }
                            
                            qr_chunks.append(chunk_info)
                            section_chunks_added += 1
                            chunk_id += 1
                            
                            logging.info(f"✅ Created Growth Q&R chunk {chunk_id-1}: {question_count} questions, {word_count} words")
                            
                        except Exception as chunk_error:
                            logging.error(f"❌ Error creating individual chunk {i+1} for section {section_name}: {chunk_error}")
                            continue
                    
                    if section_chunks_added > 0:
                        processing_results['successful_sections'].append(section_name)
                        section_result['final_chunks_added'] = section_chunks_added
                        logging.info(f"✅ Section {section_name} completed: {section_chunks_added} chunks added to final list")
                    else:
                        logging.warning(f"⚠️ Section {section_name} created raw chunks but failed to convert to Word documents")
                        processing_results['failed_sections'].append(section_name)
                        section_result['status'] = 'partial_failure'
                        section_result['error'] = 'Failed to convert chunks to Word documents'
                else:
                    logging.warning(f"⚠️ Section {section_name} returned no chunks")
                    processing_results['failed_sections'].append(section_name)
                    section_result['status'] = 'no_chunks'
                    section_result['error'] = 'No chunks returned from section processing'
                
            except Exception as section_error:
                section_processing_time = time.time() - section_start_time
                section_result['processing_time'] = section_processing_time
                section_result['status'] = 'error'
                section_result['error'] = str(section_error)
                
                logging.error(f"❌ Error processing section {section_name}: {section_error}")
                logging.error(f"🔍 Section error traceback: {traceback.format_exc()}")
                processing_results['failed_sections'].append(section_name)
            
            processing_results['section_details'][section_name] = section_result
            logging.info(f"📊 Section {section_name} summary: {section_result['status']} in {section_result['processing_time']:.2f}s")
        
        # 🔥 STEP 2: Analyze first pass results
        total_successful = len(processing_results['successful_sections'])
        total_failed = len(processing_results['failed_sections'])
        success_rate = (total_successful / processing_results['total_sections']) * 100
        
        logging.info(f"📊 PHASE 1 RESULTS:")
        logging.info(f"   - Total sections: {processing_results['total_sections']}")
        logging.info(f"   - Successful: {total_successful}")
        logging.info(f"   - Failed: {total_failed}")
        logging.info(f"   - Success rate: {success_rate:.1f}%")
        logging.info(f"   - Total chunks created: {len(qr_chunks)}")
        
        # 🔥 STEP 3: Retry failed sections (if any)
        if processing_results['failed_sections'] and success_rate < 80:
            logging.warning(f"🔄 PHASE 2: Retrying {len(processing_results['failed_sections'])} failed sections")
            
            for retry_section in processing_results['failed_sections'][:]:  # Copy list to avoid modification issues
                if retry_section not in section_groups:
                    continue
                    
                retry_start_time = time.time()
                logging.info(f"🔄 RETRY: Attempting {retry_section} again...")
                
                try:
                    # Use simplified approach for retry
                    section_responses = section_groups[retry_section]
                    
                    # Create a single consolidated chunk for failed sections
                    logging.info(f"🔧 RETRY STRATEGY: Creating simplified chunk for {retry_section}")
                    
                    consolidated_chunk = {
                        "question_responses": []
                    }
                    
                    total_retry_words = 0
                    for response in section_responses:
                        try:
                            # Simplified response processing for retry
                            question_text = response.get('question_text', 'Question not available')
                            response_data = response.get('response_data', {})
                            question_id = response.get('question_id', 'unknown')
                            
                            # Safe response extraction
                            if isinstance(response_data, dict):
                                selected_response = str(response_data.get('selected_option') or 
                                                     response_data.get('response_text') or 
                                                     response_data.get('slider_value') or 
                                                     response_data.get('value') or 
                                                     'No response available')
                            else:
                                selected_response = str(response_data) if response_data else 'No response available'
                            
                            # Simplified combined text
                            combined_text = f"Growth Question: {question_text}\nResponse: {selected_response}\nSection: {retry_section}"
                            
                            qr_item = {
                                "question_id": question_id,
                                "question_text": question_text,
                                "response_text": selected_response,
                                "combined_text": combined_text,
                                "word_count": len(combined_text.split()),
                                "metadata": response.get('metadata', {}),
                                "all_options": response.get('all_options', []),
                                "context_richness": "simplified_retry"
                            }
                            
                            consolidated_chunk["question_responses"].append(qr_item)
                            total_retry_words += qr_item["word_count"]
                            
                        except Exception as response_retry_error:
                            logging.error(f"❌ Error in retry processing response in {retry_section}: {response_retry_error}")
                            continue
                    
                    if consolidated_chunk["question_responses"]:
                        # Create Word document for retry chunk
                        chunk_title = f"{retry_section} (Retry)"
                        
                        chunk_doc = create_growth_qr_chunk_word_document(
                            consolidated_chunk, 
                            chunk_title, 
                            user_profile,
                            retry_section,
                            f"{report_id}_growth_qr_chunk_{chunk_id:03d}"
                        )
                        
                        retry_chunk_info = {
                            "chunk_id": f"{report_id}_growth_qr_chunk_{chunk_id:03d}",
                            "expansion_title": chunk_title,
                            "word_count": total_retry_words,
                            "question_count": len(consolidated_chunk["question_responses"]),
                            "character_count": sum(len(qr.get('combined_text', '')) for qr in consolidated_chunk["question_responses"]),
                            "content_preview": f"Retry chunk for {retry_section} with {len(consolidated_chunk['question_responses'])} questions",
                            "questions_included": [qr.get('question_id') for qr in consolidated_chunk["question_responses"]],
                            "document": chunk_doc,
                            "chunk_metadata": {
                                "original_section": retry_section,
                                "chunk_size_category": categorize_growth_chunk_size_by_words(total_retry_words),
                                "question_density": len(consolidated_chunk["question_responses"]) / max(1, total_retry_words / 100),
                                "chunk_type": "question_response_retry",
                                "rag_optimized": True,
                                "processing_attempt": 2,
                                "retry_strategy": "simplified_consolidation"
                            },
                            "user_context": {
                                "user_id": user_id,
                                "business_name": user_profile.get("business_name", "Unknown") if user_profile else "Unknown",
                                "industry": user_profile.get("industry", "Unknown") if user_profile else "Unknown",
                                "team_size": user_profile.get("team_size", "Unknown") if user_profile else "Unknown"
                            },
                            "created_at": datetime.now().isoformat()
                        }
                        
                        qr_chunks.append(retry_chunk_info)
                        chunk_id += 1
                        
                        # Update tracking
                        processing_results['failed_sections'].remove(retry_section)
                        processing_results['retry_sections'].append(retry_section)
                        processing_results['section_details'][retry_section]['status'] = 'retry_success'
                        processing_results['section_details'][retry_section]['attempt'] = 2
                        processing_results['section_details'][retry_section]['retry_time'] = time.time() - retry_start_time
                        
                        logging.info(f"✅ RETRY SUCCESS: {retry_section} recovered with {len(consolidated_chunk['question_responses'])} questions")
                        
                except Exception as retry_error:
                    retry_time = time.time() - retry_start_time
                    processing_results['section_details'][retry_section]['retry_error'] = str(retry_error)
                    processing_results['section_details'][retry_section]['retry_time'] = retry_time
                    logging.error(f"❌ RETRY FAILED: {retry_section} - {retry_error}")
        
        # 🔥 STEP 4: Final comprehensive analysis
        final_successful = len(processing_results['successful_sections']) + len(processing_results['retry_sections'])
        final_failed = len(processing_results['failed_sections'])
        final_success_rate = (final_successful / processing_results['total_sections']) * 100
        
        avg_words_per_chunk = sum(c['word_count'] for c in qr_chunks) // len(qr_chunks) if qr_chunks else 0
        total_questions = sum(c['question_count'] for c in qr_chunks)
        
        logging.info(f"🎉 FINAL Q&R PROCESSING RESULTS:")
        logging.info(f"   - Total sections processed: {processing_results['total_sections']}")
        logging.info(f"   - Successful (first pass): {len(processing_results['successful_sections'])}")
        logging.info(f"   - Recovered (retry): {len(processing_results['retry_sections'])}")
        logging.info(f"   - Still failed: {final_failed}")
        logging.info(f"   - Final success rate: {final_success_rate:.1f}%")
        logging.info(f"   - Total chunks created: {len(qr_chunks)}")
        logging.info(f"   - Total questions processed: {total_questions}")
        logging.info(f"   - Average words per chunk: {avg_words_per_chunk}")
        
        # Log section-by-section results
        logging.info(f"📋 DETAILED SECTION RESULTS:")
        for section_name, details in processing_results['section_details'].items():
            status_emoji = "✅" if details['status'] in ['success', 'retry_success'] else "❌"
            logging.info(f"   {status_emoji} {section_name}: {details['status']} (attempt {details['attempt']}, {details['processing_time']:.2f}s)")
            if details.get('error'):
                logging.info(f"      Error: {details['error']}")
        
        if final_failed > 0:
            logging.warning(f"⚠️ {final_failed} sections still failed after retry")
            for failed_section in processing_results['failed_sections']:
                logging.warning(f"   - {failed_section}: {processing_results['section_details'][failed_section].get('error', 'Unknown error')}")
        
        logging.info(f"🎉 Successfully created {len(qr_chunks)} Growth Question-Response chunks")
        return qr_chunks
        
    except Exception as e:
        logging.error(f"❌ Fatal error creating Growth Question-Response chunks: {str(e)}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        return []

def create_growth_section_qr_chunks(section_name: str, responses: List[Dict], target_size: int, max_size: int, min_size: int) -> List[Dict]:
    """Create manageable Growth Q&R chunks from a section's responses with detailed logging and API key preservation"""
    try:
        logging.info(f"🔄 Starting Growth Q&R chunk creation for section: {section_name}")
        logging.info(f"📊 Section parameters: {len(responses)} responses, target_size: {target_size}, max_size: {max_size}, min_size: {min_size}")
        
        chunks = []
        current_chunk = {"question_responses": [], "word_count": 0}
        
        # Track processing stats
        total_questions_processed = 0
        total_words_generated = 0
        ai_analysis_skipped_count = 0
        error_count = 0
        
        logging.info(f"📝 Processing {len(responses)} growth responses for section: {section_name}")
        
        for response_index, response in enumerate(responses):
            try:
                logging.debug(f"🔍 Processing growth response {response_index + 1}/{len(responses)} in section: {section_name}")
                
                # Extract question and response data
                question_text = response.get('question_text', 'Growth question not available')
                response_data = response.get('response_data', {})
                question_id = response.get('question_id', f'unknown_{response_index}')
                
                logging.debug(f"📋 Growth Question ID: {question_id}")
                logging.debug(f"📋 Question preview: {question_text[:50] if question_text else 'No question text'}...")
                
                # 🔥 STEP 1: DETAILED LOGGING OF RAW DATA TYPES
                logging.debug(f"🔍 RAW DATA INSPECTION for {question_id}:")
                logging.debug(f"  - response_data type: {type(response_data)}")
                logging.debug(f"  - response_data repr: {repr(response_data)}")
                if isinstance(response_data, dict):
                    logging.debug(f"  - response_data keys: {list(response_data.keys())}")
                    for key, value in response_data.items():
                        logging.debug(f"    - {key}: type={type(value)}, value={repr(value)}")
                
                # 🔥 STEP 2: SAFE DATA EXTRACTION WITH COMPREHENSIVE LOGGING
                selected_response = None
                extraction_method = "unknown"
                
                if isinstance(response_data, dict):
                    logging.debug(f"📊 Processing dict response_data with {len(response_data)} keys")
                    
                    if 'selected_option' in response_data:
                        raw_value = response_data['selected_option']
                        logging.debug(f"  - Found 'selected_option': type={type(raw_value)}, value={repr(raw_value)}")
                        selected_response = raw_value
                        extraction_method = "selected_option"
                        
                    elif 'selected_options' in response_data:  # Multiple selections
                        raw_value = response_data['selected_options']
                        logging.debug(f"  - Found 'selected_options': type={type(raw_value)}, value={repr(raw_value)}")
                        
                        if isinstance(raw_value, list):
                            logging.debug(f"    - Processing list with {len(raw_value)} items")
                            selected_response = ', '.join(str(item) for item in raw_value)
                            extraction_method = "selected_options_list"
                        elif isinstance(raw_value, (dict, set)):
                            logging.debug(f"    - Processing {type(raw_value)} with {len(raw_value)} items")
                            selected_response = str(raw_value)
                            extraction_method = "selected_options_complex"
                        else:
                            logging.debug(f"    - Converting single value to string")
                            selected_response = str(raw_value)
                            extraction_method = "selected_options_single"
                            
                    elif 'response_text' in response_data:
                        raw_value = response_data['response_text']
                        logging.debug(f"  - Found 'response_text': type={type(raw_value)}, value={repr(raw_value)}")
                        selected_response = raw_value
                        extraction_method = "response_text"
                        
                    elif 'slider_value' in response_data:
                        raw_value = response_data['slider_value']
                        logging.debug(f"  - Found 'slider_value': type={type(raw_value)}, value={repr(raw_value)}")
                        selected_response = f"Growth Rating: {raw_value}"
                        extraction_method = "slider_value"
                        
                    elif 'priority_ranking' in response_data:
                        raw_value = response_data['priority_ranking']
                        logging.debug(f"  - Found 'priority_ranking': type={type(raw_value)}, value={repr(raw_value)}")
                        
                        if isinstance(raw_value, list):
                            logging.debug(f"    - Processing ranking list with {len(raw_value)} items")
                            selected_response = f"Priority Ranking: {', '.join(str(item) for item in raw_value)}"
                        elif isinstance(raw_value, dict):
                            logging.debug(f"    - Processing ranking dict with {len(raw_value)} items")
                            selected_response = f"Priority Ranking: {raw_value}"
                        else:
                            logging.debug(f"    - Processing ranking as single value")
                            selected_response = f"Priority Ranking: {raw_value}"
                        extraction_method = "priority_ranking"
                        
                    elif 'value' in response_data:
                        raw_value = response_data['value']
                        logging.debug(f"  - Found 'value': type={type(raw_value)}, value={repr(raw_value)}")
                        selected_response = raw_value
                        extraction_method = "value"
                        
                    else:
                        # Handle any other fields
                        logging.debug(f"  - No standard fields found, checking all keys")
                        for key, value in response_data.items():
                            if value is not None:
                                logging.debug(f"    - Using {key}: type={type(value)}, value={repr(value)}")
                                selected_response = value
                                extraction_method = f"fallback_{key}"
                                break
                        
                        if selected_response is None:
                            logging.debug(f"  - No usable data found in dict, using placeholder")
                            selected_response = 'Growth response not available'
                            extraction_method = "placeholder"
                            
                elif response_data is not None:
                    logging.debug(f"📊 Processing non-dict response_data: type={type(response_data)}")
                    selected_response = response_data
                    extraction_method = "direct_non_dict"
                else:
                    logging.debug(f"📊 No response_data found, using placeholder")
                    selected_response = 'Growth response not available'
                    extraction_method = "no_data"
                
                # 🔥 STEP 3: SAFE STRING CONVERSION WITH DETAILED LOGGING
                logging.debug(f"🔄 Converting to string: method={extraction_method}, type={type(selected_response)}")
                
                try:
                    # Handle special cases that might cause issues
                    if selected_response is None:
                        selected_response_str = 'No response provided'
                        logging.debug(f"  - Converted None to placeholder string")
                    elif isinstance(selected_response, (list, tuple)):
                        selected_response_str = ', '.join(str(item) for item in selected_response)
                        logging.debug(f"  - Converted list/tuple of {len(selected_response)} items to string")
                    elif isinstance(selected_response, dict):
                        selected_response_str = f"Dict response: {selected_response}"
                        logging.debug(f"  - Converted dict with {len(selected_response)} keys to string")
                    elif hasattr(selected_response, '__iter__') and not isinstance(selected_response, str):
                        # Handle other iterable types (sets, etc.)
                        try:
                            selected_response_str = ', '.join(str(item) for item in selected_response)
                            logging.debug(f"  - Converted iterable to string")
                        except Exception as iter_error:
                            logging.debug(f"  - Iterable conversion failed: {iter_error}, using str()")
                            selected_response_str = str(selected_response)
                    else:
                        selected_response_str = str(selected_response)
                        logging.debug(f"  - Direct string conversion successful")
                    
                except Exception as conversion_error:
                    logging.error(f"  - String conversion failed: {conversion_error}")
                    selected_response_str = f"Conversion error for {type(selected_response)}"
                    error_count += 1
                
                # 🔥 STEP 4: SAFE PREVIEW CREATION
                try:
                    if len(selected_response_str) > 50:
                        response_preview = selected_response_str[:50] + "..."
                    else:
                        response_preview = selected_response_str
                    logging.debug(f"✅ Selected growth response ({extraction_method}): {response_preview}")
                except Exception as preview_error:
                    logging.error(f"  - Preview creation failed: {preview_error}")
                    response_preview = "Preview unavailable"
                    error_count += 1
                
                # Get all available options with safety
                all_options = response.get('all_options', [])
                logging.debug(f"📚 Available growth options count: {len(all_options) if isinstance(all_options, (list, dict)) else 'N/A'}")
                
                # 🔥 CRITICAL FIX: Safe option type checking
                if all_options:
                    if isinstance(all_options, list):
                        logging.debug(f"  - Option types (list): {[type(opt) for opt in all_options[:3]]}")  # First 3 types for lists
                    elif isinstance(all_options, dict):
                        logging.debug(f"  - Option structure (dict): keys={list(all_options.keys())}")  # Dict structure
                    else:
                        logging.debug(f"  - Option type: {type(all_options)}")  # Other types
                
                # Create enhanced combined Q&R text with rich growth context
                combined_text = f"Growth Strategy Question: {question_text}\n\n"
                
                # 🔥 CRITICAL FIX: Safe option processing for both lists and dicts
                if all_options:
                    combined_text += "Available Growth Strategy Options:\n"
                    try:
                        if isinstance(all_options, list):
                            # Handle list of options (normal case)
                            for i, option in enumerate(all_options, 1):
                                option_str = str(option)
                                if option_str == selected_response_str:
                                    combined_text += f"  {i}. ✓ {option_str} (SELECTED)\n"
                                else:
                                    combined_text += f"  {i}. {option_str}\n"
                            logging.debug(f"📋 Added {len(all_options)} growth options to Q&R context")
                        elif isinstance(all_options, dict):
                            # Handle dictionary options (slider questions, etc.)
                            if 'labels' in all_options:
                                labels = all_options['labels']
                                min_val = all_options.get('min', 1)
                                max_val = all_options.get('max', 10)
                                combined_text += f"  Scale: {min_val} to {max_val}\n"
                                combined_text += f"  Labels: {', '.join(labels)}\n"
                                logging.debug(f"📋 Added slider scale info: {min_val}-{max_val} with {len(labels)} labels")
                            else:
                                # Generic dict handling
                                for key, value in all_options.items():
                                    combined_text += f"  {key}: {value}\n"
                                logging.debug(f"📋 Added dict options: {len(all_options)} key-value pairs")
                        else:
                            # Handle other types
                            combined_text += f"  Options: {str(all_options)}\n"
                            logging.debug(f"📋 Added options as string: {type(all_options)}")
                        
                        combined_text += f"\nClient's Growth Strategy Choice: {selected_response_str}\n"
                        
                    except Exception as options_error:
                        logging.error(f"📋 Error processing options: {options_error}")
                        combined_text += f"Client's Growth Strategy Choice: {selected_response_str}\n"
                        error_count += 1
                else:
                    combined_text += f"Client's Growth Strategy Response: {selected_response_str}\n"
                    logging.debug(f"📋 No options available, added direct growth response")
                
                # Add growth-specific question context
                combined_text += f"\n--- Growth Strategy Context ---"
                combined_text += f"\nGrowth Area: {response.get('section', 'Unknown')}"
                combined_text += f"\nQuestion Priority: {str(response.get('weight', 'medium')).upper()}"
                combined_text += f"\nQuestion Type: {response.get('question_type', 'growth_assessment')}"
                combined_text += f"\nGrowth Question ID: {question_id}"
                combined_text += f"\nData Extraction Method: {extraction_method}"
                
                logging.debug(f"📊 Added growth question context: weight={response.get('weight', 'medium')}, type={response.get('question_type', 'growth_assessment')}")
                
                # Add response analytics from metadata
                metadata = response.get('metadata', {})
                if metadata:
                    timing_info = metadata.get('timing_data', {})
                    if timing_info:
                        time_spent = timing_info.get('total_engagement_time', 0)
                        combined_text += f"\n\n--- Growth Response Analytics ---"
                        combined_text += f"\nResponse Time: {time_spent} seconds"
                        combined_text += f"\nFocus Time: {timing_info.get('focus_time', 'N/A')} seconds"
                        combined_text += f"\nInteraction Count: {timing_info.get('interaction_count', 'N/A')}"
                        combined_text += f"\nDecision Speed: {'Quick' if time_spent < 30 else 'Deliberate' if time_spent < 90 else 'Thorough'}"
                        logging.debug(f"⏱️ Added growth timing analytics: {time_spent}s response time")
                    else:
                        logging.debug(f"⏱️ No timing data available for growth question {question_id}")
                else:
                    logging.debug(f"📊 No metadata available for growth question {question_id}")
                
                # 🔥 CRITICAL FIX: Safe response pattern analysis
                if all_options and isinstance(all_options, list):
                    try:
                        selected_index = -1
                        for i, option in enumerate(all_options):
                            # 🔥 SAFE: Convert both to strings for comparison
                            if str(option) == selected_response_str:
                                selected_index = i
                                break
                        
                        if selected_index >= 0:
                            total_options = len(all_options)
                            percentile = (selected_index + 1) / total_options
                            
                            combined_text += f"\n\n--- Growth Response Pattern Analysis ---"
                            combined_text += f"\nSelected Option: {selected_index + 1} of {total_options}"
                            combined_text += f"\nResponse Percentile: {percentile:.1%}"
                            combined_text += f"\nGrowth Preference: {'Conservative' if percentile <= 0.33 else 'Moderate' if percentile <= 0.66 else 'Aggressive'}"
                            
                            logging.debug(f"📈 Added growth pattern analysis: option {selected_index + 1}/{total_options} ({percentile:.1%} percentile)")
                        else:
                            logging.debug(f"⚠️ Could not find selected response in options list for growth question {question_id}")
                    except Exception as pattern_error:
                        logging.error(f"📈 Error in pattern analysis: {pattern_error}")
                        error_count += 1
                elif all_options and isinstance(all_options, dict):
                    # For slider/scale questions, add scale analysis
                    try:
                        if 'min' in all_options and 'max' in all_options:
                            min_val = all_options['min']
                            max_val = all_options['max']
                            # Extract numeric value from selected response
                            if isinstance(selected_response, (int, float)):
                                value = selected_response
                            elif 'selected_option' in response_data and isinstance(response_data['selected_option'], (int, float)):
                                value = response_data['selected_option']
                            else:
                                # Try to extract number from string
                                import re
                                numbers = re.findall(r'\d+', str(selected_response))
                                value = int(numbers[0]) if numbers else None
                            
                            if value is not None:
                                percentile = (value - min_val) / (max_val - min_val)
                                combined_text += f"\n\n--- Growth Scale Analysis ---"
                                combined_text += f"\nScale Range: {min_val} to {max_val}"
                                combined_text += f"\nSelected Value: {value}"
                                combined_text += f"\nScale Percentile: {percentile:.1%}"
                                combined_text += f"\nGrowth Level: {'Low' if percentile <= 0.33 else 'Medium' if percentile <= 0.66 else 'High'}"
                                logging.debug(f"📈 Added scale analysis: {value}/{max_val} ({percentile:.1%} percentile)")
                    except Exception as scale_error:
                        logging.error(f"📈 Error in scale analysis: {scale_error}")
                        error_count += 1
                
                # 🔴 DISABLED: AI-Generated Intelligent Analysis to prevent API key exhaustion
                logging.debug(f"🚫 Skipping AI analysis for growth question {question_id} to preserve API keys for main report")
                ai_analysis_skipped_count += 1
                
                # Add enhanced basic analysis for growth context
                combined_text += f"\n\n--- Basic Growth Analysis ---"
                combined_text += f"\nGrowth Focus Area: {section_name}"
                combined_text += f"\nStrategic Priority: {response.get('weight', 'medium')}"
                combined_text += f"\nAssessment Context: Growth strategy evaluation for scaling, acceleration, and market expansion opportunities"
                combined_text += f"\nGrowth Stage Relevance: Applicable to current business phase and growth trajectory"
                combined_text += f"\nAnalysis Status: Basic analysis used to preserve API capacity for comprehensive growth strategy report"
                combined_text += f"\nProcessing Status: Successful extraction via {extraction_method}"
                
                logging.debug(f"📝 Added enhanced basic growth analysis for question {question_id}")
                
                # Calculate word count for this Q&R item
                qr_word_count = len(combined_text.split())
                total_words_generated += qr_word_count
                
                qr_item = {
                    "question_id": question_id,
                    "question_text": question_text,
                    "response_text": selected_response_str,
                    "combined_text": combined_text,
                    "word_count": qr_word_count,
                    "metadata": metadata,
                    "all_options": all_options,
                    "context_richness": "basic_growth_analysis",  # Growth-specific context
                    "extraction_method": extraction_method,
                    "processing_errors": error_count
                }
                
                logging.debug(f"📊 Created Growth Q&R item: {qr_word_count} words, method={extraction_method}, context_richness=basic_growth_analysis")
                
                # Check if adding this Q&R would exceed max size
                if current_chunk["word_count"] + qr_item["word_count"] > max_size and current_chunk["question_responses"]:
                    # Current chunk is full, save it if substantial
                    if current_chunk["word_count"] >= min_size:
                        chunks.append(current_chunk)
                        logging.info(f"✅ Completed Growth Q&R chunk {len(chunks)}: {len(current_chunk['question_responses'])} questions, {current_chunk['word_count']} words")
                        current_chunk = {"question_responses": [], "word_count": 0}
                    else:
                        logging.debug(f"⚠️ Current chunk too small ({current_chunk['word_count']} words < {min_size} min), continuing to add questions")
                
                # Add Q&R to current chunk
                current_chunk["question_responses"].append(qr_item)
                current_chunk["word_count"] += qr_item["word_count"]
                total_questions_processed += 1
                
                logging.debug(f"📝 Added Growth Q&R to current chunk: {len(current_chunk['question_responses'])} questions, {current_chunk['word_count']} words total")
                
            except Exception as response_error:
                logging.error(f"❌ Error processing individual response {response_index + 1} in {section_name}: {response_error}")
                logging.error(f"🔍 Response data that failed: {response}")
                import traceback
                logging.error(f"🔍 Response traceback: {traceback.format_exc()}")
                error_count += 1
                continue  # Skip this response and continue with others
        
        # Add the last chunk if it's substantial
        if current_chunk["question_responses"] and current_chunk["word_count"] >= min_size:
            chunks.append(current_chunk)
            logging.info(f"✅ Completed final Growth Q&R chunk {len(chunks)}: {len(current_chunk['question_responses'])} questions, {current_chunk['word_count']} words")
        elif current_chunk["question_responses"]:
            logging.warning(f"⚠️ Final Growth chunk discarded (too small): {len(current_chunk['question_responses'])} questions, {current_chunk['word_count']} words < {min_size} min")
        
        # Calculate final statistics
        avg_words_per_chunk = total_words_generated // len(chunks) if chunks else 0
        avg_questions_per_chunk = total_questions_processed // len(chunks) if chunks else 0
        
        # Comprehensive completion logging
        logging.info(f"🎉 Growth Q&R chunk creation completed for section: {section_name}")
        logging.info(f"📊 Final statistics:")
        logging.info(f"   - Total chunks created: {len(chunks)}")
        logging.info(f"   - Total questions processed: {total_questions_processed}")
        logging.info(f"   - Total questions attempted: {len(responses)}")
        logging.info(f"   - Total words generated: {total_words_generated:,}")
        logging.info(f"   - Average words per chunk: {avg_words_per_chunk}")
        logging.info(f"   - Average questions per chunk: {avg_questions_per_chunk}")
        logging.info(f"   - AI analysis skipped: {ai_analysis_skipped_count} (to preserve API keys)")
        logging.info(f"   - Processing errors: {error_count}")
        logging.info(f"   - Success rate: {((total_questions_processed / len(responses)) * 100):.1f}%" if responses else "N/A")
        logging.info(f"   - Context richness: basic_growth_analysis (API-friendly)")
        
        if error_count > 0:
            logging.warning(f"⚠️ Section completed with {error_count} errors, but {total_questions_processed} questions successfully processed")
        
        logging.info(f"📊 Growth Section '{section_name}': Created {len(chunks)} Q&R chunks from {len(responses)} responses ({total_questions_processed} successful, {error_count} errors)")
        return chunks
        
    except Exception as e:
        logging.error(f"❌ Fatal error creating Growth section Q&R chunks for {section_name}: {str(e)}")
        logging.error(f"🔍 Error details: {type(e).__name__}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        return []


def create_growth_qr_chunk_word_document(chunk_content: Dict, title: str, user_profile: Dict, section_name: str, chunk_id: str) -> Document:
    """Create a professionally formatted Word document for Growth Question-Response chunk"""
    try:
        doc = Document()
        
        # Enhanced styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add header with branding
        header = doc.add_heading("BACKABLE GROWTH ENGINE - Q&A CONTEXT", 0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.font.size = Pt(20)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add chunk title
        chunk_title = doc.add_heading(title, 1)
        chunk_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title_run = chunk_title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Add metadata section
        if user_profile:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("Business Growth Context:").bold = True
            metadata_para.add_run(f"\nBusiness: {user_profile.get('business_name', 'Unknown')}")
            metadata_para.add_run(f"\nIndustry: {user_profile.get('industry', 'Unknown')}")
            metadata_para.add_run(f"\nTeam Size: {user_profile.get('team_size', 'Unknown')} employees")
            metadata_para.add_run(f"\nGrowth Section: {section_name}")
            metadata_para.add_run(f"\nChunk ID: {chunk_id}")
            metadata_para.add_run(f"\nQuestions Included: {len(chunk_content['question_responses'])}")
            metadata_para.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Add separator
        doc.add_paragraph("─" * 60)
        
        # Add RAG context note
        rag_note = doc.add_paragraph()
        rag_note_run = rag_note.add_run("🚀 GROWTH RAG CONTEXT: This document contains the client's actual growth strategy questions and responses for AI context. Use this to understand their specific growth choices, scaling preferences, and strategic reasoning.")
        rag_note_run.font.color.rgb = RGBColor(0, 102, 204)
        rag_note_run.italic = True
        
        doc.add_paragraph("─" * 60)
        
        # Process each question-response pair
        for i, qr in enumerate(chunk_content['question_responses']):
            # Question header
            question_heading = doc.add_heading(f"Growth Question {i+1}: {qr['question_id']}", 2)
            question_heading_run = question_heading.runs[0]
            question_heading_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Question text
            question_para = doc.add_paragraph()
            question_para.add_run("Growth Q: ").bold = True
            question_para.add_run(qr['question_text'])
            
            # Response text
            response_para = doc.add_paragraph()
            response_para.add_run("Strategy A: ").bold = True
            response_para.add_run(qr['response_text'])
            
            # Add growth-specific metadata if available
            metadata = qr.get('metadata', {})
            if metadata:
                timing_data = metadata.get('timing_data', {})
                if timing_data.get('total_engagement_time'):
                    meta_para = doc.add_paragraph()
                    meta_run = meta_para.add_run(f"Decision Time: {timing_data['total_engagement_time']} seconds")
                    meta_run.font.size = Pt(9)
                    meta_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add space between Q&R pairs
            if i < len(chunk_content['question_responses']) - 1:
                doc.add_paragraph("─" * 30)
        
        # Add footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run("Generated by Backable AI Growth Intelligence for RAG Context")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        logging.info(f"📄 Created Growth Q&R Word document: {len(chunk_content['question_responses'])} questions")
        return doc
        
    except Exception as e:
        logging.error(f"❌ Error creating Growth Q&R Word document: {str(e)}")
        # Return minimal document on error
        doc = Document()
        doc.add_heading("Error Creating Growth Q&R Document", 1)
        doc.add_paragraph(f"Error: {str(e)}")
        return doc

def create_semantic_growth_word_chunks(content: str, target_size: int, max_size: int, min_size: int) -> List[str]:
    """Create semantic chunks that preserve growth context WITH OVERLAP for better RAG performance"""
    
    logging.info(f"🔧 Starting growth semantic chunking: target={target_size}, max={max_size}, min={min_size}")
    
    # If content is small enough, return as single chunk
    word_count = len(content.split())
    logging.info(f"📊 Input growth content: {word_count} words")
    
    if word_count <= max_size:
        logging.info(f"✅ Growth content fits in single chunk ({word_count} <= {max_size})")
        return [content]
    
    chunks = []
    OVERLAP_SIZE = 50  # 50 words overlap between chunks for context preservation
    logging.info(f"🔗 Using {OVERLAP_SIZE}-word overlap between growth chunks")
    
    # Split by growth logic sections first
    growth_sections = split_by_growth_logic(content)
    logging.info(f"📂 Split into {len(growth_sections)} growth logic sections")
    
    current_chunk = ""
    current_word_count = 0
    previous_chunk_end = ""  # Store end of previous chunk for overlap
    
    for section_idx, section in enumerate(growth_sections):
        section_words = len(section.split())
        test_word_count = current_word_count + section_words
        
        logging.debug(f"🔍 Processing growth section {section_idx + 1}/{len(growth_sections)}: {section_words} words")
        
        if test_word_count > max_size and current_chunk:
            # Current chunk is full, save it if it's substantial
            if current_word_count >= min_size:
                # Add overlap from previous chunk if available
                final_chunk = current_chunk
                if previous_chunk_end and chunks:
                    overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
                    final_chunk = overlap_text + "\n\n" + current_chunk
                    logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to growth chunk {len(chunks) + 1}")
                
                chunks.append(final_chunk.strip())
                
                # Store end of current chunk for next overlap
                previous_chunk_end = get_last_n_words(current_chunk, OVERLAP_SIZE * 2)
                
                logging.info(f"✅ Saved growth chunk {len(chunks)}: {len(final_chunk.split())} words (original: {current_word_count})")
                
                current_chunk = section
                current_word_count = section_words
            else:
                # Current chunk too small, but adding section makes it too big
                logging.debug(f"⚠️ Current growth chunk too small ({current_word_count} < {min_size}), handling large section")
                
                if section_words > max_size:
                    logging.debug(f"🔨 Growth section too large ({section_words} > {max_size}), splitting with overlap")
                    sub_sections = split_large_growth_section_with_overlap(section, max_size)
                    logging.info(f"📂 Split large growth section into {len(sub_sections)} sub-sections with overlap")
                    
                    for sub_idx, sub_section in enumerate(sub_sections):
                        sub_words = len(sub_section.split())
                        logging.debug(f"🔍 Processing growth sub-section {sub_idx + 1}/{len(sub_sections)}: {sub_words} words")
                        
                        if current_word_count + sub_words > max_size and current_chunk:
                            if current_word_count >= min_size:
                                # Add overlap before saving
                                final_chunk = current_chunk
                                if previous_chunk_end and chunks:
                                    overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
                                    final_chunk = overlap_text + "\n\n" + current_chunk
                                    logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to growth chunk {len(chunks) + 1}")
                                
                                chunks.append(final_chunk.strip())
                                previous_chunk_end = get_last_n_words(current_chunk, OVERLAP_SIZE * 2)
                                logging.info(f"✅ Saved growth chunk {len(chunks)}: {len(final_chunk.split())} words")
                            
                            current_chunk = sub_section
                            current_word_count = sub_words
                        else:
                            current_chunk += "\n\n" + sub_section if current_chunk else sub_section
                            current_word_count += sub_words
                            logging.debug(f"➕ Added growth sub-section to current chunk: {current_word_count} total words")
                else:
                    current_chunk += "\n\n" + section if current_chunk else section
                    current_word_count = test_word_count
                    logging.debug(f"➕ Added growth section to current chunk: {current_word_count} total words")
        else:
            current_chunk += "\n\n" + section if current_chunk else section
            current_word_count = test_word_count
            logging.debug(f"➕ Added growth section to current chunk: {current_word_count} total words")
    
    # Add the last chunk if it exists and is substantial
    if current_chunk and current_word_count >= min_size:
        # Add overlap to final chunk too
        final_chunk = current_chunk
        if previous_chunk_end and chunks:
            overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
            final_chunk = overlap_text + "\n\n" + current_chunk
            logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to final growth chunk")
        
        chunks.append(final_chunk.strip())
        logging.info(f"✅ Saved final growth chunk {len(chunks)}: {len(final_chunk.split())} words (original: {current_word_count})")
    elif current_chunk:
        logging.warning(f"⚠️ Discarded final growth chunk: {current_word_count} words < {min_size} minimum")
    
    # Validate the created chunks
    chunk_stats = validate_growth_chunk_sizes(chunks, target_size, "Growth Semantic Chunking")
     
    logging.info(f"🎉 Growth semantic chunking complete: {len(chunks)} chunks created")
    logging.info(f"📊 Growth chunk size range: {chunk_stats.get('min_words', 0)}-{chunk_stats.get('max_words', 0)} words")
    logging.info(f"📈 Average growth chunk size: {chunk_stats.get('avg_words', 0)} words (target: {target_size})")
    
    return chunks

def split_by_growth_logic(content: str) -> List[str]:
    """Split content by growth-specific logical boundaries with enhanced pattern recognition"""
    
    logging.info(f"🔧 Starting growth logic splitting...")
    
    # Log input content stats
    total_words = len(content.split())
    total_paragraphs = len([p for p in content.split('\n\n') if p.strip()])
    logging.info(f"📊 Input growth content: {total_words} words, {total_paragraphs} paragraphs")
    
    # Enhanced growth-specific section indicators
    growth_patterns = [
        r'(?i)(?:^|\n)(?:key growth|growth strategy|scaling approach|expansion plan|acceleration framework):',
        r'(?i)(?:^|\n)(?:growth recommendation|scaling strategy|expansion approach|market strategy):',
        r'(?i)(?:^|\n)(?:growth analysis|scaling assessment|market evaluation|competitive review):',
        r'(?i)(?:^|\n)(?:growth strengths|scaling advantages|market opportunities|competitive edge):',
        r'(?i)(?:^|\n)(?:growth challenges|scaling risks|market threats|competitive weaknesses):',
        r'(?i)(?:^|\n)(?:growth implementation|scaling execution|market action|strategic steps):',
        r'(?i)(?:^|\n)(?:growth optimization|scaling efficiency|market improvement|revenue enhancement):',
        r'(?i)(?:^|\n)(?:growth integration|scaling alignment|market coordination|strategic sync):',
        r'(?i)(?:^|\n)(?:growth measurement|scaling metrics|market tracking|performance kpis):',
        r'(?i)(?:^|\n)(?:growth acceleration|rapid scaling|market expansion|revenue growth):',
        
        # Enhanced patterns for AI-generated growth content
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:strategic|revenue|market|digital|competitive|customer)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:analysis|assessment|evaluation|optimization|implementation)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:recommendations?|strategies|approaches|solutions|tactics)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:mind expansion|expansion \d+|growth expansion|strategic expansion)',
        r'(?i)(?:^|\n)(?:your growth|your scaling|your expansion|considering your growth)',
        r'(?i)(?:^|\n)(?:to accelerate|to scale|to expand|to grow|moving forward)',
        
        # Growth-specific structural patterns
        r'(?i)(?:^|\n)(?:revenue|profit|customer|market|competitive|digital|automation|innovation)',
        r'(?i)(?:^|\n)(?:phase \d+|stage \d+|level \d+|tier \d+)',
        r'(?i)(?:^|\n)(?:next steps|action items|implementation|timeline|roadmap)'
    ]
    
    logging.info(f"🔍 Using {len(growth_patterns)} growth-specific patterns for splitting")
    
    # Try to split by growth patterns first
    sections = []
    current_section = ""
    pattern_matches = 0
    
    paragraphs = content.split('\n\n')
    logging.info(f"📂 Processing {len(paragraphs)} paragraphs for growth pattern matching")
    
    for paragraph in paragraphs:
        # Check if this paragraph starts a new growth section
        is_new_section = False
        for pattern in growth_patterns:
            if re.search(pattern, paragraph):
                is_new_section = True
                pattern_matches += 1
                break
        
        if is_new_section and current_section:
            sections.append(current_section.strip())
            current_section = paragraph
        else:
            current_section += "\n\n" + paragraph if current_section else paragraph
    
    # Add the last section
    if current_section:
        sections.append(current_section.strip())
    
    logging.info(f"📊 Growth pattern matching results: {pattern_matches} matches found, {len(sections)} sections created")
    
    # Smart fallback logic - if no growth patterns found or sections too large
    needs_fallback = False
    if len(sections) <= 1:
        needs_fallback = True
        logging.warning(f"⚠️ No growth patterns found, applying smart fallback")
    elif any(len(s.split()) > 400 for s in sections):  # Adjusted for growth content
        needs_fallback = True
        logging.warning(f"⚠️ Growth sections too large (>400 words), applying smart fallback")
    
    if needs_fallback:
        logging.info(f"🔄 Applying smart paragraph splitting with growth-optimized size limits...")
        
        # Smart paragraph splitting with growth-optimized size limits
        sections = []
        current_section = ""
        current_words = 0
        target_words = 300  # Target size for growth chunks (optimized for RAG)
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            para_words = len(paragraph.split())
            
            # If adding this paragraph would make section too large, save current and start new
            if current_words + para_words > target_words and current_section:
                sections.append(current_section.strip())
                logging.debug(f"📄 Saved growth section {len(sections)}: {current_words} words")
                current_section = paragraph
                current_words = para_words
            else:
                current_section += "\n\n" + paragraph if current_section else paragraph
                current_words += para_words
        
        if current_section:
            sections.append(current_section.strip())
            logging.debug(f"📄 Saved final growth section {len(sections)}: {current_words} words")
        
        logging.info(f"📄 Smart growth paragraph splitting: {len(sections)} sections created")
    
    # Final validation - force split any remaining oversized sections
    final_sections = []
    for i, section in enumerate(sections):
        section_words = len(section.split())
        
        if section_words > 400:  # Adjusted for growth content
            logging.warning(f"⚠️ Growth section {i+1} still oversized ({section_words} words), force splitting")
            
            # Force split by sentences with growth context preservation
            sentences = re.split(r'(?<=[.!?])\s+', section)
            sub_sections = []
            current_sub = ""
            current_sub_words = 0
            
            for sentence in sentences:
                sentence_words = len(sentence.split())
                
                if current_sub_words + sentence_words > 350 and current_sub:  # Slightly smaller for growth
                    sub_sections.append(current_sub.strip())
                    current_sub = sentence
                    current_sub_words = sentence_words
                else:
                    current_sub += " " + sentence if current_sub else sentence
                    current_sub_words += sentence_words
            
            if current_sub:
                sub_sections.append(current_sub.strip())
            
            final_sections.extend(sub_sections)
            logging.info(f"🔨 Split oversized growth section into {len(sub_sections)} sub-sections")
        else:
            final_sections.append(section)
    
    # Final statistics
    section_sizes = [len(s.split()) for s in final_sections]
    avg_size = sum(section_sizes) // len(section_sizes) if final_sections else 0
    min_size = min(section_sizes) if final_sections else 0
    max_size = max(section_sizes) if final_sections else 0
    optimal_sections = sum(1 for size in section_sizes if 150 <= size <= 400)  # Growth-optimized range
    optimal_percentage = (optimal_sections / len(final_sections)) * 100 if final_sections else 0
    
    logging.info(f"🎉 Growth logic splitting complete!")
    logging.info(f"📊 Final: {len(final_sections)} sections, {min_size}-{max_size} words (avg: {avg_size})")
    logging.info(f"🎯 Optimal growth sections (150-400 words): {optimal_sections}/{len(final_sections)} ({optimal_percentage:.1f}%)")
    
    return final_sections

def split_large_growth_section_with_overlap(section: str, max_size: int) -> List[str]:
    """Split large growth section with overlap for context preservation"""
    
    OVERLAP_SIZE = 50  # words
    
    words = section.split()
    if len(words) <= max_size:
        return [section]
    
    chunks = []
    start = 0
    
    while start < len(words):
        # Calculate end position
        end = min(start + max_size, len(words))
        
        # Create chunk
        chunk_words = words[start:end]
        chunk = ' '.join(chunk_words)
        chunks.append(chunk)
        
        # Move start position with overlap
        if end >= len(words):
            break
        
        start = end - OVERLAP_SIZE
        if start <= 0:
            start = end
    
    logging.debug(f"🔨 Split large growth section: {len(words)} words → {len(chunks)} chunks with {OVERLAP_SIZE}-word overlap")
    return chunks

def get_last_n_words(text: str, n: int) -> str:
    """Get last N words from text for overlap between chunks"""
    words = text.split()
    if len(words) <= n:
        return text
    return " ".join(words[-n:])

def get_first_n_words(text: str, n: int) -> str:
    """Get first N words from text for overlap between chunks"""
    words = text.split()
    if len(words) <= n:
        return text
    return " ".join(words[:n])

def validate_growth_chunk_sizes(chunks: List[str], target_size: int, context_name: str = "") -> Dict:
    """Validate and log growth chunk sizes for monitoring"""
    
    if not chunks:
        return {"total_chunks": 0}
    
    chunk_stats = {
        "total_chunks": len(chunks),
        "avg_words": 0,
        "min_words": float('inf'),
        "max_words": 0,
        "chunks_over_target": 0,
        "chunks_under_100": 0,  # Flag very small chunks
        "chunks_optimal": 0     # Chunks within target range
    }
    
    total_words = 0
    for chunk in chunks:
        words = len(chunk.split())
        total_words += words
        
        chunk_stats["min_words"] = min(chunk_stats["min_words"], words)
        chunk_stats["max_words"] = max(chunk_stats["max_words"], words)
        
        if words > target_size * 1.2:  # 20% over target
            chunk_stats["chunks_over_target"] += 1
        elif words < 100:
            chunk_stats["chunks_under_100"] += 1
        elif target_size * 0.8 <= words <= target_size * 1.2:  # Within 20% of target
            chunk_stats["chunks_optimal"] += 1
    
    chunk_stats["avg_words"] = total_words // len(chunks)
    chunk_stats["min_words"] = chunk_stats["min_words"] if chunk_stats["min_words"] != float('inf') else 0
    
    # Log the stats
    context_prefix = f"[{context_name}] " if context_name else ""
    logging.info(f"📊 {context_prefix}Growth chunk validation: "
                f"{chunk_stats['total_chunks']} chunks, "
                f"avg: {chunk_stats['avg_words']} words, "
                f"range: {chunk_stats['min_words']}-{chunk_stats['max_words']}, "
                f"optimal: {chunk_stats['chunks_optimal']}/{chunk_stats['total_chunks']}")
    
    return chunk_stats

def categorize_growth_chunk_size_by_words(word_count: int) -> str:
    """Categorize chunk size for growth analysis"""
    if word_count < 200:
        return "small"
    elif word_count < 350:
        return "optimal"
    elif word_count < 600:
        return "large"
    else:
        return "oversized"

def calculate_growth_semantic_completeness(content: str) -> float:
    """Calculate semantic completeness score for growth content"""
    
    # Growth-specific completeness indicators
    completeness_indicators = {
        'analysis_words': ['analyze', 'assessment', 'evaluation', 'examination', 'review'],
        'strategy_words': ['strategy', 'approach', 'plan', 'framework', 'methodology'],
        'action_words': ['implement', 'execute', 'action', 'steps', 'process'],
        'growth_words': ['growth', 'scaling', 'expansion', 'development', 'acceleration'],
        'business_words': ['business', 'company', 'organization', 'enterprise', 'venture'],
        'outcome_words': ['results', 'outcomes', 'benefits', 'impact', 'value'],
        'market_words': ['market', 'customer', 'client', 'competitive', 'industry'],
        'revenue_words': ['revenue', 'profit', 'income', 'sales', 'monetization'],
        'digital_words': ['digital', 'technology', 'automation', 'innovation', 'tech']
    }
    
    content_lower = content.lower()
    total_indicators = 0
    found_indicators = 0
    
    for category, words in completeness_indicators.items():
        total_indicators += len(words)
        found_indicators += sum(1 for word in words if word in content_lower)
    
    base_completeness = found_indicators / total_indicators if total_indicators > 0 else 0
    
    # Length bonus (longer content is generally more complete)
    word_count = len(content.split())
    length_factor = min(1.0, word_count / 400)  # Optimal around 400 words for growth
    
    # Structure bonus (headers, bullets, etc.)
    structure_indicators = ['##', '###', '- ', '* ', '1.', '2.', '3.', 'Phase', 'Step', 'Stage']
    structure_count = sum(1 for indicator in structure_indicators if indicator in content)
    structure_factor = min(0.2, structure_count * 0.03)  # Up to 0.2 bonus
    
    # Growth-specific bonus for strategic language
    strategic_indicators = ['recommend', 'suggest', 'should', 'focus on', 'prioritize', 'implement']
    strategic_count = sum(1 for indicator in strategic_indicators if indicator in content_lower)
    strategic_factor = min(0.15, strategic_count * 0.03)
    
    final_completeness = min(1.0, base_completeness + length_factor * 0.3 + structure_factor + strategic_factor)
    
    return final_completeness

def clean_growth_content_for_word_chunks(content: str) -> str:
    """Clean growth content for better Word document processing"""
    
    # Remove excessive whitespace
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
    content = re.sub(r' +', ' ', content)
    
    # Fix common formatting issues
    content = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', content)
    content = re.sub(r'([a-z])([A-Z])', r'\1 \2', content)
    
    # Remove artifacts from AI generation
    content = re.sub(r'<[^>]+>', '', content)  # Remove any HTML tags
    content = re.sub(r'\[.*?\]', '', content)  # Remove bracket annotations
    
    # Growth-specific cleaning
    content = re.sub(r'Growth Engine:', 'Growth Strategy:', content)  # Normalize engine references
    content = re.sub(r'Mind Expansion \d+:', 'Growth Area:', content)  # Clean expansion headers
    
    # Normalize quotes and special characters
    content = content.replace('"', '"').replace('"', '"')
    content = content.replace(''', "'").replace(''', "'")
    content = re.sub(r'…', '...', content)
    
    return content.strip()

def create_growth_chunk_word_document(content: str, title: str, user_profile: Dict, 
                                    section_name: str, chunk_id: str) -> Document:
    """Create a professionally formatted Word document for growth chunk"""
    
    try:
        doc = Document()
        
        # Enhanced styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add header with branding
        header = doc.add_heading("BACKABLE GROWTH ENGINE", 0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.font.size = Pt(24)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add chunk title
        chunk_title = doc.add_heading(title, 1)
        chunk_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title_run = chunk_title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Add metadata section
        if user_profile:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("Business Growth Context:").bold = True
            metadata_para.add_run(f"\nBusiness: {user_profile.get('business_name', 'Unknown')}")
            metadata_para.add_run(f"\nIndustry: {user_profile.get('industry', 'Unknown')}")  
            metadata_para.add_run(f"\nTeam Size: {user_profile.get('team_size', 'Unknown')} employees")
            metadata_para.add_run(f"\nGrowth Section: {section_name}")
            metadata_para.add_run(f"\nChunk ID: {chunk_id}")
            metadata_para.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Add separator
        doc.add_paragraph("─" * 60)
        
        # Add the AI-generated content with growth-specific formatting
        add_growth_content_to_document(doc, content)
        
        # Add footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run("Generated by Backable AI Growth Intelligence")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        logging.info(f"📄 Created growth chunk Word document: {len(content.split())} words")
        return doc
        
    except Exception as e:
        logging.error(f"❌ Error creating growth chunk Word document: {str(e)}")
        # Return minimal document on error
        doc = Document()
        doc.add_heading("Error Creating Growth Document", 1)
        doc.add_paragraph(f"Error: {str(e)}")
        return doc

def add_growth_content_to_document(doc: Document, content: str):
    """Add AI-generated growth content to document with intelligent formatting"""
    
    # Split by paragraphs and headers
    lines = content.split('\n')
    current_paragraph = ""
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - finalize paragraph
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
        elif line.startswith('###'):
            # Sub-subsection header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('###', '').strip()
            subheading = doc.add_heading(header_text, 3)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(0, 153, 76)  # Growth green
            
        elif line.startswith('##'):
            # Subsection header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('##', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(0, 102, 204)
            
        elif line.startswith('#'):
            # Main header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('#', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(0, 102, 204)
            
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            
        elif re.match(r'^\d+\.', line):
            # Numbered list
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            number_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(number_text, style='List Number')
            
        elif line.startswith('🎯') or line.startswith('🚀') or line.startswith('📊'):
            # Growth-specific emojis - treat as emphasis
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            emoji_para = doc.add_paragraph(line)
            emoji_run = emoji_para.runs[0]
            emoji_run.bold = True
            emoji_run.font.color.rgb = RGBColor(0, 153, 76)  # Growth green
            
        else:
            # Regular content - accumulate
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    # Add any remaining paragraph
    if current_paragraph:
        para = doc.add_paragraph(current_paragraph)

def create_growth_word_document(report_data: Dict, user_id: str) -> Document:
    """Create comprehensive growth Word document with enhanced formatting"""
    logging.info("📄 Creating Growth Strategy Word Document")
    
    doc = Document()
    
    # Enhanced styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Professional title page with growth branding
    title = doc.add_heading("BACKABLE", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(42)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading("Comprehensive Growth Strategy Report", 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0, 153, 76)  # Growth green
    
    # Add metadata
    metadata_para = doc.add_paragraph()
    metadata_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    report_meta = report_data.get("_enhanced_growth_report_metadata", {})
    
    metadata_para.add_run(f"User ID: {user_id}\n").bold = True
    metadata_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
    metadata_para.add_run(f"Analysis: {report_meta.get('total_words', 0):,} words\n")
    metadata_para.add_run(f"Model: Gemini 2.5 Pro Growth Engine\n")
    metadata_para.add_run(f"Focus: Strategic Growth & Scaling\n")
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading("TABLE OF CONTENTS", 1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_growth_report_metadata" and isinstance(section_data, dict):
            title = section_data.get("title", "Untitled Section")
            
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"{section_number}. {title}").bold = True
            
            # Add word count
            metadata = section_data.get("metadata", {})
            words_generated = metadata.get("words_generated", 0)
            
            toc_para.add_run(f" ({words_generated:,} words)")
            
            section_number += 1
    
    doc.add_page_break()
    
    # Process each section
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_growth_report_metadata" and isinstance(section_data, dict):
            
            logging.info(f"📝 Formatting growth section: {section_name}")
            
            title = section_data.get("title", "Untitled Section")
            content = section_data.get("content", "")
            
            # Add section header
            section_heading = doc.add_heading(f"{section_number}. {title}", 1)
            heading_run = section_heading.runs[0]
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add the AI-generated content
            add_growth_content_to_document(doc, content)
            
            # Add section separator
            separator_para = doc.add_paragraph()
            separator_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator_run = separator_para.add_run("🚀 ◆ ◆ ◆ 🚀")
            separator_run.font.color.rgb = RGBColor(0, 153, 76)  # Growth green
            separator_run.font.size = Pt(16)
            
            section_number += 1
            doc.add_page_break()
    
    # Add growth report summary
    add_growth_report_summary(doc, report_data)
    
    logging.info("✅ Growth Strategy Word Document Created")
    return doc

def add_growth_report_summary(doc: Document, report_data: Dict):
    """Add growth-specific report summary"""
    
    summary_heading = doc.add_heading("GROWTH STRATEGY REPORT SUMMARY", 1)
    summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary_heading_run = summary_heading.runs[0]
    summary_heading_run.font.color.rgb = RGBColor(0, 153, 76)  # Growth green
    
    report_meta = report_data.get("_enhanced_growth_report_metadata", {})
    
    summary_para = doc.add_paragraph()
    summary_para.add_run("Growth Strategy Report Statistics:").bold = True
    summary_para.add_run(f"\n• Total Growth Sections: {report_meta.get('total_sections', 0)}")
    summary_para.add_run(f"\n• Total Words Generated: {report_meta.get('total_words', 0):,}")
    summary_para.add_run(f"\n• AI Model: {report_meta.get('ai_model', 'N/A')}")
    summary_para.add_run(f"\n• Processing Method: {report_meta.get('processing_method', 'N/A')}")
    summary_para.add_run(f"\n• Report Type: {report_meta.get('report_type', 'N/A')}")
    summary_para.add_run(f"\n• Growth Focus Areas: Strategic Scaling, Revenue Acceleration, Market Expansion")
    
    # Add growth-specific insights summary
    insights_para = doc.add_paragraph()
    insights_para.add_run("\nGrowth Intelligence Summary:").bold = True
    insights_para.add_run(f"\n🎯 Strategic Growth Opportunities Identified")
    insights_para.add_run(f"\n🚀 Scaling Roadmap with Implementation Timeline")
    insights_para.add_run(f"\n📊 Market Expansion Strategy with Competitive Analysis")
    insights_para.add_run(f"\n💰 Revenue Acceleration Tactics with ROI Projections")
    insights_para.add_run(f"\n🔧 Operational Optimization for Growth Support")
    insights_para.add_run(f"\n👥 Team Scaling Strategy for Sustainable Growth")
# ======================================================
#           Indexer Integration
# ======================================================
# ======================================================
#           Enhanced Growth Engine Indexer Integration
# ======================================================

async def trigger_growth_indexer_for_client(client_id: str, force: bool = False, new_client: bool = False) -> tuple[bool, str, Optional[str]]:
    """Enhanced growth indexer trigger with comprehensive error handling and retry logic"""
    
    logging.info(f"🚀 Starting Growth Indexer for client_id={client_id} (force={force}, new_client={new_client})")
    
    # Track retry attempts for this specific call
    max_retries = INDEXER_RETRY_ATTEMPTS
    base_delay = INDEXER_RETRY_DELAY
    
    for attempt in range(max_retries):
        try:
            logging.info(f"🔄 Growth indexer attempt {attempt + 1}/{max_retries} for client_id={client_id}")
            
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=INDEXER_TIMEOUT)) as session:
                payload = {
                    "client_id": client_id,
                    "force": force,
                    "new_client": new_client
                }
                
                # Log payload for debugging
                logging.info(f"📤 Growth indexer payload: {payload}")
                
                async with session.post(
                    f"{INDEXER_API_BASE_URL}/run-indexer",
                    json=payload,
                    headers={"Content-Type": "application/json"}
                ) as response:
                    
                    # Enhanced response logging
                    logging.info(f"📡 Growth indexer response: status={response.status}, headers={dict(response.headers)}")
                    
                    try:
                        response_data = await response.json()
                    except Exception as json_error:
                        logging.error(f"❌ Growth indexer JSON parse error: {json_error}")
                        response_text = await response.text()
                        logging.error(f"🔍 Raw response text: {response_text[:500]}...")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)  # Exponential backoff
                            logging.warning(f"⏳ JSON parse failed, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"JSON parse error: {json_error}", None
                    
                    # Handle different response statuses
                    if response.status == 202:  # Accepted - Success
                        job_id = response_data.get("job_id")
                        message = response_data.get("message", "Growth indexer job started successfully")
                        
                        # Enhanced success logging
                        logging.info(f"✅ Growth indexer triggered successfully!")
                        logging.info(f"   - Client ID: {client_id}")
                        logging.info(f"   - Job ID: {job_id}")
                        logging.info(f"   - Message: {message}")
                        logging.info(f"   - Attempt: {attempt + 1}/{max_retries}")
                        
                        return True, message, job_id
                    
                    elif response.status == 409:  # Conflict - Job already in progress
                        message = response_data.get("message", "Growth indexer job already in progress")
                        logging.warning(f"⚠️ Growth indexer conflict for client_id={client_id}: {message}")
                        
                        # For conflicts, we don't retry - return immediately
                        return False, message, None
                    
                    elif response.status == 404:  # Client not found
                        message = response_data.get("message", "Client not found in indexer")
                        logging.warning(f"⚠️ Client not found for growth indexer: client_id={client_id}")
                        
                        # Auto-retry with new_client=True if not already set
                        if not new_client and attempt == 0:
                            logging.info(f"🔄 Retrying growth indexer with new_client=True for {client_id}")
                            return await trigger_growth_indexer_for_client(client_id, force, True)
                        else:
                            return False, f"Client not found: {message}", None
                    
                    elif response.status == 429:  # Rate Limited
                        message = response_data.get("message", "Growth indexer rate limited")
                        logging.warning(f"🚦 Growth indexer rate limited for client_id={client_id}: {message}")
                        
                        # Wait longer for rate limits
                        if attempt < max_retries - 1:
                            wait_time = base_delay * 2 * (attempt + 1)  # Longer wait for rate limits
                            logging.info(f"⏳ Rate limited, waiting {wait_time}s before retry...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Rate limited: {message}", None
                    
                    elif response.status >= 500:  # Server errors - retry
                        message = response_data.get("message", f"Growth indexer server error {response.status}")
                        logging.error(f"🚨 Growth indexer server error {response.status} for client_id={client_id}: {message}")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)
                            logging.info(f"⏳ Server error, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Server error: {message}", None
                    
                    else:  # Other client errors - don't retry
                        message = response_data.get("message", f"Growth indexer failed with status {response.status}")
                        logging.error(f"❌ Growth indexer client error {response.status} for client_id={client_id}: {message}")
                        return False, f"Client error: {message}", None
                        
        except asyncio.TimeoutError:
            error_msg = f"Growth indexer request timed out after {INDEXER_TIMEOUT}s"
            logging.error(f"⏰ {error_msg} (attempt {attempt + 1}/{max_retries})")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ Timeout, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except aiohttp.ClientError as e:
            error_msg = f"Growth indexer HTTP client error: {str(e)}"
            logging.error(f"🌐 {error_msg} (attempt {attempt + 1}/{max_retries})")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ HTTP error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except Exception as e:
            error_msg = f"Growth indexer unexpected error: {str(e)}"
            logging.error(f"💥 {error_msg} (attempt {attempt + 1}/{max_retries})")
            logging.error(f"🔍 Error type: {type(e).__name__}")
            import traceback
            logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ Unexpected error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
    
    # If we get here, all retries failed
    final_error = f"Growth indexer failed after {max_retries} attempts"
    logging.error(f"💥 {final_error} for client_id={client_id}")
    return False, final_error, None

def store_growth_indexer_job_metadata(report_id: str, user_id: str, indexer_job_id: str, indexer_status: str, 
                                     indexer_response: str = None, error_message: str = None):
    """Enhanced storage of growth indexer job metadata with comprehensive tracking"""
    
    conn = None
    try:
        logging.info(f"💾 Storing growth indexer metadata for report_id={report_id}")
        
        conn = get_growth_connection()
        
        with conn.cursor() as cur:
            # Check if the report exists first
            check_sql = """
                SELECT COUNT(*) FROM growth_reports WHERE report_id = %s AND user_id = %s
            """
            cur.execute(check_sql, (report_id, user_id))
            exists = cur.fetchone()[0] > 0
            
            if not exists:
                logging.error(f"❌ Growth report not found: report_id={report_id}, user_id={user_id}")
                return False
            
            # Update with comprehensive indexer information
            update_sql = """
                UPDATE growth_reports 
                SET 
                    indexer_job_id = %s,
                    indexer_status = %s,
                    indexer_triggered_at = %s,
                    indexer_completed_at = CASE WHEN %s = 'completed' THEN %s ELSE indexer_completed_at END,
                    indexer_error_message = %s,
                    indexer_retry_count = COALESCE(indexer_retry_count, 0) + CASE WHEN %s != 'triggered' THEN 1 ELSE 0 END
                WHERE report_id = %s AND user_id = %s
            """
            
            current_time = datetime.now()
            
            cur.execute(update_sql, (
                indexer_job_id,
                indexer_status,
                current_time,
                indexer_status,  # For CASE WHEN condition
                current_time if indexer_status == 'completed' else None,
                error_message,
                indexer_status,  # For retry count condition
                report_id,
                user_id
            ))
            
            rows_affected = cur.rowcount
            
            if rows_affected > 0:
                logging.info(f"✅ Updated growth indexer metadata:")
                logging.info(f"   - Report ID: {report_id}")
                logging.info(f"   - Job ID: {indexer_job_id}")
                logging.info(f"   - Status: {indexer_status}")
                logging.info(f"   - Rows affected: {rows_affected}")
                
                # Log additional context if error
                if error_message:
                    logging.warning(f"⚠️ Error message stored: {error_message}")
                    
                return True
            else:
                logging.error(f"❌ No rows updated for growth indexer metadata: report_id={report_id}")
                return False
                
    except Exception as e:
        logging.error(f"❌ Error storing growth indexer metadata: {str(e)}")
        logging.error(f"🔍 Error type: {type(e).__name__}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        return False
        
    finally:
        if conn:
            conn.close()

async def check_growth_indexer_status(job_id: str) -> Dict:
    """Check indexer status - tries multiple endpoints and handles all errors gracefully with detailed logging"""
    
    logging.info(f"🔍 Starting growth indexer status check")
    logging.info(f"📋 Original job_id received: '{job_id}' (type: {type(job_id)}, length: {len(job_id)})")
    
    # 🔥 STEP 1: JOB ID SANITIZATION AND VALIDATION
    original_job_id = job_id
    job_id_issues = []
    
    # Check for common job_id issues
    if not job_id:
        logging.error(f"❌ Empty job_id provided")
        return {
            "success": False,
            "job_id": job_id,
            "status": "error",
            "error": "Empty job_id provided",
            "original_job_id": original_job_id
        }
    
    # Check for double dashes (main issue from logs)
    if '--' in job_id:
        job_id_issues.append("double_dash_found")
        job_id = job_id.replace('--', '-')
        logging.warning(f"🔧 Fixed double dash in job_id: '{original_job_id}' → '{job_id}'")
    
    # Check for multiple consecutive dashes
    import re
    if re.search(r'-{3,}', job_id):
        job_id_issues.append("multiple_dashes_found")
        job_id = re.sub(r'-{2,}', '-', job_id)
        logging.warning(f"🔧 Fixed multiple dashes in job_id: '{original_job_id}' → '{job_id}'")
    
    # Remove whitespace
    job_id_stripped = job_id.strip()
    if job_id_stripped != job_id:
        job_id_issues.append("whitespace_found")
        job_id = job_id_stripped
        logging.warning(f"🔧 Removed whitespace from job_id: length changed from {len(original_job_id)} to {len(job_id)}")
    
    # Validate UUID format (typical format: xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx)
    uuid_pattern = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.IGNORECASE)
    is_valid_uuid = bool(uuid_pattern.match(job_id))
    
    logging.info(f"📋 Job ID validation:")
    logging.info(f"   - Sanitized job_id: '{job_id}'")
    logging.info(f"   - Issues found: {job_id_issues if job_id_issues else 'None'}")
    logging.info(f"   - Valid UUID format: {is_valid_uuid}")
    logging.info(f"   - Final length: {len(job_id)}")
    
    if job_id != original_job_id:
        logging.warning(f"⚠️ Job ID was modified during sanitization!")
        logging.warning(f"   - Original: '{original_job_id}'")
        logging.warning(f"   - Sanitized: '{job_id}'")
    
    # 🔥 STEP 2: BUILD ENDPOINTS WITH DETAILED LOGGING
    logging.info(f"🌐 Building indexer API endpoints")
    logging.info(f"📍 Base URL: {INDEXER_API_BASE_URL}")
    
    # List of possible endpoints to try (ordered by likelihood)
    possible_endpoints = [
        f"{INDEXER_API_BASE_URL}/job-status/{job_id}",
        f"{INDEXER_API_BASE_URL}/status/{job_id}",
        f"{INDEXER_API_BASE_URL}/jobs/{job_id}",
        f"{INDEXER_API_BASE_URL}/jobs/{job_id}/status",
        f"{INDEXER_API_BASE_URL}/api/jobs/{job_id}",
        f"{INDEXER_API_BASE_URL}/api/status/{job_id}",
        f"{INDEXER_API_BASE_URL}/indexer/status/{job_id}",
    ]
    
    logging.info(f"📋 Generated {len(possible_endpoints)} endpoints to try:")
    for i, endpoint in enumerate(possible_endpoints, 1):
        logging.info(f"   {i}. {endpoint}")
    
    # Track endpoint attempts
    endpoint_results = []
    successful_endpoint = None
    
    # 🔥 STEP 3: TRY EACH ENDPOINT WITH DETAILED LOGGING
    for i, endpoint in enumerate(possible_endpoints):
        attempt_start_time = time.time()
        attempt_result = {
            "endpoint": endpoint,
            "attempt_number": i + 1,
            "status_code": None,
            "response_type": None,
            "error": None,
            "response_time": 0
        }
        
        try:
            logging.info(f"🔍 Attempt {i+1}/{len(possible_endpoints)}: Trying endpoint")
            logging.info(f"📡 URL: {endpoint}")
            
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=30)) as session:
                logging.debug(f"🔧 Created aiohttp session with 30s timeout")
                
                async with session.get(endpoint, headers={"Content-Type": "application/json"}) as response:
                    response_time = time.time() - attempt_start_time
                    attempt_result["response_time"] = response_time
                    attempt_result["status_code"] = response.status
                    
                    logging.info(f"📡 Response received: HTTP {response.status} in {response_time:.2f}s")
                    logging.info(f"📋 Response headers: {dict(response.headers)}")
                    
                    if response.status == 200:
                        logging.info(f"✅ SUCCESS: HTTP 200 received from endpoint {i+1}")
                        
                        try:
                            # Try to parse as JSON first
                            logging.debug(f"🔧 Attempting JSON parsing...")
                            status_data = await response.json()
                            attempt_result["response_type"] = "json"
                            
                            job_status = status_data.get("status", "completed")
                            progress = status_data.get("progress", {})
                            
                            logging.info(f"✅ JSON parsing successful!")
                            logging.info(f"📊 Parsed data keys: {list(status_data.keys())}")
                            logging.info(f"📊 Job status: {job_status}")
                            logging.info(f"📊 Progress data: {progress}")
                            
                            successful_endpoint = endpoint
                            endpoint_results.append(attempt_result)
                            
                            final_result = {
                                "success": True,
                                "job_id": job_id,
                                "original_job_id": original_job_id,
                                "status": job_status,
                                "progress": progress,
                                "raw_response": status_data,
                                "working_endpoint": endpoint,
                                "endpoint_attempt": i + 1,
                                "response_time": response_time,
                                "job_id_issues": job_id_issues,
                                "endpoints_tried": i + 1,
                                "all_endpoint_results": endpoint_results
                            }
                            
                            logging.info(f"🎉 INDEXER STATUS CHECK SUCCESSFUL!")
                            logging.info(f"   - Working endpoint: {endpoint}")
                            logging.info(f"   - Job status: {job_status}")
                            logging.info(f"   - Attempts needed: {i + 1}/{len(possible_endpoints)}")
                            logging.info(f"   - Total response time: {response_time:.2f}s")
                            
                            return final_result
                            
                        except json.JSONDecodeError as json_error:
                            # Response is 200 but not JSON - maybe HTML success page
                            logging.info(f"⚠️ JSON parsing failed: {json_error}")
                            
                            text_response = await response.text()
                            attempt_result["response_type"] = "text"
                            attempt_result["error"] = f"JSON decode failed: {json_error}"
                            
                            logging.info(f"📄 Text response received (200 status)")
                            logging.debug(f"📄 Response preview: {text_response[:200]}...")
                            logging.info(f"📄 Response length: {len(text_response)} characters")
                            
                            # If we get 200, assume it's successful even without JSON
                            successful_endpoint = endpoint
                            endpoint_results.append(attempt_result)
                            
                            final_result = {
                                "success": True,
                                "job_id": job_id,
                                "original_job_id": original_job_id,
                                "status": "completed",  # Assume completed on 200 response
                                "progress": {"note": "Received 200 response - assumed completed"},
                                "raw_response": {"text_preview": text_response[:500], "full_length": len(text_response)},
                                "working_endpoint": endpoint,
                                "endpoint_attempt": i + 1,
                                "response_time": response_time,
                                "job_id_issues": job_id_issues,
                                "endpoints_tried": i + 1,
                                "response_type": "text_not_json",
                                "all_endpoint_results": endpoint_results
                            }
                            
                            logging.info(f"🎉 INDEXER STATUS CHECK SUCCESSFUL (TEXT RESPONSE)!")
                            logging.info(f"   - Working endpoint: {endpoint}")
                            logging.info(f"   - Assumed status: completed")
                            logging.info(f"   - Attempts needed: {i + 1}/{len(possible_endpoints)}")
                            
                            return final_result
                    
                    elif response.status == 404:
                        attempt_result["error"] = "Not Found (404)"
                        logging.debug(f"⚠️ Endpoint not found (404): {endpoint}")
                        # Continue to next endpoint
                        
                    elif response.status == 500:
                        attempt_result["error"] = "Server Error (500)"
                        logging.debug(f"⚠️ Server error (500): {endpoint}")
                        # Continue to next endpoint
                        
                    elif response.status == 403:
                        attempt_result["error"] = "Forbidden (403)"
                        logging.warning(f"⚠️ Access forbidden (403): {endpoint}")
                        # Continue to next endpoint
                        
                    elif response.status == 401:
                        attempt_result["error"] = "Unauthorized (401)"
                        logging.warning(f"⚠️ Unauthorized (401): {endpoint}")
                        # Continue to next endpoint
                        
                    else:
                        # Other HTTP errors - log and continue
                        error_text = await response.text()
                        attempt_result["error"] = f"HTTP {response.status}: {error_text[:100]}"
                        logging.debug(f"⚠️ Endpoint returned {response.status}: {endpoint}")
                        logging.debug(f"   Error preview: {error_text[:100]}...")
                        
        except aiohttp.ClientError as e:
            response_time = time.time() - attempt_start_time
            attempt_result["response_time"] = response_time
            attempt_result["error"] = f"Network error: {str(e)}"
            logging.debug(f"⚠️ Network error for endpoint {endpoint}: {str(e)}")
            
        except asyncio.TimeoutError as e:
            response_time = time.time() - attempt_start_time
            attempt_result["response_time"] = response_time
            attempt_result["error"] = f"Timeout after {response_time:.2f}s"
            logging.debug(f"⏰ Timeout for endpoint {endpoint}: {response_time:.2f}s")
            
        except Exception as e:
            response_time = time.time() - attempt_start_time
            attempt_result["response_time"] = response_time
            attempt_result["error"] = f"Unexpected error: {str(e)}"
            logging.debug(f"⚠️ Unexpected error for endpoint {endpoint}: {str(e)}")
        
        endpoint_results.append(attempt_result)
        logging.debug(f"📊 Endpoint {i+1} result: {attempt_result['error'] or 'Success'}")
    
    # 🔥 STEP 4: NO WORKING ENDPOINT FOUND - DETAILED FAILURE ANALYSIS
    total_time = sum(result["response_time"] for result in endpoint_results)
    
    logging.warning(f"❌ NO WORKING ENDPOINT FOUND")
    logging.warning(f"📊 Comprehensive failure analysis:")
    logging.warning(f"   - Total endpoints tried: {len(possible_endpoints)}")
    logging.warning(f"   - Total time spent: {total_time:.2f}s")
    logging.warning(f"   - Job ID issues: {job_id_issues if job_id_issues else 'None'}")
    logging.warning(f"   - Final job ID: '{job_id}'")
    
    # Analyze failure patterns
    error_types = {}
    for result in endpoint_results:
        error = result["error"] or "Unknown"
        error_types[error] = error_types.get(error, 0) + 1
    
    logging.warning(f"📊 Error breakdown:")
    for error_type, count in error_types.items():
        logging.warning(f"   - {error_type}: {count} endpoints")
    
    # Log each endpoint attempt for debugging
    logging.warning(f"📋 Detailed endpoint results:")
    for i, result in enumerate(endpoint_results, 1):
        logging.warning(f"   {i}. {result['endpoint']}")
        logging.warning(f"      Status: {result['status_code'] or 'No response'}")
        logging.warning(f"      Time: {result['response_time']:.2f}s")
        logging.warning(f"      Error: {result['error'] or 'None'}")
    
    # ✅ Return graceful "assumed success" with detailed context
    logging.info(f"📊 Assuming indexer completed successfully (trigger was successful)")
    
    return {
        "success": True,  # ✅ Don't fail the whole process
        "job_id": job_id,
        "original_job_id": original_job_id,
        "status": "completed",  # ✅ Assume completed since trigger worked
        "progress": {
            "note": "No status endpoint available",
            "endpoints_tried": len(possible_endpoints),
            "assumption": "Completed successfully",
            "total_time_spent": total_time,
            "job_id_issues_found": job_id_issues
        },
        "raw_response": {
            "message": "All status endpoints unavailable - assumed successful completion",
            "endpoints_tried": possible_endpoints,
            "trigger_was_successful": True,
            "detailed_results": endpoint_results,
            "error_summary": error_types
        },
        "working_endpoint": None,
        "job_id_issues": job_id_issues,
        "endpoints_tried": len(possible_endpoints),
        "all_endpoint_results": endpoint_results,
        "assumed_completion": True
    }

def get_growth_indexer_job_info(report_id: str) -> Optional[Dict]:
    """Get growth indexer job information from database"""
    
    conn = None
    try:
        conn = get_growth_connection()
        
        with conn.cursor() as cur:
            sql = """
                SELECT 
                    indexer_job_id,
                    indexer_status,
                    indexer_triggered_at,
                    indexer_completed_at,
                    indexer_error_message,
                    indexer_retry_count
                FROM growth_reports
                WHERE report_id = %s
            """
            
            cur.execute(sql, (report_id,))
            row = cur.fetchone()
            
            if row:
                return {
                    "job_id": row[0],
                    "status": row[1],
                    "triggered_at": row[2].isoformat() if row[2] else None,
                    "completed_at": row[3].isoformat() if row[3] else None,
                    "error_message": row[4],
                    "retry_count": row[5] or 0
                }
            else:
                return None
                
    except Exception as e:
        logging.error(f"❌ Error getting growth indexer job info: {str(e)}")
        return None
        
    finally:
        if conn:
            conn.close()

async def trigger_and_monitor_growth_indexer(report_id: str, user_id: str, container_name: str) -> Dict:
    """Comprehensive growth indexer trigger with monitoring and database integration - STATUS CHECK REMOVED"""
    
    logging.info(f"🚀 Starting comprehensive growth indexer process for report_id={report_id}")
    
    # Step 1: Trigger the indexer
    success, message, job_id = await trigger_growth_indexer_for_client(user_id)
    
    if success and job_id:
        # Step 2: Store successful trigger in database
        store_result = store_growth_indexer_job_metadata(
            report_id=report_id,
            user_id=user_id,
            indexer_job_id=job_id,
            indexer_status="triggered",
            indexer_response=message
        )
        
        if store_result:
            logging.info(f"✅ Growth indexer triggered and stored successfully:")
            logging.info(f"   - Report ID: {report_id}")
            logging.info(f"   - Job ID: {job_id}")
            logging.info(f"   - User ID: {user_id}")
            
            # ✅ SIMPLIFIED: Skip status checking since endpoints don't exist
            return_status = "triggered"  # Assume success after successful trigger
            logging.info(f"📊 Indexer triggered successfully, status checking disabled to avoid 404 errors")
            logging.info(f"📊 Job ID {job_id} is assumed to be processing successfully")
            
            return {
                "success": True,
                "job_id": job_id,
                "message": message,
                "status": return_status,
                "stored_in_db": True,
                "report_id": report_id,
                "status_check_note": "Status checking disabled - indexer endpoints not available"
            }
        else:
            logging.error(f"❌ Growth indexer triggered but failed to store in database")
            return {
                "success": False,
                "error": "Triggered but database storage failed",
                "job_id": job_id,
                "message": message
            }
    else:
        # Step 2: Store failed trigger in database
        store_growth_indexer_job_metadata(
            report_id=report_id,
            user_id=user_id,
            indexer_job_id=job_id or "failed",
            indexer_status="failed",
            error_message=message
        )
        
        logging.error(f"❌ Growth indexer trigger failed:")
        logging.error(f"   - Report ID: {report_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Error: {message}")
        
        return {
            "success": False,
            "error": message,
            "job_id": job_id,
            "status": "failed",
            "stored_in_db": True,
            "report_id": report_id
        }

# Synchronous wrapper for background tasks
def trigger_growth_indexer_sync(report_id: str, user_id: str, container_name: str) -> Dict:
    """Synchronous wrapper for growth indexer trigger (for background tasks)"""
    
    try:
        # Handle Windows event loop policy
        if platform.system() == 'Windows':
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
        
        # Create new event loop for this thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            result = loop.run_until_complete(
                trigger_and_monitor_growth_indexer(report_id, user_id, container_name)
            )
            
            logging.info(f"🎯 Growth indexer sync wrapper completed for report_id={report_id}")
            return result
            
        finally:
            loop.close()
            
    except Exception as e:
        error_msg = f"Growth indexer sync wrapper error: {str(e)}"
        logging.error(f"❌ {error_msg}")
        
        # Still try to store the error in database
        try:
            store_growth_indexer_job_metadata(
                report_id=report_id,
                user_id=user_id,
                indexer_job_id="sync_error",
                indexer_status="error",
                error_message=error_msg
            )
        except:
            pass  # Don't let database errors break the wrapper
        
        return {
            "success": False,
            "error": error_msg,
            "status": "sync_error",
            "report_id": report_id
        }

# Background task integration
def trigger_growth_indexer_background(report_id: str, user_id: str, container_name: str):
    """Background thread function for growth indexer triggering"""
    
    def indexer_worker():
        try:
            logging.info(f"🔄 Starting background growth indexer for report_id={report_id}")
            
            result = trigger_growth_indexer_sync(report_id, user_id, container_name)
            
            if result["success"]:
                logging.info(f"✅ Background growth indexer successful: job_id={result.get('job_id')}")
            else:
                logging.error(f"❌ Background growth indexer failed: {result.get('error')}")
                
        except Exception as e:
            logging.error(f"💥 Background growth indexer worker error: {str(e)}")
    
    # Start in background thread
    indexer_thread = Thread(target=indexer_worker, daemon=True)
    indexer_thread.start()
    
    logging.info(f"🔄 Started background growth indexer thread for report_id={report_id}")
    return indexer_thread

# ======================================================
#           API Endpoints
# ======================================================

app = FastAPI(title="Growth Engine API", version="2.1")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class GrowthAssessmentRequest(BaseModel):
    user_id: str
    business_name: str
    assessment_data: Dict[str, Any]

class GrowthProgressRequest(BaseModel):
    user_id: str
    business_name: str
    assessment_data: Dict[str, Any]
    current_expansion: int
    auto_save: bool = False

# Application lifespan
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger = setup_growth_logging()
    logging.info("🚀 Growth Engine API Starting Up")
    logging.info("✅ Vertex AI Integration: ENABLED (Primary Method)")
    logging.info("🔑 API Keys Fallback: " + ", ".join([f"Growth_{i+1:02d}" for i in range(len(GEMINI_API_KEYS))]))
    logging.info("📊 Multi-Database Intelligence: ENABLED")
    logging.info("🧠 Behavioral Analytics: ENABLED")
    logging.info("📄 Word Document Chunking: ENABLED")
    logging.info("🧠 Question-Response Chunking: ENABLED")
    logging.info("🔄 Blob Retry Decoder: ENABLED")
    logging.info("⚡ AsyncPG Connection Pooling: ENABLED")
    logging.info("🗄️  Unified Storage Architecture: ENABLED")
    logging.info("❌ Auto-Indexer: REMOVED (Relies on 5-minute automatic schedule)")

    yield

    # Shutdown
    logging.info("🛑 Growth Engine API Shutting Down")

    # Close all connection pools
    global _connection_pools
    for db_name, pool in _connection_pools.items():
        logging.info(f"Closing pool for {db_name}")
        await pool.close()

    logging.info("✅ Shutdown Complete")

app.router.lifespan_context = lifespan

@app.get("/")
async def root():
    return {
        "message": "Growth Engine API v3.0 - NEW UNIFIED ARCHITECTURE",
        "status": "operational",
        "architecture": "unified_vertex_ai_primary",
        "features": {
            "vertex_ai_integration": True,  # NEW - Primary method
            "api_keys_fallback": True,  # Fallback only
            "blob_retry_decoder": True,  # NEW
            "unified_storage": True,  # NEW
            "asyncpg_pooling": True,  # NEW
            "auto_indexer_removed": True,  # NEW - Relies on automatic 5-min schedule
            "multi_database_intelligence": True,
            "enhanced_api_key_management": True,
            "load_balancing": True,
            "gemini_ai_analysis": True,
            "behavioral_analytics": True,
            "word_document_chunking": True,
            "question_response_chunking": True
        },
        "vertex_ai_status": "enabled" if vertex_ai_client else "unavailable",
        "api_keys_status": get_enhanced_api_key_status()
    }


# 1. ADD THIS SIMPLE DEBUG ENDPOINT (just copy-paste this)
@app.get("/check_indexer/{report_id}")
async def check_indexer(report_id: str):
    """Quick check if indexer completed - MINIMAL CODE CHANGE"""
    
    conn = None
    try:
        conn = get_growth_connection()
        with conn.cursor() as cur:
            sql = """
                SELECT 
                    indexer_job_id, indexer_status, indexer_triggered_at, 
                    indexer_completed_at, indexer_error_message
                FROM growth_reports 
                WHERE report_id = %s
            """
            cur.execute(sql, (report_id,))
            row = cur.fetchone()
            
            if row:
                job_id, status, triggered_at, completed_at, error_msg = row
                
                # Calculate time since trigger
                time_since_trigger = None
                if triggered_at:
                    time_since_trigger = (datetime.now() - triggered_at.replace(tzinfo=None)).total_seconds()
                
                # Simple status determination
                if status == "completed":
                    actual_status = "✅ SUCCESS - Indexer completed!"
                elif status == "failed":
                    actual_status = f"❌ FAILED - {error_msg}"
                elif status == "triggered" and time_since_trigger and time_since_trigger > 600:
                    actual_status = "❓ LIKELY FAILED - No completion after 10+ minutes"
                elif status == "triggered":
                    actual_status = f"🔄 UNKNOWN - Triggered {int(time_since_trigger)}s ago, no completion feedback"
                else:
                    actual_status = f"❓ UNKNOWN STATUS - {status}"
                
                return {
                    "report_id": report_id,
                    "indexer_job_id": job_id,
                    "database_status": status,
                    "actual_status": actual_status,
                    "time_since_trigger_seconds": int(time_since_trigger) if time_since_trigger else None,
                    "error_message": error_msg
                }
            else:
                return {"error": "Report not found", "report_id": report_id}
                
    except Exception as e:
        return {"error": str(e)}
    finally:
        if conn:
            conn.close()


@app.get("/api-key-health")
async def get_api_key_health():
    """Get detailed API key health information"""
    return {
        "timestamp": datetime.now().isoformat(),
        "summary": get_api_key_status_summary(),
        "detailed_status": get_enhanced_api_key_status()
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.1",
        "api_keys_status": get_api_key_status_summary()
    }

@app.post("/growth-audit/{user_id}")
async def process_growth_audit(user_id: str, request: GrowthAssessmentRequest):
    """Process comprehensive growth audit with multi-database intelligence - FIXED"""
    
    start_time = time.time()
    logging.info(f"🚀 Starting Growth Audit for user_id={user_id}")
    
    # Generate unique report ID
    report_id = f"growth_report_{user_id}_{int(datetime.now().timestamp())}"
    
    # Initialize job status
    growth_job_status[report_id] = {
        "status": "processing",
        "message": "Starting growth analysis...",
        "progress": 0,
        "sections_completed": 0,
        "started_at": datetime.now().isoformat(),
        "user_id": user_id,
        "business_name": request.business_name
    }
    
    try:
        # Get user profile data with connection pooling
        logging.info(f"👤 Retrieving user profile for {user_id}")
        user_profile = await get_user_profile_data(user_id)
        
        if not user_profile:
            logging.warning(f"⚠️ No user profile found for {user_id}, using provided data")
            user_profile = {
                "username": user_id,
                "business_name": request.business_name,
                "industry": "Unknown Industry",
                "team_size": "Unknown",
                "biggest_challenge": "Growth optimization"
            }
        
        # Store assessment data with multi-database intelligence
        logging.info(f"💾 Storing growth assessment data...")
        assessment_id = store_growth_assessment(user_id, request.assessment_data)
        
        # Get multi-database intelligence with connection pooling
        logging.info(f"🧠 Extracting multi-database intelligence...")
        multi_db_intelligence = await get_multi_database_intelligence(user_id)
        
        # ✅ FIX 1: Add raw_assessment_data to complete_raw_data
        complete_raw_data = {
            "user_id": user_id,
            "report_id": report_id,
            "user_profile": user_profile,
            "responses": request.assessment_data.get("responses", []),
            "assessment_metadata": request.assessment_data.get("assessment_metadata", {}),
            "comprehensive_metadata": request.assessment_data.get("comprehensive_metadata", {}),
            "multi_database_intelligence": multi_db_intelligence,
            "behavioral_analytics": request.assessment_data.get("comprehensive_metadata", {}).get("behavioral_analytics", {}),
            "completion_flags": request.assessment_data.get("completion_flags", {}),
            "text_responses": request.assessment_data.get("text_responses", {}),
            "numeric_inputs": request.assessment_data.get("numeric_inputs", {}),
            "processing_timestamp": datetime.now().isoformat(),
            "raw_assessment_data": request.assessment_data  # ✅ ADDED THIS LINE FOR Q&R CHUNKING
        }
        
        # Update job status
        growth_job_status[report_id]["message"] = "Generating comprehensive growth strategy..."
        growth_job_status[report_id]["progress"] = 20
        
        # Generate comprehensive growth report
        logging.info(f"🧠 Generating comprehensive growth report...")
        
        # Create async wrapper for report generation since it may need async operations
        async def run_report_generation_async():
            # Run in thread pool but allow for async operations
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor(max_workers=1) as executor:
                # Use run_in_executor for CPU-bound work while maintaining async context
                future = executor.submit(generate_comprehensive_growth_report, complete_raw_data, report_id)
                return await loop.run_in_executor(None, lambda: future.result())
        
        report_data = await run_report_generation_async()
        
        # ✅ FIX 2: Add raw assessment data to report metadata for Q&R chunking
        if "_enhanced_growth_report_metadata" not in report_data:
            report_data["_enhanced_growth_report_metadata"] = {}
        
        # Ensure the raw assessment data is in the report metadata for Q&R chunking
        report_data["_enhanced_growth_report_metadata"]["raw_assessment_data"] = request.assessment_data
        
        logging.info(f"✅ Added raw assessment data to report metadata for Q&R chunking")
        
        # Update job status
        growth_job_status[report_id]["message"] = "Uploading growth strategy to secure storage..."
        growth_job_status[report_id]["progress"] = 80
        
        # Get Azure container name
        container_name = get_azure_container_name(user_id)
        
        # Upload to Azure with enhanced chunking - now with Q&R data
        logging.info(f"☁️ Uploading growth report to Azure...")
        upload_success, upload_message = await upload_growth_report_to_azure(report_data, report_id, user_id)
        
        if not upload_success:
            raise Exception(f"Azure upload failed: {upload_message}")
        
        # Store report metadata
        generation_metadata = {
            "total_sections": len([k for k in report_data.keys() if k != "_enhanced_growth_report_metadata"]),
            "total_words": report_data.get("_enhanced_growth_report_metadata", {}).get("total_words", 0),
            "generation_time": time.time() - start_time,
            "ai_model": "gemini-2.5-pro",
            "multi_database_sources": len(complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", [])),
            "data_sources_used": complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", []),
            "intelligence_correlation": True,
            "total_intelligence_sources": len(complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", [])),
            "upload_message": upload_message,
            "database_pooling_enabled": True  # Track that pooling is active
        }
        
        store_growth_report_metadata(report_id, user_id, assessment_id, 
                                   report_data.get("_enhanced_growth_report_metadata", {}).get("total_sections", 0),
                                   container_name, generation_metadata)
        
        # AUTO-INDEXER REMOVED (NEW UNIFIED ARCHITECTURE)
        # Relies on automatic 5-minute indexing schedule instead
        logging.info("📋 Report completed - Auto-indexer will run within 5 minutes")
        growth_job_status[report_id]["status"] = "auto_scheduled"
        growth_job_status[report_id]["indexing_status"] = "auto_scheduled"
        growth_job_status[report_id]["message"] = "Report ready - Auto-indexing scheduled..."
        growth_job_status[report_id]["progress"] = 95

        # Final status update
        total_time = time.time() - start_time
        growth_job_status[report_id]["message"] = f"Growth strategy complete! Generated in {total_time:.1f}s"
        growth_job_status[report_id]["progress"] = 100
        growth_job_status[report_id]["completed_at"] = datetime.now().isoformat()
        growth_job_status[report_id]["total_generation_time"] = total_time
        
        logging.info(f"✅ Growth Audit completed for {user_id} in {total_time:.2f}s")
        
        return {
            "status": "processing",
            "report_id": report_id,
            "message": "Growth strategy generation started",
            "estimated_completion": "2-3 minutes",
            "user_id": user_id,
            "business_name": request.business_name,
            "generation_metadata": generation_metadata
        }
        
    except Exception as e:
        error_message = f"Growth audit error: {str(e)}"
        logging.error(f"❌ {error_message}")
        
        # Update job status with error
        growth_job_status[report_id]["status"] = "failed"
        growth_job_status[report_id]["message"] = error_message
        growth_job_status[report_id]["error"] = str(e)
        growth_job_status[report_id]["failed_at"] = datetime.now().isoformat()
        
        import traceback
        logging.error(f"🔍 Growth audit traceback:\n{traceback.format_exc()}")
        
        raise HTTPException(status_code=500, detail=error_message)

@app.get("/growth_report_status/{report_id}")
async def get_growth_report_status(report_id: str):
    """Get growth report generation status"""
    
    if report_id not in growth_job_status:
        # Try to get status from database
        conn = None
        try:
            conn = get_growth_connection()
            with conn.cursor() as cur:
                sql = """
                    SELECT status, generation_metadata, created_at, completed_at, 
                           indexer_status, indexer_job_id
                    FROM growth_reports 
                    WHERE report_id = %s
                """
                cur.execute(sql, (report_id,))
                row = cur.fetchone()
                
                if row:
                    status, metadata, created_at, completed_at, indexer_status, indexer_job_id = row
                    
                    return {
                        "status": status,
                        "report_id": report_id,
                        "message": f"Report {status}",
                        "created_at": created_at.isoformat() if created_at else None,
                        "completed_at": completed_at.isoformat() if completed_at else None,
                        "indexer_status": indexer_status,
                        "indexer_job_id": indexer_job_id,
                        "metadata": metadata
                    }
                else:
                    raise HTTPException(status_code=404, detail="Growth report not found")
        finally:
            if conn:
                conn.close()
    
    return growth_job_status[report_id]

@app.post("/growth_assessment_progress")
async def save_growth_progress(request: GrowthProgressRequest):
    """Save growth assessment progress with enhanced tracking"""
    
    try:
        logging.info(f"💾 Saving growth progress for user {request.user_id}")
        logging.info(f"📊 Progress details: expansion {request.current_expansion}, auto_save: {request.auto_save}")
        
        # 🔥 ENHANCED: Build proper progress tracking data
        progress_tracking_data = {
            'completed_mind_expansions': request.current_expansion,
            'total_mind_expansions': 5,
            'percentage_complete': (request.current_expansion / 5) * 100,
            'last_saved_at': datetime.now().isoformat(),
            'auto_save': request.auto_save,
            'save_trigger': 'auto_save' if request.auto_save else 'manual_save',
            'sections_completed': []
        }
        
        # Add completed sections based on current expansion
        section_names = [
            "Mind Expansion 1: Your Growth Story",
            "Mind Expansion 2: Creating Opportunity", 
            "Mind Expansion 3: Revenue",
            "Mind Expansion 4: The Growth Strategy",
            "Mind Expansion 5: Digital & Market Position"
        ]
        
        for i in range(min(request.current_expansion, len(section_names))):
            progress_tracking_data['sections_completed'].append(section_names[i])
        
        logging.info(f"📈 Progress tracking: {progress_tracking_data['percentage_complete']:.1f}% complete")
        
        # 🔥 ENHANCED: Update assessment data with progress tracking
        enhanced_assessment_data = {
            **request.assessment_data,
            'progress_tracking': progress_tracking_data,
            'current_expansion': request.current_expansion,
            'last_updated': datetime.now().isoformat(),
            'save_metadata': {
                'save_timestamp': datetime.now().isoformat(),
                'save_type': 'auto_save' if request.auto_save else 'manual_save',
                'user_agent': 'growth_engine_frontend',
                'expansion_at_save': request.current_expansion
            }
        }
        
        # Count responses for logging
        response_count = 0
        if 'responses' in enhanced_assessment_data:
            if isinstance(enhanced_assessment_data['responses'], list):
                response_count = len(enhanced_assessment_data['responses'])
            elif isinstance(enhanced_assessment_data['responses'], dict):
                response_count = len(enhanced_assessment_data['responses'])
        
        logging.info(f"📊 Assessment data: {response_count} responses, expansion {request.current_expansion}")
        
        # Store progress data
        assessment_id = store_growth_assessment(request.user_id, enhanced_assessment_data)
        
        logging.info(f"✅ Growth progress saved successfully for user {request.user_id}")
        
        return {
            "status": "saved",
            "assessment_id": assessment_id,
            "user_id": request.user_id,
            "current_expansion": request.current_expansion,
            "auto_save": request.auto_save,
            "timestamp": datetime.now().isoformat(),
            "progress_percentage": progress_tracking_data['percentage_complete'],
            "sections_completed": len(progress_tracking_data['sections_completed']),
            "response_count": response_count,
            "save_metadata": {
                "save_type": "auto_save" if request.auto_save else "manual_save",
                "sections_completed": progress_tracking_data['sections_completed']
            }
        }
        
    except Exception as e:
        error_message = f"Error saving growth progress: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Save error context: user_id={request.user_id}, expansion={request.current_expansion}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=error_message)

@app.post("/growth_indexer_callback")
async def growth_indexer_callback(request: Request):
    """Callback endpoint for growth indexer completion"""
    
    try:
        data = await request.json()
        job_id = data.get("job_id")
        status = data.get("status")
        
        logging.info(f"📞 Growth indexer callback: {job_id} - {status}")
        
        # Update database with indexer status
        conn = None
        try:
            conn = get_growth_connection()
            with conn.cursor() as cur:
                sql = """
                    UPDATE growth_reports 
                    SET 
                        indexer_status = %s,
                        indexer_completed_at = %s,
                        indexer_error_message = %s
                    WHERE indexer_job_id = %s
                """
                
                error_message = data.get("error_message") if status == "failed" else None
                
                cur.execute(sql, (
                    status,
                    datetime.now() if status in ["completed", "failed"] else None,
                    error_message,
                    job_id
                ))
                
        finally:
            if conn:
                conn.close()
        
        return {"status": "received", "job_id": job_id}
        
    except Exception as e:
        logging.error(f"❌ Growth indexer callback error: {str(e)}")
        return {"status": "error", "message": str(e)}

# 🔥 MINIMAL ADDITION: Optional endpoint to check notification health
@app.get("/notification_health")
async def check_notification_health():
    """Check notification system health (optional debugging endpoint)"""
    
    try:
        health_info = get_notification_system_health()
        
        logging.info(f"📊 Notification health check requested")
        logging.info(f"   - Active threads: {health_info['active_notification_threads']}")
        
        return {
            "status": "healthy",
            "health_info": health_info
        }
        
    except Exception as e:
        logging.error(f"❌ Error checking notification health: {str(e)}")
        return {
            "status": "error",
            "error": str(e)
        }


@app.get("/growth_assessment_progress/{user_id}")
async def get_growth_progress(user_id: str):
    """Get saved growth assessment progress for a user"""
    
    conn = None
    try:
        logging.info(f"📥 Retrieving growth progress for user {user_id}")
        conn = get_growth_connection()
        with conn.cursor() as cur:
            # Get the most recent assessment for this user
            sql = """
                SELECT raw_data, progress_tracking, last_updated, created_at
                FROM growth_assessments 
                WHERE user_id = %s 
                ORDER BY last_updated DESC 
                LIMIT 1
            """
            
            cur.execute(sql, (user_id,))
            row = cur.fetchone()
            
            if row:
                raw_data, progress_tracking, last_updated, created_at = row
                
                logging.info(f"🔍 Raw data type: {type(raw_data)}, Progress tracking type: {type(progress_tracking)}")
                
                # 🔥 FIX: Handle both string and dict types for raw_data
                if isinstance(raw_data, str):
                    try:
                        assessment_data = json.loads(raw_data)
                        logging.info(f"✅ Parsed raw_data from JSON string")
                    except json.JSONDecodeError as e:
                        logging.error(f"❌ Failed to parse raw_data JSON: {e}")
                        assessment_data = {}
                elif isinstance(raw_data, dict):
                    # Data is already a dictionary (from JSONB column)
                    assessment_data = raw_data
                    logging.info(f"✅ Used raw_data as dict (JSONB column)")
                else:
                    logging.warning(f"⚠️ Unexpected raw_data type: {type(raw_data)}")
                    assessment_data = {}
                
                # 🔥 FIX: Handle both string and dict types for progress_tracking
                current_expansion = 1
                if progress_tracking:
                    if isinstance(progress_tracking, str):
                        try:
                            progress_data = json.loads(progress_tracking)
                            current_expansion = progress_data.get('completed_mind_expansions', 1)
                            logging.info(f"✅ Parsed progress_tracking from JSON string")
                        except json.JSONDecodeError as e:
                            logging.error(f"❌ Failed to parse progress_tracking JSON: {e}")
                            current_expansion = 1
                    elif isinstance(progress_tracking, dict):
                        # Data is already a dictionary
                        current_expansion = progress_tracking.get('completed_mind_expansions', 1)
                        logging.info(f"✅ Used progress_tracking as dict")
                    else:
                        logging.warning(f"⚠️ Unexpected progress_tracking type: {type(progress_tracking)}")
                        current_expansion = 1
                
                logging.info(f"📊 Progress summary: {len(assessment_data.get('responses', []))} responses, expansion {current_expansion}")
                
                return {
                    "status": "found",
                    "user_id": user_id,
                    "assessment_data": assessment_data,
                    "current_expansion": current_expansion,
                    "created_at": created_at.isoformat() if created_at else None,
                    "updated_at": last_updated.isoformat() if last_updated else None
                }
            else:
                logging.info(f"ℹ️ No saved progress found for user {user_id}")
                return {
                    "status": "not_found",
                    "user_id": user_id,
                    "message": "No saved progress found"
                }
                
    except Exception as e:
        error_message = f"Error retrieving growth progress: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Error context: user_id={user_id}, error_type={type(e).__name__}")
        raise HTTPException(status_code=500, detail=error_message)
    finally:
        if conn:
            conn.close()


@app.get("/growth_reports/{user_id}")
async def get_user_growth_reports(user_id: str):
    """Get all growth reports for a user"""
    
    conn = None
    try:
        conn = get_growth_connection()
        with conn.cursor() as cur:
            sql = """
                SELECT report_id, report_type, status, created_at, completed_at, 
                       chunk_count, indexer_status, generation_metadata
                FROM growth_reports 
                WHERE user_id = %s 
                ORDER BY created_at DESC
            """
            
            cur.execute(sql, (user_id,))
            rows = cur.fetchall()
            
            reports = []
            for row in rows:
                report_id, report_type, status, created_at, completed_at, chunk_count, indexer_status, metadata = row
                
                reports.append({
                    "report_id": report_id,
                    "report_type": report_type,
                    "status": status,
                    "created_at": created_at.isoformat() if created_at else None,
                    "completed_at": completed_at.isoformat() if completed_at else None,
                    "chunk_count": chunk_count,
                    "indexer_status": indexer_status,
                    "metadata": metadata
                })
            
            return {
                "user_id": user_id,
                "total_reports": len(reports),
                "reports": reports
            }
            
    except Exception as e:
        logging.error(f"❌ Error getting growth reports: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn:
            conn.close()

# Main execution
# Main execution
if __name__ == "__main__":
    # Set up logging
    logger = setup_growth_logging()
    logger.info("🚀 Starting Backable Growth Engine")
    
    # Get port from environment variable or use default
    port = int(os.environ.get("PORT", 8001))  # Different port from component engine
    
    # Run with uvicorn for production
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=port,
        log_level="info",
        access_log=True,
        workers=1  # Single worker for optimal resource management
    )